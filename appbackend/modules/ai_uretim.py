"""AI içerik üretimi modülü (/ai/soru-uret, /ai/hikaye*, /ai/materyal/*, /ai/mini-oyun*, /ai/scaffold/*, /ai/post-reading, /ai/icerik-kalite-skoru, /ai/kelime-listesi, /ai/okuma-parcalari).

server.py'dan BİREBİR taşındı; yollar ve davranış değişmedi. AI çağrıları
core.ai üzerinden yapılır; modül DB/auth/ayar erişimini yalnızca core'dan alır.
"""
import os
import io
import re
import json
import base64
import uuid
import asyncio
import logging
import random
import math
from datetime import datetime, timedelta, timezone
from typing import Optional, List, Dict, Any

import httpx
from fastapi import APIRouter, Depends, HTTPException, Request, Body, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field

from core.db import db, prepare_for_mongo, parse_from_mongo
from core.auth import get_current_user, require_role, UserRole
from core.config import (
    GEMINI_API_KEY, GEMINI_API_KEY_2, GEMINI_API_KEY_3, GEMINI_MODELS,
    AI_MODEL, AI_DEFAULT_MODEL, AI_HAIKU_MODEL, AI_MAX_DAILY_REQUESTS,
    AI_CACHE_HOURS, YANDEX_DISK_TOKEN, ANTHROPIC_API_KEY,
)
from core.sistem import (
    get_xp_tablosu, get_puan_ayarlari, get_lig_esikleri,
    XP_TABLOSU_DEFAULT, LIG_ESIKLERI_DEFAULT, LIG_SIRA,
)
from core.ai import _gemini_call, call_claude, _mock_bilgi_tabani_response, get_ogrenci_ai_verileri

router = APIRouter()


AI_SORU_SYSTEM_PROMPT = """Sen Türkçe dil eğitimi uzmanısın. Bloom taksonomisini ve MEB müfredatını biliyorsun.
Verilen metinden çoktan seçmeli sorular üreteceksin. Her soru 4 şıklı olacak.
FORMAT: SADECE JSON array ver: [{"soru":"...", "secenekler":["A","B","C","D"], "dogru_cevap":0, "taksonomi":"bilgi|kavrama|uygulama|analiz|sentez|degerlendirme"}]"""

AI_HIKAYE_SYSTEM_PROMPT = """Sen çocuk kitabı yazarısın. MEB Türkçe müfredatını biliyorsun.
Verilen sınıf seviyesi, tema ve hedef kelimeleri kullanarak kısa bir okuma parçası yazacaksın.
Hedef kelimelerin TÜMÜNÜ metnin içinde doğal şekilde kullan.
FORMAT: SADECE JSON ver: {"baslik":"...", "metin":"...", "kelime_sayisi":0, "kullanilan_kelimeler":[], "sorular":[5 Bloom sorusu]}"""


@router.post("/ai/soru-uret")
async def ai_soru_uret(payload: dict, current_user=Depends(get_current_user)):
    """Metin + sınıf → Bloom taksonomili 5-10 soru üretimi."""
    metin = payload.get("metin", "")
    sinif = payload.get("sinif", 4)
    soru_sayisi = payload.get("soru_sayisi", 5)

    if len(metin) < 50:
        raise HTTPException(status_code=400, detail="Metin en az 50 karakter olmalı")

    user_msg = f"""Sınıf: {sinif}
Soru sayısı: {soru_sayisi}
Her Bloom basamağından en az 1 soru olsun (bilgi, kavrama, uygulama, analiz, sentez, degerlendirme).

METİN:
{metin[:3000]}

SADECE JSON array döndür: [{{"soru":"...", "secenekler":["A","B","C","D"], "dogru_cevap":0, "taksonomi":"bilgi"}}]"""

    result = await call_claude(AI_SORU_SYSTEM_PROMPT, user_msg, model="sonnet", max_tokens=2000)

    if result.get("error"):
        raise HTTPException(status_code=503, detail=result["error"])

    sorular = result.get("parsed") or []
    if isinstance(sorular, dict):
        sorular = sorular.get("sorular", [])

    return {"sorular": sorular, "token": result.get("tokens", 0), "maliyet": result.get("maliyet", 0)}


@router.post("/ai/hikaye-uret")
async def ai_hikaye_uret(payload: dict, current_user=Depends(get_current_user)):
    """Sınıf + tema + hedef kelimeler → kişisel hikâye + 5 Bloom sorusu."""
    sinif = payload.get("sinif", 3)
    tema = payload.get("tema", "Doğa ve Evren")
    kelimeler = payload.get("kelimeler", [])
    kelime_sayisi = payload.get("kelime_sayisi", 150)

    user_msg = f"""Sınıf: {sinif}. sınıf
Tema: {tema}
Hedef kelimeler: {', '.join(kelimeler) if kelimeler else 'pusula, keşif, macera, mevsim, göç'}
Kelime sayısı: ~{kelime_sayisi}
Cümle uzunluğu: max {8 + sinif} kelime

SADECE JSON döndür:
{{"baslik":"...", "metin":"...", "kelime_sayisi":0, "kullanilan_kelimeler":[], "sorular":[{{"soru":"...", "secenekler":["A","B","C","D"], "dogru_cevap":0, "taksonomi":"bilgi"}}]}}"""

    result = await call_claude(AI_HIKAYE_SYSTEM_PROMPT, user_msg, model="sonnet", max_tokens=2500)

    if result.get("error"):
        raise HTTPException(status_code=503, detail=result["error"])

    return {"hikaye": result.get("parsed") or result.get("text", ""), "token": result.get("tokens", 0), "maliyet": result.get("maliyet", 0)}


@router.get("/ai/kelime-listesi")
async def ai_kelime_listesi(sinif: int = 0, current_user=Depends(get_current_user)):
    """Kelime haritasındaki tüm kelimeleri listeler."""
    filtre = {}
    if sinif > 0:
        filtre["sinif"] = sinif
    kelimeler = await db.meb_kelime_haritasi.find(filtre).sort("kelime", 1).to_list(length=500)
    for k in kelimeler:
        k.pop("_id", None)
    return kelimeler


@router.get("/ai/okuma-parcalari")
async def ai_okuma_parcalari_listesi(current_user=Depends(get_current_user)):
    """Tüm AI okuma parçalarını listeler."""
    parcalar = await db.ai_okuma_parcalari.find({}).sort("tarih", -1).to_list(length=100)
    for p in parcalar:
        p.pop("_id", None)
    return parcalar


@router.post("/ai/mini-oyun")
async def ai_mini_oyun(payload: dict, current_user=Depends(get_current_user)):
    """Kelimelerden mini oyun üretir (kelime avı, eşleştirme, boşluk doldurma)."""
    oyun_turu = payload.get("tur", "eslestirme")  # eslestirme, kelime_avi, bosluk_doldurma, cumle_kurma
    kelimeler = payload.get("kelimeler", [])  # [{"kelime": "...", "anlam": "..."}]
    sinif = payload.get("sinif", 3)

    if not kelimeler:
        # Rastgele kelime al
        rastgele = await db.meb_kelime_haritasi.find({"sinif": sinif}).to_list(length=10)
        kelimeler = [{"kelime": k.get("kelime", ""), "anlam": k.get("anlam", ""), "ornek_cumle": k.get("ornek_cumle", "")} for k in rastgele]

    if not kelimeler:
        return {"oyun": None, "mesaj": "Kelime bulunamadı"}

    if oyun_turu == "eslestirme":
        # Kelime-anlam eşleştirme (Claude gerektirmez)
        karisik_kelimeler = kelimeler[:8]
        random.shuffle(karisik_kelimeler)
        karisik_anlamlar = [k.get("anlam", "") for k in karisik_kelimeler]
        random.shuffle(karisik_anlamlar)
        return {"oyun": {
            "tur": "eslestirme",
            "baslik": "🎲 Kelime Eşleştirme",
            "aciklama": "Her kelimeyi doğru anlamıyla eşleştir!",
            "kelimeler": [k.get("kelime", "") for k in karisik_kelimeler],
            "anlamlar": [k.get("anlam", "") for k in karisik_kelimeler],  # doğru sıralama (frontend karıştıracak)
            "xp": 5,
        }}

    elif oyun_turu == "bosluk_doldurma":
        # Cümledeki kelimeyi bul
        sorular = []
        for k in kelimeler[:6]:
            cumle = k.get("ornek_cumle", "")
            if cumle and k.get("kelime", ""):
                bos = cumle.replace(k["kelime"], "___").replace(k["kelime"].capitalize(), "___")
                if "___" in bos:
                    sorular.append({"cumle_bos": bos, "dogru": k["kelime"], "secenekler": []})
        # Şıklar ekle
        tum_kelimeler = [k.get("kelime", "") for k in kelimeler]
        for s in sorular:
            yanlis = [w for w in tum_kelimeler if w != s["dogru"]][:3]
            secenekler = [s["dogru"]] + yanlis
            random.shuffle(secenekler)
            s["secenekler"] = secenekler
        return {"oyun": {
            "tur": "bosluk_doldurma",
            "baslik": "⬜ Boşluk Doldur",
            "aciklama": "Cümledeki boşluğa uygun kelimeyi bul!",
            "sorular": sorular,
            "xp": 5,
        }}

    elif oyun_turu == "kelime_avi":
        # Kelime avı grid üretimi (AI)
        prompt = f"""Şu kelimelerden 8x8 harf gridi oluştur (kelime avı oyunu):
Kelimeler: {', '.join([k.get('kelime','') for k in kelimeler[:6]])}
SADECE JSON: {{"grid": [["A","B",...], ...], "kelimeler": ["kelime1", ...], "yonler": ["sağa","aşağı",...] }}"""

        result = await call_claude("Sen kelime oyunu tasarımcısısın.", prompt, model="haiku", max_tokens=500)
        if result.get("parsed"):
            return {"oyun": {"tur": "kelime_avi", "baslik": "🔍 Kelime Avı", "aciklama": "Gizli kelimeleri bul!", **result["parsed"], "xp": 7}}
        else:
            return {"oyun": {"tur": "kelime_avi", "baslik": "🔍 Kelime Avı", "aciklama": "Kelimeleri bul!", "kelimeler": [k.get("kelime","") for k in kelimeler[:6]], "xp": 7}}

    elif oyun_turu == "cumle_kurma":
        sorular = []
        for k in kelimeler[:5]:
            cumle = k.get("ornek_cumle", "")
            if cumle:
                kelime_listesi = cumle.split()
                karisik = kelime_listesi.copy()
                random.shuffle(karisik)
                sorular.append({"karisik": karisik, "dogru": kelime_listesi, "hedef_kelime": k.get("kelime", "")})
        return {"oyun": {
            "tur": "cumle_kurma",
            "baslik": "📝 Cümle Kurma",
            "aciklama": "Karışık kelimeleri doğru sıraya diz!",
            "sorular": sorular,
            "xp": 5,
        }}

    return {"oyun": None, "mesaj": "Bilinmeyen oyun türü"}


@router.post("/ai/mini-oyun/tamamla")
async def ai_mini_oyun_tamamla(payload: dict, current_user=Depends(get_current_user)):
    """Mini oyun tamamlama — XP ver."""
    oyun_turu = payload.get("tur", "")
    dogru_sayisi = payload.get("dogru", 0)
    toplam = payload.get("toplam", 1)
    basari = round(dogru_sayisi / max(toplam, 1) * 100)

    xp = 3 if basari < 50 else 5 if basari < 80 else 7 if basari < 100 else 10

    try:
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"toplam_xp": xp}})
    except:
        pass

    return {"xp": xp, "basari": basari, "mesaj": "Harika!" if basari >= 80 else "İyi gidiyor!" if basari >= 50 else "Tekrar deneyelim!"}


@router.post("/ai/scaffold/olustur")
async def scaffold_olustur(req: Request, current_user=Depends(get_current_user)):
    """Seçilen kitap için DNA'ya göre 3 zorluk seviyesi üretir."""
    data = await req.json()
    kitap_adi = data.get("kitap_adi", "")
    kitap_id  = data.get("kitap_id", "")
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    sinif = data.get("sinif", 3)

    # DNA bak
    dna = await db.okuma_dna.find_one({"ogrenci_id": ogrenci_id})
    seviye_skoru = 5  # varsayılan orta
    if dna:
        b = dna.get("boyutlar", {})
        seviye_skoru = round((b.get("anlama_derinligi", 50) + b.get("kelime_gucu", 50) + b.get("akicilik", 50)) / 30)

    # Cache
    cache = await db.scaffold_cache.find_one({"kitap_id": kitap_id, "ogrenci_id": ogrenci_id})
    if cache:
        # Cache'deki metin yeterince uzunsa kullan
        orta_metin = cache.get("seviyeler", {}).get("orta", {}).get("metin", "")
        if len(orta_metin.split()) >= 80:
            cache.pop("_id", None)
            return cache
        else:
            # Kısa metin cache'i — sil ve yeniden üret
            await db.scaffold_cache.delete_one({"kitap_id": kitap_id, "ogrenci_id": ogrenci_id})
            logging.info(f"[SCAFFOLD] Kısa cache silindi, yeniden üretilecek")

    prompt = f"""Sen bir çocuk edebiyatı uzmanısın. "{kitap_adi}" kitabı için {sinif}. sınıf öğrencisine uygun 3 seviyeli scaffold okuma materyali oluştur.

Öğrencinin DNA seviyesi: {seviye_skoru}/10
Önerilen seviye: {"Kolay" if seviye_skoru <= 3 else "Orta" if seviye_skoru <= 6 else "Orijinal"}

🔴 ZORUNLU KURALLAR:
- Her seviye için EN AZ 200 kelimelik, tercihen 250-300 kelimelik bir metin yaz
- Metinler kitabın gerçek karakterlerini, mekanlarını ve olaylarını içermeli
- Kısa, tek cümlelik özetler KABUL EDİLMEZ
- Her seviye tam bir sahne veya bölüm gibi okunabilir olmalı

1. KOLAY (200-250 kelime): Çok basit ve kısa cümleler. Günlük hayatta kullanılan kelimeler. Olaylar net bir sırayla anlatılır. Zor kelimeler yerine basit alternatifleri kullanılır.

2. ORTA (220-270 kelime): Orta uzunlukta cümleler. Bazı edebi ifadeler ve mecazlar var. Dolaylı anlatım ve diyaloglar kullanılır. Kelime dağarcığı biraz genişletilir.

3. ORİJİNAL (250-300 kelime): Kitabın gerçek yazarının üslubuna yakın. Zengin dil, soyut kavramlar, karmaşık cümle yapıları. Edebi sanatlar (benzetme, kişileştirme) kullanılır.

Yanıtı SADECE JSON olarak ver, başka hiçbir şey yazma:
{{
  "kitap_adi": "{kitap_adi}",
  "onerilen_seviye": "kolay|orta|orijinal",
  "seviyeler": {{
    "kolay": {{
      "baslik": "Kolay Versiyon",
      "metin": "buraya en az 200 kelimelik metin...",
      "kelime_sayisi": 220,
      "anahtar_kelimeler": ["kelime1", "kelime2", "kelime3"]
    }},
    "orta": {{
      "baslik": "Orta Versiyon",
      "metin": "buraya en az 220 kelimelik metin...",
      "kelime_sayisi": 250,
      "anahtar_kelimeler": ["kelime1", "kelime2", "kelime3"]
    }},
    "orijinal": {{
      "baslik": "Orijinal Üslup",
      "metin": "buraya en az 250 kelimelik metin...",
      "kelime_sayisi": 280,
      "anahtar_kelimeler": ["kelime1", "kelime2", "kelime3"]
    }}
  }},
  "zpd_aciklama": "Öğrencinin neden bu seviyede başlaması gerektiğinin kısa açıklaması"
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=3500)
        raw = raw.strip()
        import json as _json, re as _re
        # ``` temizle
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        # { } içini al
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        # Trailing comma düzelt
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
        # Metin uzunluğu kontrolü — çok kısaysa yeniden üret
        for sev in ["kolay","orta","orijinal"]:
            metin = result.get("seviyeler",{}).get(sev,{}).get("metin","")
            if len(metin.split()) < 80:
                logging.warning(f"[SCAFFOLD] {sev} metni çok kısa ({len(metin.split())} kelime), yeniden denenecek")
                raise Exception(f"{sev} metni çok kısa")
    except Exception as e:
        logging.error(f"[SCAFFOLD] Hata: {e} — mock döndürülüyor")
        # Mock — kısa değil gerçek metin
        result = {
            "kitap_adi": kitap_adi,
            "onerilen_seviye": "orta" if seviye_skoru <= 6 else "orijinal",
            "seviyeler": {
                "kolay": {"baslik": "Kolay Versiyon", "metin": f"'{kitap_adi}' adlı kitapta harika bir hikâye anlatılıyor. Bu kitapta bir ana karakter var. O, çok cesur ve iyi kalpli biri. Her gün yeni maceralar yaşıyor. Bazen zorluklarla karşılaşıyor ama hiç pes etmiyor. Arkadaşları ona yardım ediyor. Birlikte çalışıyorlar. Sonunda her şey güzel bir şekilde bitiyor. Bu kitap bize çok önemli bir şey öğretiyor: İyi olmak, çalışmak ve arkadaşlarına güvenmek her zaman işe yarıyor. Okurken çok heyecanlandım. Sen de okursan çok beğeneceğini düşünüyorum. Her sayfada yeni bir şey öğreniyorsun.", "kelime_sayisi": 100, "anahtar_kelimeler": ["cesaret", "arkadaşlık", "macera"]},
                "orta": {"baslik": "Orta Versiyon", "metin": f"'{kitap_adi}', okuyucusunu büyüleyici bir yolculuğa davet eden güçlü bir eser. Ana karakter, hayatının en zor döneminde bile umudunu kaybetmiyor ve içindeki gücü keşfediyor. Yazarın kalemi, her sayfada bizi farklı duygularla buluşturuyor. Kimi zaman güldürüyor, kimi zaman düşündürüyor. Karakterlerin birbirleriyle kurduğu ilişkiler, dostluğun ve güvenin ne kadar değerli olduğunu gözler önüne seriyor. Olaylar hızlı bir tempoda gelişirken, her sahne okuyucuyu bir sonrakine çekiyor. Bu kitabı elinize aldığınızda bırakmak istemeyeceksiniz. İnsan ilişkileri, cesaret ve doğruluk temaları üzerine kurulu bu yapıt, her yaştan okura farklı mesajlar veriyor.", "kelime_sayisi": 120, "anahtar_kelimeler": ["umut", "dostluk", "keşif"]},
                "orijinal": {"baslik": "Orijinal Üslup", "metin": f"'{kitap_adi}' sayfaları arasında soluk soluğa ilerleyen bu anlatı, okuyucuyu gerçeklikle kurgunun bulanık sınırında bırakır. Yazar, kelimelerini özenle seçerek her cümlede anlam katmanları oluşturmuş; yüzeyin altında akan güçlü bir duygu seli, metni her okunuşta yeniden keşfettiriyor. Ana karakterin iç dünyasına yapılan bu derin yolculuk, aslında hepimizin yaşadığı evrensel sorgulamaların bir yansımasıdır. Toplumsal baskılar, bireysel özgürlük arayışı ve kimlik mücadelesi —tüm bunlar, yazarın ustalıklı kalemi aracılığıyla doğal bir akışla birbirine örülmüş. Eserin dili, bazen şiirsel imgelerle süslenirken bazen de sert gerçeklerin çıplak ifadesine bürünüyor. Bu kontrast, okuyucuyu hem zihinsel hem duygusal düzeyde zorluyor.", "kelime_sayisi": 130, "anahtar_kelimeler": ["kimlik", "özgürlük", "anlam"]}
            },
            "zpd_aciklama": f"DNA profiline göre (seviye {seviye_skoru}/10) {'kolay' if seviye_skoru <= 3 else 'orta' if seviye_skoru <= 6 else 'orijinal'} seviyeden başlamanı öneriyoruz."
        }

    result["ogrenci_id"] = ogrenci_id
    result["kitap_id"] = kitap_id
    result["dna_seviye"] = seviye_skoru
    result["tarih"] = datetime.utcnow().isoformat()
    await db.scaffold_cache.update_one({"kitap_id": kitap_id, "ogrenci_id": ogrenci_id}, {"$set": result}, upsert=True)
    return result


@router.post("/ai/scaffold/seviye-ilerleme")
async def scaffold_seviye_ilerleme(req: Request, current_user=Depends(get_current_user)):
    """Öğrenci okumayı tamamladı — bir üst seviyeye geç veya tebrik et."""
    data = await req.json()
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    mevcut_seviye = data.get("mevcut_seviye", "kolay")
    dogru_oran = data.get("dogru_oran", 0.7)

    siradaki = {"kolay": "orta", "orta": "orijinal", "orijinal": None}
    sonraki = siradaki.get(mevcut_seviye)

    if dogru_oran >= 0.7 and sonraki:
        mesaj = f"Harika! {'Orta' if sonraki == 'orta' else 'Orijinal'} seviyeye geçmeye hazırsın! 🎉"
        xp = 15
    elif dogru_oran >= 0.7 and not sonraki:
        mesaj = "Tebrikler! Kitabı tüm seviyelerde tamamladın! 🏆"
        xp = 30
    else:
        mesaj = "Biraz daha pratik yapman iyi olur. Aynı seviyeyi tekrar dene."
        xp = 5
        sonraki = mevcut_seviye

    await db.xp_logs.insert_one({"ogrenci_id": ogrenci_id, "xp": xp, "kaynak": "scaffold", "tarih": datetime.utcnow().isoformat()})
    return {"sonraki_seviye": sonraki, "mesaj": mesaj, "xp": xp, "ilerledi": dogru_oran >= 0.7}


@router.post("/ai/materyal/uret")
async def materyal_uret(req: Request, current_user=Depends(get_current_user)):
    """Kitap / metin için çalışma materyali üret (soru seti, kelime listesi, etkinlik)."""
    data = await req.json()
    kitap_adi  = data.get("kitap_adi", "")
    yazar      = data.get("yazar", "")
    metin_id   = data.get("metin_id", "")
    tur        = data.get("tur", "soru_seti")
    sinif      = data.get("sinif", 3)
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    icerik_id  = data.get("icerik_id", "")
    metin_icerigi = data.get("metin_icerigi", "")  # Frontend'den doğrudan gelen metin

    # ── Metin toplama: öncelik sırası ──────────────────────────────────────
    # 1. Frontend'den gönderilen metin (manuel yapıştırma)
    # 2. icerik_id → gelisim_icerik koleksiyonundan (okuma_metni veya dosya_b64)
    # 3. kitap_adi → ai_okuma_parcalari koleksiyonundan (parçaları birleştir)
    # 4. kitap_adi → ai_yuklemeler koleksiyonundan (AI yükleme metni)
    # 5. kitap_adi → ai_uretilen_sorular koleksiyonundan (soru metinleri ipucu)
    # 6. Hiçbiri yoksa: Gemini kitap adını bilerek kendi bilgisiyle üretir

    metin_ek = ""
    metin_kaynak = "yok"

    # 1. Frontend'den metin geldiyse kullan
    if metin_icerigi and metin_icerigi.strip():
        metin_ek = f"\n\nKİTAP/METİN İÇERİĞİ (kullanıcı girişi):\n{metin_icerigi.strip()[:6000]}"
        metin_kaynak = "frontend"
        logging.info(f"[MATERYAL] Kaynak: frontend, {len(metin_icerigi)} karakter")

    # 2. icerik_id varsa gelisim_icerik'ten çek
    if not metin_ek and icerik_id:
        try:
            icerik_doc = await db.gelisim_icerik.find_one({"id": icerik_id})
            if not icerik_doc:
                icerik_doc = await db.gelisim_icerik.find_one({"_id": icerik_id})
            if icerik_doc:
                if icerik_doc.get("okuma_metni"):
                    metin_ek = f"\n\nKİTAP METNİ:\n{icerik_doc['okuma_metni'][:6000]}"
                    metin_kaynak = "gelisim_icerik.okuma_metni"
                elif icerik_doc.get("dosya_b64"):
                    import base64, io
                    try:
                        raw_bytes = base64.b64decode(icerik_doc["dosya_b64"])
                        dosya_turu = icerik_doc.get("dosya_turu", "")
                        if dosya_turu == "pdf":
                            from pypdf import PdfReader
                            reader = PdfReader(io.BytesIO(raw_bytes))
                            metin = " ".join(p.extract_text() or "" for p in reader.pages[:15])
                            metin_ek = f"\n\nKİTAP METNİ (PDF):\n{metin[:6000]}"
                            metin_kaynak = "gelisim_icerik.pdf"
                        elif dosya_turu in ("docx", "doc"):
                            import docx as _docx
                            doc2 = _docx.Document(io.BytesIO(raw_bytes))
                            metin = "\n".join(p.text for p in doc2.paragraphs if p.text.strip())
                            metin_ek = f"\n\nKİTAP METNİ (Word):\n{metin[:6000]}"
                            metin_kaynak = "gelisim_icerik.docx"
                    except Exception as b64_e:
                        logging.error(f"[MATERYAL] Dosya decode hatası: {b64_e}")
            logging.info(f"[MATERYAL] icerik_id kaynağı: {metin_kaynak}")
        except Exception as db_e:
            logging.error(f"[MATERYAL] gelisim_icerik çekme hatası: {db_e}")

    # 3. kitap_adi'na göre ai_okuma_parcalari'ndan parçaları birleştir
    if not metin_ek and kitap_adi:
        try:
            parcalar = await db.ai_okuma_parcalari.find(
                {"kitap_adi": {"$regex": kitap_adi[:30], "$options": "i"}}
            ).sort("bolum", 1).to_list(length=20)
            if parcalar:
                parca_metinler = []
                for p in parcalar:
                    m = p.get("metin_kesit") or p.get("metin") or p.get("icerik") or ""
                    if m:
                        parca_metinler.append(m)
                if parca_metinler:
                    birlesik = "\n\n".join(parca_metinler)
                    metin_ek = f"\n\nKİTAP METNİ (okuma parçaları, {len(parcalar)} bölüm):\n{birlesik[:7000]}"
                    metin_kaynak = f"ai_okuma_parcalari ({len(parcalar)} parça)"
                    logging.info(f"[MATERYAL] {len(parcalar)} okuma parçası bulundu, toplam {len(birlesik)} karakter")
        except Exception as e:
            logging.error(f"[MATERYAL] ai_okuma_parcalari hatası: {e}")

    # 4. ai_yuklemeler koleksiyonundan metin ara
    if not metin_ek and kitap_adi:
        try:
            yukleme = await db.ai_yuklemeler.find_one(
                {"kitap_adi": {"$regex": kitap_adi[:30], "$options": "i"}}
            )
            if yukleme:
                m = yukleme.get("metin") or yukleme.get("icerik") or yukleme.get("ozet") or ""
                if m:
                    metin_ek = f"\n\nKİTAP METNİ (yükleme):\n{m[:6000]}"
                    metin_kaynak = "ai_yuklemeler"
                    logging.info(f"[MATERYAL] ai_yuklemeler'den metin bulundu: {len(m)} karakter")
                # dosya_b64 varsa oku
                elif yukleme.get("dosya_b64"):
                    import base64, io
                    try:
                        raw_bytes = base64.b64decode(yukleme["dosya_b64"])
                        dosya_turu = yukleme.get("dosya_turu", "")
                        if dosya_turu == "pdf":
                            from pypdf import PdfReader
                            reader = PdfReader(io.BytesIO(raw_bytes))
                            metin = " ".join(p.extract_text() or "" for p in reader.pages[:15])
                            if metin.strip():
                                metin_ek = f"\n\nKİTAP METNİ (yükleme PDF):\n{metin[:6000]}"
                                metin_kaynak = "ai_yuklemeler.pdf"
                                logging.info(f"[MATERYAL] Yükleme PDF okundu: {len(metin)} karakter")
                        elif dosya_turu in ("docx", "doc"):
                            import docx as _docx
                            doc3 = _docx.Document(io.BytesIO(raw_bytes))
                            metin = "\n".join(p.text for p in doc3.paragraphs if p.text.strip())
                            if metin.strip():
                                metin_ek = f"\n\nKİTAP METNİ (yükleme Word):\n{metin[:6000]}"
                                metin_kaynak = "ai_yuklemeler.docx"
                    except Exception as yk_e:
                        logging.error(f"[MATERYAL] Yükleme dosya hatası: {yk_e}")
        except Exception as e:
            logging.error(f"[MATERYAL] ai_yuklemeler hatası: {e}")

    # 5. ai_uretilen_sorular'dan ipucu topla (metin parçaları varsa)
    if not metin_ek and kitap_adi:
        try:
            sorular_db = await db.ai_uretilen_sorular.find(
                {"kitap_adi": {"$regex": kitap_adi[:30], "$options": "i"}}
            ).to_list(length=30)
            if sorular_db:
                soru_ipuclari = []
                for s in sorular_db[:10]:
                    metin_ref = s.get("metin_ref") or s.get("paragraf") or ""
                    if metin_ref:
                        soru_ipuclari.append(metin_ref)
                if soru_ipuclari:
                    metin_ek = f"\n\nKİTAP METİN PARÇALARI (daha önce üretilmiş sorulardan):\n" + "\n".join(soru_ipuclari[:5])
                    metin_kaynak = "ai_uretilen_sorular.metin_ref"
                else:
                    # En azından soru metinlerini ipucu olarak ver
                    mevcut_sorular = [s.get("soru","") for s in sorular_db[:5] if s.get("soru")]
                    if mevcut_sorular:
                        metin_ek = f"\n\nNOT: Bu kitap için daha önce üretilmiş {len(sorular_db)} soru var. Benzer içerik ve zorluk seviyesini koru ama FARKLI sorular üret:\n" + "\n".join(f"- {s}" for s in mevcut_sorular)
                        metin_kaynak = "ai_uretilen_sorular.ipucu"
                logging.info(f"[MATERYAL] ai_uretilen_sorular'dan {len(sorular_db)} kayıt bulundu")
        except Exception as e:
            logging.error(f"[MATERYAL] ai_uretilen_sorular hatası: {e}")

    logging.info(f"[MATERYAL] Toplam metin kaynağı: {metin_kaynak}, metin_ek uzunluğu: {len(metin_ek)}")

    metin_bilgi = f" (Yazar: {yazar})" if yazar else ""
    has_metin = bool(metin_ek.strip())

    # Metin yoksa: Gemini'den önce kitap hakkında bağlam üret
    if not has_metin and GEMINI_API_KEY:
        try:
            baglam_prompt = (
                f"Türkçe çocuk kitabı: '{kitap_adi}'{metin_bilgi}\n"
                f"Bu kitap hakkında bildiklerini yaz. Eğer bu kitabı bilmiyorsan, "
                f"kitap adından ve yazarından çıkarabileceğin ipuçlarıyla tahmini bir içerik özeti yaz.\n"
                f"Şunları belirt: Ana karakter kim? Nerede geçiyor? Temel olay nedir? Önemli yan karakterler? "
                f"Kitabın mesajı/teması nedir? Dikkat çekici sahneler?\n"
                f"Kısa ve net yaz (200-300 kelime)."
            )
            kitap_baglaami = await _gemini_call(baglam_prompt, max_tokens=600)
            metin_ek = f"\n\nKİTAP HAKKINDA BİLGİ (AI analizi):\n{kitap_baglaami.strip()}"
            metin_kaynak = "gemini_baglam"
            has_metin = True
            logging.info(f"[MATERYAL] Gemini'den kitap bağlamı alındı: {len(kitap_baglaami)} karakter")
        except Exception as baglam_e:
            logging.error(f"[MATERYAL] Bağlam üretme hatası: {baglam_e}")

    # Metin varsa: metne özgü talimatlar
    if has_metin:
        metin_bağlam = metin_ek
        metin_odak = (
            "\n🔴 ZORUNLU: Aşağıdaki kitap bilgisini kullan. Sorular MUTLAKA:\n"
            "- Kitaptaki GERÇEK karakter isimlerini kullan\n"
            "- Kitaptaki GERÇEK olayları, mekanları, detayları sor\n"
            "- 'Karakterin özelliği nedir?' gibi SOYUT/JENERİK sorular KESİNLİKLE YASAK\n"
            "- Doğru cevap açıkça bulunabilmeli, yanlış şıklar inandırıcı ama yanlış olsun\n"
        )
    else:
        metin_bağlam = ""
        metin_odak = (
            "\n🔴 ZORUNLU: Kitap adından yola çıkarak ÖZGİN sorular üret. YASAK sorular:\n"
            "- 'Kitabın ana karakterinin özelliği nedir?' → YASAK\n"
            "- 'Kitap hangi türdedir?' → YASAK  \n"
            "- 'Olaylar hangi ortamda geçer?' → YASAK\n"
            "Bunların yerine kitabın adından ve içeriğinden tahmin edilen özgün sorular sor.\n"
        )

    tur_prompts = {
        "soru_seti": (
            f"'{kitap_adi}'{metin_bilgi} için {sinif}. sınıf düzeyinde TAM 10 soruluk anlama testi oluştur.\n"
            f"{metin_odak}"
            f"Bloom taksonomisinin 6 basamağından dengeli sorular ekle:\n"
            f"- 2 Bilgi sorusu (metinde doğrudan geçen bilgi)\n"
            f"- 2 Kavrama sorusu (olayı kendi sözlerinle anlat)\n"
            f"- 2 Uygulama sorusu (sen olsaydın ne yapardın?)\n"
            f"- 1 Analiz sorusu (neden-sonuç ilişkisi)\n"
            f"- 1 Sentez/Değerlendirme sorusu\n"
            f"- 2 Yaratıcı/Eleştirel düşünme sorusu\n"
            f"Her soru için 4 seçenek olsun (A,B,C,D). Yanlış seçenekler inandırıcı olsun.{metin_bağlam}\n\n"
            f"JSON formatı (tam 10 soru): {{\"baslik\": \"string\", \"sorular\": ["
            f"{{\"soru\": \"string\", \"secenekler\": [\"A...\",\"B...\",\"C...\",\"D...\"], \"dogru\": \"A...\", \"bloom_basamak\": \"string\"}}"
            f"]}}"
        ),
        "kelime_listesi": (
            f"'{kitap_adi}'{metin_bilgi} kitabından {sinif}. sınıf için 10 önemli kelime seç.{metin_odak}{metin_bağlam}\n\n"
            f"SADECE JSON döndür:\n"
            f"{{\"baslik\": \"string\", \"kelimeler\": ["
            f"{{\"kelime\": \"string\", \"anlam\": \"kısa anlam\", \"cumle\": \"kısa örnek cümle\", \"zorluk\": 1}}"
            f"]}}"
        ),
        "etkinlik": (
            f"'{kitap_adi}'{metin_bilgi} için sınıf içi grup etkinliği tasarla. {sinif}. sınıf, 20-30 dk.{metin_odak}{metin_bağlam}\n\n"
            f"JSON: {{\"baslik\": \"string\", \"sure_dk\": 25, \"grup_sayisi\": 4, "
            f"\"adimlar\": [\"string\"], \"malzemeler\": [\"string\"], \"kazanimlar\": [\"string\"]}}"
        ),
        "tahmin": (
            f"'{kitap_adi}'{metin_bilgi} okuma öncesi TAM 8 tahmin sorusu oluştur ({sinif}. sınıf).{metin_bağlam}\n\n"
            f"JSON: {{\"baslik\": \"string\", \"giris\": \"string\", \"sorular\": ["
            f"{{\"soru\": \"string\", \"ipucu\": \"string\"}}"
            f"]}}"
        ),
    }

    prompt = tur_prompts.get(tur, tur_prompts["soru_seti"]) + "\n\nSADECE JSON döndür, başka hiçbir şey yazma. Tüm anahtarlar çift tırnak içinde olsun."


    import json as _json

    if not GEMINI_API_KEY:
        logging.warning("GEMINI_API_KEY yok — mock döndürülüyor")
    else:
        logging.info(f"Gemini çağrısı başlıyor: tur={tur}, kitap={kitap_adi}, metin_len={len(metin_ek)}")

    async def _parse_gemini_json(raw_text: str):
        """Gemini yanıtından JSON çıkar — hata toleranslı."""
        import re as _re
        raw_text = raw_text.strip()

        # ``` bloklarını temizle
        if "```" in raw_text:
            match = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw_text)
            if match:
                raw_text = match.group(1).strip()
            else:
                raw_text = _re.sub(r"```(?:json)?", "", raw_text).replace("```", "").strip()

        # { ... } içini al
        brace_start = raw_text.find("{")
        brace_end = raw_text.rfind("}")
        if brace_start != -1 and brace_end != -1:
            raw_text = raw_text[brace_start:brace_end+1]

        # Önce direkt parse dene
        try:
            return _json.loads(raw_text)
        except _json.JSONDecodeError:
            pass

        # Trailing comma düzelt: ,] ve ,} → ] ve }
        cleaned = _re.sub(r",\s*([}\]])", r"\1", raw_text)
        try:
            return _json.loads(cleaned)
        except _json.JSONDecodeError:
            pass

        # Tek tırnak → çift tırnak
        cleaned2 = cleaned.replace("'", '"')
        try:
            return _json.loads(cleaned2)
        except _json.JSONDecodeError:
            pass

        # Kontrol karakterlerini temizle
        cleaned3 = _re.sub(r'[\x00-\x1f\x7f]', ' ', cleaned)
        cleaned3 = _re.sub(r',\s*([}\]])', r'\1', cleaned3)
        return _json.loads(cleaned3)


    _debug = {
        "metin_kaynak": metin_kaynak,
        "metin_uzunluk": len(metin_ek),
        "gemini_key_var": bool(GEMINI_API_KEY),
        "gemini_key_uzunluk": len(GEMINI_API_KEY) if GEMINI_API_KEY else 0,
    }

    try:
        raw = await _gemini_call(prompt, max_tokens=4000)
        logging.info(f"Gemini yanıt alındı: {len(raw)} karakter")
        result = await _parse_gemini_json(raw)
        result["tur"] = tur
        result["kitap_adi"] = kitap_adi
        result["_debug"] = {**_debug, "deneme": 1, "hata": None}
        return result
    except Exception as e:
        hata1 = f"{type(e).__name__}: {e}"
        logging.error(f"Gemini materyal hatası (1. deneme): {hata1}")
        _debug["hata1"] = hata1
        # 2. deneme: daha basit prompt ile tekrar dene
        if GEMINI_API_KEY:
            try:
                basit_prompt = (
                    f"'{kitap_adi}' kitabı için {sinif}. sınıf öğrencisine 10 anlama sorusu üret.\n"
                    f"Bu kitabın GERÇEK karakterlerini, mekanlarını ve olaylarını kullan.\n"
                    f"Her soru kitabı okumayan birinin bilemeyeceği kadar özgün olsun.\n"
                    f"{'Kitap metni: ' + metin_ek[:3000] if metin_ek else ''}\n\n"
                    f"SADECE JSON döndür:\n"
                    f"{{\"baslik\": \"string\", \"sorular\": ["
                    f"{{\"soru\": \"string\", \"secenekler\": [\"A) ...\",\"B) ...\",\"C) ...\",\"D) ...\"], \"dogru\": \"A) ...\", \"bloom_basamak\": \"Bilgi\"}}"
                    f"]}}"
                )
                raw2 = await _gemini_call(basit_prompt, max_tokens=3000)
                logging.info(f"Gemini 2. deneme yanıtı: {len(raw2)} karakter")
                result = await _parse_gemini_json(raw2)
                result["tur"] = tur
                result["kitap_adi"] = kitap_adi
                result["_debug"] = {**_debug, "deneme": 2, "hata": None}
                return result
            except Exception as e2:
                hata2 = str(e2)
                logging.error(f"Gemini 2. deneme de başarısız: {hata2}")
                _debug["hata2"] = hata2

        # Son çare: Gemini'den kitap adına göre özgün fallback al
        if GEMINI_API_KEY:
            try:
                fallback_prompt = (
                    f"Türk ilkokul öğrencileri için '{kitap_adi}' kitabı hakkında 10 soru üret. "
                    f"Sorular bu kitaba ÖZGÜN olsun, genel sorular olmasın. "
                    f"JSON: {{\"baslik\": \"string\", \"sorular\": [{{\"soru\": \"string\", \"secenekler\": [\"A) x\",\"B) x\",\"C) x\",\"D) x\"], \"dogru\": \"A) x\", \"bloom_basamak\": \"Bilgi\"}}]}}"
                )
                raw3 = await _gemini_call(fallback_prompt, max_tokens=2000)
                result = await _parse_gemini_json(raw3)
                result["tur"] = tur
                result["kitap_adi"] = kitap_adi
                result["_debug"] = {**_debug, "deneme": 3, "hata": None}
                return result
            except Exception as e3:
                hata3 = str(e3)
                logging.error(f"Gemini fallback da başarısız: {hata3}")
                _debug["hata3"] = hata3

        # Hiçbir şey çalışmadıysa son çare statik mock (artık sadece gerçekten API yoksa)
        logging.error("TÜM Gemini denemeleri başarısız — statik mock kullanılıyor")
        bloom_list = ["Bilgi","Kavrama","Uygulama","Analiz","Kavrama","Uygulama","Bilgi","Sentez","Değerlendirme","Yaratma"]
        if tur == "soru_seti":
            sorular = []
            sorular_meta = [
                (f"'{kitap_adi}' kitabı hangi türdedir?", ["Roman","Şiir","Masal","Deneme"], "Roman", "Bilgi"),
                (f"Kitabın ana karakterinin özelliği nedir?", ["Cesur","Korkak","Tembel","Kıskanç"], "Cesur", "Kavrama"),
                (f"Sen bu karakterin yerinde olsaydın ne yapardın?", ["Aynısını yapardım","Farklı davranırdım","Kaçardım","Yardım isterdim"], "Farklı davranırdım", "Uygulama"),
                (f"Kitabın ana teması nedir?", ["Arkadaşlık","Cesaret","Dürüstlük","Merak"], "Cesaret", "Kavrama"),
                (f"Olaylar hangi ortamda geçmektedir?", ["Şehirde","Köyde","Ormanda","Deniz kıyısında"], "Köyde", "Bilgi"),
                (f"Karakterin en büyük sorunu neydi?", ["Yalnızlık","Maddi sıkıntı","Güven eksikliği","Hastalık"], "Güven eksikliği", "Analiz"),
                (f"Kitabın sonu nasıl bitmektedir?", ["Mutlu son","Hüzünlü son","Açık uçlu","Sürpriz son"], "Mutlu son", "Bilgi"),
                (f"Bu kitap sana göre hangi değeri en iyi anlatıyor?", ["Dürüstlük","Sabır","Cesaret","Yardımseverlik"], "Cesaret", "Değerlendirme"),
                (f"Yazar bu kitabı neden yazmış olabilir?", ["Eğlendirmek için","Ders vermek için","Duygu aktarmak için","Belgelemek için"], "Ders vermek için", "Sentez"),
                (f"Kitabı okuduktan sonra hayatında ne değiştirebilirsin?", ["Daha cesur olabilirim","Daha sabırlı olabilirim","Daha çok kitap okurum","Arkadaşlarıma yardım ederim"], "Daha cesur olabilirim", "Yaratma"),
            ]
            for soru, sec, dogru, bloom in sorular_meta:
                sorular.append({"soru": soru, "secenekler": sec, "dogru": dogru, "bloom_basamak": bloom})
            result = {"baslik": f"{kitap_adi} — Anlama Testi", "sorular": sorular}
        elif tur == "kelime_listesi":
            kelimeler_ornek = ["macera","cesaret","yolculuk","dürüstlük","arkadaşlık","merak","umut","sabır","özgürlük","kahramanlık","sadakat","iyilik","azim","fedakarlık","başarı"]
            result = {"baslik": f"{kitap_adi} — Anahtar Kelimeler", "kelimeler": [
                {"kelime": k, "anlam": f"{k.capitalize()} kavramının anlamı", "cumle": f"Kitapta {k} teması işlendi.", "zorluk": (i%3)+1}
                for i, k in enumerate(kelimeler_ornek)
            ]}
        elif tur == "etkinlik":
            result = {"baslik": f"{kitap_adi} — Sınıf Etkinliği", "sure_dk": 25, "grup_sayisi": 4,
                "adimlar": ["Sınıfı 4 gruba ayırın","Her grup kitabın farklı bölümünü tartışsın","Karakterleri analiz edin","Gruplar sunum yapsın","Sınıf tartışması yapın"],
                "malzemeler": ["Kağıt","Kalem","Post-it","Renkli kalemler"],
                "kazanimlar": ["Eleştirel düşünme","İşbirliği","Sözlü ifade","Empati kurma"]}
        else:
            result = {"baslik": f"{kitap_adi} — Okuma Öncesi Tahmin", "giris": "Kitabı okumadan önce düşüncelerini paylaş!", "sorular": [
                {"soru": "Bu kitap ne hakkında olabilir?", "ipucu": "Kapak resmine bak"},
                {"soru": "Ana karakter nasıl biri olabilir?", "ipucu": "Başlığı düşün"},
                {"soru": "Olaylar nerede geçiyor olabilir?", "ipucu": "Kapak resmindeki ortamı incele"},
                {"soru": "Kitabın sonu nasıl bitebilir?", "ipucu": "Başlığa göre tahmin et"},
                {"soru": "Hangi sorunlarla karşılaşılacak?", "ipucu": "Türüne bak"},
                {"soru": "Bu kitaptan ne öğrenebiliriz?", "ipucu": "Yazarı araştır"},
                {"soru": "Favorin olan karakter kim olabilir?", "ipucu": "Kitabın adına bak"},
                {"soru": "Sence bu kitabın mesajı ne olacak?", "ipucu": "Türü ve konusu hakkında düşün"},
            ]}

    result["tur"] = tur
    result["kitap_adi"] = kitap_adi
    result["sinif"] = sinif
    result["tarih"] = datetime.utcnow().isoformat()
    result["_debug"] = {**_debug, "deneme": 0, "mock": True}  # mock kullanıldı

    # Kaydet + XP
    await db.ai_materyal_log.insert_one({"ogrenci_id": ogrenci_id, "kitap_adi": kitap_adi, "tur": tur, "tarih": datetime.utcnow().isoformat()})
    await db.xp_logs.insert_one({"ogrenci_id": ogrenci_id, "xp": 5, "kaynak": "materyal_uret", "tarih": datetime.utcnow().isoformat()})
    return result


@router.get("/ai/materyal/gecmis/{ogrenci_id}")
async def materyal_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    kayitlar = await db.ai_materyal_log.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=20)
    for k in kayitlar: k.pop("_id", None)
    return kayitlar


@router.post("/ai/hikaye/olustur")
async def hikaye_olustur(req: Request, current_user=Depends(get_current_user)):
    """DNA + ilgi alanı → kişisel hikâye + 5 Bloom sorusu + kelime kartları."""
    data = await req.json()
    ogrenci_id  = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    ilgi_alani  = data.get("ilgi_alani", "macera")   # macera, hayvan, uzay, tarih, spor, arkadaşlık
    kahraman_ad = data.get("kahraman_ad", "Kahraman")
    sinif       = int(data.get("sinif", 3))
    sure_dk     = data.get("sure_dk", 5)  # hedef okuma süresi

    # DNA profili
    dna = await db.okuma_dna.find_one({"ogrenci_id": ogrenci_id})
    boyutlar = dna.get("boyutlar", {}) if dna else {}
    kelime_gucu = boyutlar.get("kelime_gucu", 50)
    anlama = boyutlar.get("anlama_derinligi", 50)
    profil_tipi = dna.get("profil_tipi", "dengeli_okuyucu") if dna else "dengeli_okuyucu"

    # MEB kelime havuzu (sınıfa göre mini örnek)
    meb_kelimeler = {
        1: ["ev", "okul", "anne", "baba", "arkadaş"],
        2: ["macera", "cesur", "yardım", "dürüst", "başarı"],
        3: ["merak", "keşif", "sorumluluk", "empati", "özgüven"],
        4: ["analiz", "strateji", "iletişim", "liderlik", "çözüm"],
        5: ["eleştiri", "perspektif", "hipotez", "kanıt", "sentez"],
    }.get(sinif, ["merak", "keşif", "arkadaşlık"])

    kelime_zorluk = "basit" if kelime_gucu < 40 else "orta" if kelime_gucu < 70 else "zengin"
    hikaye_uzunluk = sure_dk * 100  # kelime

    prompt = f"""Sen çocuklar için kişiselleştirilmiş hikâyeler yazan yaratıcı bir yazarsın.

Öğrenci profili:
- Ad: {kahraman_ad}
- Sınıf: {sinif}
- İlgi alanı: {ilgi_alani}
- Okuma seviyesi: {profil_tipi}
- Kelime zenginliği: {kelime_zorluk}
- Anlama derinliği: {anlama}/100

Hikâyeye şu MEB kelimelerini doğal biçimde dahil et: {", ".join(meb_kelimeler)}

Kurallar:
- Baş karakter adı: {kahraman_ad}
- Tema: {ilgi_alani}
- Uzunluk: yaklaşık {hikaye_uzunluk} kelime
- Dil: {kelime_zorluk} kelime zenginliği
- Mutlu son zorunlu
- MEB Erdemler: sabır, dürüstlük, merak, empati, cesaret değerlerinden en az 2'si tema olsun

SADECE JSON döndür:
{{
  "baslik": "Hikâye başlığı",
  "hikaye": "Tam hikâye metni ({hikaye_uzunluk} kelime)",
  "kullanilan_meb_kelimeleri": ["kelime1", "kelime2"],
  "kazanilan_deger": "sabır|dürüstlük|merak|empati|cesaret",
  "bloom_sorulari": [
    {{"basamak": "Hatırlama", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Kavrama", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Uygulama", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Analiz", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Değerlendirme", "soru": "...", "ipucu": "..."}}
  ],
  "kelime_kartlari": [
    {{"kelime": "...", "anlam": "...", "cumle": "..."}}
  ]
}}"""

    try:
        import json as _json
        raw = await _gemini_call(prompt, max_tokens=2500)
        raw = raw.strip()
        if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:-1])
        result = _json.loads(raw)
    except Exception as e:
        result = {
            "baslik": f"{kahraman_ad}'ın {ilgi_alani.capitalize()} Macerası",
            "hikaye": f"{kahraman_ad} bir gün ormanda yürürken sihirli bir kapıyla karşılaştı. Kapının arkasında harika bir dünya vardı. Merakla içeri girdi ve yeni arkadaşlar edindi. Birlikte zorlukları aşmayı öğrendiler. Cesaret ve dürüstlükle her engeli geçtiler. Sonunda eve döndüklerinde çok şey öğrenmişlerdi.",
            "kullanilan_meb_kelimeleri": meb_kelimeler[:2],
            "kazanilan_deger": "merak",
            "bloom_sorulari": [
                {"basamak": "Hatırlama", "soru": f"{kahraman_ad} ormanda ne buldu?", "ipucu": "Kapı ile ilgili düşün."},
                {"basamak": "Kavrama", "soru": "Karakterler ne öğrendi?", "ipucu": "Birlikte ne yaptılar?"},
                {"basamak": "Uygulama", "soru": "Sen olsaydın ne yapardın?", "ipucu": "Kendi deneyiminden düşün."},
                {"basamak": "Analiz", "soru": "Hikâyedeki ana sorun neydi?", "ipucu": "Zorlukları düşün."},
                {"basamak": "Değerlendirme", "soru": "Hikâye sana ne öğretti?", "ipucu": "Değerleri düşün."}
            ],
            "kelime_kartlari": [{"kelime": meb_kelimeler[0], "anlam": "Önemli bir değer", "cumle": f"{kahraman_ad} {meb_kelimeler[0]} gösterdi."}]
        }

    result["ogrenci_id"] = ogrenci_id
    result["ilgi_alani"] = ilgi_alani
    result["kahraman_ad"] = kahraman_ad
    result["sinif"] = sinif
    result["tarih"] = datetime.utcnow().isoformat()

    hikaye_id = str(__import__("uuid").uuid4())[:8]
    result["hikaye_id"] = hikaye_id
    await db.ai_hikaye_log.insert_one({**result})
    await db.xp_logs.insert_one({"ogrenci_id": ogrenci_id, "xp": 3, "kaynak": "adaptif_hikaye", "tarih": datetime.utcnow().isoformat()})
    return result


@router.get("/ai/hikaye/gecmis/{ogrenci_id}")
async def hikaye_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    kayitlar = await db.ai_hikaye_log.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=10)
    for k in kayitlar: k.pop("_id", None)
    return kayitlar


@router.post("/ai/post-reading")
async def post_reading_ai(req: Request, current_user=Depends(get_current_user)):
    """Kitap/içerik tamamlanınca derinlik soruları + MEB Erdem değeri üretir."""
    data = await req.json()
    kitap_adi = data.get("kitap_adi", "")
    yazar = data.get("yazar", "")
    sinif = data.get("sinif", 3)
    icerik_id = data.get("icerik_id", "")
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))

    # Cache kontrol (24 saat)
    cache = await db.post_reading_cache.find_one({"icerik_id": icerik_id, "ogrenci_id": ogrenci_id})
    if cache and cache.get("tarih"):
        import dateutil.parser
        try:
            sure = (datetime.utcnow() - dateutil.parser.parse(cache["tarih"])).total_seconds()
            if sure < 86400:
                cache.pop("_id", None)
                return cache
        except:
            pass

    # Kitap metnini al (varsa)
    metin_ek = ""
    icerik = await db.gelisim_icerik.find_one({"id": icerik_id})
    if icerik and icerik.get("okuma_metni"):
        metin_ek = f"\n\nKitap/metin içeriği (ilk 1500 kelime):\n{icerik['okuma_metni'][:3000]}"

    prompt = f"""Sen bir Türkçe eğitim uzmanısın. "{kitap_adi}"{f" ({yazar})" if yazar else ""} adlı kitabı/metni {sinif}. sınıf öğrencisi tamamladı.{metin_ek}

Aşağıdaki JSON formatında derinlik analizi üret:

1. Ana fikir sorusu: Öğrenciyi düşündüren açık uçlu 1 soru
2. MEB Erdem değeri: Bu kitaptan çıkarılabilecek 1 erdem (Erdemler: sabır, dürüstlük, merak, cesaret, sorumluluk, sevgi, saygı, adalet, yardımseverlik, vefa)
3. Bloom soruları: 3 basamak (Kavrama, Analiz, Yaratma) için birer soru
4. Hayat bağlantısı: "Bu kitap senin hayatında neyi değiştirir?" sorusu
5. Öneri kitaplar: Bu kitabı seven birine 2 kitap önerisi

SADECE JSON döndür:
{{
  "ana_fikir_sorusu": "...",
  "meb_erdem": {{
    "erdem": "sabır",
    "aciklama": "Bu kitap sabırlı olmayı şöyle gösteriyor: ..."
  }},
  "bloom_sorulari": [
    {{"basamak": "Kavrama", "soru": "...", "emoji": "🔍"}},
    {{"basamak": "Analiz", "soru": "...", "emoji": "🧩"}},
    {{"basamak": "Yaratma", "soru": "...", "emoji": "✨"}}
  ],
  "hayat_baglantisi": "...",
  "oneri_kitaplar": [
    {{"baslik": "...", "yazar": "...", "neden": "..."}},
    {{"baslik": "...", "yazar": "...", "neden": "..."}}
  ],
  "ozet_cumle": "Bu kitabın özü: ..."
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=1500)
        import json as _json, re as _re
        raw = raw.strip()
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
    except Exception as e:
        logging.error(f"[POST-READING] Hata: {e}")
        result = {
            "ana_fikir_sorusu": f"'{kitap_adi}' seni en çok hangi konuda düşündürdü? Neden?",
            "meb_erdem": {"erdem": "merak", "aciklama": "Bu kitap merak etmenin ve soru sormanın önemini gösteriyor."},
            "bloom_sorulari": [
                {"basamak": "Kavrama", "soru": "Kitabın ana karakteri hangi zorluklarla karşılaştı?", "emoji": "🔍"},
                {"basamak": "Analiz", "soru": "Kitaptaki olaylar neden bu sırayla gerçekleşti?", "emoji": "🧩"},
                {"basamak": "Yaratma", "soru": "Kitabın sonunu farklı yazsan nasıl bitirirdin?", "emoji": "✨"}
            ],
            "hayat_baglantisi": "Bu kitaptaki hangi duygu veya düşünce sana tanıdık geldi?",
            "oneri_kitaplar": [
                {"baslik": "Pollyanna", "yazar": "Eleanor H. Porter", "neden": "Benzer temalar ve umutlu bir bakış açısı"},
                {"baslik": "Küçük Prens", "yazar": "Antoine de Saint-Exupéry", "neden": "Derin düşünceler, çocuksu merak"}
            ],
            "ozet_cumle": f"'{kitap_adi}' sana yeni bir bakış açısı kazandırdı."
        }

    result["kitap_adi"] = kitap_adi
    result["icerik_id"] = icerik_id
    result["ogrenci_id"] = ogrenci_id
    result["tarih"] = datetime.utcnow().isoformat()
    await db.post_reading_cache.update_one(
        {"icerik_id": icerik_id, "ogrenci_id": ogrenci_id},
        {"$set": result}, upsert=True
    )
    return result


@router.post("/ai/icerik-kalite-skoru")
async def icerik_kalite_skoru(req: Request, current_user=Depends(get_current_user)):
    """İçeriği AI ile değerlendirir: 0-100 skor → 80+ otomatik onay, 50-79 peer review, 0-49 red."""
    data = await req.json()
    icerik_id = data.get("icerik_id", "")
    baslik = data.get("baslik", "")
    aciklama = data.get("aciklama", "")
    tur = data.get("tur", "kitap")
    sinif = data.get("sinif", 3)
    metin = data.get("metin", "")

    prompt = f"""Bir Türkçe eğitim içeriğini değerlendir. 0-100 arası kalite skoru ver.

Başlık: {baslik}
Tür: {tur}
Sınıf: {sinif}
Açıklama: {aciklama[:500] if aciklama else "Yok"}
{f"İçerik metni (ilk 500 kelime): {metin[:1000]}" if metin else ""}

Değerlendirme kriterleri:
1. Yaş/sınıf uygunluğu (0-25 puan)
2. Eğitimsel değer ve MEB uyumu (0-25 puan)
3. İçerik kalitesi ve özgünlük (0-25 puan)
4. Dil doğruluğu ve anlaşılırlık (0-25 puan)

SADECE JSON döndür:
{{
  "toplam_skor": 75,
  "alt_skorlar": {{
    "yas_uygunlugu": 20,
    "egitimsel_deger": 18,
    "icerik_kalitesi": 17,
    "dil_kalitesi": 20
  }},
  "guclu_yonler": ["..."],
  "zayif_yonler": ["..."],
  "red_nedeni": null,
  "oneri": "..."
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=800)
        import json as _json, re as _re
        raw = raw.strip()
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
    except Exception as e:
        logging.error(f"[HIBRIT-ONAY] AI skor hatası: {e}")
        result = {
            "toplam_skor": 65,
            "alt_skorlar": {"yas_uygunlugu": 17, "egitimsel_deger": 16, "icerik_kalitesi": 16, "dil_kalitesi": 16},
            "guclu_yonler": ["İçerik uygun görünüyor"],
            "zayif_yonler": ["Otomatik değerlendirme yapılamadı"],
            "red_nedeni": None,
            "oneri": "Manuel inceleme önerilir"
        }

    skor = result.get("toplam_skor", 65)

    # Karar
    if skor >= 80:
        karar = "otomatik_onayla"
        karar_label = "✅ Otomatik Onay"
        karar_renk = "green"
    elif skor >= 50:
        karar = "peer_review"
        karar_label = "🔍 Peer Review Gerekli"
        karar_renk = "orange"
    else:
        karar = "red"
        karar_label = "❌ Reddedildi"
        karar_renk = "red"

    result["karar"] = karar
    result["karar_label"] = karar_label
    result["karar_renk"] = karar_renk
    result["icerik_id"] = icerik_id

    # Otomatik onay uygula
    if karar == "otomatik_onayla" and icerik_id:
        try:
            await db.gelisim_icerik.update_one(
                {"id": icerik_id},
                {"$set": {"durum": "yayinda", "ai_skor": skor, "ai_karar": karar, "ai_karar_tarihi": datetime.utcnow().isoformat()}}
            )
        except: pass
    elif icerik_id:
        try:
            await db.gelisim_icerik.update_one(
                {"id": icerik_id},
                {"$set": {"ai_skor": skor, "ai_karar": karar, "ai_karar_tarihi": datetime.utcnow().isoformat()}}
            )
        except: pass

    return result
