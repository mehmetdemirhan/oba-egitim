"""AI bilgi tabanı modülü (/ai/bilgi-tabani/*, /ai/sorular, /ai/maliyet-ozet, /ai/demo-yukle).

server.py'dan BİREBİR taşındı; yollar ve davranış değişmedi. AI çağrıları
core.ai üzerinden yapılır; modül DB/auth/ayar erişimini yalnızca core'dan alır.
"""
import os
import io
import re
import json
import base64
import uuid
import asyncio
import logging
import random
import math
from datetime import datetime, timedelta, timezone
from typing import Optional, List, Dict, Any

import httpx
from fastapi import APIRouter, Depends, HTTPException, Request, Body, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field

from core.db import db, prepare_for_mongo, parse_from_mongo
from core.auth import get_current_user, require_role, UserRole
from core.config import (
    GEMINI_API_KEY, GEMINI_API_KEY_2, GEMINI_API_KEY_3, GEMINI_MODELS,
    AI_MODEL, AI_DEFAULT_MODEL, AI_HAIKU_MODEL, AI_MAX_DAILY_REQUESTS,
    AI_CACHE_HOURS, YANDEX_DISK_TOKEN, ANTHROPIC_API_KEY,
)
from core.sistem import (
    get_xp_tablosu, get_puan_ayarlari, get_lig_esikleri,
    XP_TABLOSU_DEFAULT, LIG_ESIKLERI_DEFAULT, LIG_SIRA,
)
from core.ai import _gemini_call, call_claude, _mock_bilgi_tabani_response, get_ogrenci_ai_verileri

router = APIRouter()


AI_EGITIM_PUANLARI = {
    "pdf_yukle": 20,
    "docx_yukle": 20,
    "onaylandi": 10,       # admin onayı sonrası ek puan
    "kelime_zengin": 5,    # 50+ kelime çıkarılan yükleme bonusu
    "soru_zengin": 5,      # 20+ soru çıkarılan yükleme bonusu
}

DESTEKLENEN_FORMATLAR = [".pdf", ".docx", ".doc"]

# ── "Tüm kelimeleri tara" modu ──
# AI yalnızca 5-15 hedef kelime seçer; bu mod metindeki TÜM benzersiz kelimeleri
# (anlamsız/ham olarak) meb_kelime_haritasi'na kaydeder. AI'ın anlam ürettiği
# kelimeler zaten kayıtlıdır ve ÜZERİNE YAZILMAZ (kaynak_tip="tam_tarama" işaretli).
TUM_KELIME_MAKS = 8000  # tek yüklemede güvenlik tavanı

_TR_BUYUK = "ABCÇDEFGĞHIİJKLMNOÖPRSŞTUÜVYZ"
_TR_KUCUK = "abcçdefgğhıijklmnoöprsştuüvyz"
_TR_CEV = {b: k for b, k in zip(_TR_BUYUK, _TR_KUCUK)}


def _bt_tr_kucuk(s: str) -> str:
    return "".join(_TR_CEV.get(ch, ch.lower()) for ch in (s or ""))


# ── Hafif Türkçe kök bulma (kural-tabanlı çekim eki soyma) ──
# Tam morfolojik çözümleme DEĞİL. Yaygın İSİM çekim eklerini (çoğul + hâl) güvenli
# soyar; kök en az 3 harf kalır, ≤4 harfli kelimeler korunur. Over-stemming'i
# önlemek için tek-ünlü ekler (-a/-e/-ı/-i) ve -ya/-yı gibi riskli ekler SOYULMAZ
# (ör. "papatya", "dünya", "kelime" korunur). En fazla 2 kat ek soyulur.
_KOK_EKLER = sorted([
    "larından", "lerinden", "larında", "lerinde", "sından", "sinden",
    "larını", "lerini", "ların", "lerin", "ları", "leri",
    "sının", "sinin", "sunun", "sünün", "sını", "sini", "sunu", "sünü",
    "ından", "inden", "undan", "ünden", "ında", "inde", "unda", "ünde",
    "nın", "nin", "nun", "nün",
    "dan", "den", "tan", "ten",
    "ler", "lar",
    "da", "de", "ta", "te",
], key=len, reverse=True)


_KOK_UNSUZ = set("bcçdfgğhjklmnprsştvyz")


def _turkce_kok(kelime: str) -> str:
    k = _bt_tr_kucuk(kelime or "")
    if len(k) <= 4:
        return k
    for _ in range(2):
        soyuldu = False
        for ek in _KOK_EKLER:
            if not k.endswith(ek):
                continue
            kalan = len(k) - len(ek)
            if len(ek) == 2:
                # da/de/ta/te (hâl eki) — yalnızca önü ünsüzse ve kök >= 3 harf
                # (salata/harita/oda korunur; kitapta→kitap, bahçede→bahçe)
                if kalan < 3 or k[kalan - 1] not in _KOK_UNSUZ:
                    continue
            elif kalan < 2:
                continue
            k = k[:kalan]
            soyuldu = True
            break
        if not soyuldu:
            break
    return k


def _tum_kelimeleri_cikar(metin: str) -> list:
    """Metindeki benzersiz Türkçe kelime KÖKLERİNİ döndürür (küçük harf, 2-20 harf).

    Kelimeler köke indirilir (yansımasını→yansıma, kelebekler→kelebek) ve köke göre
    tekilleştirilir; böylece aynı kelimenin çekimli hâlleri tek kayıt olur.
    """
    kucuk = _bt_tr_kucuk(metin or "")
    ham = re.findall(r"[a-zçğıöşü]+", kucuk)
    gorulen, out = set(), []
    for w in ham:
        if not (2 <= len(w) <= 20):
            continue
        k = _turkce_kok(w)
        if len(k) < 2 or k in gorulen:
            continue
        gorulen.add(k)
        out.append(k)
        if len(out) >= TUM_KELIME_MAKS:
            break
    return out


async def _tam_kelime_kaydet(ham_metin: str, sinif: int, kitap_adi: str, yukleyen_id: str) -> int:
    """Metindeki tüm benzersiz kelimeleri meb_kelime_haritasi'na toplu ekler.

    Zaten kayıtlı (kelime, sinif) atlanır → AI'ın anlam ürettiği kelimeler korunur.
    Dönüş: yeni eklenen ham kelime sayısı.
    """
    kelimeler = _tum_kelimeleri_cikar(ham_metin)
    if not kelimeler:
        return 0
    mevcut = set(await db.meb_kelime_haritasi.distinct("kelime", {"sinif": sinif}))
    now = datetime.utcnow().isoformat()
    yeni = []
    for k in kelimeler:
        if k in mevcut:
            continue
        mevcut.add(k)
        yeni.append({
            "id": str(uuid.uuid4()), "sinif": sinif, "kelime": k,
            "anlam": "", "ornek_cumle": "", "zorluk": 5,
            "kaynak": kitap_adi, "kaynak_tip": "tam_tarama",
            "yukleyen_id": yukleyen_id, "tarih": now,
        })
    if yeni:
        await db.meb_kelime_haritasi.insert_many(yeni)
    return len(yeni)


# ── Ham kelimelere AI ile anlam üretimi (batch + dedupe) ──
AI_HARITA_BATCH = 20       # tek promptta benzersiz kelime
AI_HARITA_BEKLEME = 2.0    # batch arası bekleme (kota)
AI_HARITA_DENEME = 3
_harita_ai_aktif: set = set()  # (sinif,) eşzamanlılık kilidi


def _harita_anlam_prompt(items: list) -> tuple:
    system = ("Sen ilkokul/ortaokul öğretmeni asistanısın. Çocuk dostu, TDK uyumlu, "
              "kısa ve net tanımlar üretirsin.")
    satirlar = "\n".join(f'- {it["kelime"]} ({it.get("sinif", 3)}. sınıf)' for it in items)
    user = ("Aşağıdaki her kelime için (1) çocuk dostu Türkçe anlam (en fazla 15 kelime), "
            "(2) sınıf seviyesine uygun kısa örnek cümle, (3) zorluk (1-10) üret.\n"
            f"Kelimeler:\n{satirlar}\n"
            'SADECE şu JSON DİZİSİNİ döndür: '
            '[{"kelime":"...","anlam":"...","ornek_cumle":"...","zorluk":5}]\n'
            "Markdown, kod bloğu veya ek açıklama EKLEME.")
    return system, user


async def _harita_ai_batch(items: list) -> dict:
    system, user = _harita_anlam_prompt(items)
    for _ in range(AI_HARITA_DENEME):
        try:
            res = await call_claude(system, user, model="sonnet", max_tokens=3500)
            parsed = res.get("parsed")
            lst = parsed if isinstance(parsed, list) else (
                (parsed.get("sonuclar") or parsed.get("kelimeler")) if isinstance(parsed, dict) else None)
            if isinstance(lst, list) and lst:
                out = {}
                for s in lst:
                    if not isinstance(s, dict):
                        continue
                    k = _bt_tr_kucuk(str(s.get("kelime", "")).strip())
                    anlam = str(s.get("anlam", "")).strip()
                    if k and anlam:
                        out[k] = {
                            "anlam": anlam,
                            "ornek_cumle": str(s.get("ornek_cumle", "")).strip(),
                            "zorluk": s.get("zorluk", 5),
                        }
                if out:
                    return out
        except Exception as ex:
            logging.warning(f"[harita_anlam] AI batch hatası: {ex}")
        await asyncio.sleep(AI_HARITA_BEKLEME)
    return {}


async def _harita_anlam_uret(sinif=None):
    """meb_kelime_haritasi'ndaki anlamı BOŞ kelimelere batch+dedupe AI ile anlam üretir.

    Aynı kelime birden çok sınıfta boşsa TEK AI isteğiyle üretilir ve o kelimenin
    boş TÜM kayıtlarına yazılır. Tam başarısızlıkta (kota) durur; kaldığı yerden
    sonraki tetiklemede devam eder (üretilmiş kelime tekrar sorulmaz).
    """
    anahtar = sinif
    if anahtar in _harita_ai_aktif:
        return
    _harita_ai_aktif.add(anahtar)
    denenen: set = set()
    try:
        while True:
            sorgu = {"anlam": {"$in": [None, ""]}}
            if sinif is not None:
                sorgu["sinif"] = int(sinif)
            bekleyen = await db.meb_kelime_haritasi.find(sorgu).limit(400).to_list(length=400)
            gruplar: dict = {}
            for b in bekleyen:
                k = b.get("kelime")
                if not k or k in denenen:
                    continue
                gruplar.setdefault(k, []).append(b)
            if not gruplar:
                break
            secilen = list(gruplar.keys())[:AI_HARITA_BATCH]
            items = [{"kelime": k, "sinif": gruplar[k][0].get("sinif", 3)} for k in secilen]
            harita = await _harita_ai_batch(items)
            if not harita:
                break
            now = datetime.utcnow().isoformat()
            for k in secilen:
                denenen.add(k)
                veri = harita.get(k)
                if not (veri and veri.get("anlam")):
                    continue
                await db.meb_kelime_haritasi.update_many(
                    {"kelime": k, "anlam": {"$in": [None, ""]}},
                    {"$set": {
                        "anlam": veri["anlam"],
                        "ornek_cumle": veri.get("ornek_cumle", ""),
                        "zorluk": veri.get("zorluk", 5),
                        "ai_anlam_tarihi": now,
                    }},
                )
            await asyncio.sleep(AI_HARITA_BEKLEME)
    except Exception as ex:
        logging.warning(f"[harita_anlam] kuyruk hatası (s{sinif}): {ex}")
    finally:
        _harita_ai_aktif.discard(anahtar)


@router.post("/ai/bilgi-tabani/dosya-onizle")
async def ai_bilgi_tabani_dosya_onizle(
    dosya: UploadFile = File(...),
    current_user=Depends(get_current_user)
):
    """Yüklenen dosyayı parse edip tam metni döner — DB'ye kaydetmez."""
    import os, io
    ext = os.path.splitext(dosya.filename)[1].lower()
    if ext not in [".pdf", ".docx", ".doc", ".txt"]:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {ext}")

    icerik = await dosya.read()

    bolumler = []  # [{"baslik": str, "metin": str, "sayfa": int}]
    ham_metin = ""

    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(stream=icerik, filetype="pdf")
            for i, page in enumerate(doc):
                sayfa_metni = page.get_text()
                ham_metin += sayfa_metni + "\n"
                bolumler.append({
                    "baslik": f"Sayfa {i+1}",
                    "metin": sayfa_metni.strip(),
                    "sayfa": i+1
                })
            doc.close()

        elif ext in [".docx", ".doc"]:
            from docx import Document as DocxDocument
            doc_obj = DocxDocument(io.BytesIO(icerik))
            mevcut_baslik = "Başlangıç"
            mevcut_metin = []
            sayfa = 1

            for para in doc_obj.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                try:
                    style = para.style.name if para.style else "Normal"
                except:
                    style = "Normal"

                if "Heading" in style or "Başlık" in style:
                    if mevcut_metin:
                        bolumler.append({
                            "baslik": mevcut_baslik,
                            "metin": "\n".join(mevcut_metin),
                            "sayfa": sayfa
                        })
                        if len("\n".join(mevcut_metin)) > 2000:
                            sayfa += 1
                    mevcut_baslik = text
                    mevcut_metin = []
                else:
                    mevcut_metin.append(text)
                    ham_metin += text + "\n"

            if mevcut_metin:
                bolumler.append({
                    "baslik": mevcut_baslik,
                    "metin": "\n".join(mevcut_metin),
                    "sayfa": sayfa
                })

        elif ext == ".txt":
            ham_metin = icerik.decode("utf-8", errors="ignore")
            satirlar = ham_metin.split("\n")
            chunk = []
            for i, satir in enumerate(satirlar):
                chunk.append(satir)
                if len(chunk) >= 50 or i == len(satirlar) - 1:
                    bolumler.append({
                        "baslik": f"Bölüm {len(bolumler)+1}",
                        "metin": "\n".join(chunk),
                        "sayfa": len(bolumler) + 1
                    })
                    chunk = []

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya okunamadı: {str(e)[:200]}")

    # Boş bölümleri filtrele
    bolumler = [b for b in bolumler if len(b["metin"].strip()) > 20]

    return {
        "dosya_adi": dosya.filename,
        "ext": ext,
        "toplam_bolum": len(bolumler),
        "toplam_kelime": len(ham_metin.split()),
        "bolumler": bolumler,
    }


@router.post("/ai/bilgi-tabani/yandex-kaydet/{yukleme_id}")
async def ai_bilgi_tabani_yandex_kaydet(
    yukleme_id: str,
    dosya: UploadFile = File(...),
    current_user=Depends(get_current_user)
):
    """Eski kitabı Yandex Disk'e kaydet ve yukleme kaydını güncelle."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    if not YANDEX_DISK_TOKEN:
        raise HTTPException(status_code=500, detail="YANDEX_DISK_TOKEN tanımlı değil")

    icerik = await dosya.read()
    import os
    ext = os.path.splitext(dosya.filename)[1].lower()

    temiz_ad = f"{yukleme_id}_{dosya.filename}"
    yandex_url = await yandex_disk_yukle(icerik, temiz_ad, ext)

    await db.ai_yuklemeler.update_one(
        {"id": yukleme_id},
        {"$set": {"yandex_url": yandex_url, "dosya_format": ext, "dosya_boyut": len(icerik)}}
    )
    return {"ok": True, "yandex_url": yandex_url, "mesaj": "✅ Yandex Disk'e kaydedildi"}


@router.get("/ai/bilgi-tabani/kitabi-ac/{yukleme_id}")
async def ai_bilgi_tabani_kitabi_ac(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yandex Disk'ten dosyayı indir, parse et ve tam metni döner."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id}, {"dosya_b64": 0, "_id": 0})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")

    yandex_url = yukleme.get("yandex_url")
    if not yandex_url:
        raise HTTPException(status_code=400, detail="Bu kitap Yandex Disk'e kaydedilmemiş. Lütfen yeniden yükleyin.")

    # Yandex'ten indir
    try:
        icerik = await yandex_disk_indir(yandex_url)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya indirilemedi: {str(e)[:200]}")

    ext = yukleme.get("dosya_format", ".pdf")
    import io

    bolumler = []
    ham_metin = ""

    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(stream=icerik, filetype="pdf")
            for i, page in enumerate(doc):
                sayfa_metni = page.get_text()
                ham_metin += sayfa_metni + "\n"
                if sayfa_metni.strip():
                    bolumler.append({"baslik": f"Sayfa {i+1}", "metin": sayfa_metni.strip(), "sayfa": i+1})
            doc.close()

        elif ext in [".docx", ".doc"]:
            from docx import Document as DocxDocument
            doc_obj = DocxDocument(io.BytesIO(icerik))
            mevcut_baslik = "Başlangıç"
            mevcut_metin = []
            sayfa = 1
            for para in doc_obj.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                try:
                    style = para.style.name if para.style else "Normal"
                except:
                    style = "Normal"
                if "Heading" in style or "Başlık" in style:
                    if mevcut_metin:
                        bolumler.append({"baslik": mevcut_baslik, "metin": "\n".join(mevcut_metin), "sayfa": sayfa})
                        if len("\n".join(mevcut_metin)) > 2000:
                            sayfa += 1
                    mevcut_baslik = text
                    mevcut_metin = []
                else:
                    mevcut_metin.append(text)
                    ham_metin += text + "\n"
            if mevcut_metin:
                bolumler.append({"baslik": mevcut_baslik, "metin": "\n".join(mevcut_metin), "sayfa": sayfa})

        elif ext == ".txt":
            ham_metin = icerik.decode("utf-8", errors="ignore")
            satirlar = ham_metin.split("\n")
            chunk = []
            for i, satir in enumerate(satirlar):
                chunk.append(satir)
                if len(chunk) >= 50 or i == len(satirlar) - 1:
                    bolumler.append({"baslik": f"Bölüm {len(bolumler)+1}", "metin": "\n".join(chunk), "sayfa": len(bolumler)+1})
                    chunk = []

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya okunamadı: {str(e)[:200]}")

    bolumler = [b for b in bolumler if len(b["metin"].strip()) > 20]

    return {
        "dosya_adi": yukleme.get("dosya_adi", ""),
        "ext": ext,
        "toplam_bolum": len(bolumler),
        "toplam_kelime": len(ham_metin.split()),
        "bolumler": bolumler,
        "yukleme": yukleme,
    }


@router.get("/ai/sorular")
async def ai_sorular_listesi(current_user=Depends(get_current_user)):
    """Tüm AI üretilen soruları listeler."""
    sorular = await db.ai_uretilen_sorular.find({}).sort("tarih", -1).to_list(length=200)
    for s in sorular:
        s.pop("_id", None)
    return sorular


@router.get("/ai/maliyet-ozet")
async def ai_maliyet_ozet(current_user=Depends(require_role(UserRole.ADMIN))):
    """Admin: AI maliyet özeti."""
    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    bu_ay = datetime.utcnow().strftime("%Y-%m")

    gunluk = await db.ai_request_log.find({"tarih": {"$regex": f"^{bugun}"}}).to_list(length=None)
    aylik = await db.ai_request_log.find({"tarih": {"$regex": f"^{bu_ay}"}}).to_list(length=None)

    return {
        "gunluk": {"istek": len(gunluk), "maliyet_usd": round(sum(r.get("maliyet_usd", 0) for r in gunluk), 4)},
        "aylik": {"istek": len(aylik), "maliyet_usd": round(sum(r.get("maliyet_usd", 0) for r in aylik), 4)},
        "gunluk_limit": AI_MAX_DAILY_REQUESTS,
    }


@router.post("/ai/demo-yukle")
async def ai_demo_yukle(current_user=Depends(get_current_user)):
    """Admin: Tüm AI demo verilerini oluştur/yenile."""
    if current_user.get("role") not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yetkiniz yok")

    simdi = datetime.utcnow()

    # Tüm öğrencileri bul
    tum_ogrenciler = await db.students.find({}).to_list(length=50)
    if not tum_ogrenciler:
        tum_ogrenciler = await db.users.find({"role": "student"}).to_list(length=50)
    if not tum_ogrenciler:
        raise HTTPException(status_code=404, detail="Hiç öğrenci bulunamadı")

    ogretmen = await db.users.find_one({"role": "teacher"})
    ogretmen_id = ogretmen["id"] if ogretmen else current_user["id"]
    ogretmen_ad = f"{ogretmen.get('ad','')} {ogretmen.get('soyad','')}".strip() if ogretmen else "Admin"

    sonuc = {"dna": 0, "kocluk": 0, "kelime": 0, "kelime_tekrar": 0, "yukleme": 0, "parca": 0, "soru": 0, "socratic": 0}

    # Eski demo verileri temizle
    await db.okuma_dna.delete_many({})
    await db.ai_kocluk_cache.delete_many({"model": "demo"})
    await db.kelime_tekrar.delete_many({})
    await db.ai_socratic_log.delete_many({})

    profil_havuzu = [
        {"tip": "hayalci_okuyucu", "label": "🌈 Hayalci Okuyucu", "b": {"kelime_gucu": 72, "akicilik": 45, "anlama_derinligi": 80, "dikkat_suresi": 55, "zorluk_toleransi": 60, "kelime_tekrar_ihtiyaci": 35, "okuma_psikolojisi": "keşifçi"}},
        {"tip": "hızlı_okuyucu", "label": "⚡ Hızlı Okuyucu", "b": {"kelime_gucu": 58, "akicilik": 78, "anlama_derinligi": 42, "dikkat_suresi": 70, "zorluk_toleransi": 45, "kelime_tekrar_ihtiyaci": 50, "okuma_psikolojisi": "güvenli"}},
        {"tip": "başlangıç_okuyucu", "label": "🌱 Başlangıç Okuyucu", "b": {"kelime_gucu": 35, "akicilik": 30, "anlama_derinligi": 55, "dikkat_suresi": 25, "zorluk_toleransi": 30, "kelime_tekrar_ihtiyaci": 75, "okuma_psikolojisi": "kararsız"}},
        {"tip": "dengeli_okuyucu", "label": "⚖️ Dengeli Okuyucu", "b": {"kelime_gucu": 65, "akicilik": 62, "anlama_derinligi": 68, "dikkat_suresi": 60, "zorluk_toleransi": 55, "kelime_tekrar_ihtiyaci": 40, "okuma_psikolojisi": "güvenli"}},
        {"tip": "analitik_okuyucu", "label": "🔍 Analitik Okuyucu", "b": {"kelime_gucu": 80, "akicilik": 50, "anlama_derinligi": 85, "dikkat_suresi": 75, "zorluk_toleransi": 70, "kelime_tekrar_ihtiyaci": 25, "okuma_psikolojisi": "keşifçi"}},
    ]

    for i, ogr in enumerate(tum_ogrenciler):
        oid = ogr["id"]
        ad = ogr.get("ad", f"Öğrenci{i}")
        sinif = ogr.get("sinif", 3)
        p = profil_havuzu[i % len(profil_havuzu)]

        # Hafif randomize
        boyutlar = {}
        for k, v in p["b"].items():
            if isinstance(v, int):
                boyutlar[k] = max(5, min(100, v + random.randint(-10, 10)))
            else:
                boyutlar[k] = v

        # DNA
        await db.okuma_dna.update_one({"ogrenci_id": oid}, {"$set": {
            "ogrenci_id": oid, "boyutlar": boyutlar, "profil_tipi": p["tip"],
            "profil_label": p["label"], "sinif": sinif, "son_guncelleme": simdi.isoformat(),
        }}, upsert=True)
        sonuc["dna"] += 1

        # Koçluk cache
        await db.ai_kocluk_cache.update_one({"ogrenci_id": oid}, {"$set": {
            "id": str(uuid.uuid4()), "ogrenci_id": oid,
            "dna": {"profil_tipi": p["tip"], "profil_label": p["label"], "boyutlar": boyutlar},
            "ai_analiz": {
                "durum_degerlendirmesi": {
                    "guclu_yonler": random.sample(["Anlama kapasitesi yüksek", "Düzenli okuyor", "Hayal gücü gelişmiş", "Meraklı", "Kelime hazinesi iyi", "Cesur kitap seçimleri", "Hızlı okuma", "Dikkatli dinleme"], 3),
                    "gelisim_alanlari": random.sample(["Okuma hızı artırılmalı", "Anlama derinliği geliştirilmeli", "Kelime hazinesi genişletilmeli", "Dikkat süresi kısa", "Bloom üst basamakları zayıf", "Streak tutarsız"], 2),
                },
                "risk_analizi": {"seviye": random.choice(["düşük", "orta", "yüksek"]), "faktorler": [random.choice(["Streak kırılma riski", "Kelime gücü düşük", "Dikkat süresi kısa", "Zor metinlerden kaçınma"])], "aciliyet": random.choice(["Takip yeterli", "Haftalık kontrol", "Acil müdahale"])},
                "mudahale_plani": {"hafta_1": "Günlük 10 dk sesli okuma", "hafta_2": "Tekrarlı okuma + kelime çalışması", "hafta_3": "Bloom soru çözme pratiği", "hafta_4": "Bağımsız okuma + özet yazma"},
                "veliye_mesaj": f"Sayın Veli, {ad} okuma gelişiminde ilerleme kaydediyor. Evde günlük 10-15 dakika birlikte okuma yapmanız gelişimini hızlandıracaktır. Kitap önerilerimizi takip edebilirsiniz.",
                "haftalik_gorevler": [
                    {"gun": "Pazartesi", "gorev": "Sevdiği kitaptan 2 sayfa sesli oku", "bloom": "uygulama"},
                    {"gun": "Salı", "gorev": "5 yeni kelime öğren ve cümle kur", "bloom": "uygulama"},
                    {"gun": "Çarşamba", "gorev": "Okuduğu bölümün özetini yaz", "bloom": "sentez"},
                    {"gun": "Perşembe", "gorev": "Karakterin motivasyonunu analiz et", "bloom": "analiz"},
                    {"gun": "Cuma", "gorev": "Hikâyeye alternatif son yaz", "bloom": "yaratma"},
                ],
                "kitap_tavsiyeleri": random.sample([
                    {"ad": "Charlie'nin Çikolata Fabrikası", "yazar": "Roald Dahl", "neden": "Hayal gücü yüksek, kısa bölümler"},
                    {"ad": "Küçük Prens", "yazar": "Saint-Exupéry", "neden": "Felsefi derinlik, kısa paragraflar"},
                    {"ad": "Pollyanna", "yazar": "E.H. Porter", "neden": "Pozitif bakış, karakter gelişimi"},
                    {"ad": "Kaşağı", "yazar": "Ömer Seyfettin", "neden": "Kısa öykü, değer eğitimi"},
                    {"ad": "Martı", "yazar": "Richard Bach", "neden": "Cesaret ve azim teması"},
                ], 3),
                "motivasyon_mesaji": random.choice([
                    f"Harika gidiyorsun {ad}! Her gün biraz daha güçleniyorsun 🌟",
                    f"Süpersin {ad}! Okumaya devam et, başarı senin hakkın 💪",
                    f"Merhaba {ad}! Bugün yeni bir maceraya hazır mısın? 📚",
                    f"{ad}, senin gibisi az bulunur! Her kelime bir adım 🚀",
                ]),
                "kelime_mudahale": "Günlük 5 yeni kelime + görsel kartlarla çalışma. Spaced repetition tekrarı.",
                "metin_recetesi": {"paragraf_uzunlugu": "Orta (80-120 kelime)", "soyutluk": random.choice(["Düşük", "Orta", "Yüksek"]), "aksiyon": random.choice(["Düşük", "Orta", "Yüksek"]), "hedef_kelime_orani": "%70 bilinen + %30 yeni"},
            },
            "ai_ham_metin": "", "model": "demo", "token": 0, "maliyet": 0, "tarih": simdi.isoformat(),
        }}, upsert=True)
        sonuc["kocluk"] += 1

        # Kelime tekrar (Spaced Repetition)
        demo_kelimeler_tekrar = [
            {"kelime": "macera", "anlam": "Tehlikeli ve heyecan verici olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı."},
            {"kelime": "keşif", "anlam": "Bilinmeyen bir şeyi ilk kez bulma", "ornek_cumle": "Bilim insanı yeni bir keşif yaptı."},
            {"kelime": "pusula", "anlam": "Yön bulmaya yarayan araç", "ornek_cumle": "Kaşif pusulasıyla yolunu buldu."},
            {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme gücü", "ornek_cumle": "Küçük kız büyük cesaret gösterdi."},
            {"kelime": "merak", "anlam": "Bir şeyi bilmek isteme duygusu", "ornek_cumle": "Merak eden çocuk her şeyi sorar."},
            {"kelime": "sabır", "anlam": "Bekleyebilme gücü", "ornek_cumle": "Bahçıvan sabırla bekledi."},
            {"kelime": "göç", "anlam": "Toplu taşınma", "ornek_cumle": "Kuşlar sıcak ülkelere göç eder."},
            {"kelime": "dürüstlük", "anlam": "Doğruyu söyleme", "ornek_cumle": "Dürüstlük en değerli erdemdir."},
        ]
        for kt in demo_kelimeler_tekrar:
            kutu = random.randint(1, 5)
            gun_sonra = {1:0, 2:1, 3:5, 4:14, 5:30}[kutu]
            await db.kelime_tekrar.insert_one({
                "id": str(uuid.uuid4()), "ogrenci_id": oid, "sinif": sinif, "kutu": kutu,
                "tekrar_sayisi": random.randint(1, 8), "dogru_sayisi": random.randint(0, 6),
                "son_gosterim": (simdi - timedelta(days=random.randint(1, 7))).isoformat(),
                "sonraki_gosterim": (simdi + timedelta(days=gun_sonra)).isoformat() if gun_sonra > 0 else simdi.isoformat(),
                "tarih": (simdi - timedelta(days=14)).isoformat(), **kt,
            })
            sonuc["kelime_tekrar"] += 1

        # Socratic log
        socratic_sorular = [
            "Bu bölümde en çok ne dikkatini çekti?",
            "Karakter neden böyle davrandı sence?",
            "Sen olsaydın ne yapardın?",
            "Bu hikâyenin sana öğrettiği bir şey var mı?",
            "Hikâyenin sonu farklı olabilir miydi?",
        ]
        for j in range(random.randint(2, 4)):
            await db.ai_socratic_log.insert_one({
                "id": str(uuid.uuid4()), "ogrenci_id": oid,
                "kitap_adi": random.choice(["Ormanın Sırrı", "Dürüst Çocuk", "Göçmen Kuşlar", "Takım Çalışması"]),
                "bolum": f"Bölüm {random.randint(1,4)}", "soru": random.choice(socratic_sorular),
                "bloom": random.choice(["kavrama", "analiz", "sentez", "degerlendirme"]),
                "puan": random.randint(3, 5), "tarih": (simdi - timedelta(days=random.randint(0, 5))).isoformat(),
            })
            sonuc["socratic"] += 1

    # Kelimeler (meb_kelime_haritasi)
    demo_kelimeler = [
        {"kelime": "macera", "anlam": "Tehlikeli ve heyecan verici olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı.", "zorluk": 4, "sinif": 3},
        {"kelime": "keşif", "anlam": "Bilinmeyen bir şeyi ilk kez bulma", "ornek_cumle": "Bilim insanı yeni bir keşif yaptı.", "zorluk": 5, "sinif": 3},
        {"kelime": "pusula", "anlam": "Yön bulmaya yarayan araç", "ornek_cumle": "Kaşif pusulasıyla yolunu buldu.", "zorluk": 6, "sinif": 3},
        {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme gücü", "ornek_cumle": "Küçük kız büyük cesaret gösterdi.", "zorluk": 4, "sinif": 3},
        {"kelime": "merak", "anlam": "Bir şeyi bilmek isteme duygusu", "ornek_cumle": "Merak eden çocuk her şeyi sorar.", "zorluk": 3, "sinif": 3},
        {"kelime": "sabır", "anlam": "Bekleyebilme ve dayanma gücü", "ornek_cumle": "Bahçıvan sabırla çiçeklerin büyümesini bekledi.", "zorluk": 4, "sinif": 3},
        {"kelime": "hayal gücü", "anlam": "Zihinde yeni şeyler oluşturabilme", "ornek_cumle": "Hayal gücü güçlü olan çocuklar iyi yazar.", "zorluk": 5, "sinif": 3},
        {"kelime": "fedakarlık", "anlam": "Başkaları için vazgeçme", "ornek_cumle": "Anneler büyük fedakarlıklar yapar.", "zorluk": 7, "sinif": 4},
        {"kelime": "dürüstlük", "anlam": "Doğruyu söyleme", "ornek_cumle": "Dürüstlük en değerli erdemdir.", "zorluk": 5, "sinif": 4},
        {"kelime": "azim", "anlam": "Kararlılıkla sürdürme", "ornek_cumle": "Azimli öğrenci başarıya ulaştı.", "zorluk": 6, "sinif": 4},
        {"kelime": "empati", "anlam": "Başkalarının duygularını anlama", "ornek_cumle": "Empati kurabilen iyi arkadaş olur.", "zorluk": 7, "sinif": 5},
        {"kelime": "göç", "anlam": "Toplu taşınma", "ornek_cumle": "Kuşlar sıcak ülkelere göç eder.", "zorluk": 4, "sinif": 3},
    ]
    await db.meb_kelime_haritasi.delete_many({"kaynak": {"$regex": "Demo"}})
    for k in demo_kelimeler:
        await db.meb_kelime_haritasi.update_one({"kelime": k["kelime"]}, {"$set": {
            "id": str(uuid.uuid4()), "kaynak": "Demo - MEB Türkçe",
            "yukleyen_id": ogretmen_id, "tarih": simdi.isoformat(), **k,
        }}, upsert=True)
        sonuc["kelime"] += 1

    # Demo yükleme
    await db.ai_yuklemeler.delete_many({"kitap_adi": {"$regex": "Demo"}})
    await db.ai_okuma_parcalari.delete_many({"kitap_adi": {"$regex": "Demo"}})
    await db.ai_uretilen_sorular.delete_many({"kitap_adi": {"$regex": "Demo"}})

    demo_yuk_id = str(uuid.uuid4())
    await db.ai_yuklemeler.insert_one({
        "id": demo_yuk_id, "dosya_adi": "turkce_3_ders_kitabi.pdf", "dosya_boyut": 4500000,
        "dosya_format": ".pdf", "dosya_hash": f"demo_{uuid.uuid4().hex[:8]}", "dosya_b64": "",
        "sinif": 3, "tur": "ders_kitabi", "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)",
        "yazar": "MEB", "temalar": ["Erdemler", "Doğa ve Evren", "Çocuk Dünyası"],
        "yukleyen_id": ogretmen_id, "yukleyen_ad": ogretmen_ad, "yukleyen_rol": "teacher",
        "durum": "tamamlandi", "onayli": True, "ilerleme": 100,
        "sonuc": {"sayfa_sayisi": 180, "kelime_sayisi": 45000, "chunk_sayisi": 8, "cikarilan_kelime": 12, "eklenen_kelime": 12, "okuma_parcasi": 4, "uretilen_soru": 10, "bonus_puan": 10},
        "guven_skoru": {"toplam": 92, "seviye": "yuksek", "detay": {"icindekiler": 90, "dil_uygunlugu": 95, "bloom_dagilimi": 88}}, "okuma_seviyesi": "3. Sınıf", "versiyon": 1, "tarih": (simdi - timedelta(days=3)).isoformat(),
    })
    sonuc["yukleme"] += 1

    demo_parcalar = [
        {"baslik": "Ormanın Sırrı", "ozet": "Küçük Ali ormanda kaybolur, konuşan hayvanlarla arkadaş olur.", "tema": "Doğa ve Evren", "metin_kesit": "Ağaçların arasından süzülen güneş ışığı ormanın derinliklerini aydınlatıyordu. Küçük Ali ilk kez bu kadar içerilere gelmişti..."},
        {"baslik": "Dürüst Çocuk", "ozet": "Pazarda para bulan Elif'in dürüstlük hikâyesi.", "tema": "Erdemler", "metin_kesit": "Elif pazarda yerde parlayan bir şey gördü. Eğilip baktığında bunun bir cüzdan olduğunu anladı. 'Bunu sahibine vermem lazım' dedi..."},
        {"baslik": "Göçmen Kuşlar", "ozet": "Leyleklerin göç yolculuğu ve yol bulma yetenekleri.", "tema": "Doğa ve Evren", "metin_kesit": "Her sonbaharda leylekler uzun bir yolculuğa çıkar. Binlerce kilometre uçarak sıcak ülkelere göç ederler..."},
        {"baslik": "Takım Çalışması", "ozet": "Sınıftaki öğrencilerin birlikte proje hazırlama hikâyesi.", "tema": "Çocuk Dünyası", "metin_kesit": "Öğretmen sınıfa bir proje verdi: 'Hayalinizdeki şehri tasarlayın.' Herkes tek başına yapmak istedi ama çok zordu..."},
    ]
    for j, p in enumerate(demo_parcalar):
        await db.ai_okuma_parcalari.insert_one({"id": str(uuid.uuid4()), "yukleme_id": demo_yuk_id, "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)", "sinif": 3, "bolum": j+1, **p, "kelime_sayisi": len(p["metin_kesit"].split()), "tarih": simdi.isoformat()})
        sonuc["parca"] += 1

    demo_sorular = [
        {"soru": "Ali ormanda kime rastladı?", "secenekler": ["Sincap", "Ayı", "Balık", "Kuş"], "dogru_cevap": 0, "taksonomi": "bilgi", "bolum": 1},
        {"soru": "Hayvanlar Ali'ye nasıl yardım etti?", "secenekler": ["Yemek verdiler", "Yol gösterdiler", "Şarkı söylediler", "Uyuttular"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 1},
        {"soru": "Elif cüzdanı neden sahibine verdi?", "secenekler": ["Korktu", "Dürüsttü", "Paraya ihtiyacı yoktu", "Annesi gördü"], "dogru_cevap": 1, "taksonomi": "analiz", "bolum": 2},
        {"soru": "Sen Elif'in yerinde olsan ne yapardın?", "secenekler": ["Sahibine verirdim", "Saklardım", "Polise götürürdüm", "Bilmiyorum"], "dogru_cevap": 0, "taksonomi": "degerlendirme", "bolum": 2},
        {"soru": "Leylekler neden göç eder?", "secenekler": ["Sıcak ülke", "Arkadaş bulma", "Yeni yuva", "Uçma öğrenme"], "dogru_cevap": 0, "taksonomi": "bilgi", "bolum": 3},
        {"soru": "Kuşların yol bulma yeteneğine ne denir?", "secenekler": ["GPS", "İçgüdü", "Harita", "Rüzgar"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 3},
        {"soru": "Takım çalışması neden önemlidir?", "secenekler": ["Hızlı biter", "Herkes öğrenir", "Eğlenceli", "Hepsi"], "dogru_cevap": 3, "taksonomi": "sentez", "bolum": 4},
    ]
    for s in demo_sorular:
        await db.ai_uretilen_sorular.insert_one({"id": str(uuid.uuid4()), "yukleme_id": demo_yuk_id, "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)", "sinif": 3, **s, "tarih": simdi.isoformat()})
        sonuc["soru"] += 1

    return {"mesaj": "✅ AI demo verileri oluşturuldu!", "sonuc": sonuc}


def hesapla_guven_skoru(analiz_sonuc, sinif):
    """AI Güven Skoru: Yükleme kalitesini 0-100 puanlar"""
    skor = 0
    detay = {}

    # 1. Kelime Çeşitliliği (max 25 puan)
    kelimeler = analiz_sonuc.get("hedef_kelimeler", [])
    kelime_sayisi = len(kelimeler)
    if kelime_sayisi >= 30:
        detay["kelime_cesitliligi"] = 25
    elif kelime_sayisi >= 20:
        detay["kelime_cesitliligi"] = 20
    elif kelime_sayisi >= 10:
        detay["kelime_cesitliligi"] = 15
    elif kelime_sayisi >= 5:
        detay["kelime_cesitliligi"] = 10
    else:
        detay["kelime_cesitliligi"] = 3
    skor += detay["kelime_cesitliligi"]

    # 2. Soru Kalitesi — Bloom dağılımı (max 25 puan)
    sorular = analiz_sonuc.get("sorular", [])
    bloom_dagilim = set()
    for s in sorular:
        b = s.get("bloom_basamagi") or s.get("taksonomi", "")
        if b:
            bloom_dagilim.add(b)
    if len(bloom_dagilim) >= 5:
        detay["soru_kalitesi"] = 25
    elif len(bloom_dagilim) >= 4:
        detay["soru_kalitesi"] = 20
    elif len(bloom_dagilim) >= 3:
        detay["soru_kalitesi"] = 15
    elif len(bloom_dagilim) >= 2:
        detay["soru_kalitesi"] = 10
    else:
        detay["soru_kalitesi"] = 5 if len(sorular) > 0 else 0
    skor += detay["soru_kalitesi"]

    # 3. Zorluk Uyumu — metnin zorluk puanı sınıf seviyesine uygun mu? (max 25 puan)
    zorluk = analiz_sonuc.get("zorluk_puani", 5)
    beklenen_zorluk = sinif  # 1. sınıf → 1, 4. sınıf → 4 civarı
    fark = abs(zorluk - beklenen_zorluk)
    if fark <= 1:
        detay["zorluk_uyumu"] = 25
    elif fark <= 2:
        detay["zorluk_uyumu"] = 18
    elif fark <= 3:
        detay["zorluk_uyumu"] = 10
    else:
        detay["zorluk_uyumu"] = 3
    skor += detay["zorluk_uyumu"]

    # 4. Okuma Seviyesi Analizi — Grade Level Score varlığı (max 25 puan)
    grade_level = analiz_sonuc.get("grade_level_score", {})
    if grade_level:
        gl_sinif = grade_level.get("tahmini_sinif", 0)
        gl_fark = abs(gl_sinif - sinif)
        if gl_fark <= 0.5:
            detay["seviye_uyumu"] = 25
        elif gl_fark <= 1:
            detay["seviye_uyumu"] = 20
        elif gl_fark <= 2:
            detay["seviye_uyumu"] = 12
        else:
            detay["seviye_uyumu"] = 5
    else:
        detay["seviye_uyumu"] = 10  # Grade level yoksa orta puan
    skor += detay["seviye_uyumu"]

    return {"skor": min(skor, 100), "detay": detay, "seviye": "yuksek" if skor >= 75 else "orta" if skor >= 50 else "dusuk"}


def hesapla_grade_level(metin):
    """Okuma Seviyesi Analizi — Grade Level Score"""
    kelimeler = metin.split()
    kelime_sayisi = len(kelimeler)
    if kelime_sayisi == 0:
        return {"kelime_sayisi": 0, "ort_kelime_uzunlugu": 0, "ort_cumle_uzunlugu": 0, "tahmini_sinif": 1, "zorluk_puani": 1}

    # Ortalama kelime uzunluğu
    ort_kelime_uzunlugu = sum(len(k) for k in kelimeler) / kelime_sayisi

    # Cümle sayısı
    import re
    cumleler = re.split(r'[.!?]+', metin)
    cumle_sayisi = max(len([c for c in cumleler if c.strip()]), 1)
    ort_cumle_uzunlugu = kelime_sayisi / cumle_sayisi

    # Uzun kelime oranı (7+ harf)
    uzun_kelime_orani = sum(1 for k in kelimeler if len(k) >= 7) / kelime_sayisi

    # Soyut kelime tahmini (basit heuristik: -lık, -lik, -sel, -sal, -cilik gibi ekler)
    soyut_ekler = ["lık", "lik", "luk", "lük", "sel", "sal", "cilik", "çilik", "sızlık", "sizlik"]
    soyut_sayisi = sum(1 for k in kelimeler if any(k.lower().endswith(e) for e in soyut_ekler))
    soyutluk_orani = soyut_sayisi / kelime_sayisi

    # Grade Level hesaplama (Türkçe uyarlamalı basit formül)
    # Temel: ort_cumle_uzunlugu * 0.3 + ort_kelime_uzunlugu * 0.5 + uzun_kelime_orani * 10 + soyutluk_orani * 15
    ham_skor = (ort_cumle_uzunlugu * 0.3) + (ort_kelime_uzunlugu * 0.5) + (uzun_kelime_orani * 10) + (soyutluk_orani * 15)

    # Sınıf seviyesine çevirme (1-8 arası)
    tahmini_sinif = max(1, min(8, round(ham_skor / 2)))
    zorluk_puani = max(1, min(10, round(ham_skor)))

    return {
        "kelime_sayisi": kelime_sayisi,
        "cumle_sayisi": cumle_sayisi,
        "ort_kelime_uzunlugu": round(ort_kelime_uzunlugu, 1),
        "ort_cumle_uzunlugu": round(ort_cumle_uzunlugu, 1),
        "uzun_kelime_orani": round(uzun_kelime_orani * 100, 1),
        "soyutluk_orani": round(soyutluk_orani * 100, 1),
        "tahmini_sinif": tahmini_sinif,
        "zorluk_puani": zorluk_puani,
    }


async def yandex_disk_yukle(icerik: bytes, dosya_adi: str, ext: str) -> str:
    """Dosyayı Yandex Disk'e yükle, indirme URL'sini döner."""
    import aiohttp, urllib.parse, re
    if not YANDEX_DISK_TOKEN:
        raise Exception("YANDEX_DISK_TOKEN tanımlı değil")

    # Güvenli dosya adı oluştur
    temiz_ad = re.sub(r'[^\w\-.]', '_', dosya_adi)
    yol = f"/OBA_Egitim/{temiz_ad}"
    headers = {"Authorization": f"OAuth {YANDEX_DISK_TOKEN}"}

    async with aiohttp.ClientSession() as session:
        # Klasör oluştur (yoksa)
        await session.put(
            "https://cloud-api.yandex.net/v1/disk/resources",
            params={"path": "/OBA_Egitim"},
            headers=headers
        )

        # Upload URL al
        async with session.get(
            "https://cloud-api.yandex.net/v1/disk/resources/upload",
            params={"path": yol, "overwrite": "true"},
            headers=headers
        ) as r:
            if r.status != 200:
                raise Exception(f"Yandex upload URL alınamadı: {r.status}")
            data = await r.json()
            upload_url = data["href"]

        # Dosyayı yükle
        async with session.put(upload_url, data=icerik) as r:
            if r.status not in [200, 201]:
                raise Exception(f"Yandex yükleme hatası: {r.status}")

        # İndirme linki al
        async with session.get(
            "https://cloud-api.yandex.net/v1/disk/resources/download",
            params={"path": yol},
            headers=headers
        ) as r:
            if r.status != 200:
                raise Exception(f"Yandex indirme URL alınamadı: {r.status}")
            data = await r.json()
            return data["href"]


async def yandex_disk_indir(url: str) -> bytes:
    """Yandex Disk'ten dosyayı indir."""
    import aiohttp
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as r:
            if r.status != 200:
                raise Exception(f"Yandex indirme hatası: {r.status}")
            return await r.read()


@router.post("/ai/bilgi-tabani/yukle")
async def ai_bilgi_tabani_yukle(
    dosya: UploadFile = File(...),
    sinif: int = Form(...),
    tur: str = Form("ders_kitabi"),
    kitap_adi: str = Form(""),
    yazar: str = Form(""),
    temalar: str = Form(""),
    ders_adi: str = Form(""),
    basim_yili: str = Form(""),
    tam_tarama: str = Form("true"),
    current_user=Depends(get_current_user)
):
    import os, hashlib, io
    tam_tarama_aktif = str(tam_tarama).lower() in ("true", "1", "evet", "on")
    ext = os.path.splitext(dosya.filename)[1].lower()
    if ext not in DESTEKLENEN_FORMATLAR:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {ext}. Desteklenen: {', '.join(DESTEKLENEN_FORMATLAR)}")

    icerik = await dosya.read()

    # Duplicate kontrolü (sadece hash ile — dosya içeriği saklanmaz)
    dosya_hash = hashlib.sha256(icerik).hexdigest()
    mevcut = await db.ai_yuklemeler.find_one({"dosya_hash": dosya_hash})
    if mevcut:
        raise HTTPException(status_code=409, detail=f"Bu dosya daha önce yüklenmiş: '{mevcut.get('kitap_adi', '')}' ({mevcut.get('tarih', '')[:10]}, {mevcut.get('yukleyen_ad', '')})")

    gercek_kitap_adi = kitap_adi or dosya.filename.replace(ext, "")
    yukleme_id = str(uuid.uuid4())

    # Yandex Disk'e yükle (token varsa)
    yandex_url = None
    if YANDEX_DISK_TOKEN:
        try:
            temiz_ad = f"{yukleme_id}_{dosya.filename}"
            yandex_url = await yandex_disk_yukle(icerik, temiz_ad, ext)
        except Exception as e:
            logging.warning(f"[YANDEX] Yükleme başarısız: {e}")

    # Kayıt oluştur
    yukleme = {
        "id": yukleme_id,
        "dosya_adi": dosya.filename,
        "dosya_boyut": len(icerik),
        "dosya_format": ext,
        "dosya_hash": dosya_hash,
        "yandex_url": yandex_url,  # Yandex Disk indirme URL'si
        "sinif": sinif,
        "tur": tur,
        "kitap_adi": gercek_kitap_adi,
        "yazar": yazar,
        "ders_adi": ders_adi,
        "basim_yili": basim_yili,
        "temalar": [t.strip() for t in temalar.split(",") if t.strip()] if temalar else [],
        "yukleyen_id": current_user["id"],
        "yukleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "yukleyen_rol": current_user.get("role", ""),
        "durum": "isleniyor",
        "ilerleme": 5,
        "tam_tarama": tam_tarama_aktif,
        "onayli": current_user.get("role") in ["admin", "coordinator"],
        "guven_skoru": None,
        "okuma_seviyesi": None,
        "sonuc": {},
        "versiyon": 1,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.ai_yuklemeler.insert_one(yukleme)

    # Puan ver
    puan = AI_EGITIM_PUANLARI.get(f"{ext.replace('.', '')}_yukle", 20)
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": puan}})
    await db.ai_egitim_puanlari.insert_one({
        "id": str(uuid.uuid4()),
        "kullanici_id": current_user["id"],
        "eylem": "dosya_yukle",
        "dosya_adi": dosya.filename,
        "sinif": sinif,
        "puan": puan,
        "tarih": datetime.utcnow().isoformat(),
    })

    # ── Dosyayı direkt işle (bellekte tut, MongoDB'ye kaydetme) ──
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "metin_cikariliyor", "ilerleme": 10}})

    ham_metin = ""
    sayfa_sayisi = 0
    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(stream=icerik, filetype="pdf")
            sayfa_sayisi = len(doc)
            for page in doc:
                ham_metin += page.get_text() + "\n"
            doc.close()
        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document as DocxDocument
                doc_obj = DocxDocument(io.BytesIO(icerik))
                # Heading-aware metin çıkarma — başlıkları ayraç olarak kullan
                for para in doc_obj.paragraphs:
                    if not para.text.strip():
                        continue
                    try:
                        style = para.style.name if para.style else "Normal"
                    except:
                        style = "Normal"
                    # Başlıkları çift newline ile ayır — chunking için doğal sınır
                    if "Heading" in style or "Başlık" in style:
                        ham_metin += f"\n\n=== {para.text.strip()} ===\n\n"
                    else:
                        ham_metin += para.text + "\n"
                sayfa_sayisi = max(1, len(ham_metin) // 2000)
            except:
                ham_metin = icerik.decode("utf-8", errors="ignore")
                sayfa_sayisi = max(1, len(ham_metin) // 2000)
        elif ext == ".txt":
            ham_metin = icerik.decode("utf-8", errors="ignore")
            sayfa_sayisi = max(1, len(ham_metin) // 2000)
    except Exception as e:
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "hata", "sonuc": {"hata": str(e)[:200]}}})
        raise HTTPException(status_code=500, detail=f"Metin çıkarma hatası: {str(e)[:200]}")

    if len(ham_metin.strip()) < 100:
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "hata", "sonuc": {"hata": "Yeterli metin çıkarılamadı"}}})
        raise HTTPException(status_code=400, detail="Yeterli metin çıkarılamadı (min 100 karakter)")

    kelime_sayisi = len(ham_metin.split())
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "ai_analiz", "ilerleme": 30}})

    # Chunking — başlık sınırlarına saygılı
    chunk_boyut = {1:200, 2:300, 3:400, 4:500, 5:600, 6:700, 7:800, 8:900}.get(sinif, 500)

    # Önce === başlık === ayraçlarına göre doğal bölümlere ayır
    import re
    bolumler = re.split(r'\n\n=== .+ ===\n\n', ham_metin)
    basliklar = re.findall(r'=== (.+) ===', ham_metin)

    chunks = []
    for bi, bolum in enumerate(bolumler):
        baslik = basliklar[bi - 1] if bi > 0 and bi - 1 < len(basliklar) else ""
        paragraflar = [p.strip() for p in bolum.split("\n") if len(p.strip()) > 30]
        if not paragraflar:
            continue
        # Bölüm çok uzunsa kendi içinde parçala
        mevcut_chunk = f"{baslik}\n" if baslik else ""
        for p in paragraflar:
            if len(mevcut_chunk.split()) + len(p.split()) > chunk_boyut:
                if len(mevcut_chunk.strip().split()) > 30:
                    chunks.append(mevcut_chunk.strip())
                mevcut_chunk = p
            else:
                mevcut_chunk += "\n" + p
        if mevcut_chunk.strip():
            chunks.append(mevcut_chunk.strip())


        chunks.append(mevcut_chunk.strip())
    if not chunks:
        chunks = [ham_metin[:2000]]
    chunks = chunks[:10]

    tum_kelimeler, tum_parcalar, tum_sorular = [], [], []

    for i, chunk in enumerate(chunks):
        ilerleme = 30 + int((i / len(chunks)) * 55)
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": ilerleme}})

        ai_prompt = f"""Kitap: {gercek_kitap_adi}
Sınıf: {sinif}. sınıf
Bölüm {i+1}/{len(chunks)}

METİN:
{chunk[:2000]}

Şu JSON'u üret:
{{
  "hedef_kelimeler": [
    {{"kelime": "...", "anlam": "...", "ornek_cumle": "...", "zorluk": 1-10}}
  ],
  "okuma_parcasi": {{
    "baslik": "Bu bölümün kısa başlığı",
    "ozet": "2-3 cümle özet",
    "tema": "MEB teması",
    "kelime_sayisi": 0
  }},
  "sorular": [
    {{"soru": "...", "secenekler": ["A","B","C","D"], "dogru_cevap": 0, "taksonomi": "bilgi"}}
  ]
}}

Kurallar:
- Hedef kelimeler: {sinif}. sınıf öğrencisinin ÖĞRENMESİ gereken 5-15 kelime
- Her kelimeye çocuğun anlayacağı basit anlam ve örnek cümle yaz
- 3-5 soru üret (farklı Bloom basamaklarından)
- SADECE JSON döndür"""

        result = await call_claude(
            "Sen MEB Türkçe müfredatını bilen bir dil eğitimcisisin.",
            ai_prompt, model="sonnet", max_tokens=2000
        )

        if result.get("parsed"):
            p = result["parsed"]
            for k in p.get("hedef_kelimeler", []):
                k["kaynak_kitap"] = gercek_kitap_adi
                k["sinif"] = sinif
                k["bolum"] = i + 1
                tum_kelimeler.append(k)
            parca = p.get("okuma_parcasi", {})
            if parca:
                parca["kaynak_kitap"] = gercek_kitap_adi
                parca["bolum"] = i + 1
                parca["metin_kesit"] = chunk[:500]
                tum_parcalar.append(parca)
            for s in p.get("sorular", []):
                s["kaynak_kitap"] = gercek_kitap_adi
                s["sinif"] = sinif
                s["bolum"] = i + 1
                tum_sorular.append(s)

    # Veritabanına kaydet
    eklenen_kelime = 0
    for k in tum_kelimeler:
        mevcut_k = await db.meb_kelime_haritasi.find_one({"kelime": k.get("kelime", "").lower(), "sinif": sinif})
        if not mevcut_k:
            await db.meb_kelime_haritasi.insert_one({
                "id": str(uuid.uuid4()), "sinif": sinif,
                "kelime": k.get("kelime", "").lower(), "anlam": k.get("anlam", ""),
                "ornek_cumle": k.get("ornek_cumle", ""), "zorluk": k.get("zorluk", 5),
                "kaynak": gercek_kitap_adi, "yukleyen_id": current_user["id"],
                "tarih": datetime.utcnow().isoformat(),
            })
            eklenen_kelime += 1

    # ── TÜM KELİMELERİ TARA (opsiyonel): metindeki her benzersiz kelimeyi hafızaya al ──
    tam_kelime = 0
    if tam_tarama_aktif:
        tam_kelime = await _tam_kelime_kaydet(ham_metin, sinif, gercek_kitap_adi, current_user["id"])
        if tam_kelime:
            asyncio.create_task(_harita_anlam_uret(sinif))  # ham kelimelere arka planda anlam üret

    for p in tum_parcalar:
        await db.ai_okuma_parcalari.insert_one({
            "id": str(uuid.uuid4()), "yukleme_id": yukleme_id,
            "kitap_adi": gercek_kitap_adi, "sinif": sinif, "bolum": p.get("bolum", 0),
            "baslik": p.get("baslik", ""), "ozet": p.get("ozet", ""),
            "tema": p.get("tema", ""), "metin_kesit": p.get("metin_kesit", ""),
            "kelime_sayisi": p.get("kelime_sayisi", 0), "tarih": datetime.utcnow().isoformat(),
        })

    for s in tum_sorular:
        await db.ai_uretilen_sorular.insert_one({
            "id": str(uuid.uuid4()), "yukleme_id": yukleme_id,
            "kitap_adi": gercek_kitap_adi, "sinif": sinif, "bolum": s.get("bolum", 0),
            "soru": s.get("soru", ""), "secenekler": s.get("secenekler", []),
            "dogru_cevap": s.get("dogru_cevap", 0), "taksonomi": s.get("taksonomi", "kavrama"),
            "tarih": datetime.utcnow().isoformat(),
        })

    bonus = 0
    if eklenen_kelime >= 50:
        bonus += 5
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 5}})
    if len(tum_sorular) >= 20:
        bonus += 5
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 5}})

    sonuc = {
        "sayfa_sayisi": sayfa_sayisi, "kelime_sayisi": kelime_sayisi,
        "chunk_sayisi": len(chunks), "cikarilan_kelime": len(tum_kelimeler),
        "eklenen_kelime": eklenen_kelime, "tam_kelime": tam_kelime,
        "okuma_parcasi": len(tum_parcalar),
        "uretilen_soru": len(tum_sorular), "bonus_puan": bonus,
        "mock": not bool(GEMINI_API_KEY),
        "kelimeler": tum_kelimeler, "parcalar": tum_parcalar, "sorular": tum_sorular,
    }

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {
        "durum": "tamamlandi", "ilerleme": 100,
        "sonuc": {k: v for k, v in sonuc.items() if k not in ["kelimeler", "parcalar", "sorular"]},
    }})

    tam_metin_mesaj = f" • 📚 {tam_kelime} ham kelime hafızaya alındı" if tam_kelime else ""
    yukleme.pop("_id", None)
    return {
        "yukleme": {**yukleme, "durum": "tamamlandi"},
        "puan_kazanilan": puan,
        "mesaj": f"✅ +{puan} puan! AI öğrendi: {eklenen_kelime} kelime, {len(tum_parcalar)} parça, {len(tum_sorular)} soru.{tam_metin_mesaj}",
        **sonuc
    }


@router.post("/ai/bilgi-tabani/isle/{yukleme_id}")
async def ai_bilgi_tabani_isle(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yüklenen dosyayı AI ile işle: metin çıkar → kelimeler + okuma parçaları + sorular üret."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    if not yukleme.get("dosya_b64"):
        raise HTTPException(status_code=400, detail="Dosya verisi bulunamadı")

    import base64, gzip
    dosya_raw = base64.b64decode(yukleme["dosya_b64"])
    # Sıkıştırılmış veriyi aç
    if yukleme.get("dosya_gzip"):
        dosya_bytes = gzip.decompress(dosya_raw)
    else:
        dosya_bytes = dosya_raw
    ext = yukleme.get("dosya_format", ".pdf")
    sinif = yukleme.get("sinif", 3)
    kitap_adi = yukleme.get("kitap_adi", "Bilinmeyen")

    # ── AŞAMA 1: METİN ÇIKARMA ──
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "metin_cikariliyor", "ilerleme": 10}})

    ham_metin = ""
    sayfa_sayisi = 0
    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(stream=dosya_bytes, filetype="pdf")
            sayfa_sayisi = len(doc)
            for page in doc:
                ham_metin += page.get_text() + "\n"
            doc.close()
        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document as DocxDocument
                doc_file = io.BytesIO(dosya_bytes)
                doc_obj = DocxDocument(doc_file)
                for para in doc_obj.paragraphs:
                    if para.text.strip():
                        ham_metin += para.text + "\n"
                sayfa_sayisi = max(1, len(ham_metin) // 2000)
            except:
                ham_metin = dosya_bytes.decode("utf-8", errors="ignore")
                sayfa_sayisi = max(1, len(ham_metin) // 2000)
    except Exception as e:
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "hata", "sonuc": {"hata": f"Metin çıkarma hatası: {str(e)[:200]}"}}})
        raise HTTPException(status_code=500, detail=f"Metin çıkarma hatası: {str(e)[:200]}")

    if len(ham_metin.strip()) < 100:
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "hata", "sonuc": {"hata": "Yeterli metin çıkarılamadı (min 100 karakter)"}}})
        raise HTTPException(status_code=400, detail="Yeterli metin çıkarılamadı")

    kelime_sayisi = len(ham_metin.split())
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "ai_analiz", "ilerleme": 30, "sonuc": {"sayfa_sayisi": sayfa_sayisi, "kelime_sayisi": kelime_sayisi}}})

    # ── AŞAMA 2: CHUNKING ──
    chunk_boyut = {1:200, 2:300, 3:400, 4:500, 5:600, 6:700, 7:800, 8:900}.get(sinif, 500)
    paragraflar = [p.strip() for p in ham_metin.split("\n") if len(p.strip()) > 30]
    chunks = []
    mevcut_chunk = ""
    for p in paragraflar:
        if len(mevcut_chunk.split()) + len(p.split()) > chunk_boyut:
            if mevcut_chunk.strip():
                chunks.append(mevcut_chunk.strip())
            mevcut_chunk = p
        else:
            mevcut_chunk += "\n" + p
    if mevcut_chunk.strip():
        chunks.append(mevcut_chunk.strip())

    if not chunks:
        chunks = [ham_metin[:2000]]

    # Max 10 chunk (maliyet kontrolü)
    chunks = chunks[:10]

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": 40, "sonuc.chunk_sayisi": len(chunks)}})

    # ── AŞAMA 3: AI ANALİZİ (her chunk için) ──
    tum_kelimeler = []
    tum_parcalar = []
    tum_sorular = []

    for i, chunk in enumerate(chunks):
        ilerleme = 40 + int((i / len(chunks)) * 40)
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": ilerleme}})

        ai_prompt = f"""Kitap: {kitap_adi}
Sınıf: {sinif}. sınıf
Bölüm {i+1}/{len(chunks)}

METİN:
{chunk[:2000]}

Şu JSON'u üret:
{{
  "hedef_kelimeler": [
    {{"kelime": "...", "anlam": "...", "ornek_cumle": "...", "zorluk": 1-10}}
  ],
  "okuma_parcasi": {{
    "baslik": "Bu bölümün kısa başlığı",
    "ozet": "2-3 cümle özet",
    "tema": "MEB teması",
    "kelime_sayisi": 0
  }},
  "sorular": [
    {{"soru": "...", "secenekler": ["A","B","C","D"], "dogru_cevap": 0, "taksonomi": "bilgi"}}
  ]
}}

Kurallar:
- Hedef kelimeler: {sinif}. sınıf öğrencisinin ÖĞRENMESİ gereken 5-15 kelime
- Her kelimeye çocuğun anlayacağı basit anlam ve örnek cümle yaz
- Okuma parçası özeti çocuğun merak edeceği şekilde olsun
- 3-5 soru üret (farklı Bloom basamaklarından)
- SADECE JSON döndür"""

        result = await call_claude(
            "Sen MEB Türkçe müfredatını bilen bir dil eğitimcisisin. Çocuklara uygun kelime ve soru üretirsin.",
            ai_prompt, model="sonnet", max_tokens=2000
        )

        if result.get("parsed"):
            p = result["parsed"]
            # Kelimeler
            for k in p.get("hedef_kelimeler", []):
                k["kaynak_kitap"] = kitap_adi
                k["sinif"] = sinif
                k["bolum"] = i + 1
                tum_kelimeler.append(k)
            # Okuma parçası
            parca = p.get("okuma_parcasi", {})
            if parca:
                parca["kaynak_kitap"] = kitap_adi
                parca["bolum"] = i + 1
                parca["metin_kesit"] = chunk[:500]
                tum_parcalar.append(parca)
            # Sorular
            for s in p.get("sorular", []):
                s["kaynak_kitap"] = kitap_adi
                s["sinif"] = sinif
                s["bolum"] = i + 1
                tum_sorular.append(s)

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": 85}})

    # ── AŞAMA 4: VERİTABANINA KAYDET ──
    # Kelimeleri meb_kelime_haritasi'na ekle
    eklenen_kelime = 0
    for k in tum_kelimeler:
        mevcut = await db.meb_kelime_haritasi.find_one({"kelime": k.get("kelime", "").lower(), "sinif": sinif})
        if not mevcut:
            await db.meb_kelime_haritasi.insert_one({
                "id": str(uuid.uuid4()),
                "sinif": sinif,
                "kelime": k.get("kelime", "").lower(),
                "anlam": k.get("anlam", ""),
                "ornek_cumle": k.get("ornek_cumle", ""),
                "zorluk": k.get("zorluk", 5),
                "kaynak": kitap_adi,
                "yukleyen_id": current_user["id"],
                "tarih": datetime.utcnow().isoformat(),
            })
            eklenen_kelime += 1

    # ── TÜM KELİMELERİ TARA (opsiyonel): metindeki her benzersiz kelimeyi hafızaya al ──
    tam_kelime = 0
    if yukleme.get("tam_tarama", True):
        tam_kelime = await _tam_kelime_kaydet(ham_metin, sinif, kitap_adi, current_user["id"])
        if tam_kelime:
            asyncio.create_task(_harita_anlam_uret(sinif))  # ham kelimelere arka planda anlam üret

    # Okuma parçalarını kaydet
    for p in tum_parcalar:
        await db.ai_okuma_parcalari.insert_one({
            "id": str(uuid.uuid4()),
            "yukleme_id": yukleme_id,
            "kitap_adi": kitap_adi,
            "sinif": sinif,
            "bolum": p.get("bolum", 0),
            "baslik": p.get("baslik", ""),
            "ozet": p.get("ozet", ""),
            "tema": p.get("tema", ""),
            "metin_kesit": p.get("metin_kesit", ""),
            "kelime_sayisi": p.get("kelime_sayisi", 0),
            "tarih": datetime.utcnow().isoformat(),
        })

    # Soruları kaydet
    for s in tum_sorular:
        await db.ai_uretilen_sorular.insert_one({
            "id": str(uuid.uuid4()),
            "yukleme_id": yukleme_id,
            "kitap_adi": kitap_adi,
            "sinif": sinif,
            "bolum": s.get("bolum", 0),
            "soru": s.get("soru", ""),
            "secenekler": s.get("secenekler", []),
            "dogru_cevap": s.get("dogru_cevap", 0),
            "taksonomi": s.get("taksonomi", "kavrama"),
            "tarih": datetime.utcnow().isoformat(),
        })

    # Bonus puanlar
    bonus = 0
    if eklenen_kelime >= 50:
        bonus += 5
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 5}})
    if len(tum_sorular) >= 20:
        bonus += 5
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 5}})

    # Sonucu güncelle
    sonuc = {
        "sayfa_sayisi": sayfa_sayisi,
        "kelime_sayisi": kelime_sayisi,
        "chunk_sayisi": len(chunks),
        "cikarilan_kelime": len(tum_kelimeler),
        "eklenen_kelime": eklenen_kelime,
        "tam_kelime": tam_kelime,
        "okuma_parcasi": len(tum_parcalar),
        "uretilen_soru": len(tum_sorular),
        "bonus_puan": bonus,
        "mock": not bool(GEMINI_API_KEY),
        "kelimeler": tum_kelimeler,
        "parcalar": tum_parcalar,
        "sorular": tum_sorular,
    }

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {
        "durum": "tamamlandi",
        "ilerleme": 100,
        "sonuc": {k: v for k, v in sonuc.items() if k not in ["kelimeler", "parcalar", "sorular"]},
    }})

    return sonuc


@router.get("/ai/bilgi-tabani/sonuc/{yukleme_id}")
async def ai_bilgi_tabani_sonuc(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yükleme sonuçlarını getir: kelimeler, okuma parçaları, sorular."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id}, {"dosya_b64": 0})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    yukleme.pop("_id", None)

    kelimeler = await db.meb_kelime_haritasi.find({"kaynak": yukleme.get("kitap_adi", "")}).to_list(length=None)
    for k in kelimeler: k.pop("_id", None)

    parcalar = await db.ai_okuma_parcalari.find({"yukleme_id": yukleme_id}).sort("bolum", 1).to_list(length=None)
    for p in parcalar: p.pop("_id", None)

    sorular = await db.ai_uretilen_sorular.find({"yukleme_id": yukleme_id}).sort("bolum", 1).to_list(length=None)
    for s in sorular: s.pop("_id", None)

    return {"yukleme": yukleme, "kelimeler": kelimeler, "parcalar": parcalar, "sorular": sorular}


@router.get("/ai/bilgi-tabani/ilerleme/{yukleme_id}")
async def ai_bilgi_tabani_ilerleme(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yükleme işleme ilerleme durumu."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id}, {"dosya_b64": 0, "_id": 0})
    if not yukleme:
        return {"ilerleme": 0, "durum": "bulunamadi"}
    return {"ilerleme": yukleme.get("ilerleme", 0), "durum": yukleme.get("durum", "yuklendi"), "sonuc": yukleme.get("sonuc", {})}


@router.get("/ai/bilgi-tabani/tam-metin/{yukleme_id}")
async def ai_bilgi_tabani_tam_metin(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yüklenen dosyanın tüm bölümlerini ve metin parçalarını döner (16MB altı dosyalar için)."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id}, {"dosya_b64": 0, "_id": 0})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")

    # Boyut kontrolü kaldırıldı — dosya artık DB'de saklanmıyor

    # Tüm okuma parçalarını getir
    parcalar = await db.ai_okuma_parcalari.find(
        {"yukleme_id": yukleme_id}
    ).sort("bolum", 1).to_list(length=None)

    for p in parcalar:
        p.pop("_id", None)
        sorular = await db.ai_uretilen_sorular.find(
            {"yukleme_id": yukleme_id, "bolum": p.get("bolum", 0)}
        ).to_list(length=None)
        for s in sorular:
            s.pop("_id", None)
        p["sorular"] = sorular

    # Tüm kelimeler
    kelimeler = await db.meb_kelime_haritasi.find(
        {"kaynak": yukleme.get("kitap_adi", ""), "sinif": yukleme.get("sinif")}
    ).sort("zorluk", 1).to_list(length=None)
    for k in kelimeler:
        k.pop("_id", None)

    return {
        "yukleme": yukleme,
        "parcalar": parcalar,
        "kelimeler": kelimeler,
        "toplam_parca": len(parcalar),
        "toplam_soru": sum(len(p.get("sorular", [])) for p in parcalar),
        "toplam_kelime": len(kelimeler),
    }


async def ai_bilgi_tabani_yukle_url(payload: dict, current_user=Depends(get_current_user)):
    """URL'den PDF/Word dosyası indirip yükle."""
    url = (payload.get("url") or "").strip()
    if not url or not url.startswith("http"):
        raise HTTPException(status_code=400, detail="Geçerli bir URL girin")

    sinif = payload.get("sinif", 3)
    tur = payload.get("tur", "ders_kitabi")
    kitap_adi = payload.get("kitap_adi", "")
    yazar = payload.get("yazar", "")

    # Dosyayı indir
    try:
        async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client_http:
            resp = await client_http.get(url)
            if resp.status_code != 200:
                raise HTTPException(status_code=400, detail=f"Dosya indirilemedi: HTTP {resp.status_code}")
            icerik = resp.content
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="Dosya indirme zaman aşımı (60sn)")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"İndirme hatası: {str(e)[:200]}")

    # Format tespiti
    dosya_adi = url.split("/")[-1].split("?")[0]
    ext = "." + dosya_adi.split(".")[-1].lower() if "." in dosya_adi else ""
    if ext not in [".pdf", ".docx", ".doc"]:
        # Content-Type'dan dene
        ct = resp.headers.get("content-type", "")
        if "pdf" in ct:
            ext = ".pdf"
            dosya_adi = dosya_adi or "indirilen.pdf"
        elif "word" in ct or "docx" in ct:
            ext = ".docx"
            dosya_adi = dosya_adi or "indirilen.docx"
        else:
            ext = ".pdf"  # varsayılan
            dosya_adi = dosya_adi or "indirilen.pdf"

    if ext not in DESTEKLENEN_FORMATLAR:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {ext}")

    # Duplicate kontrolü
    import hashlib
    dosya_hash = hashlib.sha256(icerik).hexdigest()
    mevcut = await db.ai_yuklemeler.find_one({"dosya_hash": dosya_hash})
    if mevcut:
        raise HTTPException(status_code=409, detail=f"Bu dosya daha önce yüklenmiş: '{mevcut.get('kitap_adi', '')}'")

    import base64, gzip
    icerik_gz = gzip.compress(icerik, compresslevel=9)
    dosya_b64 = base64.b64encode(icerik_gz).decode("utf-8")

    yukleme = {
        "id": str(uuid.uuid4()),
        "dosya_adi": dosya_adi,
        "dosya_boyut": len(icerik),
        "dosya_format": ext,
        "dosya_hash": dosya_hash,
        "dosya_gzip": True,
        "dosya_b64": dosya_b64,
        "kaynak_url": url,
        "sinif": sinif,
        "tur": tur,
        "kitap_adi": kitap_adi or dosya_adi.replace(ext, ""),
        "yazar": yazar,
        "temalar": [],
        "yukleyen_id": current_user["id"],
        "yukleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "yukleyen_rol": current_user.get("role", ""),
        "durum": "yuklendi",
        "onayli": current_user.get("role") in ["admin", "coordinator"],
        "guven_skoru": None,
        "okuma_seviyesi": None,
        "sonuc": {},
        "versiyon": 1,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.ai_yuklemeler.insert_one(yukleme)

    # Puan
    puan = 20
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": puan}})
    await db.ai_egitim_puanlari.insert_one({
        "id": str(uuid.uuid4()),
        "kullanici_id": current_user["id"],
        "eylem": "url_yukle",
        "dosya_adi": dosya_adi,
        "sinif": sinif,
        "puan": puan,
        "tarih": datetime.utcnow().isoformat(),
    })

    yukleme.pop("_id", None)
    yukleme.pop("dosya_b64", None)
    return {"yukleme": yukleme, "puan_kazanilan": puan, "mesaj": f"✅ +{puan} puan! Link'ten dosya indirildi ve yüklendi."}


@router.get("/ai/bilgi-tabani/gecmis")
async def ai_bilgi_tabani_gecmis(current_user=Depends(get_current_user)):
    filtre = {}
    if current_user.get("role") not in ["admin", "coordinator"]:
        filtre["yukleyen_id"] = current_user["id"]
    items = await db.ai_yuklemeler.find(filtre, {"dosya_b64": 0}).sort("tarih", -1).to_list(length=200)
    for i in items:
        i.pop("_id", None)
    return items


@router.get("/ai/bilgi-tabani/istatistik")
async def ai_bilgi_tabani_istatistik(current_user=Depends(get_current_user)):
    toplam_yukleme = await db.ai_yuklemeler.count_documents({})
    tamamlanan = await db.ai_yuklemeler.count_documents({"durum": "tamamlandi"})
    bekleyen = await db.ai_yuklemeler.count_documents({"onayli": False})
    toplam_kelime = await db.meb_kelime_haritasi.count_documents({})
    toplam_soru = await db.ai_uretilen_sorular.count_documents({}) + await db.kitap_sorulari.count_documents({"kaynak": "ai_egitim"})

    sinif_dagilim = {}
    for s in range(1, 9):
        sinif_dagilim[str(s)] = {
            "yukleme": await db.ai_yuklemeler.count_documents({"sinif": s}),
            "kelime": await db.meb_kelime_haritasi.count_documents({"sinif": s}),
        }

    # En çok katkı yapan öğretmenler
    pipeline = [
        {"$group": {"_id": "$kullanici_id", "toplam_puan": {"$sum": "$puan"}, "yukleme_sayisi": {"$sum": 1}}},
        {"$sort": {"toplam_puan": -1}},
        {"$limit": 10}
    ]
    top_contributors = []
    async for doc in db.ai_egitim_puanlari.aggregate(pipeline):
        user = await db.users.find_one({"id": doc["_id"]})
        if user:
            top_contributors.append({
                "ad": f"{user.get('ad', '')} {user.get('soyad', '')}".strip(),
                "puan": doc["toplam_puan"],
                "yukleme": doc["yukleme_sayisi"],
            })

    # Güven skoru istatistikleri
    guven_yuklemeler = await db.ai_yuklemeler.find({"guven_skoru": {"$ne": None}}, {"guven_skoru": 1}).to_list(length=None)
    guven_skorlari = []
    guven_dagilim = {"yuksek": 0, "orta": 0, "dusuk": 0}
    for y in guven_yuklemeler:
        gs = y.get("guven_skoru")
        if gs is None:
            continue
        # Eski format: integer (92) → dict'e çevir
        if isinstance(gs, (int, float)):
            toplam = gs
            seviye = "yuksek" if gs >= 80 else ("orta" if gs >= 60 else "dusuk")
        else:
            toplam = gs.get("toplam", 0)
            seviye = gs.get("seviye", "")
        if toplam:
            guven_skorlari.append(toplam)
        if seviye in guven_dagilim:
            guven_dagilim[seviye] += 1
    guven_ort = round(sum(guven_skorlari) / max(len(guven_skorlari), 1), 1) if guven_skorlari else 0

    # Duplicate önleme istatistikleri
    toplam_hash = await db.ai_yuklemeler.distinct("dosya_hash")
    duplicate_engellenen = toplam_yukleme - len(toplam_hash) if toplam_yukleme > len(toplam_hash) else 0

    return {
        "toplam_yukleme": toplam_yukleme,
        "tamamlanan": tamamlanan,
        "bekleyen_onay": bekleyen,
        "toplam_kelime": toplam_kelime,
        "toplam_ai_soru": toplam_soru,
        "sinif_dagilim": sinif_dagilim,
        "top_contributors": top_contributors,
        "guven_skoru": {"ortalama": guven_ort, "dagilim": guven_dagilim},
        "duplicate_engellenen": duplicate_engellenen,
    }


@router.put("/ai/bilgi-tabani/onayla/{yukleme_id}")
async def ai_bilgi_tabani_onayla(yukleme_id: str, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"onayli": True}})
    # Onay bonusu
    await db.users.update_one({"id": yukleme["yukleyen_id"]}, {"$inc": {"puan": AI_EGITIM_PUANLARI["onaylandi"]}})
    await db.ai_egitim_puanlari.insert_one({
        "id": str(uuid.uuid4()), "kullanici_id": yukleme["yukleyen_id"],
        "eylem": "onay_bonusu", "puan": AI_EGITIM_PUANLARI["onaylandi"],
        "tarih": datetime.utcnow().isoformat(),
    })
    return {"ok": True, "mesaj": "Onaylandı, yükleyene +10 bonus puan verildi"}


@router.delete("/ai/bilgi-tabani/{yukleme_id}")
async def ai_bilgi_tabani_sil(yukleme_id: str, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    await db.ai_yuklemeler.delete_one({"id": yukleme_id})
    return {"ok": True}


@router.post("/ai/bilgi-tabani/anlam-uret")
async def ai_harita_anlam_uret_endpoint(payload: dict = None, current_user=Depends(get_current_user)):
    """Kitaptan alınan ham kelimelere (anlamı boş) AI ile toplu anlam üretir (arka plan).

    Opsiyonel {sinif}. Batch+dedupe; aynı kelime tüm boş kayıtlara tek çağrıyla yazılır.
    Yetki: öğretmen/koordinatör/admin.
    """
    if current_user.get("role") not in ("teacher", "admin", "coordinator"):
        raise HTTPException(status_code=403, detail="Yetkiniz yok.")
    payload = payload or {}
    sinif = int(payload["sinif"]) if payload.get("sinif") is not None else None
    sorgu = {"anlam": {"$in": [None, ""]}}
    if sinif is not None:
        sorgu["sinif"] = sinif
    bekleyenler = await db.meb_kelime_haritasi.find(sorgu, {"kelime": 1}).to_list(length=None)
    benzersiz = len({b.get("kelime") for b in bekleyenler if b.get("kelime")})
    toplam_batch = max(0, -(-benzersiz // AI_HARITA_BATCH))
    if bekleyenler:
        asyncio.create_task(_harita_anlam_uret(sinif))
    return {
        "kuyruk_baslatildi": bool(bekleyenler),
        "bekleyen_kelime": len(bekleyenler),
        "benzersiz_kelime": benzersiz,
        "toplam_batch": toplam_batch,
        "tahmini_kalan_sure_sn": toplam_batch * 5,
    }


@router.get("/ai/bilgi-tabani/puanlarim")
async def ai_egitim_puanlarim(current_user=Depends(get_current_user)):
    puanlar = await db.ai_egitim_puanlari.find({"kullanici_id": current_user["id"]}).sort("tarih", -1).to_list(length=100)
    for p in puanlar:
        p.pop("_id", None)
    toplam = sum(p.get("puan", 0) for p in puanlar)
    return {"toplam": toplam, "detay": puanlar}


# AI Güven Skoru hesaplama (yükleme sonrası çağrılır)
@router.post("/ai/bilgi-tabani/guven-skoru/{yukleme_id}")
async def ai_guven_skoru_hesapla(yukleme_id: str, payload: dict, current_user=Depends(get_current_user)):
    """
    AI işleme sonrası güven skoru hesaplar.
    payload: { kelime_sayisi, benzersiz_kelime, soru_sayisi, bloom_dagilim{}, zorluk_puani, sinif }
    """
    kelime_sayisi = payload.get("kelime_sayisi", 0)
    benzersiz = payload.get("benzersiz_kelime", 0)
    soru_sayisi = payload.get("soru_sayisi", 0)
    bloom = payload.get("bloom_dagilim", {})
    zorluk = payload.get("zorluk_puani", 5)
    sinif = payload.get("sinif", 3)

    # 1. Kelime çeşitliliği skoru (0-25)
    if kelime_sayisi == 0:
        kelime_skor = 0
    else:
        cesitlilik = benzersiz / max(kelime_sayisi, 1)
        kelime_skor = min(25, round(cesitlilik * 50))  # %50 çeşitlilik = 25 puan

    # 2. Soru kalitesi skoru (0-25)
    bloom_turleri = len([v for v in bloom.values() if v > 0])
    soru_skor = 0
    if soru_sayisi >= 5:
        soru_skor += 10
    elif soru_sayisi >= 3:
        soru_skor += 5
    soru_skor += min(15, bloom_turleri * 3)  # 5 bloom = 15 puan

    # 3. Zorluk uyumu skoru (0-25)
    # İdeal zorluk: sınıf seviyesine yakın (sinif * 1.2 civarı)
    ideal_zorluk = min(10, sinif * 1.2)
    zorluk_fark = abs(zorluk - ideal_zorluk)
    zorluk_skor = max(0, 25 - round(zorluk_fark * 5))

    # 4. İçerik zenginliği skoru (0-25)
    icerik_skor = 0
    if kelime_sayisi >= 500:
        icerik_skor += 10
    elif kelime_sayisi >= 200:
        icerik_skor += 5
    if benzersiz >= 50:
        icerik_skor += 8
    elif benzersiz >= 20:
        icerik_skor += 4
    if soru_sayisi >= 10:
        icerik_skor += 7
    elif soru_sayisi >= 5:
        icerik_skor += 3

    toplam = kelime_skor + soru_skor + zorluk_skor + icerik_skor
    seviye = "yuksek" if toplam >= 70 else "orta" if toplam >= 40 else "dusuk"

    guven = {
        "toplam": toplam,
        "seviye": seviye,
        "detay": {
            "kelime_cesitliligi": {"skor": kelime_skor, "max": 25, "aciklama": f"{benzersiz} benzersiz / {kelime_sayisi} toplam kelime"},
            "soru_kalitesi": {"skor": soru_skor, "max": 25, "aciklama": f"{soru_sayisi} soru, {bloom_turleri} Bloom basamağı"},
            "zorluk_uyumu": {"skor": zorluk_skor, "max": 25, "aciklama": f"Zorluk {zorluk}/10, {sinif}. sınıf için ideal ~{round(ideal_zorluk, 1)}"},
            "icerik_zenginligi": {"skor": icerik_skor, "max": 25, "aciklama": f"{kelime_sayisi} kelime, {benzersiz} benzersiz, {soru_sayisi} soru"},
        }
    }

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"guven_skoru": guven}})
    return guven


# Okuma Seviyesi Analizi (Grade Level Score)
@router.post("/ai/bilgi-tabani/okuma-seviyesi")
async def ai_okuma_seviyesi_hesapla(payload: dict, current_user=Depends(get_current_user)):
    """
    Metin zorluk analizi — adaptif motor için kritik.
    payload: { metin, sinif_hedef }
    Döner: Grade Level Score (1-8 sınıf eşdeğeri)
    """
    metin = payload.get("metin", "")
    sinif_hedef = payload.get("sinif_hedef", 3)

    if not metin or len(metin) < 50:
        return {"hata": "Metin en az 50 karakter olmalı"}

    # Metrik hesaplama
    cumleler = [c.strip() for c in metin.replace("!", ".").replace("?", ".").split(".") if c.strip()]
    kelimeler = metin.split()
    toplam_kelime = len(kelimeler)
    toplam_cumle = max(len(cumleler), 1)
    toplam_hece = sum(hece_say(k) for k in kelimeler)

    # Ort cümle uzunluğu (kelime/cümle)
    ort_cumle = toplam_kelime / toplam_cumle
    # Ort kelime uzunluğu (harf)
    ort_kelime_uzunluk = sum(len(k) for k in kelimeler) / max(toplam_kelime, 1)
    # Ort hece/kelime
    ort_hece = toplam_hece / max(toplam_kelime, 1)

    # Ateşman Okunabilirlik Formülü (Türkçe uyarlaması)
    # Okunabilirlik = 198.825 – 40.175 × (hece/kelime) – 2.610 × (kelime/cümle)
    atesman = 198.825 - (40.175 * ort_hece) - (2.610 * ort_cumle)
    atesman = max(0, min(100, round(atesman, 1)))

    # Grade level eşdeğeri
    if atesman >= 90: grade = 1
    elif atesman >= 80: grade = 2
    elif atesman >= 70: grade = 3
    elif atesman >= 60: grade = 4
    elif atesman >= 50: grade = 5
    elif atesman >= 40: grade = 6
    elif atesman >= 30: grade = 7
    else: grade = 8

    # Zorluk puanı (1-10)
    zorluk_puan = max(1, min(10, round((100 - atesman) / 10)))

    # Sınıf uyumu
    uyum = "uygun" if abs(grade - sinif_hedef) <= 1 else "kolay" if grade < sinif_hedef - 1 else "zor"

    return {
        "atesman_skoru": atesman,
        "grade_level": grade,
        "zorluk_puani": zorluk_puan,
        "sinif_uyumu": uyum,
        "metrikler": {
            "toplam_kelime": toplam_kelime,
            "toplam_cumle": toplam_cumle,
            "ort_cumle_uzunlugu": round(ort_cumle, 1),
            "ort_kelime_uzunlugu": round(ort_kelime_uzunluk, 1),
            "ort_hece_kelime": round(ort_hece, 1),
        },
        "yorum": f"Ateşman skoru {atesman} → {grade}. sınıf seviyesi. Hedef {sinif_hedef}. sınıf için {uyum}."
    }


def hece_say(kelime):
    """Türkçe hece sayma — sesli harf sayısı"""
    sesliler = set("aeıioöuüAEIİOÖUÜ")
    return max(1, sum(1 for h in kelime if h in sesliler))
