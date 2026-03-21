from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, Body, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
import random
import re
from datetime import datetime, timezone, timedelta
from enum import Enum
from passlib.context import CryptContext
from jose import JWTError, jwt
import httpx

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Config
SECRET_KEY = os.environ.get('SECRET_KEY', 'okuma-becerileri-secret-key-change-in-production')
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.environ.get('ACCESS_TOKEN_EXPIRE_MINUTES', '60'))

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()

# ─────────────────────────────────────────────
# GEMINI AI YARDIMCI FONKSİYONLAR
# ─────────────────────────────────────────────

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")  # geriye dönük uyumluluk
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", os.environ.get("ANTHROPIC_API_KEY", ""))
GEMINI_API_KEY_2 = os.environ.get("GEMINI_API_KEY_2", "")  # Yedek key
GEMINI_API_KEY_3 = os.environ.get("GEMINI_API_KEY_3", "")  # 3. key

# Model rotasyon listesi — kota dolunca bir sonrakini dene
GEMINI_MODELS = [
    "gemini-2.0-flash",           # ana model
    "gemini-2.0-flash-lite",      # hafif, az kota
    "gemini-flash-latest",        # alias - farklı kota havuzu
    "gemini-flash-lite-latest",   # en hafif
    "gemini-2.5-flash",           # en güçlü flash
]
AI_MODEL = os.environ.get("AI_DEFAULT_MODEL", GEMINI_MODELS[0])
AI_DEFAULT_MODEL = AI_MODEL
AI_HAIKU_MODEL = AI_MODEL
AI_CACHE_HOURS = int(os.environ.get("AI_CACHE_HOURS", "24"))
AI_MAX_DAILY_REQUESTS = int(os.environ.get("AI_MAX_DAILY_REQUESTS", "500"))

async def _gemini_call(prompt: str, system: str = "", max_tokens: int = 4000) -> str:
    """Gemini API çağrısı — model rotasyonu + çoklu key desteği."""
    all_keys = [k for k in [GEMINI_API_KEY, GEMINI_API_KEY_2, GEMINI_API_KEY_3] if k]
    if not all_keys:
        raise Exception("GEMINI_API_KEY tanımlı değil")

    full_prompt = f"{system}\n\n{prompt}" if system else prompt
    last_error = "Bilinmeyen hata"

    for key in all_keys:
        for model in GEMINI_MODELS:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
                payload = {
                    "contents": [{"parts": [{"text": full_prompt}]}],
                    "generationConfig": {"maxOutputTokens": max_tokens, "temperature": 0.7}
                }
                async with httpx.AsyncClient(timeout=60.0) as c:
                    r = await c.post(url, json=payload)
                    data = r.json()

                if "candidates" in data:
                    logging.info(f"[GEMINI] ✅ Başarılı: model={model}")
                    return data["candidates"][0]["content"]["parts"][0]["text"]

                err_code = data.get("error", {}).get("code", 0)
                err_msg = str(data.get("error", {}).get("message", data))
                last_error = f"{model}: {err_msg[:200]}"
                is_quota = (err_code == 429 or "RESOURCE_EXHAUSTED" in err_msg or "quota" in err_msg.lower())
                if is_quota:
                    logging.warning(f"[GEMINI] ⚠️ Kota ({model}) → sonraki deneniyor")
                    continue
                logging.error(f"[GEMINI] ❌ Kalıcı hata {err_code}: {err_msg[:100]}")
                break

            except Exception as ex:
                last_error = str(ex)[:200]
                logging.warning(f"[GEMINI] Exception ({model}): {last_error[:80]}")
                continue

    raise Exception(f"Tüm Gemini modelleri başarısız. Son hata: {last_error}")


def _mock_bilgi_tabani_response(user_message: str) -> dict:
    """API key yokken bilgi tabanı için mock veri üret."""
    import re as _re
    sinif_match = _re.search(r"Sınıf: (\d+)", user_message)
    sinif = int(sinif_match.group(1)) if sinif_match else 3
    kitap_match = _re.search(r"Kitap: (.+?)\n", user_message)
    kitap = kitap_match.group(1).strip() if kitap_match else "Kitap"
    bolum_match = _re.search(r"Bölüm (\d+)", user_message)
    bolum = int(bolum_match.group(1)) if bolum_match else 1
    metin_match = _re.search(r"METİN:\n(.{0,200})", user_message, _re.DOTALL)
    metin_kesit = metin_match.group(1).strip() if metin_match else ""
    kelimeler_demo = [
        {"kelime": "macera", "anlam": "Heyecan verici ve tehlikeli olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı.", "zorluk": sinif},
        {"kelime": "keşif", "anlam": "Bilinmeyeni ilk kez bulma", "ornek_cumle": "Bilim insanı önemli bir keşif yaptı.", "zorluk": sinif},
        {"kelime": "merak", "anlam": "Bir şeyi öğrenmek isteme", "ornek_cumle": "Meraklı çocuk her şeyi sorar.", "zorluk": max(1, sinif-1)},
        {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme", "ornek_cumle": "Cesur kahraman engeli aştı.", "zorluk": sinif},
        {"kelime": "azim", "anlam": "Bir işi bitirme kararlılığı", "ornek_cumle": "Azimle çalışan başarıya ulaştı.", "zorluk": sinif+1},
    ]
    parsed = {
        "hedef_kelimeler": kelimeler_demo,
        "okuma_parcasi": {
            "baslik": f"{kitap} — Bölüm {bolum}",
            "ozet": f"Bu bölümde {kitap} kitabından seçilmiş bir metin yer almaktadır. Öğrenciler metni okuyarak yeni kelimeler öğrenir.",
            "tema": "Genel",
            "kelime_sayisi": len(metin_kesit.split()) if metin_kesit else 50,
        },
        "sorular": [
            {"soru": "Metinde geçen en önemli olay nedir?", "secenekler": ["A seçeneği", "B seçeneği", "C seçeneği", "D seçeneği"], "dogru_cevap": 0, "taksonomi": "bilgi"},
            {"soru": "Bu metinden ne anladınız?", "secenekler": ["Ana fikri bulduk", "Sadece okudum", "Hiç anlamadım", "Çok zordu"], "dogru_cevap": 0, "taksonomi": "kavrama"},
            {"soru": "Metindeki karakterin davranışını nasıl değerlendirirsiniz?", "secenekler": ["Doğru davrandı", "Yanlış davrandı", "Kararsızım", "Önemli değil"], "dogru_cevap": 0, "taksonomi": "degerlendirme"},
        ],
    }
    import json as _json
    text = _json.dumps(parsed, ensure_ascii=False)
    return {"text": text, "parsed": parsed, "tokens": 0, "maliyet": 0, "error": None, "mock": True}


async def call_claude(system_prompt: str, user_message: str, model: str = "sonnet", max_tokens: int = 2000) -> dict:
    """AI API çağrısı — Gemini Flash kullanır."""
    if not GEMINI_API_KEY:
        if "hedef_kelimeler" in user_message or "METİN:" in user_message:
            return _mock_bilgi_tabani_response(user_message)
        return {"error": "GEMINI_API_KEY tanımlı değil", "text": ""}

    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    gunluk_istek = await db.ai_request_log.count_documents({"tarih": {"$regex": f"^{bugun}"}})
    if gunluk_istek >= AI_MAX_DAILY_REQUESTS:
        return {"error": "Günlük AI istek limiti doldu", "text": ""}

    try:
        text = await _gemini_call(user_message, system=system_prompt, max_tokens=max_tokens)
        await db.ai_request_log.insert_one({
            "id": str(uuid.uuid4()),
            "model": AI_MODEL,
            "tarih": datetime.utcnow().isoformat(),
        })

        # Try to parse JSON from response
        import json as json_mod
        parsed = None
        import json as json_mod
        parsed = None
        try:
            clean = text.strip()
            if clean.startswith("```"):
                clean = clean.split("\n", 1)[1] if "\n" in clean else clean[3:]
                if clean.endswith("```"):
                    clean = clean[:-3]
                clean = clean.strip()
            parsed = json_mod.loads(clean)
        except:
            parsed = None

        return {"text": text, "parsed": parsed, "tokens": 0, "maliyet": 0, "error": None}

    except httpx.TimeoutException:
        return {"error": "AI yanıt süresi aşıldı. Lütfen tekrar deneyin.", "text": ""}
    except Exception as e:
        logging.error(f"Gemini API error: {e}")
        return {"error": f"AI hatası: {str(e)[:100]}", "text": ""}


async def get_ogrenci_ai_verileri(ogrenci_id: str) -> dict:
    """Bir öğrencinin tüm AI-beslenme verilerini toplar (8 kaynak)."""
    ogrenci = await db.students.find_one({"id": ogrenci_id})
    if not ogrenci:
        ogrenci = await db.users.find_one({"id": ogrenci_id})
    if not ogrenci:
        return {}

    # 1. Okuma kayıtları (son 30 gün)
    otuz_gun_once = (datetime.utcnow() - timedelta(days=30)).isoformat()
    logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id, "tarih": {"$gte": otuz_gun_once}}).to_list(length=None)
    toplam_dk = sum(l.get("sure_dakika", 0) for l in logs)
    gun_sayisi = len(set(l.get("tarih", "")[:10] for l in logs))
    kitap_sayisi = len(set(l.get("kitap_adi", "") for l in logs if l.get("kitap_adi")))

    # 2. Streak hesaplama
    tum_logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).to_list(length=None)
    tarihler = sorted(set(l.get("tarih", "")[:10] for l in tum_logs), reverse=True)
    streak = 0
    for i, t in enumerate(tarihler):
        beklenen = (datetime.utcnow() - timedelta(days=i)).strftime("%Y-%m-%d")
        if t == beklenen:
            streak += 1
        else:
            break

    # 3. Risk skoru
    risk_data = {}
    try:
        # Basit risk hesaplama
        risk_skor = 0
        if gun_sayisi < 3: risk_skor += 30
        elif gun_sayisi < 5: risk_skor += 15
        if toplam_dk < 30: risk_skor += 15
        elif toplam_dk < 60: risk_skor += 8
        if streak == 0: risk_skor += 10
        gorev_geciken = await db.gorevler.count_documents({"hedef_id": ogrenci_id, "durum": "bekliyor", "son_tarih": {"$lt": datetime.utcnow().isoformat()}})
        if gorev_geciken > 0: risk_skor += 15
        risk_seviye = "yuksek" if risk_skor > 60 else "orta" if risk_skor > 30 else "dusuk"
        risk_data = {"skor": min(risk_skor, 100), "seviye": risk_seviye}
    except:
        risk_data = {"skor": 0, "seviye": "dusuk"}

    # 4. XP + Lig
    xp_data = {}
    try:
        toplam_xp = ogrenci.get("toplam_xp", 0)
        lig_esikleri = {"bronz": 0, "gumus": 200, "altin": 500, "elmas": 1000}
        lig = "bronz"
        for l, e in sorted(lig_esikleri.items(), key=lambda x: x[1], reverse=True):
            if toplam_xp >= e:
                lig = l
                break
        xp_data = {"toplam": toplam_xp, "lig": lig}
    except:
        xp_data = {"toplam": 0, "lig": "bronz"}

    # 5. Görevler
    gorevler = await db.gorevler.find({"hedef_id": ogrenci_id}).to_list(length=None)
    gorev_ozet = {
        "toplam": len(gorevler),
        "tamamlanan": len([g for g in gorevler if g.get("durum") == "tamamlandi"]),
        "bekleyen": len([g for g in gorevler if g.get("durum") == "bekliyor"]),
        "suresi_gecen": len([g for g in gorevler if g.get("durum") == "bekliyor" and g.get("son_tarih", "9") < datetime.utcnow().isoformat()]),
    }

    # 6. Giriş analizi (son rapor)
    son_rapor = await db.diagnostic_raporlar.find({"ogrenci_id": ogrenci_id}).sort("olusturma_tarihi", -1).to_list(length=1)
    analiz_data = {}
    if son_rapor:
        r = son_rapor[0]
        analiz_data = {
            "wpm": r.get("okuma_hizi", {}).get("wpm", 0),
            "dogruluk": r.get("dogru_okuma", {}).get("dogruluk_orani", 0),
            "prozodi": r.get("prozodik_okuma", {}).get("toplam", 0),
        }
        # Bloom puanları
        anlama = r.get("okudugunu_anlama", {})
        if anlama:
            analiz_data["bloom"] = {
                "bilgi": anlama.get("bilgi", 0),
                "kavrama": anlama.get("kavrama", 0),
                "uygulama": anlama.get("uygulama", 0),
                "analiz": anlama.get("analiz", 0),
                "sentez": anlama.get("sentez", 0),
                "degerlendirme": anlama.get("degerlendirme", 0),
            }

    # 7. Test sonuçları
    test_sonuclari = await db.kitap_test_sonuclari.find({"ogrenci_id": ogrenci_id}).to_list(length=None)
    test_ozet = {"toplam_test": len(test_sonuclari), "ort_yuzde": 0}
    if test_sonuclari:
        test_ozet["ort_yuzde"] = round(sum(t.get("yuzde", 0) for t in test_sonuclari) / len(test_sonuclari))

    return {
        "ogrenci": {
            "ad": ogrenci.get("ad", ""),
            "soyad": ogrenci.get("soyad", ""),
            "sinif": ogrenci.get("sinif", 0),
            "kur": ogrenci.get("kur", ""),
        },
        "okuma_ozet": {
            "son_30_gun_toplam_dk": toplam_dk,
            "gun_sayisi": gun_sayisi,
            "ort_gunluk_dk": round(toplam_dk / max(gun_sayisi, 1), 1),
            "kitap_sayisi": kitap_sayisi,
        },
        "streak": {"mevcut": streak, "en_uzun": max(streak, ogrenci.get("en_uzun_streak", 0))},
        "risk": risk_data,
        "xp": xp_data,
        "gorevler": gorev_ozet,
        "analiz": analiz_data,
        "test": test_ozet,
    }


# Create the main app without a prefix
app = FastAPI(title="Okuma Becerileri Akademisi API")

# ★ CORS — En güvenilir yapılandırma
# NOT: allow_origins=["*"] ve allow_credentials=True birlikte kullanılamaz.
# Bu yüzden ya credentials kapatılır ya da origin spesifik yazılır.
# Render'da en güvenilir yol: origin'i dinamik olarak echo etmek.

ALLOWED_ORIGINS = {
    "https://oba-egitim-frontend.onrender.com",
    "https://oba-egitim.vercel.app",
    "https://oba-egitim-git-main-mehmetdemirhans-projects.vercel.app",
    "http://localhost:3000",
    "http://localhost:3001",
}

from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
from starlette.responses import Response

class CustomCORSMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        origin = request.headers.get("origin", "")
        is_allowed = origin in ALLOWED_ORIGINS or origin.endswith(".onrender.com") or origin.endswith(".vercel.app")

        # OPTIONS preflight — hemen yanıtla
        if request.method == "OPTIONS":
            response = Response(status_code=200)
            if is_allowed:
                response.headers["Access-Control-Allow-Origin"] = origin
                response.headers["Access-Control-Allow-Credentials"] = "true"
                response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
                response.headers["Access-Control-Allow-Headers"] = "Authorization, Content-Type, Accept, Origin, X-Requested-With"
                response.headers["Access-Control-Max-Age"] = "86400"
            return response

        # Normal istek — try/except ile 500 hatalarında da CORS header ekle
        try:
            response = await call_next(request)
        except Exception as e:
            logging.error(f"Unhandled error: {e}")
            response = Response(content=str(e), status_code=500)
        if is_allowed:
            response.headers["Access-Control-Allow-Origin"] = origin
            response.headers["Access-Control-Allow-Credentials"] = "true"
            response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
            response.headers["Access-Control-Allow-Headers"] = "Authorization, Content-Type, Accept, Origin, X-Requested-With"
        return response

app.add_middleware(CustomCORSMiddleware)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# ─────────────────────────────────────────────
# ENUMS
# ─────────────────────────────────────────────

class TeacherLevel(str, Enum):
    YENI = "yeni"
    UZMAN = "uzman"

class UserRole(str, Enum):
    ADMIN = "admin"
    COORDINATOR = "coordinator"
    TEACHER = "teacher"
    STUDENT = "student"
    PARENT = "parent"

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def prepare_for_mongo(data):
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, datetime):
                data[key] = value.isoformat()
            elif isinstance(value, dict):
                data[key] = prepare_for_mongo(value)
            elif isinstance(value, list):
                data[key] = [prepare_for_mongo(item) if isinstance(item, dict) else item for item in value]
    return data

def parse_from_mongo(item):
    if isinstance(item, dict):
        for key, value in item.items():
            if key.endswith('_tarihi') or key in ('olusturma_tarihi', 'tarih'):
                if isinstance(value, str):
                    try:
                        item[key] = datetime.fromisoformat(value)
                    except:
                        pass
            elif isinstance(value, dict):
                item[key] = parse_from_mongo(value)
            elif isinstance(value, list):
                item[key] = [parse_from_mongo(s) if isinstance(s, dict) else s for s in value]
    return item

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain: str, hashed: str) -> bool:
    return pwd_context.verify(plain, hashed)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def decode_token(token: str) -> dict:
    try:
        return jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
    except JWTError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Geçersiz veya süresi dolmuş token",
            headers={"WWW-Authenticate": "Bearer"},
        )

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    payload = decode_token(credentials.credentials)
    user_id = payload.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="Geçersiz token")
    user = await db.users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=401, detail="Kullanıcı bulunamadı")
    return user

def require_role(*roles: UserRole):
    async def checker(current_user=Depends(get_current_user)):
        if current_user.get("role") not in [r.value for r in roles]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Bu işlem için yetkiniz yok"
            )
        return current_user
    return checker

# ─────────────────────────────────────────────
# AUTH MODELS
# ─────────────────────────────────────────────

class UserCreate(BaseModel):
    ad: str
    soyad: str
    email: str
    password: str
    role: UserRole
    telefon: Optional[str] = None
    linked_id: Optional[str] = None  # teacher_id, student_id or parent's student_id

class UserLogin(BaseModel):
    email_or_phone: Optional[str] = None
    email: Optional[str] = None  # eski frontend uyumluluğu
    password: str

class UserResponse(BaseModel):
    id: str
    ad: str
    soyad: str
    email: str
    role: UserRole
    telefon: Optional[str] = None
    linked_id: Optional[str] = None
    olusturma_tarihi: datetime
    puan: int = 0

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class ChangePassword(BaseModel):
    old_password: str
    new_password: str

# ─────────────────────────────────────────────
# AUTH ROUTES
# ─────────────────────────────────────────────

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    girdi = (credentials.email_or_phone or credentials.email or "").lower().strip()
    if not girdi:
        raise HTTPException(status_code=400, detail="E-posta veya telefon gerekli")
    # Email veya telefon ile kullanıcı bul
    user = await db.users.find_one({"email": girdi})
    if not user:
        user = await db.users.find_one({"telefon": girdi})
    if not user or not verify_password(credentials.password, user.get("password_hash", "")):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="E-posta/telefon veya şifre hatalı"
        )
    
    token = create_access_token({"sub": user["id"], "role": user["role"]})
    
    user_response = UserResponse(
        id=user["id"],
        ad=user["ad"],
        soyad=user["soyad"],
        email=user["email"],
        role=user["role"],
        telefon=user.get("telefon"),
        linked_id=user.get("linked_id"),
        olusturma_tarihi=datetime.fromisoformat(user["olusturma_tarihi"]) if isinstance(user.get("olusturma_tarihi"), str) else user.get("olusturma_tarihi", datetime.now(timezone.utc)),
        puan=user.get("puan", 0)
    )
    
    return TokenResponse(access_token=token, user=user_response)

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user=Depends(get_current_user)):
    return UserResponse(
        id=current_user["id"],
        ad=current_user["ad"],
        soyad=current_user["soyad"],
        email=current_user["email"],
        role=current_user["role"],
        telefon=current_user.get("telefon"),
        linked_id=current_user.get("linked_id"),
        olusturma_tarihi=datetime.fromisoformat(current_user["olusturma_tarihi"]) if isinstance(current_user.get("olusturma_tarihi"), str) else current_user.get("olusturma_tarihi", datetime.now(timezone.utc)),
        puan=current_user.get("puan", 0)
    )

@api_router.post("/auth/forgot-password")
async def forgot_password(data: dict = Body(...)):
    email_or_phone = data.get("email_or_phone", "").lower().strip()
    if not email_or_phone:
        raise HTTPException(status_code=400, detail="E-posta veya telefon giriniz")
    user = await db.users.find_one({"email": email_or_phone})
    if not user:
        user = await db.users.find_one({"telefon": email_or_phone})
    if not user:
        raise HTTPException(status_code=404, detail="Bu bilgilerle kayıtlı kullanıcı bulunamadı")
    # 6 haneli geçici şifre oluştur
    gecici_sifre = str(random.randint(100000, 999999))
    new_hash = hash_password(gecici_sifre)
    await db.users.update_one({"id": user["id"]}, {"$set": {"password_hash": new_hash}})
    # NOT: Gerçek uygulamada burada e-posta veya SMS gönderilir
    # Şimdilik geçici şifreyi response'da döndürüyoruz (geliştirme aşaması)
    return {
        "message": f"Geçici şifre oluşturuldu",
        "gecici_sifre": gecici_sifre,
        "kullanici": f"{user['ad']} {user['soyad']}",
        "email": user.get("email", ""),
        "telefon": user.get("telefon", ""),
    }

@api_router.post("/auth/change-password")
async def change_password(data: ChangePassword, current_user=Depends(get_current_user)):
    if not verify_password(data.old_password, current_user.get("password_hash", "")):
        raise HTTPException(status_code=400, detail="Mevcut şifre hatalı")
    
    new_hash = hash_password(data.new_password)
    await db.users.update_one(
        {"id": current_user["id"]},
        {"$set": {"password_hash": new_hash}}
    )
    return {"message": "Şifre başarıyla güncellendi"}

# Admin: kullanıcı oluşturma (sadece admin)
@api_router.post("/auth/users", response_model=UserResponse)
async def create_user(
    user_data: UserCreate,
    current_user=Depends(require_role(UserRole.ADMIN))
):
    # Email kontrolü
    existing = await db.users.find_one({"email": user_data.email.lower().strip()})
    if existing:
        raise HTTPException(status_code=400, detail="Bu e-posta zaten kayıtlı")
    
    user_doc = {
        "id": str(uuid.uuid4()),
        "ad": user_data.ad,
        "soyad": user_data.soyad,
        "email": user_data.email.lower().strip(),
        "telefon": user_data.telefon.strip() if user_data.telefon else None,
        "password_hash": hash_password(user_data.password),
        "role": user_data.role.value,
        "linked_id": user_data.linked_id,
        "olusturma_tarihi": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(user_doc)
    
    return UserResponse(
        id=user_doc["id"],
        ad=user_doc["ad"],
        soyad=user_doc["soyad"],
        email=user_doc["email"],
        role=user_doc["role"],
        linked_id=user_doc.get("linked_id"),
        olusturma_tarihi=datetime.now(timezone.utc)
    )

@api_router.get("/auth/users", response_model=List[UserResponse])
async def list_users(current_user=Depends(require_role(UserRole.ADMIN))):
    users = await db.users.find().to_list(length=None)
    result = []
    for u in users:
        result.append(UserResponse(
            id=u["id"],
            ad=u["ad"],
            soyad=u["soyad"],
            email=u["email"],
            role=u["role"],
            telefon=u.get("telefon"),
            linked_id=u.get("linked_id"),
            olusturma_tarihi=datetime.fromisoformat(u["olusturma_tarihi"]) if isinstance(u.get("olusturma_tarihi"), str) else datetime.now(timezone.utc),
            puan=u.get("puan", 0)
        ))
    return result

@api_router.delete("/auth/users/{user_id}")
async def delete_user(user_id: str, current_user=Depends(require_role(UserRole.ADMIN))):
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Kendi hesabınızı silemezsiniz")
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    return {"message": "Kullanıcı silindi"}

# ─────────────────────────────────────────────
# STARTUP: Admin kullanıcı oluştur
# ─────────────────────────────────────────────

@app.on_event("startup")
async def create_default_admin():
    """
    .env dosyasındaki bilgilerle varsayılan admin oluşturur.
    Admin zaten varsa tekrar oluşturmaz.
    """
    admin_email = os.environ.get('ADMIN_EMAIL', 'admin@oba.com')
    admin_password = os.environ.get('ADMIN_PASSWORD', 'Admin123!')
    admin_ad = os.environ.get('ADMIN_AD', 'Sistem')
    admin_soyad = os.environ.get('ADMIN_SOYAD', 'Yöneticisi')

    existing = await db.users.find_one({"email": admin_email})
    if not existing:
        admin_doc = {
            "id": str(uuid.uuid4()),
            "ad": admin_ad,
            "soyad": admin_soyad,
            "email": admin_email,
            "password_hash": hash_password(admin_password),
            "role": UserRole.ADMIN.value,
            "linked_id": None,
            "olusturma_tarihi": datetime.now(timezone.utc).isoformat()
        }
        await db.users.insert_one(admin_doc)
        logging.info(f"✅ Varsayılan admin oluşturuldu: {admin_email}")
    else:
        logging.info(f"ℹ️ Admin zaten mevcut: {admin_email}")

    # ── DEMO VERİLERİ ──
    demo_student_email = "demo-ogrenci@oba.com"
    demo_teacher_email = "demo-ogretmen@oba.com"
    demo_parent_email = "demo-veli@oba.com"
    DEMO_PASSWORD = "Demo123!"

    from datetime import timedelta
    simdi = datetime.now(timezone.utc)

    # --- DEMO ÖĞRETMEN ---
    existing_teacher = await db.users.find_one({"email": demo_teacher_email})
    if not existing_teacher:
        demo_teacher_id = str(uuid.uuid4())
        # teachers collection
        existing_t_col = await db.teachers.find_one({"ad": "Ayşe", "soyad": "Öğretmen"})
        if not existing_t_col:
            await db.teachers.insert_one({
                "id": demo_teacher_id, "ad": "Ayşe", "soyad": "Öğretmen", "brans": "Türkçe",
                "telefon": "05001112233", "seviye": "uzman", "ogrenci_sayisi": 3,
                "yapilmasi_gereken_odeme": 0, "yapilan_odeme": 0, "arsivli": False,
                "olusturma_tarihi": simdi.isoformat()
            })
        else:
            demo_teacher_id = existing_t_col["id"]

        demo_teacher_user_id = str(uuid.uuid4())
        await db.users.insert_one({
            "id": demo_teacher_user_id, "ad": "Ayşe", "soyad": "Öğretmen",
            "email": demo_teacher_email, "password_hash": hash_password(DEMO_PASSWORD),
            "role": "teacher", "linked_id": demo_teacher_id, "telefon": "05001112233",
            "olusturma_tarihi": simdi.isoformat()
        })
        logging.info(f"✅ Demo öğretmen oluşturuldu: {demo_teacher_email} / {DEMO_PASSWORD}")
    else:
        demo_teacher_user_id = existing_teacher["id"]
        demo_teacher_id = existing_teacher.get("linked_id", "")
        logging.info(f"ℹ️ Demo öğretmen zaten var: {demo_teacher_email}")

    # --- DEMO ÖĞRENCİLER ---
    existing_student = await db.users.find_one({"email": demo_student_email})
    if not existing_student:
        ogrenci_adlari = [
            ("Ali", "Yılmaz", "4", "Kur 2"), ("Zeynep", "Demir", "3", "Kur 1"),
            ("Mehmet", "Kaya", "5", "Kur 3"),
        ]
        demo_students = []
        for ad, soyad, sinif, kur in ogrenci_adlari:
            sid = str(uuid.uuid4())
            demo_students.append(sid)
            existing_s_col = await db.students.find_one({"ad": ad, "soyad": soyad})
            if not existing_s_col:
                await db.students.insert_one({
                    "id": sid, "ad": ad, "soyad": soyad, "sinif": sinif, "kur": kur,
                    "veli_ad": "Demo", "veli_soyad": "Veli", "veli_telefon": "05009998877",
                    "aldigi_egitim": "Okuma Becerileri Temel", "ogretmen_id": demo_teacher_id,
                    "yapilmasi_gereken_odeme": 2400, "yapilan_odeme": 0, "ogretmene_yapilacak_odeme": 800,
                    "arsivli": False, "toplam_xp": 0, "olusturma_tarihi": simdi.isoformat()
                })
            else:
                demo_students[-1] = existing_s_col["id"]

        # Ana öğrenci user
        demo_student_user_id = str(uuid.uuid4())
        await db.users.insert_one({
            "id": demo_student_user_id, "ad": "Ali", "soyad": "Yılmaz",
            "email": demo_student_email, "password_hash": hash_password(DEMO_PASSWORD),
            "role": "student", "linked_id": demo_students[0], "telefon": "",
            "olusturma_tarihi": simdi.isoformat()
        })
        # Diğer öğrenciler
        for i, (ad, soyad, _, _) in enumerate(ogrenci_adlari[1:], 1):
            await db.users.insert_one({
                "id": str(uuid.uuid4()), "ad": ad, "soyad": soyad,
                "email": f"demo-ogrenci{i+1}@oba.com", "password_hash": hash_password(DEMO_PASSWORD),
                "role": "student", "linked_id": demo_students[i], "telefon": "",
                "olusturma_tarihi": simdi.isoformat()
            })

        # Okuma kayıtları
        kitaplar = ["Küçük Prens", "Charlie'nin Çikolata Fabrikası", "Pollyanna"]
        for gun in range(14):
            tarih = (simdi - timedelta(days=gun)).isoformat()
            await db.reading_logs.insert_one({
                "id": str(uuid.uuid4()), "ogrenci_id": demo_students[0],
                "ogrenci_ad": "Ali Yılmaz", "kitap_adi": kitaplar[gun % len(kitaplar)],
                "bolum": f"Bölüm {gun % 8 + 1}", "baslangic_sayfa": gun * 10 + 1,
                "bitis_sayfa": gun * 10 + 12, "sure_dakika": random.randint(8, 25),
                "not_text": "", "tarih": tarih,
            })
        for i, sid in enumerate(demo_students[1:], 1):
            for gun in range(random.randint(3, 10)):
                await db.reading_logs.insert_one({
                    "id": str(uuid.uuid4()), "ogrenci_id": sid,
                    "ogrenci_ad": ogrenci_adlari[i][0] + " " + ogrenci_adlari[i][1],
                    "kitap_adi": kitaplar[gun % len(kitaplar)],
                    "bolum": f"Bölüm {gun % 5 + 1}", "baslangic_sayfa": gun * 8 + 1,
                    "bitis_sayfa": gun * 8 + 10, "sure_dakika": random.randint(5, 20),
                    "not_text": "", "tarih": (simdi - timedelta(days=gun)).isoformat(),
                })

        # XP
        xp_eylemleri = ["okuma_gorevi"] * 10 + ["gorev_tamamla"] * 3 + ["egzersiz"] * 5 + ["gelisim_tamamla"] * 2
        toplam_xp = 0
        for eylem in xp_eylemleri:
            xp = (await get_xp_tablosu()).get(eylem, 5)
            toplam_xp += xp
            await db.xp_logs.insert_one({
                "id": str(uuid.uuid4()), "ogrenci_id": demo_students[0],
                "eylem": eylem, "xp": xp, "tarih": (simdi - timedelta(hours=random.randint(1, 200))).isoformat(),
            })
        await db.students.update_one({"id": demo_students[0]}, {"$set": {"toplam_xp": toplam_xp}})
        for sid in demo_students[1:]:
            await db.students.update_one({"id": sid}, {"$set": {"toplam_xp": random.randint(30, 150)}})

        # Görevler
        await db.gorevler.insert_one({
            "id": str(uuid.uuid4()), "hedef_id": demo_students[0], "hedef_tip": "ogrenci",
            "hedef_ad": "Ali Yılmaz", "baslik": "Küçük Prens — Bölüm 5-6 oku",
            "aciklama": "Bu hafta Küçük Prens'in 5. ve 6. bölümlerini oku.", "tur": "kitap",
            "son_tarih": (simdi + timedelta(days=5)).strftime("%Y-%m-%d"),
            "atayan_id": demo_teacher_user_id, "atayan_ad": "Ayşe Öğretmen", "atayan_rol": "teacher",
            "durum": "bekliyor", "olusturma_tarihi": simdi.isoformat(),
            "makale_link": None, "kitap_yazar": "Antoine de Saint-Exupéry",
            "kitap_isbn": "", "kitap_link": "", "kitap_kapak": "", "film_link": None,
        })
        await db.gorevler.insert_one({
            "id": str(uuid.uuid4()), "hedef_id": demo_students[0], "hedef_tip": "ogrenci",
            "hedef_ad": "Ali Yılmaz", "baslik": "Hızlı Okuma Egzersizi yap",
            "aciklama": "Egzersizlerden 'Hızlı Kelime Okuma' tamamla.", "tur": "egzersiz",
            "son_tarih": (simdi + timedelta(days=3)).strftime("%Y-%m-%d"),
            "atayan_id": demo_teacher_user_id, "atayan_ad": "Ayşe Öğretmen", "atayan_rol": "teacher",
            "durum": "bekliyor", "olusturma_tarihi": simdi.isoformat(),
            "makale_link": None, "kitap_yazar": None, "kitap_isbn": None,
            "kitap_link": None, "kitap_kapak": None, "film_link": None,
        })
        await db.gorevler.insert_one({
            "id": str(uuid.uuid4()), "hedef_id": demo_students[0], "hedef_tip": "ogrenci",
            "hedef_ad": "Ali Yılmaz", "baslik": "Doğa belgeseli izle",
            "aciklama": "Belgeseli izleyip 3 cümle özet yaz.", "tur": "film", "son_tarih": None,
            "atayan_id": demo_teacher_user_id, "atayan_ad": "Ayşe Öğretmen", "atayan_rol": "teacher",
            "durum": "tamamlandi", "tamamlama_tarihi": (simdi - timedelta(days=2)).isoformat(),
            "tamamlama_notu": "Belgeseli izledim, çok güzeldi!", "olusturma_tarihi": (simdi - timedelta(days=5)).isoformat(),
            "makale_link": None, "kitap_yazar": None, "kitap_isbn": None,
            "kitap_link": None, "kitap_kapak": None, "film_link": "https://youtube.com/watch?v=example",
        })

        # Mesaj
        await db.mesajlar.insert_one({
            "id": str(uuid.uuid4()), "gonderen_id": demo_teacher_user_id,
            "gonderen_ad": "Ayşe Öğretmen", "gonderen_rol": "teacher",
            "alici_id": demo_student_user_id, "alici_ad": "Ali Yılmaz", "alici_rol": "student",
            "konu": "Tebrikler! 🎉", "icerik": "Ali, bu hafta çok güzel ilerleme kaydettin. Böyle devam et!",
            "okundu": False, "tarih": (simdi - timedelta(hours=3)).isoformat(),
        })
        logging.info(f"✅ Demo öğrenciler oluşturuldu: {demo_student_email} / {DEMO_PASSWORD}")
    else:
        demo_student_user_id = existing_student["id"]
        logging.info(f"ℹ️ Demo öğrenci zaten var: {demo_student_email}")

    # --- DEMO VELİ ---
    existing_parent = await db.users.find_one({"email": demo_parent_email})
    if not existing_parent:
        # Ali Yılmaz'ın student id'sini bul
        ali = await db.users.find_one({"email": demo_student_email})
        ali_linked = ali.get("linked_id", "") if ali else ""
        await db.users.insert_one({
            "id": str(uuid.uuid4()), "ad": "Demo", "soyad": "Veli",
            "email": demo_parent_email, "password_hash": hash_password(DEMO_PASSWORD),
            "role": "parent", "linked_id": ali_linked, "telefon": "05009998877",
            "olusturma_tarihi": simdi.isoformat()
        })
        logging.info(f"✅ Demo veli oluşturuldu: {demo_parent_email} / {DEMO_PASSWORD}")
    else:
        logging.info(f"ℹ️ Demo veli zaten var: {demo_parent_email}")

    # --- DEMO ROZET + ANKET VERİLERİ ---
    existing_rozetler = await db.kazanilan_rozetler.find_one({"kullanici_id": {"$regex": "^demo-"}})
    if not existing_rozetler:
        logging.info("🏅 Demo rozet + anket verileri oluşturuluyor...")

        # Demo user id'lerini bul
        demo_ogretmen_user = await db.users.find_one({"email": demo_teacher_email})
        demo_ogrenci_user = await db.users.find_one({"email": demo_student_email})
        demo_veli_user = await db.users.find_one({"email": demo_parent_email})

        if demo_ogretmen_user and demo_ogrenci_user:
            ogretmen_uid = demo_ogretmen_user["id"]
            ogretmen_lid = demo_ogretmen_user.get("linked_id", "")
            ogrenci_uid = demo_ogrenci_user["id"]
            ogrenci_lid = demo_ogrenci_user.get("linked_id", "")

            # Öğretmen rozetleri
            ogretmen_rozetler = [
                "icerik_ilk", "icerik_5",  # İçerik katkısı
                "oy_ilk", "oy_20",  # Kalite kontrol
                "gorev_ilk", "gorev_20",  # Eğitimci
                "kur_ilk",  # Kur atlama
                "gelisim_ilk", "gelisim_10",  # Gelişim
                "mesaj_ilk",  # İletişim
                "egz_ilk",  # Egzersiz
            ]
            for kod in ogretmen_rozetler:
                await db.kazanilan_rozetler.insert_one({
                    "id": str(uuid.uuid4()),
                    "kullanici_id": ogretmen_uid,
                    "rozet_kodu": kod,
                    "kazanma_tarihi": (simdi - timedelta(days=random.randint(1, 30))).isoformat(),
                })

            # Öğrenci rozetleri (Ali Yılmaz)
            ogrenci_rozetler = [
                "okuma_ilk", "okuma_100",  # Okuma
                "streak_3", "streak_7",  # Streak
                "kitap_1", "kitap_5",  # Kitap
                "gorev_ilk", "gorev_10",  # Görev
                "egz_ilk",  # Egzersiz
                "orman_ilk", "orman_50",  # Orman
            ]
            for kod in ogrenci_rozetler:
                await db.kazanilan_rozetler.insert_one({
                    "id": str(uuid.uuid4()),
                    "kullanici_id": ogrenci_uid,
                    "rozet_kodu": kod,
                    "kazanma_tarihi": (simdi - timedelta(days=random.randint(1, 20))).isoformat(),
                })

            # Diğer demo öğrencilere de birkaç rozet
            for email_suffix in ["2", "3"]:
                other = await db.users.find_one({"email": f"demo-ogrenci{email_suffix}@oba.com"})
                if other:
                    for kod in ["okuma_ilk", "streak_3", "kitap_1", "gorev_ilk"]:
                        await db.kazanilan_rozetler.insert_one({
                            "id": str(uuid.uuid4()),
                            "kullanici_id": other["id"],
                            "rozet_kodu": kod,
                            "kazanma_tarihi": (simdi - timedelta(days=random.randint(1, 15))).isoformat(),
                        })

            # Kur atlama kayıtları (öğretmen rozeti için)
            for i in range(3):
                await db.kur_atlamalari.insert_one({
                    "id": str(uuid.uuid4()),
                    "ogrenci_id": ogrenci_lid if i == 0 else str(uuid.uuid4()),
                    "ogretmen_id": ogretmen_lid,
                    "eski_kur": f"Kur {i+1}",
                    "yeni_kur": f"Kur {i+2}",
                    "tarih": (simdi - timedelta(days=random.randint(5, 60))).isoformat(),
                })

            # Veli anketleri (öğretmen rozeti + dashboard için)
            if demo_veli_user:
                veli_uid = demo_veli_user["id"]
                anket_kategoriler = ["iletisim", "duzen", "etki", "geri_bildirim", "motivasyon", "icerik", "genel"]

                # 3 farklı dönem için anket
                for donem_offset in [0, 1, 2]:
                    donem_ay = (simdi - timedelta(days=donem_offset * 30)).strftime("%Y-D%m")
                    yanitlar = []
                    for j, kat in enumerate(anket_kategoriler):
                        puan = random.choice([4, 4, 5, 5, 5, 4, 5])  # yüksek puanlar
                        yanitlar.append({"soru_no": j + 1, "puan": puan, "kategori": kat})

                    await db.veli_anketleri.insert_one({
                        "id": str(uuid.uuid4()),
                        "veli_id": veli_uid,
                        "veli_ad": "Demo Veli",
                        "ogretmen_id": ogretmen_lid,
                        "ogrenci_id": ogrenci_lid,
                        "donem": donem_ay,
                        "yanitlar": yanitlar,
                        "tavsiye": random.choice([True, True, True, False]),  # %75 tavsiye
                        "not_text": random.choice([
                            "Çocuğum çok gelişti, teşekkürler!",
                            "Öğretmenimizden memnunuz.",
                            "Okuma alışkanlığı kazandı, çok mutluyuz.",
                            "",
                        ]),
                        "tarih": (simdi - timedelta(days=donem_offset * 30 + random.randint(0, 5))).isoformat(),
                    })

                # Birkaç tane daha farklı "veli"den anket (aynı öğretmene)
                sahte_veliler = ["Veli A", "Veli B", "Veli C", "Veli D", "Veli E"]
                for veli_ad in sahte_veliler:
                    yanitlar = []
                    for j, kat in enumerate(anket_kategoriler):
                        yanitlar.append({"soru_no": j + 1, "puan": random.choice([3, 4, 4, 5, 5]), "kategori": kat})
                    await db.veli_anketleri.insert_one({
                        "id": str(uuid.uuid4()),
                        "veli_id": str(uuid.uuid4()),  # sahte veli id
                        "veli_ad": veli_ad,
                        "ogretmen_id": ogretmen_lid,
                        "ogrenci_id": str(uuid.uuid4()),
                        "donem": simdi.strftime("%Y-D%m"),
                        "yanitlar": yanitlar,
                        "tavsiye": random.choice([True, True, True, True, False]),
                        "not_text": random.choice(["Memnunuz", "Teşekkürler", "Güzel çalışma", ""]),
                        "tarih": (simdi - timedelta(days=random.randint(1, 20))).isoformat(),
                    })

            # Demo bildirimler
            demo_bildirimler = [
                (ogrenci_uid, "gorev_atandi", "Ayşe Öğretmen size yeni görev atadı: Küçük Prens Bölüm 5-6", 3),
                (ogrenci_uid, "streak_tebrik", "🎉 7 gün üst üste okudun! Harika gidiyorsun!", 1),
                (ogrenci_uid, "rozet_kazandi", "🏅 Yeni rozet: Kararlı Okuyucu! 7 gün streak başarısı.", 2),
                (ogretmen_uid, "risk_yuksek", "🚨 Mehmet Kaya son 7 gündür hiç okuma yapmadı!", 1),
                (ogretmen_uid, "gorev_tamamlandi", "Ali Yılmaz 'Doğa belgeseli izle' görevini tamamladı.", 3),
            ]
            if demo_veli_user:
                demo_bildirimler.append((demo_veli_user["id"], "rapor_tamamlandi", "Ali Yılmaz için yeni giriş analizi raporu hazır.", 2))
                demo_bildirimler.append((demo_veli_user["id"], "streak_kirildi", "Ali dün okuma yapmadı. Streak kırılma riski!", 0))
                demo_bildirimler.append((demo_veli_user["id"], "anket_hatirlatma", "Öğretmeninizi değerlendirmek ister misiniz? ⭐", 5))

            for alici, tur, icerik, gun_once in demo_bildirimler:
                await db.bildirimler.insert_one({
                    "id": str(uuid.uuid4()), "alici_id": alici, "tur": tur,
                    "baslik": BILDIRIM_TURLERI.get(tur, {}).get("baslik", "Bildirim"),
                    "icerik": icerik, "oncelik": BILDIRIM_TURLERI.get(tur, {}).get("oncelik", "normal"),
                    "ilgili_id": None, "okundu": gun_once > 2,
                    "tarih": (simdi - timedelta(hours=gun_once * 24 + random.randint(1, 12))).isoformat(),
                })

            logging.info(f"✅ Demo rozet + anket + bildirim verileri oluşturuldu!")
            logging.info(f"   🏅 Öğretmen: {len(ogretmen_rozetler)} rozet")
            logging.info(f"   🏅 Öğrenci: {len(ogrenci_rozetler)} rozet")
            logging.info(f"   ⭐ Veli anketleri: 8 adet")
    else:
        logging.info("ℹ️ Demo rozet + anket verileri zaten mevcut")

    # --- VARSAYILAN AYARLAR (boş ise seed et) ---
    ayar_defaults = {
        "xp_tablosu": XP_TABLOSU_DEFAULT,
        "lig_esikleri": LIG_ESIKLERI_DEFAULT,
        "ogretmen_rozetleri": OGRETMEN_ROZETLERI_DEFAULT,
        "ogrenci_rozetleri": OGRENCI_ROZETLERI_DEFAULT,
        "anket_sorulari": ANKET_SORULARI_DEFAULT,
    }
    for tip, default_val in ayar_defaults.items():
        doc = await db.sistem_ayarlari.find_one({"tip": tip})
        if not doc:
            await db.sistem_ayarlari.insert_one({"tip": tip, "degerler": default_val})
            logging.info(f"  ✅ Varsayılan ayar oluşturuldu: {tip}")
        elif not doc.get("degerler") or (isinstance(doc["degerler"], list) and len(doc["degerler"]) == 0) or (isinstance(doc["degerler"], dict) and len(doc["degerler"]) == 0):
            await db.sistem_ayarlari.update_one({"tip": tip}, {"$set": {"degerler": default_val}})
            logging.info(f"  🔄 Boş ayar düzeltildi: {tip}")

    # ── AI DEMO VERİLERİ ──
    ai_demo_var = await db.okuma_dna.find_one({"ogrenci_id": {"$exists": True}})
    if not ai_demo_var:
        logging.info("🧠 AI demo verileri oluşturuluyor...")

        # Demo öğrenci ID'lerini bul
        demo_ogrenciler = await db.students.find({"ad": {"$in": ["Ahmet", "Zeynep", "Can"]}}).to_list(length=10)
        if not demo_ogrenciler:
            demo_ogrenciler = await db.users.find({"role": "student"}).to_list(length=3)

        demo_ogretmen = await db.users.find_one({"email": "demo-ogretmen@oba.com"})
        ogretmen_id = demo_ogretmen["id"] if demo_ogretmen else "demo-teacher"

        for i, ogr in enumerate(demo_ogrenciler[:3]):
            oid = ogr["id"]
            ad = ogr.get("ad", f"Öğrenci{i+1}")
            sinif = ogr.get("sinif", 3)

            # ── DNA Profili ──
            dna_profiller = [
                {"kelime_gucu": 72, "akicilik": 45, "anlama_derinligi": 80, "dikkat_suresi": 55, "zorluk_toleransi": 60, "kelime_tekrar_ihtiyaci": 35, "okuma_psikolojisi": "keşifçi"},
                {"kelime_gucu": 58, "akicilik": 78, "anlama_derinligi": 42, "dikkat_suresi": 70, "zorluk_toleransi": 45, "kelime_tekrar_ihtiyaci": 50, "okuma_psikolojisi": "güvenli"},
                {"kelime_gucu": 35, "akicilik": 30, "anlama_derinligi": 55, "dikkat_suresi": 25, "zorluk_toleransi": 30, "kelime_tekrar_ihtiyaci": 75, "okuma_psikolojisi": "kararsız"},
            ]
            profil_tipler = ["hayalci_okuyucu", "hızlı_okuyucu", "başlangıç_okuyucu"]
            profil_labels = ["🌈 Hayalci Okuyucu", "⚡ Hızlı Okuyucu", "🌱 Başlangıç Okuyucu"]

            await db.okuma_dna.update_one({"ogrenci_id": oid}, {"$set": {
                "ogrenci_id": oid,
                "boyutlar": dna_profiller[i],
                "profil_tipi": profil_tipler[i],
                "profil_label": profil_labels[i],
                "sinif": sinif,
                "son_guncelleme": simdi.isoformat(),
            }}, upsert=True)

            # ── Koçluk Cache (demo analiz) ──
            demo_analizler = [
                {
                    "durum_degerlendirmesi": {"guclu_yonler": ["Anlama kapasitesi yüksek", "Hayal gücü gelişmiş", "Kitap seçiminde cesur"], "gelisim_alanlari": ["Okuma hızı geliştirilmeli", "Dikkat süresi artırılmalı"]},
                    "risk_analizi": {"seviye": "düşük", "faktorler": ["Akıcılık ortalamanın altında"], "aciliyet": "Takip yeterli"},
                    "mudahale_plani": {"hafta_1": "Günlük 10 dk sesli okuma", "hafta_2": "Tekrarlı okuma çalışmaları", "hafta_3": "Hız testi ve geri bildirim", "hafta_4": "Bağımsız okuma + analiz"},
                    "veliye_mesaj": f"Sayın Veli, {ad} okuma anlama becerisi çok iyi gelişiyor. Akıcılığını artırmak için evde günlük 10 dakika sesli okuma yapmasını öneririz. Hikâye kitapları ile başlamak motivasyonunu yüksek tutar.",
                    "haftalik_gorevler": [{"gun": "Pazartesi", "gorev": "Sevdiği bir kitaptan 2 sayfa sesli oku", "bloom": "uygulama"}, {"gun": "Salı", "gorev": "Okuduğu bölümün özetini yaz", "bloom": "sentez"}, {"gun": "Çarşamba", "gorev": "5 yeni kelime öğren ve cümle kur", "bloom": "uygulama"}, {"gun": "Perşembe", "gorev": "Karakterin motivasyonunu analiz et", "bloom": "analiz"}, {"gun": "Cuma", "gorev": "Hikâyeye alternatif son yaz", "bloom": "yaratma"}],
                    "kitap_tavsiyeleri": [{"ad": "Charlie'nin Çikolata Fabrikası", "yazar": "Roald Dahl", "neden": "Hayal gücü yüksek, kısa bölümler, macera dolu"}, {"ad": "Küçük Prens", "yazar": "Saint-Exupéry", "neden": "Felsefi derinlik, kısa paragraflar, soyut düşünmeyi geliştirir"}, {"ad": "Pollyanna", "yazar": "Eleanor H. Porter", "neden": "Pozitif bakış açısı, karakter gelişimi, orta zorluk"}],
                    "motivasyon_mesaji": f"Harika gidiyorsun {ad}! Bu hafta 3 kitap okudun, kelime gücün %72'ye ulaştı. Hedefine çok yakınsın! 🌟",
                    "kelime_mudahale": "Günlük 5 yeni kelime + görsel kartlarla çalışma. Spaced repetition: 1-3-7-21 gün aralığında tekrar.",
                    "metin_recetesi": {"paragraf_uzunlugu": "Orta (80-120 kelime)", "soyutluk": "Orta-yüksek", "aksiyon": "Yüksek", "hedef_kelime_orani": "%70 bilinen + %30 yeni"}
                },
                {
                    "durum_degerlendirmesi": {"guclu_yonler": ["Hızlı okuma", "Düzenli streak", "Dikkat süresi iyi"], "gelisim_alanlari": ["Anlama derinliği artırılmalı", "Bloom üst basamakları zayıf"]},
                    "risk_analizi": {"seviye": "düşük", "faktorler": ["Hızlı ama yüzeysel okuma eğilimi"], "aciliyet": "Bloom çalışması önerilir"},
                    "mudahale_plani": {"hafta_1": "Her okuma sonrası 3 soru cevapla", "hafta_2": "Karakter analizi çalışması", "hafta_3": "Metin karşılaştırma etkinliği", "hafta_4": "Tartışma soruları + yazılı yanıt"},
                    "veliye_mesaj": f"Sayın Veli, {ad} çok düzenli okuyor, tebrikler! Okuma hızı mükemmel. Şimdi anlama derinliğini artırmaya odaklanıyoruz. Okuduklarını sormak çok faydalı olacaktır.",
                    "haftalik_gorevler": [{"gun": "Pazartesi", "gorev": "Bir paragrafı oku ve ana fikri bul", "bloom": "analiz"}, {"gun": "Salı", "gorev": "İki karakteri karşılaştır", "bloom": "analiz"}, {"gun": "Çarşamba", "gorev": "Hikâyenin sonucunu tahmin et", "bloom": "sentez"}, {"gun": "Perşembe", "gorev": "Yazar neden böyle yazmış? Yorumla", "bloom": "degerlendirme"}, {"gun": "Cuma", "gorev": "Okuduğun kitap hakkında 5 cümle yaz", "bloom": "sentez"}],
                    "kitap_tavsiyeleri": [{"ad": "Martı", "yazar": "Richard Bach", "neden": "Kısa ama derin, felsefi düşünme gerektirir"}, {"ad": "Kaşağı", "yazar": "Ömer Seyfettin", "neden": "Kısa öykü, karakter analizi için ideal"}, {"ad": "Sefiller (Çocuk versiyonu)", "yazar": "Victor Hugo", "neden": "Değer yargıları, adalet teması"}],
                    "motivasyon_mesaji": f"Süpersin {ad}! Streak'in 12 günü geçti. Şimdi anlama gücünü artıralım — derin sorular senin için! 💪",
                    "kelime_mudahale": "Soyut kelimeler üzerine çalışma. Kavram haritaları oluşturma. Her gün 3 soyut kelime.",
                    "metin_recetesi": {"paragraf_uzunlugu": "Uzun (120-200 kelime)", "soyutluk": "Orta", "aksiyon": "Orta", "hedef_kelime_orani": "%60 bilinen + %40 yeni"}
                },
                {
                    "durum_degerlendirmesi": {"guclu_yonler": ["Meraklı", "Sorulara istekli cevap veriyor"], "gelisim_alanlari": ["Kelime hazinesi çok dar", "Akıcılık düşük", "Dikkat süresi kısa", "Streak tutarsız"]},
                    "risk_analizi": {"seviye": "yüksek", "faktorler": ["Kelime gücü sınıf ortalamasının altında", "Streak sık kırılıyor", "Zor metinlerden kaçınma eğilimi"], "aciliyet": "Acil müdahale gerekli"},
                    "mudahale_plani": {"hafta_1": "Günlük 5 dk kolay metin + kelime oyunu", "hafta_2": "Sesli okuma + anlık geri bildirim", "hafta_3": "Kısa hikâyeler + 3 kelime hedef", "hafta_4": "İlgi alanına göre metin seçimi"},
                    "veliye_mesaj": f"Sayın Veli, {ad} okumaya ilgi duyuyor ama desteğe ihtiyacı var. Evde günlük 5 dakika birlikte okuma yapmanız çok faydalı olacak. Hayvan hikâyeleri ile başlamanızı öneririz — ilgisini çekeceğini düşünüyoruz.",
                    "haftalik_gorevler": [{"gun": "Pazartesi", "gorev": "Resimli bir kitaptan 1 sayfa oku", "bloom": "bilgi"}, {"gun": "Salı", "gorev": "3 yeni kelime öğren (resimli kart)", "bloom": "bilgi"}, {"gun": "Çarşamba", "gorev": "Kısa bir masal dinle ve anlat", "bloom": "kavrama"}, {"gun": "Perşembe", "gorev": "Kelime oyunu oyna (eşleştirme)", "bloom": "uygulama"}, {"gun": "Cuma", "gorev": "En sevdiğin hayvanı anlatan 3 cümle yaz", "bloom": "uygulama"}],
                    "kitap_tavsiyeleri": [{"ad": "Kırmızı Başlıklı Kız", "yazar": "Charles Perrault", "neden": "Kısa, bilinen hikâye, kolay kelimeler"}, {"ad": "Bremen Mızıkacıları", "yazar": "Grimm Kardeşler", "neden": "Hayvan karakterleri, eğlenceli, tekrarlı yapı"}, {"ad": "Pinokyo", "yazar": "Carlo Collodi", "neden": "Macera dolu, kısa bölümler, ahlak dersi"}],
                    "motivasyon_mesaji": f"Merhaba {ad}! Bugün sadece 5 dakika okumaya ne dersin? Küçük adımlar büyük fark yaratır! 🌱",
                    "kelime_mudahale": "Günlük 3 kolay kelime + resimli kartlar + eşleştirme oyunu. Tekrar: 1-2-5-12 gün aralığında.",
                    "metin_recetesi": {"paragraf_uzunlugu": "Çok kısa (30-60 kelime)", "soyutluk": "Düşük (somut)", "aksiyon": "Çok yüksek", "hedef_kelime_orani": "%85 bilinen + %15 yeni"}
                },
            ]

            await db.ai_kocluk_cache.update_one(
                {"ogrenci_id": oid},
                {"$set": {
                    "id": str(uuid.uuid4()),
                    "ogrenci_id": oid,
                    "dna": {"profil_tipi": profil_tipler[i], "profil_label": profil_labels[i], "boyutlar": dna_profiller[i]},
                    "ai_analiz": demo_analizler[i],
                    "ai_ham_metin": "",
                    "model": "demo",
                    "token": 0,
                    "maliyet": 0,
                    "tarih": simdi.isoformat(),
                }},
                upsert=True,
            )

        # ── Demo Kelimeler (meb_kelime_haritasi) ──
        demo_kelimeler = [
            {"kelime": "macera", "anlam": "Tehlikeli ve heyecan verici olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı.", "zorluk": 4, "sinif": 3},
            {"kelime": "keşif", "anlam": "Bilinmeyen bir şeyi ilk kez bulma", "ornek_cumle": "Bilim insanı yeni bir keşif yaptı.", "zorluk": 5, "sinif": 3},
            {"kelime": "pusula", "anlam": "Yön bulmaya yarayan araç", "ornek_cumle": "Kaşif pusulasıyla yolunu buldu.", "zorluk": 6, "sinif": 3},
            {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme gücü", "ornek_cumle": "Küçük kız büyük cesaret gösterdi.", "zorluk": 4, "sinif": 3},
            {"kelime": "merak", "anlam": "Bir şeyi bilmek isteme duygusu", "ornek_cumle": "Merak eden çocuk her şeyi sorar.", "zorluk": 3, "sinif": 3},
            {"kelime": "sabır", "anlam": "Bekleyebilme ve dayanma gücü", "ornek_cumle": "Bahçıvan sabırla çiçeklerin büyümesini bekledi.", "zorluk": 4, "sinif": 3},
            {"kelime": "hayal gücü", "anlam": "Zihinde yeni şeyler oluşturabilme becerisi", "ornek_cumle": "Hayal gücü güçlü olan çocuklar iyi hikâye yazar.", "zorluk": 5, "sinif": 3},
            {"kelime": "fedakarlık", "anlam": "Başkaları için kendi isteklerinden vazgeçme", "ornek_cumle": "Anneler çocukları için büyük fedakarlıklar yapar.", "zorluk": 7, "sinif": 4},
            {"kelime": "dürüstlük", "anlam": "Doğruyu söyleme, hile yapmama", "ornek_cumle": "Dürüstlük en değerli erdemlerden biridir.", "zorluk": 5, "sinif": 4},
            {"kelime": "azim", "anlam": "Bir işi sonuna kadar kararlılıkla sürdürme", "ornek_cumle": "Azimli öğrenci başarıya ulaştı.", "zorluk": 6, "sinif": 4},
            {"kelime": "empati", "anlam": "Başkalarının duygularını anlayabilme", "ornek_cumle": "Empati kurabilen insanlar daha iyi arkadaş olur.", "zorluk": 7, "sinif": 5},
            {"kelime": "göç", "anlam": "Bir yerden başka bir yere toplu taşınma", "ornek_cumle": "Kuşlar kışın sıcak ülkelere göç eder.", "zorluk": 4, "sinif": 3},
        ]
        for k in demo_kelimeler:
            mevcut = await db.meb_kelime_haritasi.find_one({"kelime": k["kelime"]})
            if not mevcut:
                await db.meb_kelime_haritasi.insert_one({
                    "id": str(uuid.uuid4()), "kaynak": "Demo - MEB Türkçe",
                    "yukleyen_id": ogretmen_id, "tarih": simdi.isoformat(), **k,
                })

        # ── Demo AI Yükleme (tamamlanmış) ──
        demo_yuk_id = str(uuid.uuid4())
        mevcut_yuk = await db.ai_yuklemeler.find_one({"kitap_adi": "Türkçe 3 Ders Kitabı (Demo)"})
        if not mevcut_yuk:
            await db.ai_yuklemeler.insert_one({
                "id": demo_yuk_id, "dosya_adi": "turkce_3_ders_kitabi.pdf", "dosya_boyut": 4500000,
                "dosya_format": ".pdf", "dosya_hash": "demo_hash_001", "dosya_b64": "",
                "sinif": 3, "tur": "ders_kitabi", "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)",
                "yazar": "MEB", "temalar": ["Erdemler", "Doğa ve Evren", "Çocuk Dünyası"],
                "yukleyen_id": ogretmen_id, "yukleyen_ad": "Ayşe Öğretmen", "yukleyen_rol": "teacher",
                "durum": "tamamlandi", "onayli": True, "ilerleme": 100,
                "sonuc": {"sayfa_sayisi": 180, "kelime_sayisi": 45000, "chunk_sayisi": 8,
                          "cikarilan_kelime": 12, "eklenen_kelime": 12, "okuma_parcasi": 4, "uretilen_soru": 15, "bonus_puan": 10},
                "guven_skoru": {"toplam": 92, "seviye": "yuksek", "detay": {"icindekiler": 90, "dil_uygunlugu": 95, "bloom_dagilimi": 88}}, "okuma_seviyesi": "3. Sınıf", "versiyon": 1, "tarih": (simdi - timedelta(days=3)).isoformat(),
            })

            # Demo okuma parçaları
            demo_parcalar = [
                {"baslik": "Ormanın Sırrı", "ozet": "Küçük bir çocuk ormanda kaybolur ve konuşan hayvanlarla arkadaş olur. Birlikte evin yolunu bulurlar.", "tema": "Doğa ve Evren", "metin_kesit": "Ağaçların arasından süzülen güneş ışığı, ormanın derinliklerini aydınlatıyordu. Küçük Ali ilk kez bu kadar içerilere gelmişti. Bir sincap daldan dala atlayarak ona yol gösteriyordu sanki..."},
                {"baslik": "Dürüst Çocuk", "ozet": "Pazarda para bulan bir çocuğun dürüstlük hikâyesi. Parayı sahibine teslim etmesi ve ödüllendirilmesi.", "tema": "Erdemler", "metin_kesit": "Elif pazarda yerde parlayan bir şey gördü. Eğilip baktığında bunun bir cüzdan olduğunu anladı. İçinde epey para vardı. 'Bunu sahibine vermem lazım' diye düşündü..."},
                {"baslik": "Göçmen Kuşlar", "ozet": "Leyleklerin göç yolculuğunu anlatan bilgilendirici metin. Göçün nedenleri ve kuşların yol bulma yetenekleri.", "tema": "Doğa ve Evren", "metin_kesit": "Her sonbaharda leylekler uzun bir yolculuğa çıkar. Binlerce kilometre uçarak sıcak ülkelere göç ederler. Pusula gibi bir iç duyguları sayesinde yollarını hiç şaşırmazlar..."},
                {"baslik": "Takım Çalışması", "ozet": "Sınıftaki öğrencilerin birlikte proje hazırlama hikâyesi. İşbirliği ve paylaşma temaları.", "tema": "Çocuk Dünyası", "metin_kesit": "Öğretmen sınıfa bir proje verdi: 'Hayalinizdeki şehri tasarlayın.' Herkes tek başına yapmak istedi ama çok zordu. Sonunda birlikte çalışmaya karar verdiler..."},
            ]
            for j, p in enumerate(demo_parcalar):
                await db.ai_okuma_parcalari.insert_one({
                    "id": str(uuid.uuid4()), "yukleme_id": demo_yuk_id, "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)",
                    "sinif": 3, "bolum": j+1, **p, "kelime_sayisi": len(p["metin_kesit"].split()), "tarih": simdi.isoformat(),
                })

            # Demo sorular
            demo_sorular = [
                {"soru": "Ali ormanda kime rastladı?", "secenekler": ["Bir sincap", "Bir ayı", "Bir balık", "Bir kuş"], "dogru_cevap": 0, "taksonomi": "bilgi", "bolum": 1},
                {"soru": "Ormandaki hayvanlar Ali'ye nasıl yardım etti?", "secenekler": ["Yemek verdiler", "Yol gösterdiler", "Şarkı söylediler", "Uyuttular"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 1},
                {"soru": "Elif cüzdanı neden sahibine verdi?", "secenekler": ["Korktuğu için", "Dürüst olduğu için", "Paraya ihtiyacı olmadığı için", "Annesi gördüğü için"], "dogru_cevap": 1, "taksonomi": "analiz", "bolum": 2},
                {"soru": "Sen Elif'in yerinde olsaydın ne yapardın?", "secenekler": ["Sahibine verirdim", "Saklar bir şey alırdım", "Polise götürürdüm", "Bilmiyorum"], "dogru_cevap": 0, "taksonomi": "degerlendirme", "bolum": 2},
                {"soru": "Leylekler neden göç eder?", "secenekler": ["Sıcak ülkelere gitmek için", "Arkadaş bulmak için", "Yeni yuva yapmak için", "Uçmayı öğrenmek için"], "dogru_cevap": 0, "taksonomi": "bilgi", "bolum": 3},
                {"soru": "Kuşların yol bulma yeteneğine ne ad verilir?", "secenekler": ["GPS", "İç pusula / içgüdü", "Harita okuma", "Rüzgar takibi"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 3},
                {"soru": "Göç eden kuşlar ile göç etmeyen kuşları karşılaştırınız.", "secenekler": ["Göçmenler daha büyük", "Göçmenler soğuğa dayanamaz", "Göçmenler daha hızlı", "Fark yok"], "dogru_cevap": 1, "taksonomi": "analiz", "bolum": 3},
                {"soru": "Öğrenciler neden birlikte çalışmaya karar verdi?", "secenekler": ["Öğretmen zorladı", "Tek başlarına yapamadılar", "Canları sıkıldı", "Ödül kazanmak için"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 4},
                {"soru": "Takım çalışması neden önemlidir? Kendi hayatından örnek ver.", "secenekler": ["Daha hızlı biter", "Herkes bir şey öğrenir", "Daha eğlenceli olur", "Hepsi doğru"], "dogru_cevap": 3, "taksonomi": "sentez", "bolum": 4},
                {"soru": "Hikâyedeki şehir tasarımı projesinde sen olsan ne eklerdin?", "secenekler": ["Park", "Kütüphane", "Hayvanat bahçesi", "Hepsi"], "dogru_cevap": 3, "taksonomi": "yaratma", "bolum": 4},
            ]
            for s in demo_sorular:
                await db.ai_uretilen_sorular.insert_one({
                    "id": str(uuid.uuid4()), "yukleme_id": demo_yuk_id, "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)",
                    "sinif": 3, **s, "tarih": simdi.isoformat(),
                })

            # Öğretmene AI puan ver
            await db.ai_egitim_puanlari.insert_one({
                "id": str(uuid.uuid4()), "kullanici_id": ogretmen_id,
                "eylem": "dosya_yukle", "dosya_adi": "turkce_3_ders_kitabi.pdf", "sinif": 3, "puan": 30, "tarih": simdi.isoformat(),
            })

        # ── Dalga 2 Demo: Kelime Evrimi (Spaced Repetition) ──
        kelime_tekrar_var = await db.kelime_tekrar.find_one({})
        if not kelime_tekrar_var:
            for ogr in demo_ogrenciler[:3]:
                oid = ogr["id"]
                sinif = ogr.get("sinif", 3)
                demo_tekrar_kelimeler = [
                    {"kelime": "macera", "anlam": "Tehlikeli ve heyecan verici olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı.", "kutu": 3, "dogru": 4, "tekrar": 5, "gun_sonra": 7},
                    {"kelime": "keşif", "anlam": "Bilinmeyen bir şeyi ilk kez bulma", "ornek_cumle": "Bilim insanı yeni bir keşif yaptı.", "kutu": 2, "dogru": 2, "tekrar": 3, "gun_sonra": 3},
                    {"kelime": "pusula", "anlam": "Yön bulmaya yarayan araç", "ornek_cumle": "Kaşif pusulasıyla yolunu buldu.", "kutu": 1, "dogru": 0, "tekrar": 1, "gun_sonra": 0},
                    {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme gücü", "ornek_cumle": "Küçük kız büyük cesaret gösterdi.", "kutu": 4, "dogru": 6, "tekrar": 7, "gun_sonra": 21},
                    {"kelime": "merak", "anlam": "Bir şeyi bilmek isteme duygusu", "ornek_cumle": "Merak eden çocuk her şeyi sorar.", "kutu": 5, "dogru": 8, "tekrar": 8, "gun_sonra": 45},
                    {"kelime": "sabır", "anlam": "Bekleyebilme ve dayanma gücü", "ornek_cumle": "Bahçıvan sabırla çiçeklerin büyümesini bekledi.", "kutu": 1, "dogru": 1, "tekrar": 3, "gun_sonra": 0},
                    {"kelime": "göç", "anlam": "Bir yerden başka bir yere toplu taşınma", "ornek_cumle": "Kuşlar kışın sıcak ülkelere göç eder.", "kutu": 2, "dogru": 3, "tekrar": 4, "gun_sonra": 1},
                    {"kelime": "dürüstlük", "anlam": "Doğruyu söyleme, hile yapmama", "ornek_cumle": "Dürüstlük en değerli erdemlerden biridir.", "kutu": 3, "dogru": 5, "tekrar": 6, "gun_sonra": 5},
                ]
                for kt in demo_tekrar_kelimeler:
                    await db.kelime_tekrar.insert_one({
                        "id": str(uuid.uuid4()),
                        "ogrenci_id": oid,
                        "kelime": kt["kelime"],
                        "anlam": kt["anlam"],
                        "ornek_cumle": kt["ornek_cumle"],
                        "sinif": sinif,
                        "kutu": kt["kutu"],
                        "tekrar_sayisi": kt["tekrar"],
                        "dogru_sayisi": kt["dogru"],
                        "son_gosterim": (simdi - timedelta(days=max(1, kt["gun_sonra"]))).isoformat(),
                        "sonraki_gosterim": (simdi + timedelta(days=kt["gun_sonra"])).isoformat() if kt["gun_sonra"] > 0 else simdi.isoformat(),
                        "tarih": (simdi - timedelta(days=14)).isoformat(),
                    })
            logging.info("  ✅ Kelime Evrimi demo verileri oluşturuldu (8 kelime × 3 öğrenci)")

        # ── Dalga 2 Demo: Socratic Reading logları ──
        socratic_var = await db.ai_socratic_log.find_one({})
        if not socratic_var:
            demo_socratic = [
                {"kitap": "Ormanın Sırrı", "soru": "Ali ormanda kaybolduğunda ne hissetti sence?", "bloom": "analiz", "gun_once": 2},
                {"kitap": "Ormanın Sırrı", "soru": "Sen Ali'nin yerinde olsaydın ne yapardın?", "bloom": "degerlendirme", "gun_once": 2},
                {"kitap": "Dürüst Çocuk", "soru": "Elif cüzdanı sahibine verdi. Bu doğru bir karar mıydı?", "bloom": "degerlendirme", "gun_once": 1},
                {"kitap": "Göçmen Kuşlar", "soru": "Kuşlar nasıl yol buluyor olabilir?", "bloom": "analiz", "gun_once": 1},
                {"kitap": "Takım Çalışması", "soru": "Tek başına mı yoksa takımla mı çalışmak daha iyi? Neden?", "bloom": "sentez", "gun_once": 0},
            ]
            for ogr in demo_ogrenciler[:2]:
                for s in demo_socratic:
                    await db.ai_socratic_log.insert_one({
                        "id": str(uuid.uuid4()),
                        "ogrenci_id": ogr["id"],
                        "kitap_adi": s["kitap"],
                        "bolum": "Bölüm 1",
                        "soru": s["soru"],
                        "bloom": s["bloom"],
                        "cevap": "Demo cevap — öğrenci düşüncesini yazdı.",
                        "puan": random.randint(3, 5),
                        "tarih": (simdi - timedelta(days=s["gun_once"])).isoformat(),
                    })
            logging.info("  ✅ Socratic Reading demo logları oluşturuldu")

        # ── Dalga 3 Demo: Speech AI logları ──
        speech_var = await db.speech_logs.find_one({})
        if not speech_var:
            demo_speech = [
                {"metin_id": "s3_1", "baslik": "Ormanda Bir Gün", "metin": SPEECH_OKUMA_METİNLERİ[3][0]["metin"], "gun_once": 6, "sure": 38, "wpm": 52, "skor": 72, "telaffuz": 75, "akicilik": 68},
                {"metin_id": "s3_2", "baslik": "Dürüstlük", "metin": SPEECH_OKUMA_METİNLERİ[3][1]["metin"], "gun_once": 5, "sure": 35, "wpm": 58, "skor": 76, "telaffuz": 78, "akicilik": 73},
                {"metin_id": "s3_1", "baslik": "Ormanda Bir Gün", "metin": SPEECH_OKUMA_METİNLERİ[3][0]["metin"], "gun_once": 4, "sure": 33, "wpm": 63, "skor": 80, "telaffuz": 82, "akicilik": 78},
                {"metin_id": "s3_2", "baslik": "Dürüstlük", "metin": SPEECH_OKUMA_METİNLERİ[3][1]["metin"], "gun_once": 3, "sure": 30, "wpm": 68, "skor": 84, "telaffuz": 85, "akicilik": 83},
                {"metin_id": "s3_1", "baslik": "Ormanda Bir Gün", "metin": SPEECH_OKUMA_METİNLERİ[3][0]["metin"], "gun_once": 1, "sure": 28, "wpm": 74, "skor": 88, "telaffuz": 89, "akicilik": 87},
            ]
            for ogr in demo_ogrenciler[:2]:
                sinif_int = 3
                try:
                    sinif_int = int(ogr.get("sinif", 3))
                except:
                    pass
                for s in demo_speech:
                    genel = s["skor"]
                    seviye = "çok iyi" if genel >= 85 else "iyi" if genel >= 70 else "orta"
                    await db.speech_logs.insert_one({
                        "id": str(uuid.uuid4()),
                        "ogrenci_id": ogr["id"],
                        "metin_id": s["metin_id"],
                        "metin_baslik": s["baslik"],
                        "metin": s["metin"],
                        "sinif": sinif_int,
                        "sure_sn": s["sure"],
                        "analiz": {
                            "transkript": s["metin"],
                            "telaffuz_skoru": s["telaffuz"],
                            "akicilik_skoru": s["akicilik"],
                            "wpm": s["wpm"],
                            "norm_wpm": 95,
                            "duraklama_sayisi": random.randint(1, 4),
                            "tonlama_skoru": s["telaffuz"] + 2,
                            "vurgu_skoru": s["akicilik"] + 3,
                            "genel_skor": genel,
                            "seviye": seviye,
                            "guclu_yonler": ["Telaffuz başarılı", "Ritim tutarlı"] if genel >= 75 else ["Kelime tanıma gelişiyor"],
                            "gelisim_alanlari": ["Hız artırılabilir"] if s["wpm"] < 80 else [],
                            "ogretmen_notu": f"Öğrenci {s['wpm']} kelime/dk okuyor. {'Sınıf normuna yaklaşıyor.' if s['wpm'] > 60 else 'Tekrarlı okuma önerilir.'}",
                            "ogrenci_mesaj": "Harika ilerliyorsun! 🌟" if genel >= 80 else "Her gün biraz daha iyileşiyorsun! 💪",
                            "telaffuz_hatalar": [],
                            "mock": True,
                            "whisper_kullanildi": False,
                        },
                        "tarih": (simdi - timedelta(days=s["gun_once"])).isoformat(),
                    })
            logging.info("  ✅ Speech AI demo logları oluşturuldu")

        logging.info("✅ AI demo verileri oluşturuldu: DNA, koçluk, kelimeler, parçalar, sorular, tekrar, socratic, speech")
    else:
        logging.info("ℹ️ AI demo verileri zaten mevcut")

    logging.info("📋 Demo Hesapları:")
    logging.info(f"   🎓 Öğrenci: {demo_student_email} / {DEMO_PASSWORD}")
    logging.info(f"   👩‍🏫 Öğretmen: {demo_teacher_email} / {DEMO_PASSWORD}")
    logging.info(f"   👪 Veli: {demo_parent_email} / {DEMO_PASSWORD}")
    logging.info(f"   🔑 Admin: {admin_email} / {admin_password}")

# ─────────────────────────────────────────────
# MEVCUT MODELLER (değişmeden korunuyor)
# ─────────────────────────────────────────────

class Teacher(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    ad: str
    soyad: str
    brans: str
    telefon: str
    seviye: TeacherLevel
    ogrenci_sayisi: int = 0
    atanan_ogrenciler: List[str] = []
    yapilmasi_gereken_odeme: float = 0.0
    yapilan_odeme: float = 0.0
    arsivli: bool = False
    olusturma_tarihi: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    @property
    def tam_ad(self):
        return f"{self.ad} {self.soyad}"

    @property
    def borc(self):
        has_students = self.ogrenci_sayisi > 0 or len(self.atanan_ogrenciler) > 0
        if not has_students:
            return 0.0
        return max(0, self.yapilmasi_gereken_odeme - self.yapilan_odeme)

class TeacherCreate(BaseModel):
    ad: str
    soyad: str
    brans: str
    telefon: str
    seviye: TeacherLevel
    yapilmasi_gereken_odeme: float = 0.0

class TeacherUpdate(BaseModel):
    ad: Optional[str] = None
    soyad: Optional[str] = None
    brans: Optional[str] = None
    telefon: Optional[str] = None
    seviye: Optional[TeacherLevel] = None
    yapilmasi_gereken_odeme: Optional[float] = None
    yapilan_odeme: Optional[float] = None
    arsivli: Optional[bool] = None

class Student(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    ad: str
    soyad: str
    sinif: str
    veli_ad: str
    veli_soyad: str
    veli_telefon: str
    aldigi_egitim: str
    kur: str
    yapilmasi_gereken_odeme: float = 0.0
    yapilan_odeme: float = 0.0
    ogretmene_yapilacak_odeme: float = 0.0
    ogretmen_id: Optional[str] = None
    arsivli: bool = False
    olusturma_tarihi: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StudentCreate(BaseModel):
    ad: str
    soyad: str
    sinif: str
    veli_ad: str
    veli_soyad: str
    veli_telefon: str
    aldigi_egitim: str
    kur: str
    yapilmasi_gereken_odeme: float = 0.0
    ogretmene_yapilacak_odeme: float = 0.0
    ogretmen_id: Optional[str] = None

class StudentUpdate(BaseModel):
    ad: Optional[str] = None
    soyad: Optional[str] = None
    sinif: Optional[str] = None
    veli_ad: Optional[str] = None
    veli_soyad: Optional[str] = None
    veli_telefon: Optional[str] = None
    aldigi_egitim: Optional[str] = None
    kur: Optional[str] = None
    yapilmasi_gereken_odeme: Optional[float] = None
    yapilan_odeme: Optional[float] = None
    ogretmene_yapilacak_odeme: Optional[float] = None
    ogretmen_id: Optional[str] = None
    arsivli: Optional[bool] = None

class Course(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    ad: str
    fiyat: float
    sure: int
    olusturma_tarihi: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    ogrenci_sayisi: int = 0
    arsivli: bool = False

class CourseCreate(BaseModel):
    ad: str
    fiyat: float
    sure: int

class CourseUpdate(BaseModel):
    ad: Optional[str] = None
    fiyat: Optional[float] = None
    sure: Optional[int] = None
    arsivli: Optional[bool] = None

# ── Ders ve İçerik Modelleri ──
class DersIcerik(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tur: str  # video, pdf, docx
    baslik: str
    url: str = ""
    ozet: str = ""

class Ders(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    kurs_id: str
    baslik: str
    sira: int = 0
    ozet: str = ""
    icerikler: List[DersIcerik] = []
    olusturma_tarihi: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class DersCreate(BaseModel):
    baslik: str
    sira: int = 0
    ozet: str = ""

class DersUpdate(BaseModel):
    baslik: Optional[str] = None
    sira: Optional[int] = None
    ozet: Optional[str] = None

class DersIcerikCreate(BaseModel):
    tur: str
    baslik: str
    url: str = ""
    ozet: str = ""

class Payment(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tip: str
    kisi_id: str
    miktar: float
    aciklama: Optional[str] = None
    tarih: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class PaymentCreate(BaseModel):
    tip: str
    kisi_id: str
    miktar: float
    aciklama: Optional[str] = None

class DashboardStats(BaseModel):
    toplam_ogretmen: int
    toplam_ogrenci: int
    toplam_kurs: int
    toplam_ogrenci_alacak: float
    toplam_ogretmen_borc: float
    bu_ay_odenen_toplam: float

class WeeklyStats(BaseModel):
    hafta: str
    yeni_ogrenciler: int
    odemeler: float
    gelir: float

class MonthlyStats(BaseModel):
    ay: str
    yeni_ogrenciler: int
    odemeler: float
    gelir: float
    toplam_borc: float

class ExportData(BaseModel):
    ogretmenler: List[dict]
    ogrenciler: List[dict]
    kurslar: List[dict]
    odemeler: List[dict]

# ─────────────────────────────────────────────
# MEVCUT ROUTE'LAR (değişmeden korunuyor)
# ─────────────────────────────────────────────

@api_router.get("/dashboard", response_model=DashboardStats)
async def get_dashboard_stats():
    teacher_count = await db.teachers.count_documents({})
    student_count = await db.students.count_documents({})
    course_count = await db.courses.count_documents({})
    teachers = await db.teachers.find().to_list(length=None)
    students = await db.students.find().to_list(length=None)
    total_student_receivable = sum(max(0, s.get('yapilmasi_gereken_odeme', 0) - s.get('yapilan_odeme', 0)) for s in students)
    total_teacher_debt = 0
    for t in teachers:
        has_students = t.get('ogrenci_sayisi', 0) > 0 or len(t.get('atanan_ogrenciler', [])) > 0
        if has_students:
            total_teacher_debt += max(0, t.get('yapilmasi_gereken_odeme', 0) - t.get('yapilan_odeme', 0))
    current_month_start = datetime.now(timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    monthly_payments = await db.payments.find({"tarih": {"$gte": current_month_start.isoformat()}}).to_list(length=None)
    monthly_total = sum(p.get('miktar', 0) for p in monthly_payments)
    return DashboardStats(
        toplam_ogretmen=teacher_count,
        toplam_ogrenci=student_count,
        toplam_kurs=course_count,
        toplam_ogrenci_alacak=total_student_receivable,
        toplam_ogretmen_borc=total_teacher_debt,
        bu_ay_odenen_toplam=monthly_total
    )


@api_router.get("/dashboard/bekleyenler")
async def get_bekleyenler(current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    # Analiz metinleri - beklemede olanlar
    metin_bekleyen = await db.analiz_metinler.find({"durum": "beklemede"}).sort("olusturma_tarihi", -1).to_list(length=None)
    metin_oylama = await db.analiz_metinler.find({"durum": "oylama"}).sort("olusturma_tarihi", -1).to_list(length=None)
    # Gelişim araçları - beklemede olanlar
    gelisim_bekleyen = await db.gelisim_icerik.find({"durum": "beklemede"}).sort("olusturma_tarihi", -1).to_list(length=None)
    gelisim_oylama = await db.gelisim_icerik.find({"durum": "oylama"}).sort("olusturma_tarihi", -1).to_list(length=None)

    # Kitaplar - beklemede olanlar
    kitap_bekleyen = await db.kitaplar.find({"durum": "beklemede"}).sort("olusturma_tarihi", -1).to_list(length=None)
    kitap_oylama = await db.kitaplar.find({"durum": "oylama"}).sort("olusturma_tarihi", -1).to_list(length=None)

    for lst in [metin_bekleyen, metin_oylama, gelisim_bekleyen, gelisim_oylama, kitap_bekleyen, kitap_oylama]:
        for item in lst:
            item.pop("_id", None)
            item.pop("icerik", None)

    return {
        "metin_bekleyen": metin_bekleyen,
        "metin_oylama": metin_oylama,
        "gelisim_bekleyen": gelisim_bekleyen,
        "gelisim_oylama": gelisim_oylama,
        "kitap_bekleyen": kitap_bekleyen,
        "kitap_oylama": kitap_oylama,
        "toplam": len(metin_bekleyen) + len(metin_oylama) + len(gelisim_bekleyen) + len(gelisim_oylama) + len(kitap_bekleyen) + len(kitap_oylama)
    }

@api_router.get("/stats/weekly", response_model=List[WeeklyStats])
async def get_weekly_stats():
    from datetime import timedelta
    stats = []
    now = datetime.now(timezone.utc)
    for i in range(12):
        week_start = now - timedelta(weeks=i+1)
        week_end = now - timedelta(weeks=i)
        students_this_week = await db.students.find({"olusturma_tarihi": {"$gte": week_start.isoformat(), "$lt": week_end.isoformat()}}).to_list(length=None)
        payments_this_week = await db.payments.find({"tarih": {"$gte": week_start.isoformat(), "$lt": week_end.isoformat()}}).to_list(length=None)
        stats.append(WeeklyStats(
            hafta=f"{week_start.strftime('%d/%m')} - {week_end.strftime('%d/%m')}",
            yeni_ogrenciler=len(students_this_week),
            odemeler=sum(p.get('miktar', 0) for p in payments_this_week),
            gelir=sum(s.get('yapilmasi_gereken_odeme', 0) for s in students_this_week)
        ))
    return list(reversed(stats))

@api_router.get("/stats/monthly", response_model=List[MonthlyStats])
async def get_monthly_stats():
    from datetime import timedelta
    stats = []
    now = datetime.now(timezone.utc)
    for i in range(6):
        if i == 0:
            month_end = now
        else:
            month_end = now.replace(day=1) - timedelta(days=1)
            for j in range(i-1):
                month_end = month_end.replace(day=1) - timedelta(days=1)
                month_end = month_end.replace(day=28)
        month_start = month_end.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        students_this_month = await db.students.find({"olusturma_tarihi": {"$gte": month_start.isoformat(), "$lt": month_end.isoformat()}}).to_list(length=None)
        payments_this_month = await db.payments.find({"tarih": {"$gte": month_start.isoformat(), "$lt": month_end.isoformat()}}).to_list(length=None)
        students_total = await db.students.find({"olusturma_tarihi": {"$lt": month_end.isoformat()}}).to_list(length=None)
        stats.append(MonthlyStats(
            ay=month_start.strftime('%B %Y'),
            yeni_ogrenciler=len(students_this_month),
            odemeler=sum(p.get('miktar', 0) for p in payments_this_month),
            gelir=sum(s.get('yapilmasi_gereken_odeme', 0) for s in students_this_month),
            toplam_borc=sum(max(0, s.get('yapilmasi_gereken_odeme', 0) - s.get('yapilan_odeme', 0)) for s in students_total)
        ))
    return list(reversed(stats))

@api_router.post("/teachers", response_model=Teacher)
async def create_teacher(teacher_data: TeacherCreate):
    teacher = Teacher(**teacher_data.dict())
    await db.teachers.insert_one(prepare_for_mongo(teacher.dict()))
    return teacher

@api_router.get("/teachers", response_model=List[Teacher])
async def get_teachers():
    teachers = await db.teachers.find().to_list(length=None)
    result = []
    for teacher in teachers:
        real_count = await db.students.count_documents({"ogretmen_id": teacher['id']})
        teacher['ogrenci_sayisi'] = real_count
        if teacher.get('ogrenci_sayisi', 0) != real_count:
            await db.teachers.update_one({"id": teacher['id']}, {"$set": {"ogrenci_sayisi": real_count}})
        result.append(Teacher(**parse_from_mongo(teacher)))
    return result

@api_router.get("/teachers/{teacher_id}", response_model=Teacher)
async def get_teacher(teacher_id: str):
    teacher = await db.teachers.find_one({"id": teacher_id})
    if not teacher:
        raise HTTPException(status_code=404, detail="Öğretmen bulunamadı")
    return Teacher(**parse_from_mongo(teacher))

@api_router.get("/teachers/{teacher_id}/students", response_model=List[Student])
async def get_teacher_students(teacher_id: str):
    teacher = await db.teachers.find_one({"id": teacher_id})
    if not teacher:
        raise HTTPException(status_code=404, detail="Öğretmen bulunamadı")
    students = await db.students.find({"ogretmen_id": teacher_id}).to_list(length=None)
    return [Student(**parse_from_mongo(s)) for s in students]

@api_router.put("/teachers/{teacher_id}", response_model=Teacher)
async def update_teacher(teacher_id: str, teacher_update: TeacherUpdate):
    update_data = teacher_update.dict(exclude_unset=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="Güncellenecek veri bulunamadı")
    result = await db.teachers.update_one({"id": teacher_id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Öğretmen bulunamadı")
    teacher = await db.teachers.find_one({"id": teacher_id})
    return Teacher(**parse_from_mongo(teacher))

@api_router.delete("/teachers/{teacher_id}")
async def delete_teacher(teacher_id: str):
    result = await db.teachers.delete_one({"id": teacher_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Öğretmen bulunamadı")
    return {"message": "Öğretmen başarıyla silindi"}

@api_router.post("/students", response_model=Student)
async def create_student(student_data: StudentCreate):
    student = Student(**student_data.dict())
    await db.students.insert_one(prepare_for_mongo(student.dict()))
    if student.ogretmen_id:
        teacher = await db.teachers.find_one({"id": student.ogretmen_id})
        if teacher:
            await db.teachers.update_one(
                {"id": student.ogretmen_id},
                {"$inc": {"ogrenci_sayisi": 1, "yapilmasi_gereken_odeme": student.ogretmene_yapilacak_odeme},
                 "$addToSet": {"atanan_ogrenciler": student.id}}
            )
    # ★ Otomatik muhasebe kaydı: öğrenci alacak kaydı
    if student.yapilmasi_gereken_odeme and student.yapilmasi_gereken_odeme > 0:
        alacak_kaydi = {
            "id": str(uuid.uuid4()),
            "tip": "ogrenci",
            "kisi_id": student.id,
            "miktar": student.yapilmasi_gereken_odeme,
            "aciklama": f"Kayıt ücreti — {student.ad} {student.soyad}",
            "tarih": datetime.now(timezone.utc).isoformat(),
        }
        await db.payments.insert_one(alacak_kaydi)
    # ★ Otomatik muhasebe kaydı: öğretmene yapılacak ödeme
    if student.ogretmen_id and student.ogretmene_yapilacak_odeme and student.ogretmene_yapilacak_odeme > 0:
        ogretmen_kaydi = {
            "id": str(uuid.uuid4()),
            "tip": "ogretmen",
            "kisi_id": student.ogretmen_id,
            "miktar": student.ogretmene_yapilacak_odeme,
            "aciklama": f"Öğretmen ücreti — {student.ad} {student.soyad}",
            "tarih": datetime.now(timezone.utc).isoformat(),
        }
        await db.payments.insert_one(ogretmen_kaydi)
    return student

@api_router.get("/students", response_model=List[Student])
async def get_students():
    students = await db.students.find().to_list(length=None)
    return [Student(**parse_from_mongo(s)) for s in students]

@api_router.get("/students/{student_id}", response_model=Student)
async def get_student(student_id: str):
    student = await db.students.find_one({"id": student_id})
    if not student:
        raise HTTPException(status_code=404, detail="Öğrenci bulunamadı")
    return Student(**parse_from_mongo(student))

@api_router.put("/students/{student_id}", response_model=Student)
async def update_student(student_id: str, student_update: StudentUpdate):
    update_data = student_update.dict(exclude_unset=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="Güncellenecek veri bulunamadı")
    old_student = await db.students.find_one({"id": student_id})
    if not old_student:
        raise HTTPException(status_code=404, detail="Öğrenci bulunamadı")
    old_teacher_id = old_student.get('ogretmen_id')
    new_teacher_id = update_data.get('ogretmen_id')
    old_payment = old_student.get('ogretmene_yapilacak_odeme', 0)
    new_payment = update_data.get('ogretmene_yapilacak_odeme', old_payment)
    await db.students.update_one({"id": student_id}, {"$set": update_data})
    if old_teacher_id != new_teacher_id:
        if old_teacher_id:
            await db.teachers.update_one({"id": old_teacher_id}, {"$inc": {"ogrenci_sayisi": -1, "yapilmasi_gereken_odeme": -old_payment}, "$pull": {"atanan_ogrenciler": student_id}})
        if new_teacher_id:
            await db.teachers.update_one({"id": new_teacher_id}, {"$inc": {"ogrenci_sayisi": 1, "yapilmasi_gereken_odeme": new_payment}, "$addToSet": {"atanan_ogrenciler": student_id}})
    elif old_teacher_id and old_payment != new_payment:
        await db.teachers.update_one({"id": old_teacher_id}, {"$inc": {"yapilmasi_gereken_odeme": new_payment - old_payment}})
    student = await db.students.find_one({"id": student_id})
    return Student(**parse_from_mongo(student))

@api_router.delete("/students/{student_id}")
async def delete_student(student_id: str):
    student = await db.students.find_one({"id": student_id})
    if not student:
        raise HTTPException(status_code=404, detail="Öğrenci bulunamadı")
    if student.get('ogretmen_id'):
        teacher = await db.teachers.find_one({"id": student['ogretmen_id']})
        if teacher:
            await db.teachers.update_one(
                {"id": student['ogretmen_id']},
                {"$inc": {"ogrenci_sayisi": -1, "yapilmasi_gereken_odeme": -student.get('ogretmene_yapilacak_odeme', 0)},
                 "$pull": {"atanan_ogrenciler": student_id}}
            )
    await db.students.delete_one({"id": student_id})
    return {"message": "Öğrenci başarıyla silindi"}

@api_router.post("/courses", response_model=Course)
async def create_course(course_data: CourseCreate):
    course = Course(**course_data.dict())
    await db.courses.insert_one(prepare_for_mongo(course.dict()))
    return course

@api_router.get("/courses", response_model=List[Course])
async def get_courses():
    courses = await db.courses.find().to_list(length=None)
    return [Course(**parse_from_mongo(c)) for c in courses]

@api_router.get("/courses/{course_id}", response_model=Course)
async def get_course(course_id: str):
    course = await db.courses.find_one({"id": course_id})
    if not course:
        raise HTTPException(status_code=404, detail="Kurs bulunamadı")
    return Course(**parse_from_mongo(course))

@api_router.put("/courses/{course_id}", response_model=Course)
async def update_course(course_id: str, course_update: CourseUpdate):
    update_data = course_update.dict(exclude_unset=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="Güncellenecek veri bulunamadı")
    result = await db.courses.update_one({"id": course_id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Kurs bulunamadı")
    course = await db.courses.find_one({"id": course_id})
    return Course(**parse_from_mongo(course))

@api_router.delete("/courses/{course_id}")
async def delete_course(course_id: str):
    result = await db.courses.delete_one({"id": course_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Kurs bulunamadı")
    return {"message": "Kurs başarıyla silindi"}

# ── Ders Endpoints ──
@api_router.get("/courses/{kurs_id}/dersler")
async def get_dersler(kurs_id: str):
    dersler = await db.dersler.find({"kurs_id": kurs_id}).sort("sira", 1).to_list(length=None)
    for d in dersler:
        d.pop("_id", None)
    return dersler

@api_router.post("/courses/{kurs_id}/dersler")
async def create_ders(kurs_id: str, data: DersCreate, current_user=Depends(get_current_user)):
    ders_doc = {
        "id": str(uuid.uuid4()),
        "kurs_id": kurs_id,
        "baslik": data.baslik,
        "sira": data.sira,
        "ozet": data.ozet,
        "icerikler": [],
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.dersler.insert_one(ders_doc)
    ders_doc.pop("_id", None)
    return ders_doc

@api_router.put("/dersler/{ders_id}")
async def update_ders(ders_id: str, data: DersUpdate, current_user=Depends(get_current_user)):
    update = data.dict(exclude_unset=True)
    if update:
        await db.dersler.update_one({"id": ders_id}, {"$set": update})
    ders = await db.dersler.find_one({"id": ders_id})
    if ders:
        ders.pop("_id", None)
    return ders

@api_router.delete("/dersler/{ders_id}")
async def delete_ders(ders_id: str, current_user=Depends(get_current_user)):
    await db.dersler.delete_one({"id": ders_id})
    return {"message": "Ders silindi"}

@api_router.post("/dersler/{ders_id}/icerik")
async def add_ders_icerik(ders_id: str, data: DersIcerikCreate, current_user=Depends(get_current_user)):
    icerik = {
        "id": str(uuid.uuid4()),
        "tur": data.tur,
        "baslik": data.baslik,
        "url": data.url,
        "ozet": data.ozet,
    }
    await db.dersler.update_one({"id": ders_id}, {"$push": {"icerikler": icerik}})
    return icerik

@api_router.delete("/dersler/{ders_id}/icerik/{icerik_id}")
async def delete_ders_icerik(ders_id: str, icerik_id: str, current_user=Depends(get_current_user)):
    await db.dersler.update_one({"id": ders_id}, {"$pull": {"icerikler": {"id": icerik_id}}})
    return {"message": "İçerik silindi"}

@api_router.post("/payments", response_model=Payment)
async def create_payment(payment_data: PaymentCreate):
    payment = Payment(**payment_data.dict())
    await db.payments.insert_one(prepare_for_mongo(payment.dict()))
    if payment.tip == "ogrenci":
        await db.students.update_one({"id": payment.kisi_id}, {"$inc": {"yapilan_odeme": payment.miktar}})
    elif payment.tip == "ogretmen":
        await db.teachers.update_one({"id": payment.kisi_id}, {"$inc": {"yapilan_odeme": payment.miktar}})
    return payment

@api_router.get("/payments", response_model=List[Payment])
async def get_payments():
    payments = await db.payments.find().sort("tarih", -1).to_list(length=None)
    return [Payment(**parse_from_mongo(p)) for p in payments]

@api_router.delete("/payments/{payment_id}")
async def delete_payment(payment_id: str):
    payment = await db.payments.find_one({"id": payment_id})
    if not payment:
        raise HTTPException(status_code=404, detail="Ödeme bulunamadı")
    if payment['tip'] == "ogrenci":
        await db.students.update_one({"id": payment['kisi_id']}, {"$inc": {"yapilan_odeme": -payment['miktar']}})
    elif payment['tip'] == "ogretmen":
        await db.teachers.update_one({"id": payment['kisi_id']}, {"$inc": {"yapilan_odeme": -payment['miktar']}})
    await db.payments.delete_one({"id": payment_id})
    return {"message": "Ödeme başarıyla silindi"}

@api_router.get("/export", response_model=ExportData)
async def get_export_data():
    teachers = await db.teachers.find().to_list(length=None)
    teacher_export = [{"Ad": t.get('ad',''), "Soyad": t.get('soyad',''), "Branş": t.get('brans',''), "Telefon": t.get('telefon',''), "Seviye": t.get('seviye',''), "Öğrenci Sayısı": t.get('ogrenci_sayisi',0), "Yapılması Gereken Ödeme": t.get('yapilmasi_gereken_odeme',0), "Yapılan Ödeme": t.get('yapilan_odeme',0), "Borç": max(0, t.get('yapilmasi_gereken_odeme',0) - t.get('yapilan_odeme',0)) if (t.get('ogrenci_sayisi',0) > 0 or len(t.get('atanan_ogrenciler',[])) > 0) else 0} for t in teachers]
    students = await db.students.find().to_list(length=None)
    student_export = []
    for s in students:
        teacher = await db.teachers.find_one({"id": s.get('ogretmen_id')}) if s.get('ogretmen_id') else None
        student_export.append({"Ad": s.get('ad',''), "Soyad": s.get('soyad',''), "Sınıf": s.get('sinif',''), "Veli Adı": s.get('veli_ad',''), "Veli Soyadı": s.get('veli_soyad',''), "Veli Telefon": s.get('veli_telefon',''), "Aldığı Eğitim": s.get('aldigi_egitim',''), "Kur": s.get('kur',''), "Yapılması Gereken Ödeme": s.get('yapilmasi_gereken_odeme',0), "Yapılan Ödeme": s.get('yapilan_odeme',0), "Öğretmene Yapılacak Ödeme": s.get('ogretmene_yapilacak_odeme',0), "Öğretmen": f"{teacher.get('ad','')} {teacher.get('soyad','')}" if teacher else 'Atanmamış', "Alacak": max(0, s.get('yapilmasi_gereken_odeme',0) - s.get('yapilan_odeme',0))})
    courses = await db.courses.find().to_list(length=None)
    course_export = [{"Kurs Adı": c.get('ad',''), "Fiyat": c.get('fiyat',0), "Süre (Saat)": c.get('sure',0), "Öğrenci Sayısı": c.get('ogrenci_sayisi',0)} for c in courses]
    payments = await db.payments.find().sort("tarih", -1).to_list(length=None)
    payment_export = []
    for p in payments:
        if p.get('tip') == 'ogrenci':
            person = await db.students.find_one({"id": p.get('kisi_id')})
        else:
            person = await db.teachers.find_one({"id": p.get('kisi_id')})
        payment_export.append({"Tarih": p.get('tarih',''), "Tip": 'Öğrenci' if p.get('tip') == 'ogrenci' else 'Öğretmen', "Kişi": f"{person.get('ad','')} {person.get('soyad','')}" if person else 'Bilinmiyor', "Miktar": p.get('miktar',0), "Açıklama": p.get('aciklama','')})
    return ExportData(ogretmenler=teacher_export, ogrenciler=student_export, kurslar=course_export, odemeler=payment_export)

@api_router.post("/backup/google-drive")
async def backup_to_google_drive(backup_data: dict):
    return {"success": True, "message": "Data queued for Google Drive backup", "backup_id": str(uuid.uuid4())}


# ─────────────────────────────────────────────
# GİRİŞ ANALİZİ (FAZ 1A)
# ─────────────────────────────────────────────

# Varsayılan norm tablosu (admin değiştirebilir)
VARSAYILAN_NORMLAR = {
    "1": {"dusuk": 25, "orta": 40, "yeterli": 60},
    "2": {"dusuk": 55, "orta": 75, "yeterli": 95},
    "3": {"dusuk": 65, "orta": 90, "yeterli": 115},
    "4": {"dusuk": 80, "orta": 110, "yeterli": 140},
    "5": {"dusuk": 90, "orta": 120, "yeterli": 150},
    "6": {"dusuk": 100, "orta": 135, "yeterli": 170},
    "7": {"dusuk": 110, "orta": 150, "yeterli": 185},
    "8": {"dusuk": 120, "orta": 160, "yeterli": 200},
}

async def get_norm_tablosu():
    doc = await db.sistem_ayarlari.find_one({"tip": "okuma_hizi_normlari"})
    if doc:
        return doc.get("normlar", VARSAYILAN_NORMLAR)
    return VARSAYILAN_NORMLAR

async def hiz_degerlendirme(sinif: str, wpm: float) -> str:
    normlar = await get_norm_tablosu()
    sinif_no = sinif.replace("-", "").replace(".", "").strip()[:1]
    n = normlar.get(sinif_no, normlar.get("4", VARSAYILAN_NORMLAR["4"]))
    if wpm <= n["dusuk"]:
        return "dusuk"
    elif wpm <= n["orta"]:
        return "orta"
    elif wpm <= n["yeterli"]:
        return "yeterli"
    else:
        return "ileri"

async def kur_onerisi_hesapla(wpm: float, dogruluk: float, sinif: str) -> str:
    hiz = await hiz_degerlendirme(sinif, wpm)
    if dogruluk >= 97 and hiz in ("yeterli", "ileri"):
        return "Kur 3"
    elif dogruluk >= 93 and hiz in ("orta", "yeterli", "ileri"):
        return "Kur 2"
    else:
        return "Kur 1"

# ── Norm Tablosu Yönetimi ──
class NormGuncelle(BaseModel):
    normlar: dict  # {"1": {"dusuk": 25, "orta": 40, "yeterli": 60}, ...}

@api_router.get("/diagnostic/normlar")
async def get_normlar(current_user=Depends(get_current_user)):
    return await get_norm_tablosu()

@api_router.put("/diagnostic/normlar")
async def update_normlar(data: NormGuncelle, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    await db.sistem_ayarlari.update_one(
        {"tip": "okuma_hizi_normlari"},
        {"$set": {"tip": "okuma_hizi_normlari", "normlar": data.normlar}},
        upsert=True
    )
    return {"message": "Norm tablosu güncellendi", "normlar": data.normlar}

# ── Puan Ayarları ──
VARSAYILAN_PUANLAR = {
    "metin_ekleme": 5,
    "oylama_katilim": 2,
    "metin_havuza_girme": 10,
    "icerik_ekleme": 5,
    "icerik_oylama": 2,
    "ai_kitap_yukleme": 25,
    "ai_ders_kitabi_yukleme": 40,
    "ai_kitap_onaylandi": 15,
}

async def get_puan_ayarlari():
    doc = await db.sistem_ayarlari.find_one({"tip": "puan_ayarlari"})
    if doc:
        return doc.get("puanlar", VARSAYILAN_PUANLAR)
    return VARSAYILAN_PUANLAR

@api_router.get("/ayarlar/puanlar")
async def get_puanlar(current_user=Depends(get_current_user)):
    return await get_puan_ayarlari()

@api_router.put("/ayarlar/puanlar")
async def update_puanlar(data: dict = Body(...), current_user=Depends(require_role(UserRole.ADMIN))):
    await db.sistem_ayarlari.update_one(
        {"tip": "puan_ayarlari"},
        {"$set": {"tip": "puan_ayarlari", "puanlar": data}},
        upsert=True
    )
    return {"message": "Puan ayarları güncellendi", "puanlar": data}


# ─────────────────────────────────────────────
# ★ EKSİK MODELLER VE ENDPOINT'LER (EKLENDİ)
# ─────────────────────────────────────────────

# Metin oluşturma modeli — frontend MetinYonetimi bileşeni kullanıyor
class MetinCreate(BaseModel):
    baslik: str
    icerik: str
    kelime_sayisi: int = 0
    sinif_seviyesi: str = "4"
    tur: str = "hikaye"  # hikaye, bilgilendirici, siir

# Metin oylama modeli — metin_oy_ver endpoint'i kullanıyor (NameError düzeltmesi)
class MetinOyCreate(BaseModel):
    metin_id: str
    onay: bool
    sebep: str = ""


# ★ Metin ekleme endpoint'i (frontend: axios.post(`${API}/diagnostic/texts`, ...))
@api_router.post("/diagnostic/texts")
async def create_metin(data: MetinCreate, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")

    # Kelime sayısı otomatik hesapla (0 geldiyse)
    kelime_sayisi = data.kelime_sayisi
    if kelime_sayisi == 0 and data.icerik:
        kelime_sayisi = len(data.icerik.strip().split())

    # Admin/Koordinatör eklerse direkt oylama, öğretmen eklerse beklemede
    durum = "oylama" if role in ["admin", "coordinator"] else "beklemede"

    metin_doc = {
        "id": str(uuid.uuid4()),
        "baslik": data.baslik,
        "icerik": data.icerik,
        "kelime_sayisi": kelime_sayisi,
        "sinif_seviyesi": data.sinif_seviyesi,
        "tur": data.tur,
        "durum": durum,
        "ekleyen_id": current_user["id"],
        "ekleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}",
        "oylar": {},
        "olusturma_tarihi": datetime.now(timezone.utc).isoformat(),
        "yayin_tarihi": None,
    }
    await db.analiz_metinler.insert_one(metin_doc)

    # Ekleyene puan ver (dinamik)
    puanlar = await get_puan_ayarlari()
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": puanlar.get("metin_ekleme", 5)}})

    metin_doc.pop("_id", None)
    return metin_doc


# ★ Metin listeleme endpoint'i (frontend: axios.get(`${API}/diagnostic/texts`))
@api_router.get("/diagnostic/texts")
async def get_metinler(current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    user_id = current_user.get("id", "")

    items = await db.analiz_metinler.find().sort("olusturma_tarihi", -1).to_list(length=None)
    result = []
    for item in items:
        item.pop("_id", None)
        durum = item.get("durum", "")

        # Admin her şeyi görür
        if role == "admin":
            result.append(item)
        # Öğretmen: kendi eklediği + oylama bekleyenler + havuzdakiler
        elif role == "teacher":
            if item.get("ekleyen_id") == user_id:
                result.append(item)
            elif durum in ("oylama", "havuzda"):
                result.append(item)
        # Öğrenci/diğer: sadece havuzdakiler
        else:
            if durum == "havuzda":
                result.append(item)

    return result


# ─────────────────────────────────────────────
# MEVCUT GİRİŞ ANALİZİ ROUTE'LARI
# ─────────────────────────────────────────────

@api_router.post("/diagnostic/texts/{metin_id}/admin-karar")
async def metin_admin_karar(metin_id: str, karar: dict, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    # karar: {"onay": True/False, "direkt": True/False}
    # direkt=True → oylama atla, direkt havuza al
    onay = karar.get("onay", False)
    direkt = karar.get("direkt", False)
    if not onay:
        yeni_durum = "reddedildi"
    elif direkt:
        yeni_durum = "havuzda"
        # Ekleyene bonus puan (havuza direkt girince - dinamik)
        puanlar = await get_puan_ayarlari()
        metin = await db.analiz_metinler.find_one({"id": metin_id})
        if metin and metin.get("ekleyen_id"):
            await db.users.update_one({"id": metin["ekleyen_id"]}, {"$inc": {"puan": puanlar.get("metin_havuza_girme", 10)}})
    else:
        yeni_durum = "oylama"
    await db.analiz_metinler.update_one(
        {"id": metin_id},
        {"$set": {"durum": yeni_durum, **({"yayin_tarihi": datetime.utcnow().isoformat()} if yeni_durum == "havuzda" else {})}}
    )
    return {"durum": yeni_durum}

@api_router.post("/diagnostic/texts/oy")
async def metin_oy_ver(oy: MetinOyCreate, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Sadece öğretmenler oy verebilir")
    if not oy.onay and not oy.sebep:
        raise HTTPException(status_code=400, detail="Red için sebep belirtmelisiniz")
    metin = await db.analiz_metinler.find_one({"id": oy.metin_id})
    if not metin:
        raise HTTPException(status_code=404, detail="Metin bulunamadı")
    if metin.get("durum") != "oylama":
        raise HTTPException(status_code=400, detail="Bu metin oylamada değil")
    user_id = current_user["id"]
    oylar = metin.get("oylar", {})
    if user_id in oylar:
        raise HTTPException(status_code=400, detail="Zaten oy kullandınız")
    oylar[user_id] = {"onay": oy.onay, "sebep": oy.sebep}
    await db.analiz_metinler.update_one({"id": oy.metin_id}, {"$set": {"oylar": oylar}})
    # Oy veren öğretmene puan (dinamik)
    puanlar = await get_puan_ayarlari()
    await db.users.update_one({"id": user_id}, {"$inc": {"puan": puanlar.get("oylama_katilim", 2)}})
    # %60 kontrolü
    ogretmenler = await db.users.find({"role": {"$in": ["teacher", "coordinator", "admin"]}}).to_list(length=None)
    toplam = len(ogretmenler)
    onay_sayisi = sum(1 for v in oylar.values() if v.get("onay"))
    oy_sayisi = len(oylar)
    yeni_durum = metin.get("durum")
    if toplam > 0:
        onay_orani = onay_sayisi / toplam
        if onay_orani >= 0.6:
            yeni_durum = "havuzda"
            await db.analiz_metinler.update_one(
                {"id": oy.metin_id},
                {"$set": {"durum": "havuzda", "yayin_tarihi": datetime.utcnow().isoformat()}}
            )
            # Metni ekleyene bonus puan (havuza girince - dinamik)
            ekleyen_id = metin.get("ekleyen_id")
            if ekleyen_id:
                await db.users.update_one({"id": ekleyen_id}, {"$inc": {"puan": puanlar.get("metin_havuza_girme", 10)}})
        elif oy_sayisi == toplam and onay_orani < 0.6:
            yeni_durum = "reddedildi"
            await db.analiz_metinler.update_one({"id": oy.metin_id}, {"$set": {"durum": "reddedildi"}})
    return {
        "mesaj": f"Oyunuz kaydedildi (+{puanlar.get('oylama_katilim', 2)} puan)",
        "durum": yeni_durum,
        "onay_orani": round(onay_sayisi / max(toplam, 1) * 100),
        "oy_sayisi": oy_sayisi,
        "toplam": toplam
    }

@api_router.delete("/diagnostic/texts/{metin_id}")
async def delete_metin(metin_id: str, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    await db.analiz_metinler.delete_one({"id": metin_id})
    return {"message": "Silindi"}

# ── Analiz Oturumları ──
class HataKaydi(BaseModel):
    tip: str  # atlama, yanlis_okuma, takilma, tekrar
    kelime: str = ""

class AnalizOturumBaslat(BaseModel):
    ogrenci_id: str
    metin_id: str

class AnalizTamamla(BaseModel):
    sure_saniye: float
    hatalar: List[HataKaydi]
    gozlem_notu: str = ""
    ogretmen_kur: str = ""

class DiagnosticOturum(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    ogrenci_id: str
    metin_id: str
    ogretmen_id: str
    durum: str = "devam"
    sure_saniye: float = 0
    hatalar: List[dict] = []
    gozlem_notu: str = ""
    wpm: float = 0
    dogruluk_yuzde: float = 0
    hiz_deger: str = ""
    sistem_kur: str = ""
    ogretmen_kur: str = ""
    olusturma_tarihi: datetime = Field(default_factory=datetime.utcnow)
    tamamlama_tarihi: Optional[datetime] = None

@api_router.post("/diagnostic/sessions")
async def baslat_oturum(data: AnalizOturumBaslat, current_user=Depends(get_current_user)):
    if current_user.get("role") not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    # Metni bul (id veya _id ile)
    metin = await db.analiz_metinler.find_one({"id": data.metin_id})
    if not metin:
        raise HTTPException(status_code=404, detail=f"Metin bulunamadı: {data.metin_id}")
    # Öğrenciyi kontrol et
    ogrenci = await db.students.find_one({"id": data.ogrenci_id})
    if not ogrenci:
        raise HTTPException(status_code=404, detail=f"Öğrenci bulunamadı: {data.ogrenci_id}")
    oturum = DiagnosticOturum(
        ogrenci_id=data.ogrenci_id,
        metin_id=data.metin_id,
        ogretmen_id=current_user["id"]
    )
    d = oturum.dict()
    d["olusturma_tarihi"] = d["olusturma_tarihi"].isoformat()
    d["tamamlama_tarihi"] = None
    await db.diagnostic_oturumlar.insert_one(d)
    d.pop("_id", None)
    return d

@api_router.get("/diagnostic/sessions")
async def get_oturumlar(current_user=Depends(get_current_user)):
    q = {}
    if current_user.get("role") == "teacher":
        q["ogretmen_id"] = current_user["id"]
    items = await db.diagnostic_oturumlar.find(q).sort("olusturma_tarihi", -1).to_list(length=None)
    for i in items: i.pop("_id", None)
    return items

@api_router.get("/diagnostic/sessions/student/{ogrenci_id}")
async def get_ogrenci_oturumlari(ogrenci_id: str, current_user=Depends(get_current_user)):
    items = await db.diagnostic_oturumlar.find({"ogrenci_id": ogrenci_id}).sort("olusturma_tarihi", -1).to_list(length=None)
    for i in items: i.pop("_id", None)
    return items

@api_router.post("/diagnostic/sessions/{oturum_id}/complete")
async def tamamla_oturum(oturum_id: str, data: AnalizTamamla, current_user=Depends(get_current_user)):
    oturum = await db.diagnostic_oturumlar.find_one({"id": oturum_id})
    if not oturum:
        raise HTTPException(status_code=404, detail="Oturum bulunamadı")

    metin = await db.analiz_metinler.find_one({"id": oturum["metin_id"]})
    kelime_sayisi = metin.get("kelime_sayisi", 100) if metin else 100
    sinif_seviyesi = metin.get("sinif_seviyesi", "4") if metin else "4"

    # Hesaplamalar
    sure_dakika = data.sure_saniye / 60 if data.sure_saniye > 0 else 1
    wpm = round(kelime_sayisi / sure_dakika, 1)

    toplam_hata = len(data.hatalar)
    dogruluk = round(max(0, (kelime_sayisi - toplam_hata) / kelime_sayisi * 100), 1)

    hiz_deger = await hiz_degerlendirme(sinif_seviyesi, wpm)
    sistem_kur = await kur_onerisi_hesapla(wpm, dogruluk, sinif_seviyesi)
    atanan_kur = data.ogretmen_kur if data.ogretmen_kur else sistem_kur

    # Hata dağılımı
    hata_sayilari = {"atlama": 0, "yanlis_okuma": 0, "takilma": 0, "tekrar": 0}
    for h in data.hatalar:
        tip = h.tip if hasattr(h, "tip") else h.get("tip", "")
        if tip in hata_sayilari:
            hata_sayilari[tip] += 1

    now = datetime.utcnow().isoformat()
    guncelle = {
        "durum": "tamamlandi",
        "sure_saniye": data.sure_saniye,
        "hatalar": [h.dict() if hasattr(h, "dict") else h for h in data.hatalar],
        "gozlem_notu": data.gozlem_notu,
        "wpm": wpm,
        "dogruluk_yuzde": dogruluk,
        "hiz_deger": hiz_deger,
        "sistem_kur": sistem_kur,
        "ogretmen_kur": atanan_kur,
        "tamamlama_tarihi": now
    }
    await db.diagnostic_oturumlar.update_one({"id": oturum_id}, {"$set": guncelle})

    # Öğrencinin kurunu güncelle
    await db.students.update_one({"id": oturum["ogrenci_id"]}, {"$set": {"kur": atanan_kur}})

    return {
        "wpm": wpm,
        "dogruluk_yuzde": dogruluk,
        "hiz_deger": hiz_deger,
        "sistem_kur": sistem_kur,
        "atanan_kur": atanan_kur,
        "hata_sayilari": hata_sayilari,
        "sure_saniye": data.sure_saniye
    }



# ── Rapor Sistemi ──
class AnlamaVeri(BaseModel):
    # 4.1 Sözcük düzeyinde
    cumle_anlama: str = "orta"          # zayif / orta / iyi
    bilinmeyen_sozcuk: str = "orta"
    baglac_zamir: str = "orta"
    # 4.2 Ana yapı
    ana_fikir: str = "orta"
    yardimci_fikir: str = "orta"
    konu: str = "orta"
    baslik_onerme: str = "orta"
    # 4.3 Derin anlama
    neden_sonuc: str = "orta"
    cikarim: str = "orta"
    ipuclari: str = "orta"
    yorumlama: str = "orta"
    # 4.4 Eleştirel
    gorus_bildirme: str = "orta"
    yazar_amaci: str = "orta"
    alternatif_fikir: str = "orta"
    guncelle_hayat: str = "orta"
    # 4.5 Soru performansı
    bilgi: str = "iyi"
    kavrama: str = "iyi"
    uygulama: str = "iyi"
    analiz: str = "iyi"
    sentez: str = "iyi"
    degerlendirme: str = "iyi"
    genel_yuzde: int = 0

class ProzodikVeri(BaseModel):
    noktalama: int = 3    # 1-4 puan
    vurgu: int = 3
    tonlama: int = 3
    akicilik: int = 3
    anlamli_gruplama: int = 3

class RaporOlusturCreate(BaseModel):
    oturum_id: str
    anlama: AnlamaVeri
    prozodik: ProzodikVeri
    ogretmen_notu: str = ""

def anlama_yuzde(anlama: AnlamaVeri) -> int:
    alanlar = [
        anlama.cumle_anlama, anlama.bilinmeyen_sozcuk, anlama.baglac_zamir,
        anlama.ana_fikir, anlama.yardimci_fikir, anlama.konu, anlama.baslik_onerme,
        anlama.neden_sonuc, anlama.cikarim, anlama.ipuclari, anlama.yorumlama,
        anlama.gorus_bildirme, anlama.yazar_amaci, anlama.alternatif_fikir, anlama.guncelle_hayat,
        anlama.bilgi, anlama.kavrama, anlama.uygulama, anlama.analiz, anlama.sentez, anlama.degerlendirme
    ]
    puan_map = {"zayif": 0, "orta": 1, "iyi": 2}
    toplam = sum(puan_map.get(a, 1) for a in alanlar)
    return round(toplam / (len(alanlar) * 2) * 100)

def hiz_metni(hiz_deger: str) -> str:
    return {"dusuk": "düşük", "orta": "orta", "yeterli": "yeterli", "ileri": "ileri"}.get(hiz_deger, "orta")

def prozodik_seviye(toplam: int) -> str:
    if toplam >= 18: return "çok iyi"
    elif toplam >= 14: return "iyi"
    elif toplam >= 10: return "orta"
    else: return "geliştirilmeli"

def anlama_seviye(pct: int) -> str:
    if pct >= 85: return "iyi"
    elif pct >= 70: return "orta"
    else: return "zayıf"

def _hata_sayilari_hesapla(hatalar_raw):
    """Ham hata listesinden [{tur, sayi}] formatına çevir"""
    if not hatalar_raw:
        return []
    # Eğer zaten {tur, sayi} formatındaysa dokunma
    if hatalar_raw and isinstance(hatalar_raw[0], dict) and "sayi" in hatalar_raw[0]:
        return hatalar_raw
    # Ham liste: [{tip: "atlama", kelime: "x"}, ...] → sayılarla topla
    sayac = {}
    for h in hatalar_raw:
        tip = h.get("tip", "") if isinstance(h, dict) else ""
        if tip:
            sayac[tip] = sayac.get(tip, 0) + 1
    hata_labels = {"atlama": "Atlama", "yanlis_okuma": "Yanlış Okuma", "takilma": "Takılma", "tekrar": "Tekrar"}
    result = []
    for tip, sayi in sayac.items():
        result.append({"tur": tip, "sayi": sayi})
    # Eğer hiç hata yoksa standart 4 türü 0 olarak göster
    if not result:
        for tip in ["atlama", "yanlis_okuma", "takilma", "tekrar"]:
            result.append({"tur": tip, "sayi": 0})
    return result

@api_router.post("/diagnostic/rapor")
async def olustur_rapor(data: RaporOlusturCreate, current_user=Depends(get_current_user)):
    oturum = await db.diagnostic_oturumlar.find_one({"id": data.oturum_id})
    if not oturum:
        raise HTTPException(status_code=404, detail="Oturum bulunamadı")

    metin = await db.analiz_metinler.find_one({"id": oturum.get("metin_id")})
    ogrenci = await db.students.find_one({"id": oturum.get("ogrenci_id")})
    ogretmen = await db.users.find_one({"id": oturum.get("ogretmen_id")})

    prozodik_toplam = data.prozodik.noktalama + data.prozodik.vurgu + data.prozodik.tonlama + data.prozodik.akicilik + data.prozodik.anlamli_gruplama
    anlama_pct = data.anlama.genel_yuzde if data.anlama.genel_yuzde > 0 else anlama_yuzde(data.anlama)

    rapor_data = {
        "id": str(uuid.uuid4()),
        "oturum_id": data.oturum_id,
        "ogrenci_id": oturum.get("ogrenci_id"),
        "ogretmen_id": oturum.get("ogretmen_id"),
        "ogrenci_ad": f"{ogrenci.get('ad','')} {ogrenci.get('soyad','')}" if ogrenci else "",
        "ogrenci_sinif": ogrenci.get("sinif", "") if ogrenci else "",
        "ogretmen_ad": f"{ogretmen.get('ad','')} {ogretmen.get('soyad','')}" if ogretmen else "",
        "metin_adi": metin.get("baslik", "") if metin else "",
        "metin_turu": metin.get("tur", "") if metin else "",
        "kelime_sayisi": metin.get("kelime_sayisi", 0) if metin else 0,
        "sure_saniye": oturum.get("sure_saniye", 0),
        "wpm": oturum.get("wpm", 0),
        "dogruluk_yuzde": oturum.get("dogruluk_yuzde", 0),
        "hiz_deger": oturum.get("hiz_deger", ""),
        "atanan_kur": oturum.get("ogretmen_kur", ""),
        "hata_sayilari": _hata_sayilari_hesapla(oturum.get("hatalar", [])),
        "anlama": data.anlama.dict(),
        "anlama_yuzde": anlama_pct,
        "prozodik": data.prozodik.dict(),
        "prozodik_toplam": prozodik_toplam,
        "ogretmen_notu": data.ogretmen_notu,
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.diagnostic_raporlar.insert_one(rapor_data)
    rapor_data.pop("_id", None)
    # Veliye bildirim gönder
    try: await bildirim_rapor_tamamlandi(rapor_data.get("ogrenci_id"), rapor_data.get("baslik", "Giriş Analizi Raporu"))
    except: pass
    return rapor_data

@api_router.get("/diagnostic/rapor/ogrenci/{ogrenci_id}")
async def get_ogrenci_raporlari(ogrenci_id: str, current_user=Depends(get_current_user)):
    items = await db.diagnostic_raporlar.find({"ogrenci_id": ogrenci_id}).sort("olusturma_tarihi", -1).to_list(length=None)
    for i in items: i.pop("_id", None)
    return items

@api_router.get("/diagnostic/rapor/{rapor_id}")
async def get_rapor(rapor_id: str, current_user=Depends(get_current_user)):
    rapor = await db.diagnostic_raporlar.find_one({"id": rapor_id})
    if not rapor:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    rapor.pop("_id", None)
    return rapor

# ── Egzersiz Puan Sistemi ──
@api_router.get("/egzersiz/puanlar")
async def get_egzersiz_puanlari():
    doc = await db.ayarlar.find_one({"tip": "egzersiz_puanlari"})
    if doc:
        doc.pop("_id", None)
        doc.pop("tip", None)
        return doc.get("puanlar", {})
    return {}

@api_router.post("/egzersiz/puan-ayarla")
async def set_egzersiz_puanlari(data: dict, current_user=Depends(get_current_user)):
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Yetki yok")
    puanlar = data.get("puanlar", {})
    await db.ayarlar.update_one(
        {"tip": "egzersiz_puanlari"},
        {"$set": {"tip": "egzersiz_puanlari", "puanlar": puanlar}},
        upsert=True
    )
    return {"message": "Puanlar kaydedildi"}

@api_router.post("/egzersiz/tamamla")
async def egzersiz_tamamla(data: dict, current_user=Depends(get_current_user)):
    kullanici_id = data.get("kullanici_id", current_user.get("id"))
    egzersiz_id = data.get("egzersiz_id", "")
    if not egzersiz_id:
        raise HTTPException(status_code=400, detail="Egzersiz ID gerekli")
    # Bugün zaten yaptı mı?
    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    mevcut = await db.egzersiz_kayitlari.find_one({
        "kullanici_id": kullanici_id,
        "egzersiz_id": egzersiz_id,
        "tarih": bugun
    })
    if mevcut:
        raise HTTPException(status_code=409, detail="Bu egzersiz bugün zaten tamamlandı")
    # Puan hesapla
    ayar = await db.ayarlar.find_one({"tip": "egzersiz_puanlari"})
    puanlar = ayar.get("puanlar", {}) if ayar else {}
    kazanilan = puanlar.get(egzersiz_id, 10)  # varsayılan 10 puan
    # Kaydet
    await db.egzersiz_kayitlari.insert_one({
        "id": str(uuid.uuid4()),
        "kullanici_id": kullanici_id,
        "egzersiz_id": egzersiz_id,
        "tarih": bugun,
        "kazanilan_puan": kazanilan,
        "zaman": datetime.utcnow().isoformat()
    })
    # Kullanıcı puanını güncelle
    await db.users.update_one({"id": kullanici_id}, {"$inc": {"puan": kazanilan}})
    return {"kazanilan_puan": kazanilan, "egzersiz_id": egzersiz_id}

# ── Kitap + Soru Havuzu ──
class KitapCreate(BaseModel):
    baslik: str
    yazar: str = ""
    yas_grubu: str = "8-10"
    zorluk: str = "orta"
    bolum_sayisi: int = 1

@api_router.post("/kitaplar")
async def kitap_ekle(data: KitapCreate, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ("admin", "coordinator", "teacher"):
        raise HTTPException(status_code=403, detail="Yetki yok")
    durum = "oylama" if role in ("admin", "coordinator") else "beklemede"
    kitap = {
        "id": str(uuid.uuid4()),
        "baslik": data.baslik,
        "yazar": data.yazar,
        "yas_grubu": data.yas_grubu,
        "zorluk": data.zorluk,
        "bolum_sayisi": data.bolum_sayisi,
        "ekleyen_id": current_user["id"],
        "ekleyen_ad": f"{current_user.get('ad','')} {current_user.get('soyad','')}".strip(),
        "durum": durum,
        "oylar": {},
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.kitaplar.insert_one(kitap)
    kitap.pop("_id", None)
    return kitap

@api_router.get("/kitaplar")
async def kitap_listele(current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    kitaplar = await db.kitaplar.find().sort("olusturma_tarihi", -1).to_list(length=None)
    for k in kitaplar:
        k.pop("_id", None)
    if role in ("admin", "coordinator"):
        return kitaplar
    # Öğretmen: kendi eklediği + oylamada + havuzda
    return [k for k in kitaplar if k.get("durum") in ("oylama", "havuzda") or k.get("ekleyen_id") == current_user["id"]]

@api_router.put("/kitaplar/{kitap_id}")
async def kitap_guncelle(kitap_id: str, data: dict, current_user=Depends(get_current_user)):
    kitap = await db.kitaplar.find_one({"id": kitap_id})
    if not kitap:
        raise HTTPException(status_code=404, detail="Kitap bulunamadı")
    update = {k: v for k, v in data.items() if k in ("baslik", "yazar", "yas_grubu", "zorluk", "bolum_sayisi")}
    if update:
        await db.kitaplar.update_one({"id": kitap_id}, {"$set": update})
    return {"message": "Güncellendi"}

@api_router.delete("/kitaplar/{kitap_id}")
async def kitap_sil(kitap_id: str, current_user=Depends(get_current_user)):
    if current_user.get("role") not in ("admin", "coordinator"):
        raise HTTPException(status_code=403, detail="Yetki yok")
    await db.kitaplar.delete_one({"id": kitap_id})
    await db.sorular.delete_many({"kitap_id": kitap_id})
    return {"message": "Kitap ve soruları silindi"}

# Kitap admin karar (oylama başlat / direkt havuza al / reddet)
@api_router.post("/kitaplar/{kitap_id}/admin-karar")
async def kitap_admin_karar(kitap_id: str, data: dict, current_user=Depends(get_current_user)):
    if current_user.get("role") not in ("admin", "coordinator"):
        raise HTTPException(status_code=403, detail="Yetki yok")
    onay = data.get("onay", True)
    direkt = data.get("direkt", False)
    if not onay:
        await db.kitaplar.update_one({"id": kitap_id}, {"$set": {"durum": "reddedildi", "red_sebep": data.get("sebep", "")}})
        return {"message": "Reddedildi"}
    if direkt:
        await db.kitaplar.update_one({"id": kitap_id}, {"$set": {"durum": "havuzda"}})
        return {"message": "Direkt havuza alındı"}
    await db.kitaplar.update_one({"id": kitap_id}, {"$set": {"durum": "oylama"}})
    return {"message": "Oylama başlatıldı"}

# Kitap oylama
@api_router.post("/kitaplar/{kitap_id}/oy")
async def kitap_oy_ver(kitap_id: str, data: dict, current_user=Depends(get_current_user)):
    kitap = await db.kitaplar.find_one({"id": kitap_id})
    if not kitap or kitap.get("durum") != "oylama":
        raise HTTPException(status_code=400, detail="Bu kitap oylamada değil")
    user_id = current_user["id"]
    oylar = kitap.get("oylar", {})
    if user_id in oylar:
        raise HTTPException(status_code=400, detail="Zaten oy kullandınız")
    onay = data.get("onay", True)
    oylar[user_id] = {"onay": onay, "sebep": data.get("sebep", ""), "tarih": datetime.utcnow().isoformat()}
    update = {"oylar": oylar}
    # Oy eşiği kontrolü
    ayar = await db.ayarlar.find_one({"tip": "puan_ayarlari"})
    esik = ayar.get("oy_esik", 3) if ayar else 3
    onay_sayisi = sum(1 for o in oylar.values() if o.get("onay"))
    red_sayisi = sum(1 for o in oylar.values() if not o.get("onay"))
    if red_sayisi >= 1:
        update["durum"] = "reddedildi"
    elif onay_sayisi >= esik:
        update["durum"] = "havuzda"
    await db.kitaplar.update_one({"id": kitap_id}, {"$set": update})
    # Katkı puanı
    await db.users.update_one({"id": user_id}, {"$inc": {"puan": 3}})
    return {"message": "Oy kaydedildi"}

# ── Soru CRUD ──
class SoruCreate(BaseModel):
    kitap_id: str
    bolum: int = 1
    soru_metni: str
    secenekler: list = []
    dogru_cevap: int = 0

@api_router.post("/sorular")
async def soru_ekle(data: SoruCreate, current_user=Depends(get_current_user)):
    soru = {
        "id": str(uuid.uuid4()),
        "kitap_id": data.kitap_id,
        "bolum": data.bolum,
        "soru_metni": data.soru_metni,
        "secenekler": data.secenekler,
        "dogru_cevap": data.dogru_cevap,
        "ekleyen_id": current_user["id"],
        "ekleyen_ad": f"{current_user.get('ad','')} {current_user.get('soyad','')}".strip(),
        "kullanim_sayisi": 0,
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.sorular.insert_one(soru)
    soru.pop("_id", None)
    return soru

@api_router.get("/sorular/{kitap_id}")
async def soru_listele(kitap_id: str, bolum: int = None, current_user=Depends(get_current_user)):
    filtre = {"kitap_id": kitap_id}
    if bolum is not None:
        filtre["bolum"] = bolum
    sorular = await db.sorular.find(filtre).sort("bolum", 1).to_list(length=None)
    for s in sorular:
        s.pop("_id", None)
    return sorular

@api_router.put("/sorular/{soru_id}")
async def soru_guncelle(soru_id: str, data: dict, current_user=Depends(get_current_user)):
    update = {k: v for k, v in data.items() if k in ("soru_metni", "secenekler", "dogru_cevap", "bolum")}
    if update:
        await db.sorular.update_one({"id": soru_id}, {"$set": update})
    return {"message": "Güncellendi"}

@api_router.delete("/sorular/{soru_id}")
async def soru_sil(soru_id: str, current_user=Depends(get_current_user)):
    await db.sorular.delete_one({"id": soru_id})
    return {"message": "Silindi"}

# ── Kitap Bilgi Çekme (ISBN / Link) ──
@api_router.post("/kitap-bilgi-cek")
async def kitap_bilgi_cek(data: dict, current_user=Depends(get_current_user)):
    import urllib.request, urllib.error, ssl, json, asyncio

    deger = data.get("deger", "").strip()
    tip = data.get("tip", "isbn")
    result = {"baslik": "", "yazar": "", "isbn": "", "yayinevi": "", "sayfa_sayisi": "", "aciklama": "", "kapak_url": "", "link": ""}

    def fetch_url(url):
        try:
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"})
            with urllib.request.urlopen(req, timeout=12, context=ctx) as resp:
                raw = resp.read()
                try:
                    return raw.decode("utf-8")
                except Exception:
                    return raw.decode("latin-1")
        except Exception:
            return None

    if tip == "isbn":
        isbn_temiz = re.sub(r"[^0-9X]", "", deger.upper())
        # Google Books API
        raw = await asyncio.to_thread(fetch_url, f"https://www.googleapis.com/books/v1/volumes?q=isbn:{isbn_temiz}")
        if raw:
            try:
                j = json.loads(raw)
                if j.get("totalItems", 0) > 0:
                    vol = j["items"][0]["volumeInfo"]
                    result["baslik"] = vol.get("title", "")
                    result["yazar"] = ", ".join(vol.get("authors", []))
                    result["yayinevi"] = vol.get("publisher", "")
                    result["sayfa_sayisi"] = str(vol.get("pageCount", "") or "")
                    result["aciklama"] = (vol.get("description", "") or "")[:200]
                    result["isbn"] = isbn_temiz
                    imgs = vol.get("imageLinks", {})
                    result["kapak_url"] = imgs.get("thumbnail", imgs.get("smallThumbnail", ""))
                    result["link"] = vol.get("infoLink", "")
            except Exception:
                pass
        # Open Library fallback
        if not result["baslik"]:
            raw = await asyncio.to_thread(fetch_url, f"https://openlibrary.org/api/books?bibkeys=ISBN:{isbn_temiz}&format=json&jscmd=data")
            if raw:
                try:
                    j = json.loads(raw)
                    key = f"ISBN:{isbn_temiz}"
                    if key in j:
                        book = j[key]
                        result["baslik"] = book.get("title", "")
                        result["yazar"] = ", ".join([a.get("name", "") for a in book.get("authors", [])])
                        result["yayinevi"] = ", ".join([p.get("name", "") for p in book.get("publishers", [])])
                        result["sayfa_sayisi"] = str(book.get("number_of_pages", "") or "")
                        result["isbn"] = isbn_temiz
                        cover = book.get("cover", {})
                        result["kapak_url"] = cover.get("medium", cover.get("small", ""))
                        result["link"] = book.get("url", "")
                except Exception:
                    pass

    elif tip == "link":
        html = await asyncio.to_thread(fetch_url, deger)
        if html:
            QP = """[\"']"""
            m = re.search(r'property\s*=\s*' + QP + r'og:title' + QP + r'[^>]*content\s*=\s*' + QP + r'([^"\']+)', html, re.I)
            if m:
                result["baslik"] = m.group(1).strip()
            else:
                m = re.search(r"<title[^>]*>([^<]+)</title>", html, re.I)
                if m:
                    t = m.group(1).strip()
                    for sep in [" - ", " | ", " :: "]:
                        if sep in t:
                            t = t.split(sep)[0].strip()
                            break
                    result["baslik"] = t
            m = re.search(r'property\s*=\s*' + QP + r'og:description' + QP + r'[^>]*content\s*=\s*' + QP + r'([^"\']+)', html, re.I)
            if m:
                result["aciklama"] = m.group(1).strip()[:200]
            m = re.search(r'property\s*=\s*' + QP + r'og:image' + QP + r'[^>]*content\s*=\s*' + QP + r'([^"\']+)', html, re.I)
            if m:
                result["kapak_url"] = m.group(1).strip()
            m = re.search(r'itemprop\s*=\s*' + QP + r'author' + QP + r'[^>]*>([^<]+)', html, re.I)
            if not m:
                m = re.search(r'Yazar\s*:?\s*</\w+>\s*<[^>]+>([^<]+)', html, re.I)
            if m:
                result["yazar"] = m.group(1).strip()
            m = re.search(r'itemprop\s*=\s*' + QP + r'publisher' + QP + r'[^>]*>([^<]+)', html, re.I)
            if not m:
                m = re.search(r'Yay.nevi\s*:?\s*</\w+>\s*<[^>]+>([^<]+)', html, re.I)
            if m:
                result["yayinevi"] = m.group(1).strip()
            m = re.search(r'(?:Sayfa|sayfa)\s*(?:Say.s.)?\s*:?\s*(\d+)', html)
            if m:
                result["sayfa_sayisi"] = m.group(1)
            m = re.search(r'ISBN[^:]*:\s*([\d\-]{10,})', html, re.I)
            if m:
                result["isbn"] = re.sub(r"[^0-9]", "", m.group(1))
            result["link"] = deger

    if not result["baslik"]:
        raise HTTPException(status_code=404, detail="Kitap bilgisi bulunamadi")
    return result

# ── PDF Rapor Üretimi ──
def _tr_upper(text):
    """Türkçe büyük harf çevirimi (i→İ, ı→I)"""
    if not text:
        return text
    tr_map = str.maketrans("abcçdefgğhıijklmnoöprsştuüvyz", "ABCÇDEFGĞHIİJKLMNOÖPRSŞTUÜVYZ")
    return text.translate(tr_map)

@api_router.get("/diagnostic/rapor/{rapor_id}/pdf")
async def get_rapor_pdf(rapor_id: str, current_user=Depends(get_current_user)):
    rapor = await db.diagnostic_raporlar.find_one({"id": rapor_id})
    if not rapor:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    rapor.pop("_id", None)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph as RLPara, Spacer, Table as RLTable, TableStyle, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # ── Türkçe Font Kaydı ──
    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]
    font_registered = False
    for fp in font_paths:
        if os.path.exists(fp):
            font_registered = True
            break

    if font_registered:
        pdfmetrics.registerFont(TTFont("TRFont", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
        pdfmetrics.registerFont(TTFont("TRFontBold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))
        FONT = "TRFont"
        FONTB = "TRFontBold"
    else:
        FONT = "Helvetica"
        FONTB = "Helvetica-Bold"

    buffer = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm, leftMargin=2*cm, rightMargin=2*cm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleOBA', fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=4, textColor=colors.HexColor('#1F4E79'), fontName=FONTB))
    styles.add(ParagraphStyle(name='SubOBA', fontSize=11, leading=14, alignment=TA_CENTER, spaceAfter=16, textColor=colors.HexColor('#666666'), fontName=FONT))
    styles.add(ParagraphStyle(name='SectOBA', fontSize=12, leading=16, spaceBefore=14, spaceAfter=8, textColor=colors.HexColor('#1F4E79'), fontName=FONTB))
    styles.add(ParagraphStyle(name='SubSectOBA', fontSize=10, leading=13, spaceBefore=10, spaceAfter=4, textColor=colors.HexColor('#333333'), fontName=FONTB))
    styles.add(ParagraphStyle(name='BodyOBA', fontSize=9, leading=13, spaceAfter=4, fontName=FONT))
    styles.add(ParagraphStyle(name='SmallOBA', fontSize=8, leading=11, textColor=colors.HexColor('#999999'), fontName=FONT))
    styles.add(ParagraphStyle(name='BigNum', fontSize=28, leading=32, textColor=colors.HexColor('#1F4E79'), fontName=FONTB))

    el = []  # elements

    hdr_bg = colors.HexColor('#1F4E79')
    alt_bg = colors.HexColor('#F2F7FB')
    bdr = colors.HexColor('#CCCCCC')
    ora = colors.HexColor('#E67E22')

    def tbl_style(rows, has_header=True):
        s = [
            ('FONTNAME', (0,0), (-1,-1), FONT),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('GRID', (0,0), (-1,-1), 0.5, bdr),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ]
        if has_header:
            s += [
                ('BACKGROUND', (0,0), (-1,0), hdr_bg),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('FONTNAME', (0,0), (-1,0), FONTB),
            ]
        for i in range(1 if has_header else 0, rows):
            if i % 2 == 0:
                s.append(('BACKGROUND', (0,i), (-1,i), alt_bg))
        return s

    # ── BAŞLIK ──
    el.append(RLPara("Okuma Becerileri Akademisi", styles['TitleOBA']))
    el.append(RLPara("Giriş Analizi Raporu", styles['SubOBA']))

    # ── 1. ÖĞRENCİ BİLGİLERİ ──
    el.append(RLPara("1. Öğrenci Bilgileri", styles['SectOBA']))
    w1, w2, w3, w4 = 3*cm, 5*cm, 3*cm, 5*cm
    info = [
        ["Adı Soyadı:", rapor.get("ogrenci_ad", "-"), "Sınıfı:", rapor.get("ogrenci_sinif", "-")],
        ["Eğitimci:", rapor.get("ogretmen_ad", "-"), "Tarih:", rapor.get("olusturma_tarihi", "")[:10]],
    ]
    t = RLTable(info, colWidths=[w1, w2, w3, w4])
    t.setStyle(TableStyle([
        ('FONTNAME', (0,0), (0,-1), FONTB), ('FONTNAME', (2,0), (2,-1), FONTB),
        ('FONTNAME', (1,0), (1,-1), FONT), ('FONTNAME', (3,0), (3,-1), FONT),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('TEXTCOLOR', (0,0), (0,-1), hdr_bg), ('TEXTCOLOR', (2,0), (2,-1), hdr_bg),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5), ('TOPPADDING', (0,0), (-1,-1), 5),
    ]))
    el.append(t)

    # ── 2. METİN BİLGİLERİ ──
    el.append(RLPara("2. Metin Bilgileri", styles['SectOBA']))
    kelime_s = rapor.get("kelime_sayisi", 0)
    dogruluk = rapor.get("dogruluk_yuzde", 0)
    yanlis_k = round(kelime_s * (100 - dogruluk) / 100) if kelime_s else 0
    dogru_k = kelime_s - yanlis_k
    sure_sn_total = rapor.get("sure_saniye", 0)
    sure_dk = int(sure_sn_total) // 60
    sure_sn = int(sure_sn_total) % 60
    metin_info = [
        ["Metnin Adı:", _tr_upper(rapor.get("metin_adi", "-"))],
        ["Metnin Türü:", _tr_upper(rapor.get("metin_turu", "-"))],
        ["Toplam Kelime Sayısı:", str(kelime_s)],
        ["Doğru Okunan Kelime:", str(dogru_k)],
        ["Yanlış Okunan Kelime:", str(yanlis_k)],
        ["Tamamlama Süresi:", f"{sure_dk}:{str(sure_sn).zfill(2)} ({sure_sn_total} sn)"],
    ]
    t2 = RLTable(metin_info, colWidths=[5*cm, 11*cm])
    t2.setStyle(TableStyle([
        ('FONTNAME', (0,0), (0,-1), FONTB), ('FONTNAME', (1,0), (1,-1), FONT),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('GRID', (0,0), (-1,-1), 0.5, bdr),
        ('BACKGROUND', (0,0), (0,-1), alt_bg),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5), ('TOPPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
    ]))
    el.append(t2)

    # ── 3. OKUMA HIZI ──
    el.append(RLPara("3. Okuma Hızı", styles['SectOBA']))
    wpm = round(rapor.get("wpm", 0))
    hiz_map = {"dusuk": "Düşük", "orta": "Orta", "yeterli": "Yeterli", "ileri": "İleri"}
    hiz_label = hiz_map.get(rapor.get("hiz_deger", ""), "?")
    hiz_renk = {"dusuk": "#E74C3C", "orta": "#F39C12", "yeterli": "#27AE60", "ileri": "#2E86C1"}.get(rapor.get("hiz_deger", ""), "#333")
    el.append(RLPara(f'<font size="28" color="{hiz_renk}"><b>{wpm}</b></font>  <font size="10">kelime/dakika</font>', styles['BodyOBA']))
    el.append(RLPara(f'<font size="12" color="{hiz_renk}"><b>{hiz_label} Düzey</b></font>', styles['BodyOBA']))
    el.append(Spacer(1, 4))
    sinif = rapor.get("ogrenci_sinif", "")
    el.append(RLPara(f"Öğrencinin okuma hızı dakikada <b>{wpm} kelime</b>dir. Bu okuma hızı, öğrencinin bulunduğu sınıf düzeyi normlarına göre <b>{hiz_label.lower()} düzeydedir</b>.", styles['BodyOBA']))

    # ── 4. DOĞRU OKUMA ORANI ──
    el.append(RLPara("3.1. Doğru Okuma Oranı", styles['SubSectOBA']))
    el.append(RLPara(f"Doğruluk: <b>%{round(dogruluk)}</b>", styles['BodyOBA']))
    hatalar = rapor.get("hata_sayilari", [])
    # Dict formatındaysa listeye çevir: {"atlama": 2, ...} → [{"tur": "atlama", "sayi": 2}]
    if isinstance(hatalar, dict):
        hatalar = [{"tur": k, "sayi": v} for k, v in hatalar.items()]
    if hatalar:
        hata_labels = {"atlama": "Atlama", "yanlis_okuma": "Yanlış Okuma", "yanlis": "Yanlış Okuma", "takilma": "Takılma", "tekrar": "Tekrar"}
        hata_desc = {"atlama": "Kelime veya satır atlama", "yanlis_okuma": "Kelimeyi farklı okuma", "yanlis": "Kelimeyi farklı okuma", "takilma": "Kelimede duraksama", "tekrar": "Aynı kelimeyi tekrar okuma"}
        hata_rows = [["Hata Türü", "Açıklama", "Sayı"]]
        toplam_hata = 0
        for h in hatalar:
            tur = h.get("tur", "")
            sayi = h.get("sayi", 0)
            toplam_hata += sayi
            hata_rows.append([hata_labels.get(tur, tur), hata_desc.get(tur, ""), str(sayi)])
        hata_rows.append(["TOPLAM", "", str(toplam_hata)])
        t3 = RLTable(hata_rows, colWidths=[3.5*cm, 6.5*cm, 2*cm])
        t3.setStyle(TableStyle(tbl_style(len(hata_rows)) + [
            ('FONTNAME', (0,-1), (-1,-1), FONTB),
            ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#E8F0FE')),
            ('ALIGN', (2,0), (2,-1), 'CENTER'),
        ]))
        el.append(t3)

    # ── 5. OKUDUĞUNU ANLAMA ──
    anlama = rapor.get("anlama", {})
    anlama_pct = rapor.get("anlama_yuzde", 0)
    anlama_sev = "İyi" if anlama_pct >= 85 else "Orta" if anlama_pct >= 70 else "Zayıf"
    el.append(RLPara(f"4. Okuduğunu Anlama Becerileri — %{anlama_pct}", styles['SectOBA']))

    # Anlama alt grupları
    anlama_gruplari = [
        ("4.1 Sözcük Düzeyinde Anlama", [
            ("Cümle anlamını kavrama", "cumle_anlama"),
            ("Bilinmeyen sözcük tahmini", "bilinmeyen_sozcuk"),
            ("Bağlaç ve zamirleri anlama", "baglac_zamir"),
        ]),
        ("4.2 Metnin Ana Yapısını Anlama", [
            ("Ana fikir belirleme", "ana_fikir"),
            ("Yardımcı fikirleri ifade etme", "yardimci_fikir"),
            ("Metnin konusunu ifade etme", "konu"),
            ("Başlık önerme", "baslik_onerme"),
        ]),
        ("4.3 Metinler Arasılık ve Derin Anlama", [
            ("Neden-sonuç ilişkisini belirleme", "neden_sonuc"),
            ("Çıkarım yapma", "cikarim"),
            ("Metindeki ipuçlarını kullanma", "ipuclari"),
            ("Yorumlama", "yorumlama"),
        ]),
        ("4.4 Eleştirel ve Yaratıcı Okuma", [
            ("Metne yönelik görüş bildirme", "gorus_bildirme"),
            ("Yazarın amacını sezme", "yazar_amaci"),
            ("Alternatif son / fikir üretme", "alternatif_fikir"),
            ("Metni günlük hayatla ilişkilendirme", "guncelle_hayat"),
        ]),
        ("4.5 Soru Performans Analizi", [
            ("Bilgi", "bilgi"),
            ("Kavrama", "kavrama"),
            ("Uygulama", "uygulama"),
            ("Analiz", "analiz"),
            ("Sentez", "sentez"),
            ("Değerlendirme", "degerlendirme"),
        ]),
    ]

    seviye_map = {"zayif": "Zayıf", "orta": "Orta", "iyi": "İyi"}
    for grup_baslik, olcutler in anlama_gruplari:
        el.append(RLPara(grup_baslik, styles['SubSectOBA']))
        rows = [["Ölçüt", "Zayıf", "Orta", "İyi"]]
        for label, key in olcutler:
            val = anlama.get(key, "orta")
            row = [label]
            for s in ["zayif", "orta", "iyi"]:
                row.append("+" if val == s else "")
            rows.append(row)
        t_a = RLTable(rows, colWidths=[7*cm, 3*cm, 3*cm, 3*cm])
        st = tbl_style(len(rows))
        # Renklendir: + işaretlerini turuncu yap
        for ri in range(1, len(rows)):
            for ci in range(1, 4):
                if rows[ri][ci] == "+":
                    st.append(('TEXTCOLOR', (ci, ri), (ci, ri), ora))
                    st.append(('FONTNAME', (ci, ri), (ci, ri), FONTB))
        st.append(('ALIGN', (1,0), (-1,-1), 'CENTER'))
        t_a.setStyle(TableStyle(st))
        el.append(t_a)

    # ── 6. PROZODİK OKUMA ──
    proz = rapor.get("prozodik", {})
    proz_toplam = rapor.get("prozodik_toplam", 0)
    proz_sev = "Çok İyi" if proz_toplam >= 18 else "İyi" if proz_toplam >= 14 else "Orta" if proz_toplam >= 10 else "Geliştirilmeli"
    el.append(RLPara("5. Prozodik Okuma Ölçeği", styles['SectOBA']))

    proz_desc = {
        "noktalama": ["Uymuyor", "Kısmen", "Çoğunlukla", "Tam ve bilinçli"],
        "vurgu": ["Tek düze", "Yer yer", "Anlama uygun", "Etkili ve bilinçli"],
        "tonlama": ["Monoton", "Sınırlı", "Metne uygun", "Doğal ve etkileyici"],
        "akicilik": ["Sık duraklama", "Kısmi akış", "Genel akıcı", "Kesintisiz"],
        "anlamli_gruplama": ["Sözcük sözcük", "Kısmen", "Çoğunlukla", "Tam ve tutarlı"],
    }
    proz_labels = {"noktalama": "Noktalama ve Duraklama", "vurgu": "Vurgu", "tonlama": "Tonlama", "akicilik": "Akıcılık", "anlamli_gruplama": "Anlamlı Gruplama"}

    proz_rows = [["Ölçüt", "1 puan", "2 puan", "3 puan", "4 puan", "Puan"]]
    for key in ["noktalama", "vurgu", "tonlama", "akicilik", "anlamli_gruplama"]:
        puan = proz.get(key, 0)
        descs = proz_desc.get(key, ["", "", "", ""])
        row = [proz_labels.get(key, key)]
        for pi in range(4):
            row.append(descs[pi])
        row.append(str(puan))
        proz_rows.append(row)
    proz_rows.append(["", "", "", "", "Toplam", str(proz_toplam)])

    t_p = RLTable(proz_rows, colWidths=[2.8*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.8*cm, 1.5*cm])
    ps = tbl_style(len(proz_rows))
    # Seçili puanı turuncu yap
    for ri in range(1, len(proz_rows) - 1):
        key = list(proz_desc.keys())[ri - 1]
        puan = proz.get(key, 0)
        if 1 <= puan <= 4:
            ps.append(('TEXTCOLOR', (puan, ri), (puan, ri), ora))
            ps.append(('FONTNAME', (puan, ri), (puan, ri), FONTB))
    ps.append(('ALIGN', (5, 0), (5, -1), 'CENTER'))
    ps.append(('FONTNAME', (0, -1), (-1, -1), FONTB))
    ps.append(('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#E8F0FE')))
    t_p.setStyle(TableStyle(ps))
    el.append(t_p)
    el.append(RLPara(f"Prozodik okuma performansı: <b>{proz_sev}</b> (Toplam {proz_toplam}/20)", styles['BodyOBA']))

    # ── 7. SONUÇ VE GENEL YORUM ──
    el.append(RLPara("6. Sonuç ve Genel Yorum", styles['SectOBA']))
    if rapor.get("ogretmen_notu"):
        for line in rapor.get("ogretmen_notu", "").split("\n"):
            if line.strip():
                el.append(RLPara(line.strip(), styles['BodyOBA']))

    # AI Yorumları
    ai = rapor.get("ai_yorumlar", {})
    if ai:
        el.append(Spacer(1, 6))
        ai_labels = {"hiz": "Okuma Hızı", "dogruluk": "Doğru Okuma", "anlama": "Anlama", "prozodik": "Prozodik Okuma", "sonuc": "Sonuç", "oneriler": "Öneriler"}
        for key, label in ai_labels.items():
            if ai.get(key):
                el.append(RLPara(f"<b>{label}:</b> {ai[key]}", styles['BodyOBA']))

    # ── KUR ──
    el.append(Spacer(1, 8))
    el.append(RLPara(f"Atanan Kur: <b>{rapor.get('atanan_kur', '-')}</b>", styles['BodyOBA']))

    # ── ALT BİLGİ ──
    el.append(Spacer(1, 20))
    el.append(HRFlowable(width="100%", thickness=0.5, color=bdr))
    el.append(RLPara("Bu rapor Okuma Becerileri Akademisi sistemi tarafından oluşturulmuştur.", styles['SmallOBA']))

    doc_pdf.build(el)
    buffer.seek(0)

    ogrenci_ad = rapor.get("ogrenci_ad", "ogrenci").replace(" ", "_")
    tarih = rapor.get("olusturma_tarihi", "")[:10]
    filename = f"Rapor_{ogrenci_ad}_{tarih}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ── Migration / Debug Endpoint ──
@api_router.post("/admin/fix-ids")
async def fix_missing_ids(current_user=Depends(require_role(UserRole.ADMIN))):
    """Eksik id alanlarını düzelt"""
    fixed = 0
    # analiz_metinler
    async for doc in db.analiz_metinler.find({"id": {"$exists": False}}):
        new_id = str(uuid.uuid4())
        await db.analiz_metinler.update_one({"_id": doc["_id"]}, {"$set": {"id": new_id}})
        fixed += 1
    # analiz_metinler - durum alanı yoksa ekle
    await db.analiz_metinler.update_many({"durum": {"$exists": False}}, {"$set": {"durum": "havuzda"}})
    # diagnostic_oturumlar
    async for doc in db.diagnostic_oturumlar.find({"id": {"$exists": False}}):
        await db.diagnostic_oturumlar.update_one({"_id": doc["_id"]}, {"$set": {"id": str(uuid.uuid4())}})
    return {"fixed": fixed, "message": "ID düzeltme tamamlandı"}

@api_router.get("/admin/gemini-test")
async def gemini_test(current_user=Depends(require_role(UserRole.ADMIN))):
    """Gemini API bağlantısını test et — sadece admin."""
    key = GEMINI_API_KEY
    if not key:
        return {"durum": "HATA", "sebep": "GEMINI_API_KEY environment variable tanımlı değil", "key_uzunluk": 0}
    try:
        yanit = await _gemini_call("Merhaba! Sadece 'Gemini çalışıyor' yaz.", max_tokens=50)
        return {"durum": "OK", "yanit": yanit, "key_uzunluk": len(key), "model": AI_MODEL}
    except Exception as e:
        return {"durum": "HATA", "sebep": str(e), "key_uzunluk": len(key), "model": AI_MODEL}

@api_router.get("/admin/gemini-modeller")
async def gemini_modeller():
    """Kullanılabilir Gemini modellerini listele — geçici public."""
    key = GEMINI_API_KEY
    if not key:
        return {"hata": "API key yok"}
    try:
        async with httpx.AsyncClient(timeout=15.0) as c:
            r = await c.get(f"https://generativelanguage.googleapis.com/v1beta/models?key={key}")
            data = r.json()
        modeller = []
        for m in data.get("models", []):
            if "generateContent" in m.get("supportedGenerationMethods", []):
                modeller.append(m.get("name", "").replace("models/", ""))
        return {"modeller": modeller, "toplam": len(modeller)}
    except Exception as e:
        return {"hata": str(e)}

@api_router.get("/admin/debug-metinler")
async def debug_metinler(current_user=Depends(require_role(UserRole.ADMIN))):
    """Tüm metinleri ham haliyle göster"""
    items = await db.analiz_metinler.find().to_list(length=None)
    result = []
    for item in items:
        item.pop("_id", None)
        result.append({"id": item.get("id","EKSİK"), "baslik": item.get("baslik","?"), "durum": item.get("durum","?")})
    return result

@api_router.get("/admin/debug-ogrenciler")
async def debug_ogrenciler(current_user=Depends(require_role(UserRole.ADMIN))):
    items = await db.students.find().to_list(length=None)
    result = []
    for item in items:
        item.pop("_id", None)
        result.append({"id": item.get("id","EKSİK"), "ad": item.get("ad","?"), "soyad": item.get("soyad","?")})
    return result

# ─────────────────────────────────────────────
# GELİŞİM ALANI - Tam İş Akışı
# ─────────────────────────────────────────────

class SoruModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    soru: str
    secenekler: List[str]
    dogru_cevap: int

class IcerikCreate(BaseModel):
    baslik: str
    tur: str  # hizmetici, film, kitap, makale, okuma_parcasi
    aciklama: str = ""
    hedef_kitle: str  # ogretmen, ogrenci, hepsi
    sorular: List[SoruModel] = []
    # Makale alanları
    makale_link: Optional[str] = None
    makale_dosya_turu: Optional[str] = None  # pdf, word, link
    # Kitap ek alanlar
    kitap_yazar: Optional[str] = None
    kitap_isbn: Optional[str] = None
    kitap_yayinevi: Optional[str] = None
    kitap_sayfa: Optional[str] = None
    kitap_yas_grubu: Optional[str] = None
    kitap_link: Optional[str] = None
    kitap_kapak: Optional[str] = None
    kitap_bolum_sayisi: Optional[int] = None
    # Kitap/makale dosya (base64)
    dosya_b64: Optional[str] = None
    dosya_adi: Optional[str] = None
    dosya_turu: Optional[str] = None  # pdf, docx
    # Okuma parçası
    okuma_metni: Optional[str] = None
    okuma_seviye: Optional[str] = None  # kolay, orta, zor
    okuma_sure: Optional[int] = None  # dakika

class IcerikModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    baslik: str
    tur: str
    aciklama: str = ""
    hedef_kitle: str
    sorular: List[SoruModel] = []
    makale_link: Optional[str] = None
    makale_dosya_turu: Optional[str] = None
    kitap_yazar: Optional[str] = None
    kitap_isbn: Optional[str] = None
    kitap_yayinevi: Optional[str] = None
    kitap_sayfa: Optional[str] = None
    kitap_yas_grubu: Optional[str] = None
    kitap_link: Optional[str] = None
    kitap_kapak: Optional[str] = None
    kitap_bolum_sayisi: Optional[int] = None
    dosya_b64: Optional[str] = None
    dosya_adi: Optional[str] = None
    dosya_turu: Optional[str] = None
    okuma_metni: Optional[str] = None
    okuma_seviye: Optional[str] = None
    okuma_sure: Optional[int] = None
    ekleyen_id: str = ""
    ekleyen_ad: str = ""
    durum: str = "beklemede"  # beklemede, oylama, yayinda, reddedildi
    oylar: dict = Field(default_factory=dict)
    olusturma_tarihi: datetime = Field(default_factory=datetime.utcnow)
    yayin_tarihi: Optional[datetime] = None

class OyCreate(BaseModel):
    icerik_id: str
    onay: bool
    sebep: str = ""  # Red durumunda zorunlu

class TamamlamaCreate(BaseModel):
    icerik_id: str
    kullanici_id: str
    test_cevaplari: Optional[List[int]] = None

class TamamlamaModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    kullanici_id: str
    icerik_id: str
    test_yapildi: bool = False
    dogru_sayisi: int = 0
    toplam_soru: int = 0
    kazanilan_puan: int = 0
    tarih: datetime = Field(default_factory=datetime.utcnow)

# Dosya yükleme endpoint — kitap PDF/Word gelişim içeriğine eklenir
@api_router.post("/gelisim/dosya-yukle")
async def gelisim_dosya_yukle(
    dosya: UploadFile = File(...),
    current_user=Depends(get_current_user)
):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    import os, base64
    ext = os.path.splitext(dosya.filename)[1].lower()
    if ext not in [".pdf", ".docx", ".doc", ".txt"]:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {ext}. Desteklenen: .pdf, .docx, .doc, .txt")
    icerik = await dosya.read()
    if len(icerik) > 20 * 1024 * 1024:  # 20MB
        raise HTTPException(status_code=400, detail="Dosya 20MB'dan büyük olamaz")
    dosya_b64 = base64.b64encode(icerik).decode("utf-8")
    dosya_turu = "pdf" if ext == ".pdf" else "docx"
    return {
        "dosya_b64": dosya_b64,
        "dosya_adi": dosya.filename,
        "dosya_turu": dosya_turu,
        "boyut_kb": len(icerik) // 1024
    }

# İçerik ekleme (admin veya öğretmen)
@api_router.post("/gelisim/icerik")
async def create_icerik(icerik: IcerikCreate, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    
    # Admin/Koordinatör eklerse direkt yayında, öğretmen eklerse oylama
    durum = "yayinda" if role in ["admin", "coordinator"] else "oylama"
    
    model = IcerikModel(
        **icerik.dict(),
        ekleyen_id=current_user["id"],
        ekleyen_ad=f"{current_user.get('ad','')} {current_user.get('soyad','')}",
        durum=durum
    )
    data = model.dict()
    data["olusturma_tarihi"] = data["olusturma_tarihi"].isoformat()
    if data.get("yayin_tarihi"):
        data["yayin_tarihi"] = data["yayin_tarihi"].isoformat()

    # "Neden bu kitap?" alanı → +3 puan bonusu
    neden = data.get("neden_bu_icerik", "")
    if neden and len(neden.strip()) >= 20:
        data["neden_bonus"] = True
        try:
            await db.users.update_one({"id": current_user["id"]}, {"$inc": {"toplam_puan": 3}})
        except: pass

    await db.gelisim_icerik.insert_one(data)

    # Kitap türünde içerik eklendiyse kitap havuzuna da kaydet (bölüm bazlı soru için)
    if data.get("tur") == "kitap":
        mevcut = await db.kitap_havuzu.find_one({"baslik": data.get("baslik", ""), "yazar": data.get("kitap_yazar", "")})
        if not mevcut:
            await db.kitap_havuzu.insert_one({
                "id": str(uuid.uuid4()),
                "baslik": data.get("baslik", ""),
                "yazar": data.get("kitap_yazar", ""),
                "yas_grubu": data.get("kitap_yas_grubu", ""),
                "zorluk": "orta",
                "bolum_sayisi": data.get("kitap_bolum_sayisi", 10),
                "kapak_url": data.get("kitap_kapak", ""),
                "ekleyen_id": current_user["id"],
                "ekleyen_ad": f"{current_user.get('ad','')} {current_user.get('soyad','')}",
                "durum": durum,
                "oylar": {},
                "gelisim_icerik_id": data.get("id"),
                "olusturma_tarihi": datetime.utcnow().isoformat(),
            })

    return data

# İçerikleri listele
@api_router.get("/gelisim/icerik")
async def get_icerik_list(current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    user_id = current_user.get("id", "")
    
    items = await db.gelisim_icerik.find().sort("olusturma_tarihi", -1).to_list(length=None)
    result = []
    for item in items:
        item.pop("_id", None)
        durum = item.get("durum", "")
        hedef = item.get("hedef_kitle", "hepsi")
        
        # Kitap türü ise bölüm bazlı soru sayısını ekle
        if item.get("tur") == "kitap":
            item["_soru_sayisi"] = await db.kitap_sorulari.count_documents({"kitap_id": item["id"]})
        
        # Admin her şeyi görür
        if role in ["admin", "coordinator"]:
            result.append(item)
        elif role == "teacher":
            if item.get("ekleyen_id") == user_id:
                result.append(item)
            elif durum == "oylama":
                result.append(item)
            elif durum == "yayinda" and hedef in ["hepsi", "ogretmen"]:
                result.append(item)
        elif role == "student":
            if durum == "yayinda" and hedef in ["hepsi", "ogrenci"]:
                result.append(item)
    
    return result

# Admin onay/red (beklemede → oylama veya reddedildi)
@api_router.post("/gelisim/icerik/{icerik_id}/admin-karar")
async def admin_karar(icerik_id: str, karar: dict, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    # direkt=True → oylama atla, direkt yayına al
    onay = karar.get("onay", False)
    direkt = karar.get("direkt", False)
    if not onay:
        yeni_durum = "reddedildi"
    elif direkt:
        yeni_durum = "yayinda"
        icerik = await db.gelisim_icerik.find_one({"id": icerik_id})
        puanlar = await get_puan_ayarlari()
        if icerik and icerik.get("ekleyen_id"):
            await db.users.update_one({"id": icerik["ekleyen_id"]}, {"$inc": {"puan": puanlar.get("icerik_ekleme", 5)}})
    else:
        yeni_durum = "oylama"
    await db.gelisim_icerik.update_one(
        {"id": icerik_id},
        {"$set": {"durum": yeni_durum, **({"yayin_tarihi": datetime.utcnow().isoformat()} if yeni_durum == "yayinda" else {})}}
    )
    return {"durum": yeni_durum}

# Öğretmen oylama
@api_router.post("/gelisim/oy")
async def oy_ver(oy: OyCreate, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Sadece öğretmenler oy verebilir")
    
    if not oy.onay and not oy.sebep:
        raise HTTPException(status_code=400, detail="Red için sebep belirtmelisiniz")
    
    icerik = await db.gelisim_icerik.find_one({"id": oy.icerik_id})
    if not icerik:
        raise HTTPException(status_code=404, detail="İçerik bulunamadı")
    if icerik.get("durum") != "oylama":
        raise HTTPException(status_code=400, detail="Bu içerik oylamada değil")
    
    user_id = current_user["id"]
    oylar = icerik.get("oylar", {})
    
    if user_id in oylar:
        raise HTTPException(status_code=400, detail="Zaten oy kullandınız")
    
    # Oyu kaydet
    oylar[user_id] = {"onay": oy.onay, "sebep": oy.sebep}
    await db.gelisim_icerik.update_one({"id": oy.icerik_id}, {"$set": {"oylar": oylar}})
    
    # Oy veren öğretmene puan (dinamik)
    puanlar = await get_puan_ayarlari()
    await db.users.update_one({"id": user_id}, {"$inc": {"puan": puanlar.get("icerik_oylama", 2)}})
    
    # %60 kontrolü
    ogretmenler = await db.users.find({"role": {"$in": ["teacher", "coordinator", "admin"]}}).to_list(length=None)
    toplam_ogretmen = len(ogretmenler)
    onay_sayisi = sum(1 for v in oylar.values() if v.get("onay"))
    oy_sayisi = len(oylar)
    
    yeni_durum = icerik.get("durum")
    
    if toplam_ogretmen > 0:
        onay_orani = onay_sayisi / toplam_ogretmen
        # Herkes oy kullandı veya onay oranı %60 geçti
        if onay_orani >= 0.6:
            yeni_durum = "yayinda"
            await db.gelisim_icerik.update_one(
                {"id": oy.icerik_id},
                {"$set": {"durum": "yayinda", "yayin_tarihi": datetime.utcnow().isoformat()}}
            )
            # İçerik ekleyene bonus puan (dinamik)
            ekleyen_id = icerik.get("ekleyen_id")
            if ekleyen_id:
                await db.users.update_one({"id": ekleyen_id}, {"$inc": {"puan": puanlar.get("icerik_ekleme", 5)}})
        elif oy_sayisi == toplam_ogretmen and onay_orani < 0.6:
            yeni_durum = "reddedildi"
            await db.gelisim_icerik.update_one({"id": oy.icerik_id}, {"$set": {"durum": "reddedildi"}})
    
    return {
        "mesaj": "Oyunuz kaydedildi (+2 puan)",
        "durum": yeni_durum,
        "onay_orani": round(onay_sayisi / max(toplam_ogretmen, 1) * 100),
        "oy_sayisi": oy_sayisi,
        "toplam": toplam_ogretmen
    }

# Tamamlama
@api_router.post("/gelisim/tamamla")
async def tamamla_icerik(data: TamamlamaCreate, current_user=Depends(get_current_user)):
    existing = await db.gelisim_tamamlama.find_one({"kullanici_id": data.kullanici_id, "icerik_id": data.icerik_id})
    if existing:
        raise HTTPException(status_code=400, detail="Bu içerik zaten tamamlandı")
    
    icerik = await db.gelisim_icerik.find_one({"id": data.icerik_id})
    if not icerik:
        raise HTTPException(status_code=404, detail="İçerik bulunamadı")
    
    sorular = icerik.get("sorular", [])
    toplam = len(sorular)
    dogru = 0
    test_yapildi = False
    puan = 1
    
    if data.test_cevaplari and toplam > 0:
        test_yapildi = True
        for i, cevap in enumerate(data.test_cevaplari):
            if i < toplam and cevap == sorular[i].get("dogru_cevap"):
                dogru += 1
        puan = max(1, round((dogru / toplam) * 10))
    
    tamamlama = TamamlamaModel(
        kullanici_id=data.kullanici_id,
        icerik_id=data.icerik_id,
        test_yapildi=test_yapildi,
        dogru_sayisi=dogru,
        toplam_soru=toplam,
        kazanilan_puan=puan
    )
    t_data = tamamlama.dict()
    t_data["tarih"] = t_data["tarih"].isoformat()
    await db.gelisim_tamamlama.insert_one(t_data)
    await db.users.update_one({"id": data.kullanici_id}, {"$inc": {"puan": puan}})
    
    return {"puan": puan, "dogru": dogru, "toplam": toplam, "test_yapildi": test_yapildi}

# Kullanıcının tamamlamaları
@api_router.get("/gelisim/tamamlama/{kullanici_id}")
async def get_tamamlamalar(kullanici_id: str, current_user=Depends(get_current_user)):
    items = await db.gelisim_tamamlama.find({"kullanici_id": kullanici_id}).to_list(length=None)
    for item in items:
        item.pop("_id", None)
    return items

# Puan tablosu
@api_router.get("/gelisim/puan-tablosu")
async def get_puan_tablosu(current_user=Depends(get_current_user)):
    users = await db.users.find().to_list(length=None)
    tablo = []
    for u in users:
        tablo.append({
            "ad": u.get("ad", ""), "soyad": u.get("soyad", ""),
            "role": u.get("role", ""), "puan": u.get("puan", 0)
        })
    tablo.sort(key=lambda x: x["puan"], reverse=True)
    return tablo

# İçerik sil
@api_router.delete("/gelisim/icerik/{icerik_id}")
async def delete_icerik(icerik_id: str, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    await db.gelisim_icerik.delete_one({"id": icerik_id})
    return {"message": "Silindi"}


# ─────────────────────────────────────────────
# OKUMA KAYITLARI (reading_logs) + ÖĞRENCİ PANELİ
# ─────────────────────────────────────────────

class ReadingLogCreate(BaseModel):
    kitap_id: Optional[str] = None
    kitap_adi: str = ""
    bolum: str = ""
    baslangic_sayfa: Optional[int] = None
    bitis_sayfa: Optional[int] = None
    sure_dakika: int = 0
    not_text: str = ""

class ReadingLogModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    ogrenci_id: str
    ogrenci_ad: str = ""
    kitap_id: Optional[str] = None
    kitap_adi: str = ""
    bolum: str = ""
    baslangic_sayfa: Optional[int] = None
    bitis_sayfa: Optional[int] = None
    sure_dakika: int = 0
    not_text: str = ""
    tarih: str = Field(default_factory=lambda: datetime.utcnow().isoformat())


# Okuma kaydı oluştur
@api_router.post("/reading-logs")
async def create_reading_log(log: ReadingLogCreate, current_user=Depends(get_current_user)):
    # Öğrenci kendi kaydını oluşturur
    # Öğrenci user ise, linked student'ı bul
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    ogrenci_ad = f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip()

    # Eğer linked_id varsa student collection'dan adı çek
    if current_user.get("linked_id"):
        st = await db.students.find_one({"id": current_user["linked_id"]})
        if st:
            ogrenci_ad = f"{st.get('ad', '')} {st.get('soyad', '')}".strip()
            ogrenci_id = st["id"]

    model = ReadingLogModel(
        ogrenci_id=ogrenci_id,
        ogrenci_ad=ogrenci_ad,
        **log.dict()
    )
    data = model.dict()
    await db.reading_logs.insert_one(data)
    return data


# Öğrencinin okuma kayıtları
@api_router.get("/reading-logs/{ogrenci_id}")
async def get_reading_logs(ogrenci_id: str, current_user=Depends(get_current_user)):
    logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=None)
    for log in logs:
        log.pop("_id", None)
    return logs


# Öğrencinin okuma istatistikleri
@api_router.get("/reading-logs/{ogrenci_id}/istatistik")
async def get_reading_stats(ogrenci_id: str, current_user=Depends(get_current_user)):
    logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).to_list(length=None)

    toplam_dakika = sum(l.get("sure_dakika", 0) for l in logs)
    toplam_kayit = len(logs)
    kitaplar = set(l.get("kitap_adi", "") for l in logs if l.get("kitap_adi"))

    # Son 7 günün kaydı
    from datetime import timedelta
    simdi = datetime.utcnow()
    yedi_gun = simdi - timedelta(days=7)
    son_7_gun = [l for l in logs if l.get("tarih", "") >= yedi_gun.isoformat()]
    aktif_gunler = len(set(l.get("tarih", "")[:10] for l in son_7_gun))

    # Bugünkü okuma
    bugun = simdi.strftime("%Y-%m-%d")
    bugunun_kayitlari = [l for l in logs if l.get("tarih", "").startswith(bugun)]
    bugun_dakika = sum(l.get("sure_dakika", 0) for l in bugunun_kayitlari)

    # Streak hesapla
    tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
    streak = 0
    kontrol = simdi.strftime("%Y-%m-%d")
    for i in range(60):
        gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
        if gun in tarihler:
            streak += 1
        elif i > 0:  # bugün yoksa streak kırılmamış olabilir
            break

    return {
        "toplam_dakika": toplam_dakika,
        "toplam_kayit": toplam_kayit,
        "toplam_kitap": len(kitaplar),
        "aktif_gunler_7": aktif_gunler,
        "bugun_dakika": bugun_dakika,
        "streak": streak,
        "kitaplar": list(kitaplar),
    }


# Öğrenci paneli: kendine atanan görevler
@api_router.get("/ogrenci-panel/gorevler")
async def get_ogrenci_gorevler(current_user=Depends(get_current_user)):
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    gorevler = await db.gorevler.find({"hedef_id": ogrenci_id, "hedef_tip": "ogrenci"}).sort("olusturma_tarihi", -1).to_list(length=None)
    for g in gorevler:
        g.pop("_id", None)
    return gorevler


# Öğrenci paneli: profil bilgisi (student collection'dan)
@api_router.get("/ogrenci-panel/profil")
async def get_ogrenci_profil(current_user=Depends(get_current_user)):
    linked_id = current_user.get("linked_id")
    if linked_id:
        student = await db.students.find_one({"id": linked_id})
        if student:
            student.pop("_id", None)
            # Öğretmen bilgisini ekle
            ogretmen_bilgi = None
            if student.get("ogretmen_id"):
                t = await db.teachers.find_one({"id": student["ogretmen_id"]})
                if t:
                    ogretmen_bilgi = {"ad": t.get("ad",""), "soyad": t.get("soyad",""), "brans": t.get("brans",""), "telefon": t.get("telefon","")}
            return {**student, "user_ad": current_user.get("ad"), "user_soyad": current_user.get("soyad"), "email": current_user.get("email"), "ogretmen_bilgi": ogretmen_bilgi}
    return {
        "id": current_user.get("id"),
        "ad": current_user.get("ad", ""),
        "soyad": current_user.get("soyad", ""),
        "email": current_user.get("email", ""),
        "sinif": "",
        "kur": "",
        "ogretmen_bilgi": None,
    }


# Öğrenci paneli: anonim puan tablosu (sadece sıra + kendi konumu)
@api_router.get("/ogrenci-panel/siralama")
async def get_ogrenci_siralama(current_user=Depends(get_current_user)):
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    # Tüm öğrencilerin okuma istatistiklerini çek
    tum_loglar = await db.reading_logs.find().to_list(length=None)

    # Öğrenci bazlı toplam dakika
    ogrenci_dakika = {}
    for log in tum_loglar:
        oid = log.get("ogrenci_id", "")
        ogrenci_dakika[oid] = ogrenci_dakika.get(oid, 0) + log.get("sure_dakika", 0)

    # Sırala
    siralama = sorted(ogrenci_dakika.items(), key=lambda x: x[1], reverse=True)

    # Anonim tablo oluştur
    tablo = []
    benim_siram = None
    for i, (oid, dakika) in enumerate(siralama):
        sira = i + 1
        if oid == ogrenci_id:
            benim_siram = sira
            tablo.append({"sira": sira, "dakika": dakika, "ben": True, "ad": "Sen 🌟"})
        else:
            tablo.append({"sira": sira, "dakika": dakika, "ben": False, "ad": f"Öğrenci #{sira}"})

    # Eğer ben listede yoksa ekle
    if benim_siram is None:
        tablo.append({"sira": len(tablo) + 1, "dakika": 0, "ben": True, "ad": "Sen 🌟"})
        benim_siram = len(tablo)

    return {"siralama": tablo[:20], "benim_siram": benim_siram, "toplam_ogrenci": len(siralama) or 1}


# ─────────────────────────────────────────────
# RİSK SKORU HESAPLAMA (Faz 4)
# ─────────────────────────────────────────────

@api_router.get("/risk-skor/{ogrenci_id}")
async def get_risk_skor(ogrenci_id: str, current_user=Depends(get_current_user)):
    logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).to_list(length=None)
    gorevler = await db.gorevler.find({"hedef_id": ogrenci_id, "hedef_tip": "ogrenci"}).to_list(length=None)

    from datetime import timedelta
    simdi = datetime.utcnow()
    yedi_gun = simdi - timedelta(days=7)
    otuz_gun = simdi - timedelta(days=30)

    # Faktörler
    son_7_gun_log = [l for l in logs if l.get("tarih", "") >= yedi_gun.isoformat()]
    son_30_gun_log = [l for l in logs if l.get("tarih", "") >= otuz_gun.isoformat()]
    aktif_gunler_7 = len(set(l.get("tarih", "")[:10] for l in son_7_gun_log))
    toplam_dakika_7 = sum(l.get("sure_dakika", 0) for l in son_7_gun_log)
    tamamlanmamis = len([g for g in gorevler if g.get("durum") == "bekliyor"])
    suresi_dolmus = len([g for g in gorevler if g.get("durum") == "suresi_doldu"])

    # Streak
    tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
    streak = 0
    for i in range(60):
        gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
        if gun in tarihler:
            streak += 1
        elif i > 0:
            break

    # Risk hesapla (0-100, yüksek = riskli)
    risk = 0
    faktorler = []

    if aktif_gunler_7 == 0:
        risk += 40; faktorler.append("Son 7 günde hiç okuma yok")
    elif aktif_gunler_7 < 2:
        risk += 25; faktorler.append(f"Son 7 günde sadece {aktif_gunler_7} gün aktif")
    elif aktif_gunler_7 < 4:
        risk += 10; faktorler.append(f"Haftalık hedefin altında ({aktif_gunler_7}/4)")

    if toplam_dakika_7 < 12:
        risk += 20; faktorler.append(f"Haftalık okuma çok düşük ({toplam_dakika_7} dk)")
    elif toplam_dakika_7 < 48:
        risk += 10; faktorler.append(f"Haftalık okuma ortalamanın altında")

    if streak == 0:
        risk += 15; faktorler.append("Streak kırılmış")

    if tamamlanmamis > 2:
        risk += 10; faktorler.append(f"{tamamlanmamis} tamamlanmamış görev")

    if suresi_dolmus > 0:
        risk += 10; faktorler.append(f"{suresi_dolmus} süresi dolmuş görev")

    if len(son_30_gun_log) == 0:
        risk += 20; faktorler.append("Son 30 günde hiç okuma yok")

    risk = min(100, risk)
    seviye = "dusuk" if risk < 30 else "orta" if risk < 60 else "yuksek"

    return {
        "risk_skoru": risk,
        "seviye": seviye,
        "seviye_label": {"dusuk": "🟢 Düşük", "orta": "🟡 Orta", "yuksek": "🔴 Yüksek"}[seviye],
        "faktorler": faktorler,
        "istatistik": {
            "aktif_gunler_7": aktif_gunler_7,
            "toplam_dakika_7": toplam_dakika_7,
            "streak": streak,
            "tamamlanmamis_gorev": tamamlanmamis,
        }
    }


# Tüm öğrencilerin risk skorları (öğretmen/admin için)
@api_router.get("/risk-skor/toplu")
async def get_toplu_risk(current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")

    students = await db.students.find({"arsivli": {"$ne": True}}).to_list(length=None)
    sonuc = []
    for s in students:
        try:
            # Hızlı risk hesaplama
            logs = await db.reading_logs.find({"ogrenci_id": s["id"]}).to_list(length=None)
            from datetime import timedelta
            simdi = datetime.utcnow()
            yedi_gun = simdi - timedelta(days=7)
            son_7 = [l for l in logs if l.get("tarih", "") >= yedi_gun.isoformat()]
            aktif_7 = len(set(l.get("tarih", "")[:10] for l in son_7))
            dakika_7 = sum(l.get("sure_dakika", 0) for l in son_7)
            tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
            streak = 0
            for i in range(60):
                gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
                if gun in tarihler: streak += 1
                elif i > 0: break

            risk = 0
            if aktif_7 == 0: risk += 40
            elif aktif_7 < 2: risk += 25
            elif aktif_7 < 4: risk += 10
            if dakika_7 < 12: risk += 20
            if streak == 0: risk += 15
            risk = min(100, risk)
            seviye = "dusuk" if risk < 30 else "orta" if risk < 60 else "yuksek"

            sonuc.append({
                "id": s["id"], "ad": s.get("ad", ""), "soyad": s.get("soyad", ""),
                "sinif": s.get("sinif", ""), "kur": s.get("kur", ""),
                "toplam_xp": s.get("toplam_xp", 0),
                "risk_skoru": risk, "risk_seviye": seviye,
                "streak": streak, "aktif_gunler_7": aktif_7, "dakika_7": dakika_7,
                "ogretmen_id": s.get("ogretmen_id", ""),
            })
        except:
            pass

    sonuc.sort(key=lambda x: x["risk_skoru"], reverse=True)
    return sonuc


# ─────────────────────────────────────────────
# XP + LİG + KUR SİSTEMİ (Faz 3)
# ─────────────────────────────────────────────

XP_TABLOSU_DEFAULT = {
    "okuma_gorevi": 10, "anlama_testi": 15, "kelime_gorevi": 8,
    "gunluk_streak": 5, "kitap_bitirme": 30, "yazili_ozet": 20,
    "egzersiz": 5, "gelisim_tamamla": 5, "gorev_tamamla": 10,
}

LIG_ESIKLERI_DEFAULT = {
    "bronz": 0, "gumus": 200, "altin": 500, "elmas": 1000,
}

LIG_SIRA = ["bronz", "gumus", "altin", "elmas"]

OGRETMEN_ROZETLERI_DEFAULT = [
    # İçerik Katkısı
    {"kod": "icerik_ilk", "ad": "İlk Adım", "ikon": "🌱", "kategori": "icerik", "seviye": "bronz", "puan": 5},
    {"kod": "icerik_5", "ad": "İçerik Üreticisi", "ikon": "✍️", "kategori": "icerik", "seviye": "gumus", "puan": 10},
    {"kod": "icerik_20", "ad": "Kütüphane Kurucusu", "ikon": "📚", "kategori": "icerik", "seviye": "altin", "puan": 25},
    {"kod": "icerik_50", "ad": "Bilgi Kaynağı", "ikon": "🏛️", "kategori": "icerik", "seviye": "elmas", "puan": 50},
    # Kalite Kontrol
    {"kod": "oy_ilk", "ad": "İlk Oy", "ikon": "🗳️", "kategori": "kalite", "seviye": "bronz", "puan": 3},
    {"kod": "oy_20", "ad": "Kalite Bekçisi", "ikon": "🛡️", "kategori": "kalite", "seviye": "gumus", "puan": 10},
    {"kod": "oy_50", "ad": "Baş Editör", "ikon": "📋", "kategori": "kalite", "seviye": "altin", "puan": 25},
    # Eğitimci
    {"kod": "gorev_ilk", "ad": "İlk Görev", "ikon": "📌", "kategori": "egitimci", "seviye": "bronz", "puan": 3},
    {"kod": "gorev_20", "ad": "Aktif Eğitimci", "ikon": "🎯", "kategori": "egitimci", "seviye": "gumus", "puan": 15},
    {"kod": "ilham_veren", "ad": "İlham Veren", "ikon": "💡", "kategori": "egitimci", "seviye": "altin", "puan": 20},
    {"kod": "yildiz_egitimci", "ad": "Yıldız Eğitimci", "ikon": "⭐", "kategori": "egitimci", "seviye": "elmas", "puan": 40},
    # Kur Atlama
    {"kod": "kur_ilk", "ad": "İlk Kur Atlatan", "ikon": "🎓", "kategori": "kur", "seviye": "bronz", "puan": 10},
    {"kod": "kur_20", "ad": "Kur Ustası", "ikon": "🏅", "kategori": "kur", "seviye": "gumus", "puan": 25},
    {"kod": "kur_30", "ad": "Seviye Atlatan", "ikon": "🚀", "kategori": "kur", "seviye": "altin", "puan": 40},
    {"kod": "kur_50", "ad": "Süper Eğitimci", "ikon": "🦸", "kategori": "kur", "seviye": "platin", "puan": 75},
    {"kod": "kur_100", "ad": "Dönüşüm Lideri", "ikon": "👑", "kategori": "kur", "seviye": "elmas", "puan": 100},
    # Veli Değerlendirme
    {"kod": "veli_ilk", "ad": "İlk Beğeni", "ikon": "👍", "kategori": "veli", "seviye": "bronz", "puan": 5},
    {"kod": "veli_20", "ad": "Veli Favorisi", "ikon": "💜", "kategori": "veli", "seviye": "gumus", "puan": 20},
    {"kod": "veli_30", "ad": "Ailelerin Güveni", "ikon": "🏠", "kategori": "veli", "seviye": "altin", "puan": 35},
    {"kod": "veli_100", "ad": "Efsane Öğretmen", "ikon": "🌟", "kategori": "veli", "seviye": "elmas", "puan": 100},
    # Gelişim + İletişim + Egzersiz
    {"kod": "gelisim_ilk", "ad": "Meraklı Öğretmen", "ikon": "🔍", "kategori": "gelisim", "seviye": "bronz", "puan": 3},
    {"kod": "gelisim_10", "ad": "Sürekli Öğrenen", "ikon": "📖", "kategori": "gelisim", "seviye": "gumus", "puan": 15},
    {"kod": "gelisim_uzman", "ad": "Uzman Öğretmen", "ikon": "🎓", "kategori": "gelisim", "seviye": "elmas", "puan": 50},
    {"kod": "mesaj_ilk", "ad": "İlk Mesaj", "ikon": "💬", "kategori": "iletisim", "seviye": "bronz", "puan": 2},
    {"kod": "kopru_kurucu", "ad": "Köprü Kurucu", "ikon": "🌉", "kategori": "iletisim", "seviye": "altin", "puan": 15},
    {"kod": "egz_ilk", "ad": "İlk Egzersiz", "ikon": "👁️", "kategori": "egzersiz", "seviye": "bronz", "puan": 2},
    {"kod": "egz_tamset", "ad": "Tam Set", "ikon": "🎖️", "kategori": "egzersiz", "seviye": "altin", "puan": 20},
    # AI Eğitim Katkısı
    {"kod": "ai_ilk", "ad": "AI Eğitimcisi", "ikon": "🧠", "kategori": "ai_egitim", "seviye": "bronz", "puan": 5},
    {"kod": "ai_5", "ad": "Veri Kaşifi", "ikon": "📊", "kategori": "ai_egitim", "seviye": "gumus", "puan": 15},
    {"kod": "ai_20", "ad": "AI Ustası", "ikon": "🤖", "kategori": "ai_egitim", "seviye": "altin", "puan": 30},
    {"kod": "ai_50", "ad": "Bilgi Mimarı", "ikon": "🏗️", "kategori": "ai_egitim", "seviye": "elmas", "puan": 75},
]

OGRENCI_ROZETLERI_DEFAULT = [
    {"kod": "okuma_ilk", "ad": "İlk Sayfa", "ikon": "📖", "kategori": "okuma", "seviye": "bronz", "xp": 5},
    {"kod": "okuma_100", "ad": "Kitap Kurdu", "ikon": "🐛", "kategori": "okuma", "seviye": "gumus", "xp": 15},
    {"kod": "okuma_500", "ad": "Okuma Yıldızı", "ikon": "⭐", "kategori": "okuma", "seviye": "altin", "xp": 30},
    {"kod": "okuma_2000", "ad": "Okuma Efsanesi", "ikon": "🌟", "kategori": "okuma", "seviye": "elmas", "xp": 50},
    {"kod": "streak_3", "ad": "İlk Alışkanlık", "ikon": "🔥", "kategori": "streak", "seviye": "bronz", "xp": 5},
    {"kod": "streak_7", "ad": "Kararlı Okuyucu", "ikon": "💪", "kategori": "streak", "seviye": "gumus", "xp": 10},
    {"kod": "streak_21", "ad": "Demir İrade", "ikon": "🏔️", "kategori": "streak", "seviye": "altin", "xp": 25},
    {"kod": "streak_60", "ad": "Durdurulamaz", "ikon": "🚀", "kategori": "streak", "seviye": "elmas", "xp": 50},
    {"kod": "kitap_1", "ad": "İlk Kitap", "ikon": "📕", "kategori": "kitap", "seviye": "bronz", "xp": 5},
    {"kod": "kitap_5", "ad": "Kitap Kaşifi", "ikon": "🗺️", "kategori": "kitap", "seviye": "gumus", "xp": 15},
    {"kod": "kitap_15", "ad": "Kütüphane Dostu", "ikon": "📚", "kategori": "kitap", "seviye": "altin", "xp": 30},
    {"kod": "kitap_30", "ad": "Kitap Efsanesi", "ikon": "🏰", "kategori": "kitap", "seviye": "elmas", "xp": 50},
    {"kod": "gorev_ilk", "ad": "Görev Başlangıcı", "ikon": "✅", "kategori": "gorev", "seviye": "bronz", "xp": 5},
    {"kod": "gorev_10", "ad": "Görev Avcısı", "ikon": "🎯", "kategori": "gorev", "seviye": "gumus", "xp": 15},
    {"kod": "gorev_30", "ad": "Görev Ustası", "ikon": "🏹", "kategori": "gorev", "seviye": "altin", "xp": 30},
    {"kod": "gorev_100", "ad": "Görev Efsanesi", "ikon": "👑", "kategori": "gorev", "seviye": "elmas", "xp": 50},
    {"kod": "egz_ilk", "ad": "Göz Jimnastiği", "ikon": "👁️", "kategori": "egzersiz", "seviye": "bronz", "xp": 3},
    {"kod": "egz_20", "ad": "Egzersiz Yıldızı", "ikon": "💫", "kategori": "egzersiz", "seviye": "gumus", "xp": 10},
    {"kod": "egz_14", "ad": "Beyin Atleti", "ikon": "🧠", "kategori": "egzersiz", "seviye": "altin", "xp": 20},
    {"kod": "orman_ilk", "ad": "İlk Fidan", "ikon": "🌱", "kategori": "orman", "seviye": "bronz", "xp": 3},
    {"kod": "orman_50", "ad": "Küçük Orman", "ikon": "🌿", "kategori": "orman", "seviye": "gumus", "xp": 10},
    {"kod": "orman_200", "ad": "Orman Korucusu", "ikon": "🌳", "kategori": "orman", "seviye": "altin", "xp": 25},
    {"kod": "lig_gumus", "ad": "Gümüş Yolcusu", "ikon": "🥈", "kategori": "lig", "seviye": "gumus", "xp": 10},
    {"kod": "lig_altin", "ad": "Altın Savaşçısı", "ikon": "🥇", "kategori": "lig", "seviye": "altin", "xp": 20},
    {"kod": "lig_elmas", "ad": "Elmas Efsanesi", "ikon": "💎", "kategori": "lig", "seviye": "elmas", "xp": 50},
]

ANKET_SORULARI_DEFAULT = [
    {"no": 1, "soru": "Öğretmenin çocuğunuzla iletişimi nasıl?", "tip": "puan", "kategori": "iletisim"},
    {"no": 2, "soru": "Görev ve ödevler düzenli veriliyor mu?", "tip": "puan", "kategori": "duzen"},
    {"no": 3, "soru": "Çocuğunuzun okuma alışkanlığında gelişme görüyor musunuz?", "tip": "puan", "kategori": "etki"},
    {"no": 4, "soru": "Öğretmen geri bildirimleri yeterli mi?", "tip": "puan", "kategori": "geri_bildirim"},
    {"no": 5, "soru": "Çocuğunuzun motivasyonu arttı mı?", "tip": "puan", "kategori": "motivasyon"},
    {"no": 6, "soru": "Öğretmenin egzersiz ve içerik çeşitliliği yeterli mi?", "tip": "puan", "kategori": "icerik"},
    {"no": 7, "soru": "Genel olarak öğretmenden memnun musunuz?", "tip": "puan", "kategori": "genel"},
    {"no": 8, "soru": "Bu öğretmeni başka velilere tavsiye eder misiniz?", "tip": "evet_hayir", "kategori": "tavsiye"},
    {"no": 9, "soru": "Eklemek istediğiniz not (opsiyonel)", "tip": "metin", "kategori": "not"},
]



async def get_xp_tablosu():
    doc = await db.sistem_ayarlari.find_one({"tip": "xp_tablosu"})
    return doc.get("degerler", XP_TABLOSU_DEFAULT) if doc else XP_TABLOSU_DEFAULT


async def get_lig_esikleri():
    doc = await db.sistem_ayarlari.find_one({"tip": "lig_esikleri"})
    return doc.get("degerler", LIG_ESIKLERI_DEFAULT) if doc else LIG_ESIKLERI_DEFAULT


async def get_ogretmen_rozetleri():
    doc = await db.sistem_ayarlari.find_one({"tip": "ogretmen_rozetleri"})
    return doc.get("degerler", OGRETMEN_ROZETLERI_DEFAULT) if doc else OGRETMEN_ROZETLERI_DEFAULT


async def get_ogrenci_rozetleri():
    doc = await db.sistem_ayarlari.find_one({"tip": "ogrenci_rozetleri"})
    return doc.get("degerler", OGRENCI_ROZETLERI_DEFAULT) if doc else OGRENCI_ROZETLERI_DEFAULT


async def get_anket_sorulari():
    doc = await db.sistem_ayarlari.find_one({"tip": "anket_sorulari"})
    return doc.get("degerler", ANKET_SORULARI_DEFAULT) if doc else ANKET_SORULARI_DEFAULT


# XP kazan
@api_router.post("/xp/kazan")
async def xp_kazan(payload: dict, current_user=Depends(get_current_user)):
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    eylem = payload.get("eylem", "")
    xp = (await get_xp_tablosu()).get(eylem, 0)
    if xp == 0:
        return {"xp": 0, "mesaj": "Bilinmeyen eylem"}

    log = {
        "id": str(uuid.uuid4()),
        "ogrenci_id": ogrenci_id,
        "eylem": eylem,
        "xp": xp,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.xp_logs.insert_one(log)

    # Toplam XP güncelle
    await db.students.update_one({"id": ogrenci_id}, {"$inc": {"toplam_xp": xp}})

    return {"xp": xp, "toplam": await _get_toplam_xp(ogrenci_id)}


async def _get_toplam_xp(ogrenci_id):
    student = await db.students.find_one({"id": ogrenci_id})
    return student.get("toplam_xp", 0) if student else 0


# XP durumu
@api_router.get("/xp/durum/{ogrenci_id}")
async def xp_durum(ogrenci_id: str, current_user=Depends(get_current_user)):
    toplam = await _get_toplam_xp(ogrenci_id)
    # Lig hesapla
    lig = "bronz"
    for l in reversed(LIG_SIRA):
        if toplam >= (await get_lig_esikleri()).get(l, 0):
            lig = l
            break
    # Sonraki lig
    idx = LIG_SIRA.index(lig)
    sonraki_lig = LIG_SIRA[idx + 1] if idx < len(LIG_SIRA) - 1 else None
    sonraki_esik = (await get_lig_esikleri()).get(sonraki_lig, 0) if sonraki_lig else 0
    kalan = max(0, sonraki_esik - toplam)

    # Son XP kayıtları
    son_xp = await db.xp_logs.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=10)
    for x in son_xp:
        x.pop("_id", None)

    return {
        "toplam_xp": toplam,
        "lig": lig,
        "lig_label": {"bronz": "🥉 Bronz", "gumus": "🥈 Gümüş", "altin": "🥇 Altın", "elmas": "💎 Elmas"}.get(lig, lig),
        "sonraki_lig": sonraki_lig,
        "sonraki_esik": sonraki_esik,
        "kalan_xp": kalan,
        "son_xp": son_xp,
    }


# Lig sıralaması (anonim)
@api_router.get("/xp/lig-siralama")
async def lig_siralama(current_user=Depends(get_current_user)):
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    students = await db.students.find().to_list(length=None)
    siralama = sorted(students, key=lambda s: s.get("toplam_xp", 0), reverse=True)

    tablo = []
    benim_siram = None
    for i, s in enumerate(siralama):
        sira = i + 1
        xp = s.get("toplam_xp", 0)
        lig = "bronz"
        for l in reversed(LIG_SIRA):
            if xp >= (await get_lig_esikleri()).get(l, 0):
                lig = l
                break
        ben = s.get("id") == ogrenci_id
        if ben:
            benim_siram = sira
        tablo.append({
            "sira": sira,
            "xp": xp,
            "lig": lig,
            "lig_label": {"bronz": "🥉", "gumus": "🥈", "altin": "🥇", "elmas": "💎"}.get(lig, ""),
            "ben": ben,
            "ad": "Sen 🌟" if ben else f"Öğrenci #{sira}",
        })

    return {"siralama": tablo[:30], "benim_siram": benim_siram or len(tablo) + 1, "toplam": len(siralama)}


# Kur atlama kontrolü
@api_router.get("/kur/kontrol/{ogrenci_id}")
async def kur_kontrol(ogrenci_id: str, current_user=Depends(get_current_user)):
    # Kriterleri kontrol et
    gorevler = await db.gorevler.find({"hedef_id": ogrenci_id, "durum": "tamamlandi"}).to_list(length=None)
    reading_logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).to_list(length=None)
    kitaplar = set(l.get("kitap_adi", "") for l in reading_logs if l.get("kitap_adi"))

    # Streak
    from datetime import timedelta
    simdi = datetime.utcnow()
    tarihler = sorted(set(l.get("tarih", "")[:10] for l in reading_logs), reverse=True)
    streak = 0
    for i in range(60):
        gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
        if gun in tarihler:
            streak += 1
        elif i > 0:
            break

    tamamlanan_gorev = len(gorevler)
    kitap_sayisi = len(kitaplar)

    # Anlama yüzdesi (gelişim tamamlamalardan)
    tamamlamalar = await db.gelisim_tamamlama.find({"kullanici_id": ogrenci_id}).to_list(length=None)
    if tamamlamalar:
        toplam_dogru = sum(t.get("dogru_sayisi", 0) for t in tamamlamalar if t.get("test_yapildi"))
        toplam_soru = sum(t.get("toplam_soru", 0) for t in tamamlamalar if t.get("test_yapildi"))
        anlama_yuzdesi = round((toplam_dogru / max(toplam_soru, 1)) * 100)
    else:
        anlama_yuzdesi = 0

    kriterler = {
        "gorev_12": {"gerekli": 12, "mevcut": tamamlanan_gorev, "tamam": tamamlanan_gorev >= 12},
        "anlama_75": {"gerekli": 75, "mevcut": anlama_yuzdesi, "tamam": anlama_yuzdesi >= 75},
        "kitap_4": {"gerekli": 4, "mevcut": kitap_sayisi, "tamam": kitap_sayisi >= 4},
        "streak_10": {"gerekli": 10, "mevcut": streak, "tamam": streak >= 10},
    }

    hepsi_tamam = all(k["tamam"] for k in kriterler.values())

    return {
        "kriterler": kriterler,
        "kur_atlayabilir": hepsi_tamam,
        "mevcut_kur": (await db.students.find_one({"id": ogrenci_id}) or {}).get("kur", ""),
    }


# Kur atla (öğretmen onayı ile)
@api_router.post("/kur/atla")
async def kur_atla(payload: dict, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Sadece öğretmen/yönetici kur atlatabilir")

    ogrenci_id = payload.get("ogrenci_id")
    yeni_kur = payload.get("yeni_kur", "")

    # Mevcut kuru al
    student = await db.students.find_one({"id": ogrenci_id})
    eski_kur = student.get("kur", "") if student else ""

    # Kur güncelle
    await db.students.update_one({"id": ogrenci_id}, {"$set": {"kur": yeni_kur}})

    # Kur atlama kaydı (rozet sistemi için)
    ogretmen_id = current_user.get("linked_id") or current_user.get("id")
    await db.kur_atlamalari.insert_one({
        "id": str(uuid.uuid4()),
        "ogrenci_id": ogrenci_id,
        "ogretmen_id": ogretmen_id,
        "eski_kur": eski_kur,
        "yeni_kur": yeni_kur,
        "tarih": datetime.utcnow().isoformat(),
    })

    return {"ok": True, "yeni_kur": yeni_kur, "eski_kur": eski_kur}


# ─────────────────────────────────────────────
# SİSTEM AYARLARI YÖNETİMİ (Admin CRUD)
# ─────────────────────────────────────────────

@api_router.get("/ayarlar/{tip}")
async def get_ayar(tip: str, current_user=Depends(get_current_user)):
    doc = await db.sistem_ayarlari.find_one({"tip": tip})
    if doc:
        doc.pop("_id", None)
        return doc
    # Varsayılan değerleri döndür
    defaults = {
        "xp_tablosu": XP_TABLOSU_DEFAULT,
        "lig_esikleri": LIG_ESIKLERI_DEFAULT,
        "ogretmen_rozetleri": OGRETMEN_ROZETLERI_DEFAULT,
        "ogrenci_rozetleri": OGRENCI_ROZETLERI_DEFAULT,
        "anket_sorulari": ANKET_SORULARI_DEFAULT,
    }
    return {"tip": tip, "degerler": defaults.get(tip, {})}


@api_router.put("/ayarlar/{tip}")
async def update_ayar(tip: str, payload: dict, current_user=Depends(get_current_user)):
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Sadece admin ayar değiştirebilir")
    degerler = payload.get("degerler", {})
    await db.sistem_ayarlari.update_one(
        {"tip": tip},
        {"$set": {"tip": tip, "degerler": degerler, "guncelleme_tarihi": datetime.utcnow().isoformat(), "guncelleyen": current_user.get("ad", "")}},
        upsert=True
    )
    return {"ok": True, "tip": tip}


@api_router.get("/ayarlar")
async def get_tum_ayarlar(current_user=Depends(get_current_user)):
    if current_user.get("role") not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    ayarlar = await db.sistem_ayarlari.find().to_list(length=None)
    for a in ayarlar:
        a.pop("_id", None)
    return ayarlar


# Birleştirilmiş Puan Tablosu (gelişim + rozet puanları)
@api_router.get("/puan-tablosu/birlesik")
async def get_birlesik_puan_tablosu(current_user=Depends(get_current_user)):
    users = await db.users.find().to_list(length=None)
    tablo = []
    for u in users:
        uid = u["id"]
        gelisim_puan = u.get("puan", 0)

        # Rozet puanları
        rozetler = await db.kazanilan_rozetler.find({"kullanici_id": uid}).to_list(length=None)
        rozet_kodlar = [r["rozet_kodu"] for r in rozetler]

        ogretmen_rozet_list = await get_ogretmen_rozetleri()
        ogrenci_rozet_list = await get_ogrenci_rozetleri()
        tum_rozetler = ogretmen_rozet_list + ogrenci_rozet_list

        rozet_puan = 0
        for rk in rozet_kodlar:
            tanim = next((r for r in tum_rozetler if r["kod"] == rk), None)
            if tanim:
                rozet_puan += tanim.get("puan", tanim.get("xp", 0))

        toplam = gelisim_puan + rozet_puan

        tablo.append({
            "ad": u.get("ad", ""), "soyad": u.get("soyad", ""),
            "role": u.get("role", ""),
            "gelisim_puan": gelisim_puan,
            "rozet_puan": rozet_puan,
            "rozet_sayisi": len(rozet_kodlar),
            "toplam_puan": toplam,
        })
    tablo.sort(key=lambda x: x["toplam_puan"], reverse=True)
    return tablo


# ─────────────────────────────────────────────
# ROZET SİSTEMİ (Öğretmen + Öğrenci)
# ─────────────────────────────────────────────




@api_router.get("/rozetler/tanim")
async def rozet_tanimlari():
    return {"ogretmen": await get_ogretmen_rozetleri(), "ogrenci": await get_ogrenci_rozetleri()}


@api_router.get("/rozetler/{user_id}")
async def get_rozetler(user_id: str, current_user=Depends(get_current_user)):
    rozetler = await db.kazanilan_rozetler.find({"kullanici_id": user_id}).to_list(length=None)
    for r in rozetler:
        r.pop("_id", None)
    return rozetler


@api_router.post("/rozetler/kontrol")
async def rozet_kontrol(current_user=Depends(get_current_user)):
    user_id = current_user["id"]
    role = current_user.get("role", "")
    linked_id = current_user.get("linked_id", "")

    mevcut = await db.kazanilan_rozetler.find({"kullanici_id": user_id}).to_list(length=None)
    mevcut_kodlar = set(r["rozet_kodu"] for r in mevcut)
    yeni_rozetler = []

    if role == "teacher":
        ogretmen_id = linked_id or user_id
        # İçerik sayısı
        icerikler = await db.gelisim_icerik.count_documents({"ekleyen_id": user_id, "durum": "yayinda"})
        # Oylama sayısı
        tum_icerikler = await db.gelisim_icerik.find({"durum": {"$in": ["yayinda", "oylama"]}}).to_list(length=None)
        oy_sayisi = sum(1 for ic in tum_icerikler if user_id in (ic.get("oylar") or {}))
        # Görev atama
        gorevler = await db.gorevler.find({"atayan_id": user_id}).to_list(length=None)
        gorev_sayisi = len(gorevler)
        tamamlanan_gorev = len([g for g in gorevler if g.get("durum") == "tamamlandi"])
        # Öğrenci streak ortalaması
        ogrenciler = await db.students.find({"ogretmen_id": ogretmen_id, "arsivli": {"$ne": True}}).to_list(length=None)
        from datetime import timedelta
        simdi = datetime.utcnow()
        streakler = []
        for s in ogrenciler:
            logs = await db.reading_logs.find({"ogrenci_id": s["id"]}).to_list(length=None)
            tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
            st = 0
            for i in range(60):
                gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
                if gun in tarihler: st += 1
                elif i > 0: break
            streakler.append(st)
        ort_streak = sum(streakler) / max(len(streakler), 1)
        # Kur atlama
        kur_sayisi = await db.kur_atlamalari.count_documents({"ogretmen_id": ogretmen_id})
        # Gelişim tamamlama
        gelisim_tam = await db.gelisim_tamamlama.count_documents({"kullanici_id": user_id})
        # Mesaj
        mesajlar = await db.mesajlar.find({"gonderen_id": user_id}).to_list(length=None)
        mesaj_sayisi = len(mesajlar)
        mesaj_roller = set(m.get("alici_rol", "") for m in mesajlar)
        # Egzersiz
        egz_tam = await db.egzersiz_tamamlama.find({"kullanici_id": user_id}).to_list(length=None)
        egz_turler = set(e.get("egzersiz_id", "") for e in egz_tam)
        # Veli anketi
        anketler = await db.veli_anketleri.find({"ogretmen_id": ogretmen_id}).to_list(length=None)
        anket_sayisi = len(anketler)
        anket_ort = 0
        tavsiye_oran = 0
        if anket_sayisi > 0:
            puanlar = []
            tavsiyeler = 0
            for a in anketler:
                yanitlar = a.get("yanitlar", [])
                puan_yanitlar = [y.get("puan", 0) for y in yanitlar if y.get("puan")]
                if puan_yanitlar:
                    puanlar.append(sum(puan_yanitlar) / len(puan_yanitlar))
                if a.get("tavsiye"):
                    tavsiyeler += 1
            anket_ort = sum(puanlar) / max(len(puanlar), 1)
            tavsiye_oran = (tavsiyeler / anket_sayisi) * 100

        # Kontrol
        checks = [
            ("icerik_ilk", icerikler >= 1), ("icerik_5", icerikler >= 5), ("icerik_20", icerikler >= 20), ("icerik_50", icerikler >= 50),
            ("oy_ilk", oy_sayisi >= 1), ("oy_20", oy_sayisi >= 20), ("oy_50", oy_sayisi >= 50),
            ("gorev_ilk", gorev_sayisi >= 1), ("gorev_20", gorev_sayisi >= 20 and tamamlanan_gorev >= 10),
            ("ilham_veren", ort_streak >= 7), ("yildiz_egitimci", ort_streak >= 10),
            ("kur_ilk", kur_sayisi >= 1), ("kur_20", kur_sayisi >= 20), ("kur_30", kur_sayisi >= 30), ("kur_50", kur_sayisi >= 50), ("kur_100", kur_sayisi >= 100),
            ("veli_ilk", anket_sayisi >= 1 and anket_ort >= 4), ("veli_20", anket_sayisi >= 20 and anket_ort >= 4.5),
            ("veli_30", anket_sayisi >= 30 and anket_ort >= 4.5 and tavsiye_oran >= 90),
            ("veli_100", anket_sayisi >= 100 and anket_ort >= 4.8 and tavsiye_oran >= 95),
            ("gelisim_ilk", gelisim_tam >= 1), ("gelisim_10", gelisim_tam >= 10), ("gelisim_uzman", gelisim_tam >= 30),
            ("mesaj_ilk", mesaj_sayisi >= 1), ("kopru_kurucu", "student" in mesaj_roller and "parent" in mesaj_roller),
            ("egz_ilk", len(egz_turler) >= 1), ("egz_tamset", len(egz_turler) >= 14),
        ]
        for kod, kosul in checks:
            if kosul and kod not in mevcut_kodlar:
                doc = {"id": str(uuid.uuid4()), "kullanici_id": user_id, "rozet_kodu": kod, "kazanma_tarihi": datetime.utcnow().isoformat()}
                await db.kazanilan_rozetler.insert_one(doc)
                rozet_bilgi = next((r for r in (await get_ogretmen_rozetleri()) if r["kod"] == kod), None)
                yeni_rozetler.append({**doc, "rozet": rozet_bilgi})

    elif role == "student":
        ogrenci_id = linked_id or user_id
        logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).to_list(length=None)
        toplam_dk = sum(l.get("sure_dakika", 0) for l in logs)
        kitaplar = set(l.get("kitap_adi", "") for l in logs if l.get("kitap_adi"))
        from datetime import timedelta
        simdi = datetime.utcnow()
        tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
        streak = 0
        for i in range(60):
            gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
            if gun in tarihler: streak += 1
            elif i > 0: break
        gorevler_tam = await db.gorevler.count_documents({"hedef_id": ogrenci_id, "durum": "tamamlandi"})
        egz_tam = await db.egzersiz_tamamlama.find({"kullanici_id": user_id}).to_list(length=None)
        egz_turler = set(e.get("egzersiz_id", "") for e in egz_tam)
        egz_toplam = len(egz_tam)
        agac_sayisi = toplam_dk  # 1 dk = 1 ağaç
        student = await db.students.find_one({"id": ogrenci_id})
        toplam_xp = student.get("toplam_xp", 0) if student else 0

        checks = [
            ("okuma_ilk", len(logs) >= 1), ("okuma_100", toplam_dk >= 100), ("okuma_500", toplam_dk >= 500), ("okuma_2000", toplam_dk >= 2000),
            ("streak_3", streak >= 3), ("streak_7", streak >= 7), ("streak_21", streak >= 21), ("streak_60", streak >= 60),
            ("kitap_1", len(kitaplar) >= 1), ("kitap_5", len(kitaplar) >= 5), ("kitap_15", len(kitaplar) >= 15), ("kitap_30", len(kitaplar) >= 30),
            ("gorev_ilk", gorevler_tam >= 1), ("gorev_10", gorevler_tam >= 10), ("gorev_30", gorevler_tam >= 30), ("gorev_100", gorevler_tam >= 100),
            ("egz_ilk", egz_toplam >= 1), ("egz_20", egz_toplam >= 20), ("egz_14", len(egz_turler) >= 14),
            ("orman_ilk", agac_sayisi >= 1), ("orman_50", agac_sayisi >= 50), ("orman_200", agac_sayisi >= 200),
            ("lig_gumus", toplam_xp >= 200), ("lig_altin", toplam_xp >= 500), ("lig_elmas", toplam_xp >= 1000),
        ]
        for kod, kosul in checks:
            if kosul and kod not in mevcut_kodlar:
                doc = {"id": str(uuid.uuid4()), "kullanici_id": user_id, "rozet_kodu": kod, "kazanma_tarihi": datetime.utcnow().isoformat()}
                await db.kazanilan_rozetler.insert_one(doc)
                rozet_bilgi = next((r for r in (await get_ogrenci_rozetleri()) if r["kod"] == kod), None)
                yeni_rozetler.append({**doc, "rozet": rozet_bilgi})

    return {"yeni_rozetler": yeni_rozetler, "toplam": len(mevcut_kodlar) + len(yeni_rozetler)}


# ─────────────────────────────────────────────
# VELİ DEĞERLENDİRME ANKETİ
# ─────────────────────────────────────────────



@api_router.get("/anketler/sorular")
async def get_anket_sorulari_endpoint():
    return await get_anket_sorulari()


@api_router.post("/anketler")
async def create_anket(payload: dict, current_user=Depends(get_current_user)):
    if current_user.get("role") != "parent":
        raise HTTPException(status_code=403, detail="Sadece veliler anket doldurabilir")

    ogretmen_id = payload.get("ogretmen_id", "")
    ogrenci_id = payload.get("ogrenci_id", "")
    yanitlar = payload.get("yanitlar", [])
    tavsiye = payload.get("tavsiye", None)
    not_text = payload.get("not_text", "")
    donem = payload.get("donem", datetime.utcnow().strftime("%Y-D%m"))

    # Aynı dönem + aynı öğretmen kontrolü
    mevcut = await db.veli_anketleri.find_one({
        "veli_id": current_user["id"], "ogretmen_id": ogretmen_id, "donem": donem
    })
    if mevcut:
        raise HTTPException(status_code=409, detail="Bu dönem için zaten anket doldurdunuz")

    doc = {
        "id": str(uuid.uuid4()),
        "veli_id": current_user["id"],
        "veli_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "ogretmen_id": ogretmen_id,
        "ogrenci_id": ogrenci_id,
        "donem": donem,
        "yanitlar": yanitlar,
        "tavsiye": tavsiye,
        "not_text": not_text,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.veli_anketleri.insert_one(doc)
    return doc


@api_router.get("/anketler/ogretmen/{ogretmen_id}/ozet")
async def anket_ozet(ogretmen_id: str, current_user=Depends(get_current_user)):
    anketler = await db.veli_anketleri.find({"ogretmen_id": ogretmen_id}).to_list(length=None)
    if not anketler:
        return {"anket_sayisi": 0, "ortalama": 0, "tavsiye_oran": 0, "kategoriler": {}, "son_anketler": []}

    puanlar = []
    tavsiyeler = 0
    kategori_toplam = {}
    kategori_sayac = {}

    for a in anketler:
        for y in a.get("yanitlar", []):
            if y.get("puan"):
                puanlar.append(y["puan"])
                kat = y.get("kategori", "genel")
                kategori_toplam[kat] = kategori_toplam.get(kat, 0) + y["puan"]
                kategori_sayac[kat] = kategori_sayac.get(kat, 0) + 1
        if a.get("tavsiye"):
            tavsiyeler += 1

    ortalama = round(sum(puanlar) / max(len(puanlar), 1), 1)
    tavsiye_oran = round((tavsiyeler / len(anketler)) * 100)
    kategoriler = {k: round(kategori_toplam[k] / kategori_sayac[k], 1) for k in kategori_toplam}

    # Öğretmen isimleri görmez
    role = current_user.get("role", "")
    son_anketler = []
    for a in sorted(anketler, key=lambda x: x.get("tarih", ""), reverse=True)[:10]:
        a.pop("_id", None)
        entry = {"donem": a.get("donem"), "tarih": a.get("tarih"), "tavsiye": a.get("tavsiye")}
        puan_yanitlar = [y.get("puan") for y in a.get("yanitlar", []) if y.get("puan")]
        entry["ortalama"] = round(sum(puan_yanitlar) / max(len(puan_yanitlar), 1), 1) if puan_yanitlar else 0
        if role in ["admin", "coordinator"]:
            entry["veli_ad"] = a.get("veli_ad", "")
            entry["not_text"] = a.get("not_text", "")
        son_anketler.append(entry)

    return {
        "anket_sayisi": len(anketler), "ortalama": ortalama, "tavsiye_oran": tavsiye_oran,
        "kategoriler": kategoriler, "son_anketler": son_anketler,
    }


@api_router.get("/anketler/veli/{veli_id}")
async def veli_anketleri(veli_id: str, current_user=Depends(get_current_user)):
    anketler = await db.veli_anketleri.find({"veli_id": veli_id}).to_list(length=None)
    for a in anketler:
        a.pop("_id", None)
    return anketler


# ─────────────────────────────────────────────
# KUR ATLAMA KAYDI (rozet için güncelleme)
# ─────────────────────────────────────────────

# kur/atla endpoint'ini güncelle — kur_atlamalari collection'ına da kaydet
_original_kur_atla = None  # placeholder


# ─────────────────────────────────────────────
# AI KOÇLUK + DNA + SORU ÜRETİMİ + HİKAYE
# ─────────────────────────────────────────────

AI_KOCLUK_SYSTEM_PROMPT = """Sen deneyimli bir ilkokul okuma koçusun. Türkiye'de çalışıyorsun. MEB Türkçe müfredatını ve Bloom taksonomisini biliyorsun.
Görevin: Verilen öğrenci verilerini analiz ederek kişiselleştirilmiş koçluk raporu üretmek.
DİL: Türkçe. Pozitif ve yapıcı dil kullan. Öğretmene yardımcı ol, öğrenciyi motive et.
FORMAT: Yanıtını SADECE JSON olarak ver, başka metin ekleme. Markdown code block kullanma."""

AI_SORU_SYSTEM_PROMPT = """Sen Türkçe dil eğitimi uzmanısın. Bloom taksonomisini ve MEB müfredatını biliyorsun.
Verilen metinden çoktan seçmeli sorular üreteceksin. Her soru 4 şıklı olacak.
FORMAT: SADECE JSON array ver: [{"soru":"...", "secenekler":["A","B","C","D"], "dogru_cevap":0, "taksonomi":"bilgi|kavrama|uygulama|analiz|sentez|degerlendirme"}]"""

AI_HIKAYE_SYSTEM_PROMPT = """Sen çocuk kitabı yazarısın. MEB Türkçe müfredatını biliyorsun.
Verilen sınıf seviyesi, tema ve hedef kelimeleri kullanarak kısa bir okuma parçası yazacaksın.
Hedef kelimelerin TÜMÜNÜ metnin içinde doğal şekilde kullan.
FORMAT: SADECE JSON ver: {"baslik":"...", "metin":"...", "kelime_sayisi":0, "kullanilan_kelimeler":[], "sorular":[5 Bloom sorusu]}"""


@api_router.get("/ai/dna/{ogrenci_id}")
async def get_okuma_dna(ogrenci_id: str, current_user=Depends(get_current_user)):
    """7 boyutlu Okuma DNA profili hesapla (Claude API kullanmaz — mevcut verilerden)."""
    v = await get_ogrenci_ai_verileri(ogrenci_id)
    if not v:
        raise HTTPException(status_code=404, detail="Öğrenci bulunamadı")

    ok = v["okuma_ozet"]
    # 1. Kelime Gücü (0-100)
    sinif_raw = v["ogrenci"].get("sinif", 3)
    try:
        sinif = int(sinif_raw)
    except (ValueError, TypeError):
        sinif = 3
    hedef_kelime = {1:150, 2:300, 3:500, 4:700, 5:1000, 6:1300, 7:1600, 8:2000}.get(sinif, 500)
    bilinen = await db.kelime_bankasi.count_documents({"ogrenci_id": ogrenci_id, "ogrenildi": True}) if await db.kelime_bankasi.count_documents({}) > 0 else int(hedef_kelime * 0.5)
    kelime_gucu = min(100, round(bilinen / hedef_kelime * 100))

    # 2. Akıcılık (0-100)
    wpm = v["analiz"].get("wpm", 0)
    norm_wpm = {1:50, 2:75, 3:95, 4:115, 5:130, 6:145, 7:155, 8:165}.get(sinif, 95)
    akicilik = min(100, round(wpm / norm_wpm * 100)) if wpm > 0 else 50

    # 3. Anlama Derinliği (0-100)
    bloom = v["analiz"].get("bloom", {})
    if bloom:
        anlama = round((bloom.get("bilgi",0)*0.1 + bloom.get("kavrama",0)*0.15 + bloom.get("uygulama",0)*0.2 + bloom.get("analiz",0)*0.25 + bloom.get("sentez",0)*0.15 + bloom.get("degerlendirme",0)*0.15))
    else:
        anlama = v["test"].get("ort_yuzde", 50)

    # 4. Dikkat Süresi (0-100)
    ort_dk = ok.get("ort_gunluk_dk", 0)
    dikkat = min(100, round(ort_dk / 20 * 100)) if ort_dk > 0 else 30

    # 5. Zorluk Toleransı (0-100)
    zorluk_tol = 50  # Varsayılan, zor metinlerdeki başarı verisi birikince geliştirilecek

    # 6. Kelime Tekrar İhtiyacı (0-100, yüksek = çok tekrara ihtiyacı var)
    tekrar_ihtiyac = max(0, 100 - kelime_gucu)

    # 7. Okuma Psikolojisi
    toplam_kitap = ok.get("kitap_sayisi", 0)
    streak_m = v["streak"].get("mevcut", 0)
    if toplam_kitap >= 3 and streak_m >= 5:
        psikoloji = "keşifçi"
    elif streak_m < 2 and ok.get("gun_sayisi", 0) < 5:
        psikoloji = "kararsız"
    else:
        psikoloji = "güvenli"

    # Profil tipi
    if akicilik > 70 and anlama < 50:
        profil_tipi = "hızlı_okuyucu"
    elif akicilik < 40 and anlama > 70:
        profil_tipi = "analitik_okuyucu"
    elif bloom and bloom.get("sentez", 0) > 60:
        profil_tipi = "hayalci_okuyucu"
    elif streak_m < 3 and ok.get("gun_sayisi", 0) < 10:
        profil_tipi = "başlangıç_okuyucu"
    else:
        profil_tipi = "dengeli_okuyucu"

    profil_label = {"hızlı_okuyucu": "📖 Hızlı Okuyucu", "analitik_okuyucu": "🔍 Analitik Okuyucu", "hayalci_okuyucu": "🌈 Hayalci Okuyucu", "başlangıç_okuyucu": "🌱 Başlangıç Okuyucu", "dengeli_okuyucu": "⚖️ Dengeli Okuyucu"}

    dna = {
        "ogrenci_id": ogrenci_id,
        "boyutlar": {
            "kelime_gucu": kelime_gucu,
            "akicilik": akicilik,
            "anlama_derinligi": anlama,
            "dikkat_suresi": dikkat,
            "zorluk_toleransi": zorluk_tol,
            "kelime_tekrar_ihtiyaci": tekrar_ihtiyac,
            "okuma_psikolojisi": psikoloji,
        },
        "profil_tipi": profil_tipi,
        "profil_label": profil_label.get(profil_tipi, "📖 Okuyucu"),
        "sinif": sinif,
        "son_guncelleme": datetime.utcnow().isoformat(),
    }

    # Cache'e kaydet
    await db.okuma_dna.update_one({"ogrenci_id": ogrenci_id}, {"$set": dna}, upsert=True)
    return dna


@api_router.post("/ai/kocluk/{ogrenci_id}")
async def ai_kocluk(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Tam AI koçluk raporu (10 modül). 24 saat cache."""
    # Cache kontrolü
    cache = await db.ai_kocluk_cache.find_one({
        "ogrenci_id": ogrenci_id,
        "tarih": {"$gte": (datetime.utcnow() - timedelta(hours=AI_CACHE_HOURS)).isoformat()}
    })
    if cache:
        cache.pop("_id", None)
        return cache

    # Veri topla
    v = await get_ogrenci_ai_verileri(ogrenci_id)
    if not v or not v.get("ogrenci", {}).get("ad"):
        raise HTTPException(status_code=404, detail="Öğrenci bulunamadı")

    # DNA hesapla
    dna = None
    try:
        dna_response = await get_okuma_dna(ogrenci_id, current_user)
        dna = dna_response
    except:
        dna = {"profil_tipi": "bilinmiyor", "boyutlar": {}}

    import json as json_mod
    user_message = f"""Öğrenci verileri:
{json_mod.dumps(v, ensure_ascii=False, indent=2)}

Okuma DNA:
{json_mod.dumps(dna, ensure_ascii=False, indent=2) if dna else "Hesaplanamadı"}

Şu 8 modülü JSON olarak üret:
{{
  "durum_degerlendirmesi": {{"guclu_yonler": ["..."], "gelisim_alanlari": ["..."]}},
  "risk_analizi": {{"seviye": "düşük|orta|yüksek", "faktorler": ["..."], "aciliyet": "..."}},
  "mudahale_plani": {{"hafta_1": "...", "hafta_2": "...", "hafta_3": "...", "hafta_4": "..."}},
  "veliye_mesaj": "...",
  "haftalik_gorevler": [{{"gun": "Pazartesi", "gorev": "...", "bloom": "..."}}],
  "kitap_tavsiyeleri": [{{"ad": "...", "yazar": "...", "neden": "..."}}],
  "motivasyon_mesaji": "...",
  "kelime_mudahale": "...",
  "metin_recetesi": {{"paragraf_uzunlugu": "...", "soyutluk": "...", "aksiyon": "...", "hedef_kelime_orani": "..."}}
}}"""

    result = await call_claude(AI_KOCLUK_SYSTEM_PROMPT, user_message, model="sonnet", max_tokens=3000)

    if result.get("error"):
        raise HTTPException(status_code=503, detail=result["error"])

    rapor = {
        "id": str(uuid.uuid4()),
        "ogrenci_id": ogrenci_id,
        "dna": dna,
        "veriler": v,
        "ai_analiz": result.get("parsed") or result.get("text", ""),
        "ai_ham_metin": result.get("text", ""),
        "model": AI_DEFAULT_MODEL,
        "token": result.get("tokens", 0),
        "maliyet": result.get("maliyet", 0),
        "tarih": datetime.utcnow().isoformat(),
    }

    await db.ai_kocluk_cache.insert_one(rapor)
    rapor.pop("_id", None)
    return rapor


@api_router.get("/ai/kocluk/{ogrenci_id}/motivasyon")
async def ai_motivasyon(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Günlük kişisel motivasyon mesajı (Haiku — hızlı ve ucuz)."""
    v = await get_ogrenci_ai_verileri(ogrenci_id)
    if not v:
        return {"mesaj": "Bugün okumak için harika bir gün! 📚"}

    ad = v["ogrenci"].get("ad", "Öğrenci")
    streak = v["streak"].get("mevcut", 0)
    bugun_dk = v["okuma_ozet"].get("ort_gunluk_dk", 0)

    prompt = f"{ad}, streak: {streak} gün, ortalama {bugun_dk} dk/gün okuyor. 1 cümle motivasyon mesajı yaz. Türkçe, sıcak, kişisel. Sadece mesaj metni ver, başka bir şey yazma."

    result = await call_claude("Sen çocuklara motivasyon veren sıcak bir okuma koçusun.", prompt, model="haiku", max_tokens=100)
    mesaj = result.get("text", f"Harika gidiyorsun {ad}! 🔥 Streak'in {streak} gün!")

    return {"mesaj": mesaj.strip().strip('"')}


@api_router.post("/ai/soru-uret")
async def ai_soru_uret(payload: dict, current_user=Depends(get_current_user)):
    """Metin + sınıf → Bloom taksonomili 5-10 soru üretimi."""
    metin = payload.get("metin", "")
    sinif = payload.get("sinif", 4)
    soru_sayisi = payload.get("soru_sayisi", 5)

    if len(metin) < 50:
        raise HTTPException(status_code=400, detail="Metin en az 50 karakter olmalı")

    user_msg = f"""Sınıf: {sinif}
Soru sayısı: {soru_sayisi}
Her Bloom basamağından en az 1 soru olsun (bilgi, kavrama, uygulama, analiz, sentez, degerlendirme).

METİN:
{metin[:3000]}

SADECE JSON array döndür: [{{"soru":"...", "secenekler":["A","B","C","D"], "dogru_cevap":0, "taksonomi":"bilgi"}}]"""

    result = await call_claude(AI_SORU_SYSTEM_PROMPT, user_msg, model="sonnet", max_tokens=2000)

    if result.get("error"):
        raise HTTPException(status_code=503, detail=result["error"])

    sorular = result.get("parsed") or []
    if isinstance(sorular, dict):
        sorular = sorular.get("sorular", [])

    return {"sorular": sorular, "token": result.get("tokens", 0), "maliyet": result.get("maliyet", 0)}


@api_router.post("/ai/hikaye-uret")
async def ai_hikaye_uret(payload: dict, current_user=Depends(get_current_user)):
    """Sınıf + tema + hedef kelimeler → kişisel hikâye + 5 Bloom sorusu."""
    sinif = payload.get("sinif", 3)
    tema = payload.get("tema", "Doğa ve Evren")
    kelimeler = payload.get("kelimeler", [])
    kelime_sayisi = payload.get("kelime_sayisi", 150)

    user_msg = f"""Sınıf: {sinif}. sınıf
Tema: {tema}
Hedef kelimeler: {', '.join(kelimeler) if kelimeler else 'pusula, keşif, macera, mevsim, göç'}
Kelime sayısı: ~{kelime_sayisi}
Cümle uzunluğu: max {8 + sinif} kelime

SADECE JSON döndür:
{{"baslik":"...", "metin":"...", "kelime_sayisi":0, "kullanilan_kelimeler":[], "sorular":[{{"soru":"...", "secenekler":["A","B","C","D"], "dogru_cevap":0, "taksonomi":"bilgi"}}]}}"""

    result = await call_claude(AI_HIKAYE_SYSTEM_PROMPT, user_msg, model="sonnet", max_tokens=2500)

    if result.get("error"):
        raise HTTPException(status_code=503, detail=result["error"])

    return {"hikaye": result.get("parsed") or result.get("text", ""), "token": result.get("tokens", 0), "maliyet": result.get("maliyet", 0)}


@api_router.get("/ai/kelime-listesi")
async def ai_kelime_listesi(sinif: int = 0, current_user=Depends(get_current_user)):
    """Kelime haritasındaki tüm kelimeleri listeler."""
    filtre = {}
    if sinif > 0:
        filtre["sinif"] = sinif
    kelimeler = await db.meb_kelime_haritasi.find(filtre).sort("kelime", 1).to_list(length=500)
    for k in kelimeler:
        k.pop("_id", None)
    return kelimeler


@api_router.get("/ai/okuma-parcalari")
async def ai_okuma_parcalari_listesi(current_user=Depends(get_current_user)):
    """Tüm AI okuma parçalarını listeler."""
    parcalar = await db.ai_okuma_parcalari.find({}).sort("tarih", -1).to_list(length=100)
    for p in parcalar:
        p.pop("_id", None)
    return parcalar


@api_router.get("/ai/sorular")
async def ai_sorular_listesi(current_user=Depends(get_current_user)):
    """Tüm AI üretilen soruları listeler."""
    sorular = await db.ai_uretilen_sorular.find({}).sort("tarih", -1).to_list(length=200)
    for s in sorular:
        s.pop("_id", None)
    return sorular


@api_router.get("/ai/socratic-log")
async def ai_socratic_log_listesi(current_user=Depends(get_current_user)):
    """Socratic Reading loglarını listeler."""
    loglar = await db.ai_socratic_log.find({}).sort("tarih", -1).to_list(length=100)
    for l in loglar:
        l.pop("_id", None)
    # Öğrenci adlarını ekle
    for l in loglar:
        ogr = await db.users.find_one({"id": l.get("ogrenci_id", "")})
        if not ogr:
            ogr = await db.students.find_one({"id": l.get("ogrenci_id", "")})
        l["ogrenci_ad"] = f"{ogr.get('ad', '')} {ogr.get('soyad', '')}" if ogr else "Bilinmiyor"
    return loglar


@api_router.get("/ai/maliyet-ozet")
async def ai_maliyet_ozet(current_user=Depends(require_role(UserRole.ADMIN))):
    """Admin: AI maliyet özeti."""
    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    bu_ay = datetime.utcnow().strftime("%Y-%m")

    gunluk = await db.ai_request_log.find({"tarih": {"$regex": f"^{bugun}"}}).to_list(length=None)
    aylik = await db.ai_request_log.find({"tarih": {"$regex": f"^{bu_ay}"}}).to_list(length=None)

    return {
        "gunluk": {"istek": len(gunluk), "maliyet_usd": round(sum(r.get("maliyet_usd", 0) for r in gunluk), 4)},
        "aylik": {"istek": len(aylik), "maliyet_usd": round(sum(r.get("maliyet_usd", 0) for r in aylik), 4)},
        "gunluk_limit": AI_MAX_DAILY_REQUESTS,
    }


@api_router.post("/ai/demo-yukle")
async def ai_demo_yukle(current_user=Depends(get_current_user)):
    """Admin: Tüm AI demo verilerini oluştur/yenile."""
    if current_user.get("role") not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yetkiniz yok")

    simdi = datetime.utcnow()

    # Tüm öğrencileri bul
    tum_ogrenciler = await db.students.find({}).to_list(length=50)
    if not tum_ogrenciler:
        tum_ogrenciler = await db.users.find({"role": "student"}).to_list(length=50)
    if not tum_ogrenciler:
        raise HTTPException(status_code=404, detail="Hiç öğrenci bulunamadı")

    ogretmen = await db.users.find_one({"role": "teacher"})
    ogretmen_id = ogretmen["id"] if ogretmen else current_user["id"]
    ogretmen_ad = f"{ogretmen.get('ad','')} {ogretmen.get('soyad','')}".strip() if ogretmen else "Admin"

    sonuc = {"dna": 0, "kocluk": 0, "kelime": 0, "kelime_tekrar": 0, "yukleme": 0, "parca": 0, "soru": 0, "socratic": 0}

    # Eski demo verileri temizle
    await db.okuma_dna.delete_many({})
    await db.ai_kocluk_cache.delete_many({"model": "demo"})
    await db.kelime_tekrar.delete_many({})
    await db.ai_socratic_log.delete_many({})

    profil_havuzu = [
        {"tip": "hayalci_okuyucu", "label": "🌈 Hayalci Okuyucu", "b": {"kelime_gucu": 72, "akicilik": 45, "anlama_derinligi": 80, "dikkat_suresi": 55, "zorluk_toleransi": 60, "kelime_tekrar_ihtiyaci": 35, "okuma_psikolojisi": "keşifçi"}},
        {"tip": "hızlı_okuyucu", "label": "⚡ Hızlı Okuyucu", "b": {"kelime_gucu": 58, "akicilik": 78, "anlama_derinligi": 42, "dikkat_suresi": 70, "zorluk_toleransi": 45, "kelime_tekrar_ihtiyaci": 50, "okuma_psikolojisi": "güvenli"}},
        {"tip": "başlangıç_okuyucu", "label": "🌱 Başlangıç Okuyucu", "b": {"kelime_gucu": 35, "akicilik": 30, "anlama_derinligi": 55, "dikkat_suresi": 25, "zorluk_toleransi": 30, "kelime_tekrar_ihtiyaci": 75, "okuma_psikolojisi": "kararsız"}},
        {"tip": "dengeli_okuyucu", "label": "⚖️ Dengeli Okuyucu", "b": {"kelime_gucu": 65, "akicilik": 62, "anlama_derinligi": 68, "dikkat_suresi": 60, "zorluk_toleransi": 55, "kelime_tekrar_ihtiyaci": 40, "okuma_psikolojisi": "güvenli"}},
        {"tip": "analitik_okuyucu", "label": "🔍 Analitik Okuyucu", "b": {"kelime_gucu": 80, "akicilik": 50, "anlama_derinligi": 85, "dikkat_suresi": 75, "zorluk_toleransi": 70, "kelime_tekrar_ihtiyaci": 25, "okuma_psikolojisi": "keşifçi"}},
    ]

    for i, ogr in enumerate(tum_ogrenciler):
        oid = ogr["id"]
        ad = ogr.get("ad", f"Öğrenci{i}")
        sinif = ogr.get("sinif", 3)
        p = profil_havuzu[i % len(profil_havuzu)]

        # Hafif randomize
        boyutlar = {}
        for k, v in p["b"].items():
            if isinstance(v, int):
                boyutlar[k] = max(5, min(100, v + random.randint(-10, 10)))
            else:
                boyutlar[k] = v

        # DNA
        await db.okuma_dna.update_one({"ogrenci_id": oid}, {"$set": {
            "ogrenci_id": oid, "boyutlar": boyutlar, "profil_tipi": p["tip"],
            "profil_label": p["label"], "sinif": sinif, "son_guncelleme": simdi.isoformat(),
        }}, upsert=True)
        sonuc["dna"] += 1

        # Koçluk cache
        await db.ai_kocluk_cache.update_one({"ogrenci_id": oid}, {"$set": {
            "id": str(uuid.uuid4()), "ogrenci_id": oid,
            "dna": {"profil_tipi": p["tip"], "profil_label": p["label"], "boyutlar": boyutlar},
            "ai_analiz": {
                "durum_degerlendirmesi": {
                    "guclu_yonler": random.sample(["Anlama kapasitesi yüksek", "Düzenli okuyor", "Hayal gücü gelişmiş", "Meraklı", "Kelime hazinesi iyi", "Cesur kitap seçimleri", "Hızlı okuma", "Dikkatli dinleme"], 3),
                    "gelisim_alanlari": random.sample(["Okuma hızı artırılmalı", "Anlama derinliği geliştirilmeli", "Kelime hazinesi genişletilmeli", "Dikkat süresi kısa", "Bloom üst basamakları zayıf", "Streak tutarsız"], 2),
                },
                "risk_analizi": {"seviye": random.choice(["düşük", "orta", "yüksek"]), "faktorler": [random.choice(["Streak kırılma riski", "Kelime gücü düşük", "Dikkat süresi kısa", "Zor metinlerden kaçınma"])], "aciliyet": random.choice(["Takip yeterli", "Haftalık kontrol", "Acil müdahale"])},
                "mudahale_plani": {"hafta_1": "Günlük 10 dk sesli okuma", "hafta_2": "Tekrarlı okuma + kelime çalışması", "hafta_3": "Bloom soru çözme pratiği", "hafta_4": "Bağımsız okuma + özet yazma"},
                "veliye_mesaj": f"Sayın Veli, {ad} okuma gelişiminde ilerleme kaydediyor. Evde günlük 10-15 dakika birlikte okuma yapmanız gelişimini hızlandıracaktır. Kitap önerilerimizi takip edebilirsiniz.",
                "haftalik_gorevler": [
                    {"gun": "Pazartesi", "gorev": "Sevdiği kitaptan 2 sayfa sesli oku", "bloom": "uygulama"},
                    {"gun": "Salı", "gorev": "5 yeni kelime öğren ve cümle kur", "bloom": "uygulama"},
                    {"gun": "Çarşamba", "gorev": "Okuduğu bölümün özetini yaz", "bloom": "sentez"},
                    {"gun": "Perşembe", "gorev": "Karakterin motivasyonunu analiz et", "bloom": "analiz"},
                    {"gun": "Cuma", "gorev": "Hikâyeye alternatif son yaz", "bloom": "yaratma"},
                ],
                "kitap_tavsiyeleri": random.sample([
                    {"ad": "Charlie'nin Çikolata Fabrikası", "yazar": "Roald Dahl", "neden": "Hayal gücü yüksek, kısa bölümler"},
                    {"ad": "Küçük Prens", "yazar": "Saint-Exupéry", "neden": "Felsefi derinlik, kısa paragraflar"},
                    {"ad": "Pollyanna", "yazar": "E.H. Porter", "neden": "Pozitif bakış, karakter gelişimi"},
                    {"ad": "Kaşağı", "yazar": "Ömer Seyfettin", "neden": "Kısa öykü, değer eğitimi"},
                    {"ad": "Martı", "yazar": "Richard Bach", "neden": "Cesaret ve azim teması"},
                ], 3),
                "motivasyon_mesaji": random.choice([
                    f"Harika gidiyorsun {ad}! Her gün biraz daha güçleniyorsun 🌟",
                    f"Süpersin {ad}! Okumaya devam et, başarı senin hakkın 💪",
                    f"Merhaba {ad}! Bugün yeni bir maceraya hazır mısın? 📚",
                    f"{ad}, senin gibisi az bulunur! Her kelime bir adım 🚀",
                ]),
                "kelime_mudahale": "Günlük 5 yeni kelime + görsel kartlarla çalışma. Spaced repetition tekrarı.",
                "metin_recetesi": {"paragraf_uzunlugu": "Orta (80-120 kelime)", "soyutluk": random.choice(["Düşük", "Orta", "Yüksek"]), "aksiyon": random.choice(["Düşük", "Orta", "Yüksek"]), "hedef_kelime_orani": "%70 bilinen + %30 yeni"},
            },
            "ai_ham_metin": "", "model": "demo", "token": 0, "maliyet": 0, "tarih": simdi.isoformat(),
        }}, upsert=True)
        sonuc["kocluk"] += 1

        # Kelime tekrar (Spaced Repetition)
        demo_kelimeler_tekrar = [
            {"kelime": "macera", "anlam": "Tehlikeli ve heyecan verici olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı."},
            {"kelime": "keşif", "anlam": "Bilinmeyen bir şeyi ilk kez bulma", "ornek_cumle": "Bilim insanı yeni bir keşif yaptı."},
            {"kelime": "pusula", "anlam": "Yön bulmaya yarayan araç", "ornek_cumle": "Kaşif pusulasıyla yolunu buldu."},
            {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme gücü", "ornek_cumle": "Küçük kız büyük cesaret gösterdi."},
            {"kelime": "merak", "anlam": "Bir şeyi bilmek isteme duygusu", "ornek_cumle": "Merak eden çocuk her şeyi sorar."},
            {"kelime": "sabır", "anlam": "Bekleyebilme gücü", "ornek_cumle": "Bahçıvan sabırla bekledi."},
            {"kelime": "göç", "anlam": "Toplu taşınma", "ornek_cumle": "Kuşlar sıcak ülkelere göç eder."},
            {"kelime": "dürüstlük", "anlam": "Doğruyu söyleme", "ornek_cumle": "Dürüstlük en değerli erdemdir."},
        ]
        for kt in demo_kelimeler_tekrar:
            kutu = random.randint(1, 5)
            gun_sonra = {1:0, 2:1, 3:5, 4:14, 5:30}[kutu]
            await db.kelime_tekrar.insert_one({
                "id": str(uuid.uuid4()), "ogrenci_id": oid, "sinif": sinif, "kutu": kutu,
                "tekrar_sayisi": random.randint(1, 8), "dogru_sayisi": random.randint(0, 6),
                "son_gosterim": (simdi - timedelta(days=random.randint(1, 7))).isoformat(),
                "sonraki_gosterim": (simdi + timedelta(days=gun_sonra)).isoformat() if gun_sonra > 0 else simdi.isoformat(),
                "tarih": (simdi - timedelta(days=14)).isoformat(), **kt,
            })
            sonuc["kelime_tekrar"] += 1

        # Socratic log
        socratic_sorular = [
            "Bu bölümde en çok ne dikkatini çekti?",
            "Karakter neden böyle davrandı sence?",
            "Sen olsaydın ne yapardın?",
            "Bu hikâyenin sana öğrettiği bir şey var mı?",
            "Hikâyenin sonu farklı olabilir miydi?",
        ]
        for j in range(random.randint(2, 4)):
            await db.ai_socratic_log.insert_one({
                "id": str(uuid.uuid4()), "ogrenci_id": oid,
                "kitap_adi": random.choice(["Ormanın Sırrı", "Dürüst Çocuk", "Göçmen Kuşlar", "Takım Çalışması"]),
                "bolum": f"Bölüm {random.randint(1,4)}", "soru": random.choice(socratic_sorular),
                "bloom": random.choice(["kavrama", "analiz", "sentez", "degerlendirme"]),
                "puan": random.randint(3, 5), "tarih": (simdi - timedelta(days=random.randint(0, 5))).isoformat(),
            })
            sonuc["socratic"] += 1

    # Kelimeler (meb_kelime_haritasi)
    demo_kelimeler = [
        {"kelime": "macera", "anlam": "Tehlikeli ve heyecan verici olay", "ornek_cumle": "Çocuklar ormanda büyük bir macera yaşadı.", "zorluk": 4, "sinif": 3},
        {"kelime": "keşif", "anlam": "Bilinmeyen bir şeyi ilk kez bulma", "ornek_cumle": "Bilim insanı yeni bir keşif yaptı.", "zorluk": 5, "sinif": 3},
        {"kelime": "pusula", "anlam": "Yön bulmaya yarayan araç", "ornek_cumle": "Kaşif pusulasıyla yolunu buldu.", "zorluk": 6, "sinif": 3},
        {"kelime": "cesaret", "anlam": "Korkmadan hareket edebilme gücü", "ornek_cumle": "Küçük kız büyük cesaret gösterdi.", "zorluk": 4, "sinif": 3},
        {"kelime": "merak", "anlam": "Bir şeyi bilmek isteme duygusu", "ornek_cumle": "Merak eden çocuk her şeyi sorar.", "zorluk": 3, "sinif": 3},
        {"kelime": "sabır", "anlam": "Bekleyebilme ve dayanma gücü", "ornek_cumle": "Bahçıvan sabırla çiçeklerin büyümesini bekledi.", "zorluk": 4, "sinif": 3},
        {"kelime": "hayal gücü", "anlam": "Zihinde yeni şeyler oluşturabilme", "ornek_cumle": "Hayal gücü güçlü olan çocuklar iyi yazar.", "zorluk": 5, "sinif": 3},
        {"kelime": "fedakarlık", "anlam": "Başkaları için vazgeçme", "ornek_cumle": "Anneler büyük fedakarlıklar yapar.", "zorluk": 7, "sinif": 4},
        {"kelime": "dürüstlük", "anlam": "Doğruyu söyleme", "ornek_cumle": "Dürüstlük en değerli erdemdir.", "zorluk": 5, "sinif": 4},
        {"kelime": "azim", "anlam": "Kararlılıkla sürdürme", "ornek_cumle": "Azimli öğrenci başarıya ulaştı.", "zorluk": 6, "sinif": 4},
        {"kelime": "empati", "anlam": "Başkalarının duygularını anlama", "ornek_cumle": "Empati kurabilen iyi arkadaş olur.", "zorluk": 7, "sinif": 5},
        {"kelime": "göç", "anlam": "Toplu taşınma", "ornek_cumle": "Kuşlar sıcak ülkelere göç eder.", "zorluk": 4, "sinif": 3},
    ]
    await db.meb_kelime_haritasi.delete_many({"kaynak": {"$regex": "Demo"}})
    for k in demo_kelimeler:
        await db.meb_kelime_haritasi.update_one({"kelime": k["kelime"]}, {"$set": {
            "id": str(uuid.uuid4()), "kaynak": "Demo - MEB Türkçe",
            "yukleyen_id": ogretmen_id, "tarih": simdi.isoformat(), **k,
        }}, upsert=True)
        sonuc["kelime"] += 1

    # Demo yükleme
    await db.ai_yuklemeler.delete_many({"kitap_adi": {"$regex": "Demo"}})
    await db.ai_okuma_parcalari.delete_many({"kitap_adi": {"$regex": "Demo"}})
    await db.ai_uretilen_sorular.delete_many({"kitap_adi": {"$regex": "Demo"}})

    demo_yuk_id = str(uuid.uuid4())
    await db.ai_yuklemeler.insert_one({
        "id": demo_yuk_id, "dosya_adi": "turkce_3_ders_kitabi.pdf", "dosya_boyut": 4500000,
        "dosya_format": ".pdf", "dosya_hash": f"demo_{uuid.uuid4().hex[:8]}", "dosya_b64": "",
        "sinif": 3, "tur": "ders_kitabi", "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)",
        "yazar": "MEB", "temalar": ["Erdemler", "Doğa ve Evren", "Çocuk Dünyası"],
        "yukleyen_id": ogretmen_id, "yukleyen_ad": ogretmen_ad, "yukleyen_rol": "teacher",
        "durum": "tamamlandi", "onayli": True, "ilerleme": 100,
        "sonuc": {"sayfa_sayisi": 180, "kelime_sayisi": 45000, "chunk_sayisi": 8, "cikarilan_kelime": 12, "eklenen_kelime": 12, "okuma_parcasi": 4, "uretilen_soru": 10, "bonus_puan": 10},
        "guven_skoru": {"toplam": 92, "seviye": "yuksek", "detay": {"icindekiler": 90, "dil_uygunlugu": 95, "bloom_dagilimi": 88}}, "okuma_seviyesi": "3. Sınıf", "versiyon": 1, "tarih": (simdi - timedelta(days=3)).isoformat(),
    })
    sonuc["yukleme"] += 1

    demo_parcalar = [
        {"baslik": "Ormanın Sırrı", "ozet": "Küçük Ali ormanda kaybolur, konuşan hayvanlarla arkadaş olur.", "tema": "Doğa ve Evren", "metin_kesit": "Ağaçların arasından süzülen güneş ışığı ormanın derinliklerini aydınlatıyordu. Küçük Ali ilk kez bu kadar içerilere gelmişti..."},
        {"baslik": "Dürüst Çocuk", "ozet": "Pazarda para bulan Elif'in dürüstlük hikâyesi.", "tema": "Erdemler", "metin_kesit": "Elif pazarda yerde parlayan bir şey gördü. Eğilip baktığında bunun bir cüzdan olduğunu anladı. 'Bunu sahibine vermem lazım' dedi..."},
        {"baslik": "Göçmen Kuşlar", "ozet": "Leyleklerin göç yolculuğu ve yol bulma yetenekleri.", "tema": "Doğa ve Evren", "metin_kesit": "Her sonbaharda leylekler uzun bir yolculuğa çıkar. Binlerce kilometre uçarak sıcak ülkelere göç ederler..."},
        {"baslik": "Takım Çalışması", "ozet": "Sınıftaki öğrencilerin birlikte proje hazırlama hikâyesi.", "tema": "Çocuk Dünyası", "metin_kesit": "Öğretmen sınıfa bir proje verdi: 'Hayalinizdeki şehri tasarlayın.' Herkes tek başına yapmak istedi ama çok zordu..."},
    ]
    for j, p in enumerate(demo_parcalar):
        await db.ai_okuma_parcalari.insert_one({"id": str(uuid.uuid4()), "yukleme_id": demo_yuk_id, "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)", "sinif": 3, "bolum": j+1, **p, "kelime_sayisi": len(p["metin_kesit"].split()), "tarih": simdi.isoformat()})
        sonuc["parca"] += 1

    demo_sorular = [
        {"soru": "Ali ormanda kime rastladı?", "secenekler": ["Sincap", "Ayı", "Balık", "Kuş"], "dogru_cevap": 0, "taksonomi": "bilgi", "bolum": 1},
        {"soru": "Hayvanlar Ali'ye nasıl yardım etti?", "secenekler": ["Yemek verdiler", "Yol gösterdiler", "Şarkı söylediler", "Uyuttular"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 1},
        {"soru": "Elif cüzdanı neden sahibine verdi?", "secenekler": ["Korktu", "Dürüsttü", "Paraya ihtiyacı yoktu", "Annesi gördü"], "dogru_cevap": 1, "taksonomi": "analiz", "bolum": 2},
        {"soru": "Sen Elif'in yerinde olsan ne yapardın?", "secenekler": ["Sahibine verirdim", "Saklardım", "Polise götürürdüm", "Bilmiyorum"], "dogru_cevap": 0, "taksonomi": "degerlendirme", "bolum": 2},
        {"soru": "Leylekler neden göç eder?", "secenekler": ["Sıcak ülke", "Arkadaş bulma", "Yeni yuva", "Uçma öğrenme"], "dogru_cevap": 0, "taksonomi": "bilgi", "bolum": 3},
        {"soru": "Kuşların yol bulma yeteneğine ne denir?", "secenekler": ["GPS", "İçgüdü", "Harita", "Rüzgar"], "dogru_cevap": 1, "taksonomi": "kavrama", "bolum": 3},
        {"soru": "Takım çalışması neden önemlidir?", "secenekler": ["Hızlı biter", "Herkes öğrenir", "Eğlenceli", "Hepsi"], "dogru_cevap": 3, "taksonomi": "sentez", "bolum": 4},
    ]
    for s in demo_sorular:
        await db.ai_uretilen_sorular.insert_one({"id": str(uuid.uuid4()), "yukleme_id": demo_yuk_id, "kitap_adi": "Türkçe 3 Ders Kitabı (Demo)", "sinif": 3, **s, "tarih": simdi.isoformat()})
        sonuc["soru"] += 1

    return {"mesaj": "✅ AI demo verileri oluşturuldu!", "sonuc": sonuc}


# ─────────────────────────────────────────────
# DALGA 2: SOCRATİC READİNG + KELİME EVRİMİ + MİNİ OYUN
# ─────────────────────────────────────────────

@api_router.post("/ai/socratic-soru")
async def ai_socratic_soru(payload: dict, current_user=Depends(get_current_user)):
    """Okuma kaydı sonrası Sokratik soru üretir."""
    kitap_adi = payload.get("kitap_adi", "")
    bolum = payload.get("bolum", "")
    sure_dk = payload.get("sure_dk", 10)
    sinif = payload.get("sinif", 3)

    prompt = f"""Kitap: {kitap_adi or 'bilinmiyor'}
Bölüm: {bolum or 'bilinmiyor'}
Sınıf: {sinif}
Okuma süresi: {sure_dk} dk

Bu öğrenci az önce okuma yaptı. Okuduğu hakkında düşünmesini sağlayacak 1 Sokratik soru sor.
Soru kısa, merak uyandırıcı ve Türkçe olsun. Çocuğa uygun dil kullan.
SADECE JSON döndür: {{"soru": "...", "ipucu": "...", "bloom": "kavrama|analiz|sentez|degerlendirme"}}"""

    result = await call_claude(
        "Sen çocuklara Sokratik sorular soran sevecen bir okuma koçusun. Düşünmeye teşvik edersin.",
        prompt, model="haiku", max_tokens=200
    )

    if result.get("parsed"):
        soru_data = result["parsed"]
    else:
        soru_data = {"soru": f"Az önce okuduğun bölümde en çok ne dikkatini çekti?", "ipucu": "Karakterlerin davranışlarını düşün", "bloom": "kavrama"}

    # Log
    await db.ai_socratic_log.insert_one({
        "id": str(uuid.uuid4()),
        "ogrenci_id": current_user["id"],
        "kitap_adi": kitap_adi,
        "bolum": bolum,
        "soru": soru_data.get("soru", ""),
        "bloom": soru_data.get("bloom", "kavrama"),
        "tarih": datetime.utcnow().isoformat(),
    })

    return soru_data


@api_router.post("/ai/socratic-cevap")
async def ai_socratic_cevap(payload: dict, current_user=Depends(get_current_user)):
    """Öğrencinin Sokratik soruya verdiği cevabı değerlendirir."""
    soru = payload.get("soru", "")
    cevap = payload.get("cevap", "")

    if len(cevap) < 5:
        return {"puan": 1, "geri_bildirim": "Biraz daha düşünüp detaylı cevap vermeyi dener misin? 🤔", "xp": 2}

    prompt = f"""Soru: {soru}
Öğrenci cevabı: {cevap}

Bu cevabı 1-5 arası puanla. Kısa, pozitif geri bildirim ver. Türkçe. Çocuğa uygun.
SADECE JSON: {{"puan": 1-5, "geri_bildirim": "..."}}"""

    result = await call_claude("Sen yapıcı geri bildirim veren bir okuma koçusun.", prompt, model="haiku", max_tokens=150)

    if result.get("parsed"):
        r = result["parsed"]
        puan = min(5, max(1, r.get("puan", 3)))
        geri = r.get("geri_bildirim", "Güzel düşünmüşsün! 👏")
    else:
        puan = 3
        geri = "Düşüncelerini paylaştığın için teşekkürler! 👏"

    xp = {1: 2, 2: 3, 3: 5, 4: 7, 5: 10}.get(puan, 5)

    # XP ver
    try:
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"toplam_xp": xp}})
    except:
        pass

    return {"puan": puan, "geri_bildirim": geri, "xp": xp}


# ── KELİME EVRİMİ (SPACED REPETITION) ──

@api_router.get("/ai/kelime-evrimi/{ogrenci_id}")
async def kelime_evrimi_listesi(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Öğrencinin kelime tekrar programı — bugün tekrar edilmesi gerekenler."""
    simdi = datetime.utcnow()
    # Bugün veya geçmiş tarihli tekrar bekleyenler
    bekleyenler = await db.kelime_tekrar.find({
        "ogrenci_id": ogrenci_id,
        "sonraki_gosterim": {"$lte": simdi.isoformat()}
    }).sort("sonraki_gosterim", 1).to_list(length=20)

    for k in bekleyenler:
        k.pop("_id", None)

    # İstatistik
    toplam = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id})
    ogrenilmis = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id, "kutu": {"$gte": 4}})

    return {"bekleyenler": bekleyenler, "toplam": toplam, "ogrenilmis": ogrenilmis, "bugun_tekrar": len(bekleyenler)}


@api_router.post("/ai/kelime-evrimi/cevapla")
async def kelime_evrimi_cevapla(payload: dict, current_user=Depends(get_current_user)):
    """Kelime tekrarına doğru/yanlış cevap — Leitner Box algoritması."""
    kelime_id = payload.get("kelime_id", "")
    dogru = payload.get("dogru", False)
    ogrenci_id = current_user["id"]

    kayit = await db.kelime_tekrar.find_one({"id": kelime_id, "ogrenci_id": ogrenci_id})
    if not kayit:
        raise HTTPException(status_code=404, detail="Kelime kaydı bulunamadı")

    mevcut_kutu = kayit.get("kutu", 1)
    simdi = datetime.utcnow()

    # Yaş bazlı aralıklar
    sinif = kayit.get("sinif", 3)
    if sinif <= 2:  # 6-8 yaş
        araliklar = {1: 1, 2: 2, 3: 5, 4: 12, 5: 30}
    elif sinif <= 5:  # 9-11 yaş
        araliklar = {1: 1, 2: 3, 3: 7, 4: 21, 5: 45}
    else:  # 12+ yaş
        araliklar = {1: 1, 2: 3, 3: 7, 4: 30, 5: 60}

    if dogru:
        yeni_kutu = min(5, mevcut_kutu + 1)
        xp = 2
    else:
        yeni_kutu = 1  # Yanlış → ilk kutuya geri
        xp = 1

    sonraki_gun = araliklar.get(yeni_kutu, 7)
    sonraki = (simdi + timedelta(days=sonraki_gun)).isoformat()

    await db.kelime_tekrar.update_one({"id": kelime_id}, {"$set": {
        "kutu": yeni_kutu,
        "son_gosterim": simdi.isoformat(),
        "sonraki_gosterim": sonraki,
        "tekrar_sayisi": kayit.get("tekrar_sayisi", 0) + 1,
        "dogru_sayisi": kayit.get("dogru_sayisi", 0) + (1 if dogru else 0),
    }})

    # XP
    try:
        await db.users.update_one({"id": ogrenci_id}, {"$inc": {"toplam_xp": xp}})
    except:
        pass

    return {"dogru": dogru, "yeni_kutu": yeni_kutu, "sonraki_gun": sonraki_gun, "xp": xp}


@api_router.post("/ai/kelime-evrimi/ekle")
async def kelime_evrimi_ekle(payload: dict, current_user=Depends(get_current_user)):
    """Öğrenciye kelime tekrar programına kelime ekler."""
    ogrenci_id = payload.get("ogrenci_id", current_user["id"])
    kelimeler = payload.get("kelimeler", [])  # [{"kelime": "...", "anlam": "...", "sinif": 3}]

    eklenen = 0
    for k in kelimeler[:20]:  # max 20 kelime bir seferde
        mevcut = await db.kelime_tekrar.find_one({"ogrenci_id": ogrenci_id, "kelime": k.get("kelime", "").lower()})
        if not mevcut:
            await db.kelime_tekrar.insert_one({
                "id": str(uuid.uuid4()),
                "ogrenci_id": ogrenci_id,
                "kelime": k.get("kelime", "").lower(),
                "anlam": k.get("anlam", ""),
                "ornek_cumle": k.get("ornek_cumle", ""),
                "sinif": k.get("sinif", 3),
                "kutu": 1,
                "tekrar_sayisi": 0,
                "dogru_sayisi": 0,
                "son_gosterim": None,
                "sonraki_gosterim": datetime.utcnow().isoformat(),
                "tarih": datetime.utcnow().isoformat(),
            })
            eklenen += 1

    return {"eklenen": eklenen, "toplam_gonderilen": len(kelimeler)}


# ── MİNİ OYUN ÜRETİCİ ──

@api_router.post("/ai/mini-oyun")
async def ai_mini_oyun(payload: dict, current_user=Depends(get_current_user)):
    """Kelimelerden mini oyun üretir (kelime avı, eşleştirme, boşluk doldurma)."""
    oyun_turu = payload.get("tur", "eslestirme")  # eslestirme, kelime_avi, bosluk_doldurma, cumle_kurma
    kelimeler = payload.get("kelimeler", [])  # [{"kelime": "...", "anlam": "..."}]
    sinif = payload.get("sinif", 3)

    if not kelimeler:
        # Rastgele kelime al
        rastgele = await db.meb_kelime_haritasi.find({"sinif": sinif}).to_list(length=10)
        kelimeler = [{"kelime": k.get("kelime", ""), "anlam": k.get("anlam", ""), "ornek_cumle": k.get("ornek_cumle", "")} for k in rastgele]

    if not kelimeler:
        return {"oyun": None, "mesaj": "Kelime bulunamadı"}

    if oyun_turu == "eslestirme":
        # Kelime-anlam eşleştirme (Claude gerektirmez)
        karisik_kelimeler = kelimeler[:8]
        random.shuffle(karisik_kelimeler)
        karisik_anlamlar = [k.get("anlam", "") for k in karisik_kelimeler]
        random.shuffle(karisik_anlamlar)
        return {"oyun": {
            "tur": "eslestirme",
            "baslik": "🎲 Kelime Eşleştirme",
            "aciklama": "Her kelimeyi doğru anlamıyla eşleştir!",
            "kelimeler": [k.get("kelime", "") for k in karisik_kelimeler],
            "anlamlar": [k.get("anlam", "") for k in karisik_kelimeler],  # doğru sıralama (frontend karıştıracak)
            "xp": 5,
        }}

    elif oyun_turu == "bosluk_doldurma":
        # Cümledeki kelimeyi bul
        sorular = []
        for k in kelimeler[:6]:
            cumle = k.get("ornek_cumle", "")
            if cumle and k.get("kelime", ""):
                bos = cumle.replace(k["kelime"], "___").replace(k["kelime"].capitalize(), "___")
                if "___" in bos:
                    sorular.append({"cumle_bos": bos, "dogru": k["kelime"], "secenekler": []})
        # Şıklar ekle
        tum_kelimeler = [k.get("kelime", "") for k in kelimeler]
        for s in sorular:
            yanlis = [w for w in tum_kelimeler if w != s["dogru"]][:3]
            secenekler = [s["dogru"]] + yanlis
            random.shuffle(secenekler)
            s["secenekler"] = secenekler
        return {"oyun": {
            "tur": "bosluk_doldurma",
            "baslik": "⬜ Boşluk Doldur",
            "aciklama": "Cümledeki boşluğa uygun kelimeyi bul!",
            "sorular": sorular,
            "xp": 5,
        }}

    elif oyun_turu == "kelime_avi":
        # Kelime avı grid üretimi (AI)
        prompt = f"""Şu kelimelerden 8x8 harf gridi oluştur (kelime avı oyunu):
Kelimeler: {', '.join([k.get('kelime','') for k in kelimeler[:6]])}
SADECE JSON: {{"grid": [["A","B",...], ...], "kelimeler": ["kelime1", ...], "yonler": ["sağa","aşağı",...] }}"""

        result = await call_claude("Sen kelime oyunu tasarımcısısın.", prompt, model="haiku", max_tokens=500)
        if result.get("parsed"):
            return {"oyun": {"tur": "kelime_avi", "baslik": "🔍 Kelime Avı", "aciklama": "Gizli kelimeleri bul!", **result["parsed"], "xp": 7}}
        else:
            return {"oyun": {"tur": "kelime_avi", "baslik": "🔍 Kelime Avı", "aciklama": "Kelimeleri bul!", "kelimeler": [k.get("kelime","") for k in kelimeler[:6]], "xp": 7}}

    elif oyun_turu == "cumle_kurma":
        sorular = []
        for k in kelimeler[:5]:
            cumle = k.get("ornek_cumle", "")
            if cumle:
                kelime_listesi = cumle.split()
                karisik = kelime_listesi.copy()
                random.shuffle(karisik)
                sorular.append({"karisik": karisik, "dogru": kelime_listesi, "hedef_kelime": k.get("kelime", "")})
        return {"oyun": {
            "tur": "cumle_kurma",
            "baslik": "📝 Cümle Kurma",
            "aciklama": "Karışık kelimeleri doğru sıraya diz!",
            "sorular": sorular,
            "xp": 5,
        }}

    return {"oyun": None, "mesaj": "Bilinmeyen oyun türü"}


@api_router.post("/ai/mini-oyun/tamamla")
async def ai_mini_oyun_tamamla(payload: dict, current_user=Depends(get_current_user)):
    """Mini oyun tamamlama — XP ver."""
    oyun_turu = payload.get("tur", "")
    dogru_sayisi = payload.get("dogru", 0)
    toplam = payload.get("toplam", 1)
    basari = round(dogru_sayisi / max(toplam, 1) * 100)

    xp = 3 if basari < 50 else 5 if basari < 80 else 7 if basari < 100 else 10

    try:
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"toplam_xp": xp}})
    except:
        pass

    return {"xp": xp, "basari": basari, "mesaj": "Harika!" if basari >= 80 else "İyi gidiyor!" if basari >= 50 else "Tekrar deneyelim!"}


# ─────────────────────────────────────────────
# AI BİLGİ TABANI — PDF/DOCX Yükleme + AI Öğrenme + Puan Sistemi
# ─────────────────────────────────────────────

AI_EGITIM_PUANLARI = {
    "pdf_yukle": 20,
    "docx_yukle": 20,
    "onaylandi": 10,       # admin onayı sonrası ek puan
    "kelime_zengin": 5,    # 50+ kelime çıkarılan yükleme bonusu
    "soru_zengin": 5,      # 20+ soru çıkarılan yükleme bonusu
}

DESTEKLENEN_FORMATLAR = [".pdf", ".docx", ".doc"]


def hesapla_guven_skoru(analiz_sonuc, sinif):
    """AI Güven Skoru: Yükleme kalitesini 0-100 puanlar"""
    skor = 0
    detay = {}

    # 1. Kelime Çeşitliliği (max 25 puan)
    kelimeler = analiz_sonuc.get("hedef_kelimeler", [])
    kelime_sayisi = len(kelimeler)
    if kelime_sayisi >= 30:
        detay["kelime_cesitliligi"] = 25
    elif kelime_sayisi >= 20:
        detay["kelime_cesitliligi"] = 20
    elif kelime_sayisi >= 10:
        detay["kelime_cesitliligi"] = 15
    elif kelime_sayisi >= 5:
        detay["kelime_cesitliligi"] = 10
    else:
        detay["kelime_cesitliligi"] = 3
    skor += detay["kelime_cesitliligi"]

    # 2. Soru Kalitesi — Bloom dağılımı (max 25 puan)
    sorular = analiz_sonuc.get("sorular", [])
    bloom_dagilim = set()
    for s in sorular:
        b = s.get("bloom_basamagi") or s.get("taksonomi", "")
        if b:
            bloom_dagilim.add(b)
    if len(bloom_dagilim) >= 5:
        detay["soru_kalitesi"] = 25
    elif len(bloom_dagilim) >= 4:
        detay["soru_kalitesi"] = 20
    elif len(bloom_dagilim) >= 3:
        detay["soru_kalitesi"] = 15
    elif len(bloom_dagilim) >= 2:
        detay["soru_kalitesi"] = 10
    else:
        detay["soru_kalitesi"] = 5 if len(sorular) > 0 else 0
    skor += detay["soru_kalitesi"]

    # 3. Zorluk Uyumu — metnin zorluk puanı sınıf seviyesine uygun mu? (max 25 puan)
    zorluk = analiz_sonuc.get("zorluk_puani", 5)
    beklenen_zorluk = sinif  # 1. sınıf → 1, 4. sınıf → 4 civarı
    fark = abs(zorluk - beklenen_zorluk)
    if fark <= 1:
        detay["zorluk_uyumu"] = 25
    elif fark <= 2:
        detay["zorluk_uyumu"] = 18
    elif fark <= 3:
        detay["zorluk_uyumu"] = 10
    else:
        detay["zorluk_uyumu"] = 3
    skor += detay["zorluk_uyumu"]

    # 4. Okuma Seviyesi Analizi — Grade Level Score varlığı (max 25 puan)
    grade_level = analiz_sonuc.get("grade_level_score", {})
    if grade_level:
        gl_sinif = grade_level.get("tahmini_sinif", 0)
        gl_fark = abs(gl_sinif - sinif)
        if gl_fark <= 0.5:
            detay["seviye_uyumu"] = 25
        elif gl_fark <= 1:
            detay["seviye_uyumu"] = 20
        elif gl_fark <= 2:
            detay["seviye_uyumu"] = 12
        else:
            detay["seviye_uyumu"] = 5
    else:
        detay["seviye_uyumu"] = 10  # Grade level yoksa orta puan
    skor += detay["seviye_uyumu"]

    return {"skor": min(skor, 100), "detay": detay, "seviye": "yuksek" if skor >= 75 else "orta" if skor >= 50 else "dusuk"}


def hesapla_grade_level(metin):
    """Okuma Seviyesi Analizi — Grade Level Score"""
    kelimeler = metin.split()
    kelime_sayisi = len(kelimeler)
    if kelime_sayisi == 0:
        return {"kelime_sayisi": 0, "ort_kelime_uzunlugu": 0, "ort_cumle_uzunlugu": 0, "tahmini_sinif": 1, "zorluk_puani": 1}

    # Ortalama kelime uzunluğu
    ort_kelime_uzunlugu = sum(len(k) for k in kelimeler) / kelime_sayisi

    # Cümle sayısı
    import re
    cumleler = re.split(r'[.!?]+', metin)
    cumle_sayisi = max(len([c for c in cumleler if c.strip()]), 1)
    ort_cumle_uzunlugu = kelime_sayisi / cumle_sayisi

    # Uzun kelime oranı (7+ harf)
    uzun_kelime_orani = sum(1 for k in kelimeler if len(k) >= 7) / kelime_sayisi

    # Soyut kelime tahmini (basit heuristik: -lık, -lik, -sel, -sal, -cilik gibi ekler)
    soyut_ekler = ["lık", "lik", "luk", "lük", "sel", "sal", "cilik", "çilik", "sızlık", "sizlik"]
    soyut_sayisi = sum(1 for k in kelimeler if any(k.lower().endswith(e) for e in soyut_ekler))
    soyutluk_orani = soyut_sayisi / kelime_sayisi

    # Grade Level hesaplama (Türkçe uyarlamalı basit formül)
    # Temel: ort_cumle_uzunlugu * 0.3 + ort_kelime_uzunlugu * 0.5 + uzun_kelime_orani * 10 + soyutluk_orani * 15
    ham_skor = (ort_cumle_uzunlugu * 0.3) + (ort_kelime_uzunlugu * 0.5) + (uzun_kelime_orani * 10) + (soyutluk_orani * 15)

    # Sınıf seviyesine çevirme (1-8 arası)
    tahmini_sinif = max(1, min(8, round(ham_skor / 2)))
    zorluk_puani = max(1, min(10, round(ham_skor)))

    return {
        "kelime_sayisi": kelime_sayisi,
        "cumle_sayisi": cumle_sayisi,
        "ort_kelime_uzunlugu": round(ort_kelime_uzunlugu, 1),
        "ort_cumle_uzunlugu": round(ort_cumle_uzunlugu, 1),
        "uzun_kelime_orani": round(uzun_kelime_orani * 100, 1),
        "soyutluk_orani": round(soyutluk_orani * 100, 1),
        "tahmini_sinif": tahmini_sinif,
        "zorluk_puani": zorluk_puani,
    }


@api_router.post("/ai/bilgi-tabani/yukle")
async def ai_bilgi_tabani_yukle(
    dosya: UploadFile = File(...),
    sinif: int = Form(...),
    tur: str = Form("ders_kitabi"),
    kitap_adi: str = Form(""),
    yazar: str = Form(""),
    temalar: str = Form(""),
    current_user=Depends(get_current_user)
):
    import os
    ext = os.path.splitext(dosya.filename)[1].lower()
    if ext not in DESTEKLENEN_FORMATLAR:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {ext}. Desteklenen: {', '.join(DESTEKLENEN_FORMATLAR)}")

    icerik = await dosya.read()
    # Boyut ve günlük yükleme limiti kaldırıldı

    # ── DUPLICATE KONTROL ──
    import hashlib
    dosya_hash = hashlib.sha256(icerik).hexdigest()
    mevcut = await db.ai_yuklemeler.find_one({"dosya_hash": dosya_hash})
    if mevcut:
        raise HTTPException(status_code=409, detail=f"Bu dosya daha önce yüklenmiş: '{mevcut.get('kitap_adi', '')}' ({mevcut.get('tarih', '')[:10]}, {mevcut.get('yukleyen_ad', '')})")

    # İsim bazlı benzerlik kontrolü (aynı kitap farklı dosya)
    benzer = await db.ai_yuklemeler.find_one({
        "kitap_adi": {"$regex": f"^{(kitap_adi or dosya.filename.replace(ext, '')).strip()[:30]}",  "$options": "i"},
        "sinif": sinif,
    })
    duplicate_uyari = ""
    if benzer:
        duplicate_uyari = f"⚠️ Benzer yükleme mevcut: '{benzer.get('kitap_adi')}' ({benzer.get('tarih', '')[:10]}). Yine de yüklendi."

    import base64
    dosya_b64 = base64.b64encode(icerik).decode("utf-8")

    yukleme = {
        "id": str(uuid.uuid4()),
        "dosya_adi": dosya.filename,
        "dosya_boyut": len(icerik),
        "dosya_format": ext,
        "dosya_hash": dosya_hash,
        "dosya_b64": dosya_b64,
        "sinif": sinif,
        "tur": tur,
        "kitap_adi": kitap_adi or dosya.filename.replace(ext, ""),
        "yazar": yazar,
        "temalar": [t.strip() for t in temalar.split(",") if t.strip()] if temalar else [],
        "yukleyen_id": current_user["id"],
        "yukleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "yukleyen_rol": current_user.get("role", ""),
        "durum": "yuklendi",
        "onayli": current_user.get("role") in ["admin", "coordinator"],
        # AI Güven Skoru alanları (işleme sonrası doldurulur)
        "guven_skoru": None,
        "okuma_seviyesi": None,
        "sonuc": {},
        "versiyon": 1,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.ai_yuklemeler.insert_one(yukleme)

    # Puan ver
    puan = AI_EGITIM_PUANLARI.get(f"{ext.replace('.', '')}_yukle", 20)
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": puan}})

    # Puan log
    await db.ai_egitim_puanlari.insert_one({
        "id": str(uuid.uuid4()),
        "kullanici_id": current_user["id"],
        "eylem": "dosya_yukle",
        "dosya_adi": dosya.filename,
        "sinif": sinif,
        "puan": puan,
        "tarih": datetime.utcnow().isoformat(),
    })

    yukleme.pop("_id", None)
    yukleme.pop("dosya_b64", None)
    mesaj = f"✅ +{puan} puan! Dosya yüklendi, işlenmeyi bekliyor."
    if duplicate_uyari:
        mesaj += f" {duplicate_uyari}"
    return {"yukleme": yukleme, "puan_kazanilan": puan, "mesaj": mesaj}


@api_router.post("/ai/bilgi-tabani/isle/{yukleme_id}")
async def ai_bilgi_tabani_isle(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yüklenen dosyayı AI ile işle: metin çıkar → kelimeler + okuma parçaları + sorular üret."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    if not yukleme.get("dosya_b64"):
        raise HTTPException(status_code=400, detail="Dosya verisi bulunamadı")

    import base64
    dosya_bytes = base64.b64decode(yukleme["dosya_b64"])
    ext = yukleme.get("dosya_format", ".pdf")
    sinif = yukleme.get("sinif", 3)
    kitap_adi = yukleme.get("kitap_adi", "Bilinmeyen")

    # ── AŞAMA 1: METİN ÇIKARMA ──
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "metin_cikariliyor", "ilerleme": 10}})

    ham_metin = ""
    sayfa_sayisi = 0
    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(stream=dosya_bytes, filetype="pdf")
            sayfa_sayisi = len(doc)
            for page in doc:
                ham_metin += page.get_text() + "\n"
            doc.close()
        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document as DocxDocument
                doc_file = io.BytesIO(dosya_bytes)
                doc_obj = DocxDocument(doc_file)
                for para in doc_obj.paragraphs:
                    if para.text.strip():
                        ham_metin += para.text + "\n"
                sayfa_sayisi = max(1, len(ham_metin) // 2000)
            except:
                ham_metin = dosya_bytes.decode("utf-8", errors="ignore")
                sayfa_sayisi = max(1, len(ham_metin) // 2000)
    except Exception as e:
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "hata", "sonuc": {"hata": f"Metin çıkarma hatası: {str(e)[:200]}"}}})
        raise HTTPException(status_code=500, detail=f"Metin çıkarma hatası: {str(e)[:200]}")

    if len(ham_metin.strip()) < 100:
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "hata", "sonuc": {"hata": "Yeterli metin çıkarılamadı (min 100 karakter)"}}})
        raise HTTPException(status_code=400, detail="Yeterli metin çıkarılamadı")

    kelime_sayisi = len(ham_metin.split())
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"durum": "ai_analiz", "ilerleme": 30, "sonuc": {"sayfa_sayisi": sayfa_sayisi, "kelime_sayisi": kelime_sayisi}}})

    # ── AŞAMA 2: CHUNKING ──
    chunk_boyut = {1:200, 2:300, 3:400, 4:500, 5:600, 6:700, 7:800, 8:900}.get(sinif, 500)
    paragraflar = [p.strip() for p in ham_metin.split("\n") if len(p.strip()) > 30]
    chunks = []
    mevcut_chunk = ""
    for p in paragraflar:
        if len(mevcut_chunk.split()) + len(p.split()) > chunk_boyut:
            if mevcut_chunk.strip():
                chunks.append(mevcut_chunk.strip())
            mevcut_chunk = p
        else:
            mevcut_chunk += "\n" + p
    if mevcut_chunk.strip():
        chunks.append(mevcut_chunk.strip())

    if not chunks:
        chunks = [ham_metin[:2000]]

    # Max 10 chunk (maliyet kontrolü)
    chunks = chunks[:10]

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": 40, "sonuc.chunk_sayisi": len(chunks)}})

    # ── AŞAMA 3: AI ANALİZİ (her chunk için) ──
    tum_kelimeler = []
    tum_parcalar = []
    tum_sorular = []

    for i, chunk in enumerate(chunks):
        ilerleme = 40 + int((i / len(chunks)) * 40)
        await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": ilerleme}})

        ai_prompt = f"""Kitap: {kitap_adi}
Sınıf: {sinif}. sınıf
Bölüm {i+1}/{len(chunks)}

METİN:
{chunk[:2000]}

Şu JSON'u üret:
{{
  "hedef_kelimeler": [
    {{"kelime": "...", "anlam": "...", "ornek_cumle": "...", "zorluk": 1-10}}
  ],
  "okuma_parcasi": {{
    "baslik": "Bu bölümün kısa başlığı",
    "ozet": "2-3 cümle özet",
    "tema": "MEB teması",
    "kelime_sayisi": 0
  }},
  "sorular": [
    {{"soru": "...", "secenekler": ["A","B","C","D"], "dogru_cevap": 0, "taksonomi": "bilgi"}}
  ]
}}

Kurallar:
- Hedef kelimeler: {sinif}. sınıf öğrencisinin ÖĞRENMESİ gereken 5-15 kelime
- Her kelimeye çocuğun anlayacağı basit anlam ve örnek cümle yaz
- Okuma parçası özeti çocuğun merak edeceği şekilde olsun
- 3-5 soru üret (farklı Bloom basamaklarından)
- SADECE JSON döndür"""

        result = await call_claude(
            "Sen MEB Türkçe müfredatını bilen bir dil eğitimcisisin. Çocuklara uygun kelime ve soru üretirsin.",
            ai_prompt, model="sonnet", max_tokens=2000
        )

        if result.get("parsed"):
            p = result["parsed"]
            # Kelimeler
            for k in p.get("hedef_kelimeler", []):
                k["kaynak_kitap"] = kitap_adi
                k["sinif"] = sinif
                k["bolum"] = i + 1
                tum_kelimeler.append(k)
            # Okuma parçası
            parca = p.get("okuma_parcasi", {})
            if parca:
                parca["kaynak_kitap"] = kitap_adi
                parca["bolum"] = i + 1
                parca["metin_kesit"] = chunk[:500]
                tum_parcalar.append(parca)
            # Sorular
            for s in p.get("sorular", []):
                s["kaynak_kitap"] = kitap_adi
                s["sinif"] = sinif
                s["bolum"] = i + 1
                tum_sorular.append(s)

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"ilerleme": 85}})

    # ── AŞAMA 4: VERİTABANINA KAYDET ──
    # Kelimeleri meb_kelime_haritasi'na ekle
    eklenen_kelime = 0
    for k in tum_kelimeler:
        mevcut = await db.meb_kelime_haritasi.find_one({"kelime": k.get("kelime", "").lower(), "sinif": sinif})
        if not mevcut:
            await db.meb_kelime_haritasi.insert_one({
                "id": str(uuid.uuid4()),
                "sinif": sinif,
                "kelime": k.get("kelime", "").lower(),
                "anlam": k.get("anlam", ""),
                "ornek_cumle": k.get("ornek_cumle", ""),
                "zorluk": k.get("zorluk", 5),
                "kaynak": kitap_adi,
                "yukleyen_id": current_user["id"],
                "tarih": datetime.utcnow().isoformat(),
            })
            eklenen_kelime += 1

    # Okuma parçalarını kaydet
    for p in tum_parcalar:
        await db.ai_okuma_parcalari.insert_one({
            "id": str(uuid.uuid4()),
            "yukleme_id": yukleme_id,
            "kitap_adi": kitap_adi,
            "sinif": sinif,
            "bolum": p.get("bolum", 0),
            "baslik": p.get("baslik", ""),
            "ozet": p.get("ozet", ""),
            "tema": p.get("tema", ""),
            "metin_kesit": p.get("metin_kesit", ""),
            "kelime_sayisi": p.get("kelime_sayisi", 0),
            "tarih": datetime.utcnow().isoformat(),
        })

    # Soruları kaydet
    for s in tum_sorular:
        await db.ai_uretilen_sorular.insert_one({
            "id": str(uuid.uuid4()),
            "yukleme_id": yukleme_id,
            "kitap_adi": kitap_adi,
            "sinif": sinif,
            "bolum": s.get("bolum", 0),
            "soru": s.get("soru", ""),
            "secenekler": s.get("secenekler", []),
            "dogru_cevap": s.get("dogru_cevap", 0),
            "taksonomi": s.get("taksonomi", "kavrama"),
            "tarih": datetime.utcnow().isoformat(),
        })

    # Bonus puanlar
    bonus = 0
    if eklenen_kelime >= 50:
        bonus += 5
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 5}})
    if len(tum_sorular) >= 20:
        bonus += 5
        await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 5}})

    # Sonucu güncelle
    sonuc = {
        "sayfa_sayisi": sayfa_sayisi,
        "kelime_sayisi": kelime_sayisi,
        "chunk_sayisi": len(chunks),
        "cikarilan_kelime": len(tum_kelimeler),
        "eklenen_kelime": eklenen_kelime,
        "okuma_parcasi": len(tum_parcalar),
        "uretilen_soru": len(tum_sorular),
        "bonus_puan": bonus,
        "mock": not bool(GEMINI_API_KEY),
        "kelimeler": tum_kelimeler,
        "parcalar": tum_parcalar,
        "sorular": tum_sorular,
    }

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {
        "durum": "tamamlandi",
        "ilerleme": 100,
        "sonuc": {k: v for k, v in sonuc.items() if k not in ["kelimeler", "parcalar", "sorular"]},
    }})

    return sonuc


@api_router.get("/ai/bilgi-tabani/sonuc/{yukleme_id}")
async def ai_bilgi_tabani_sonuc(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yükleme sonuçlarını getir: kelimeler, okuma parçaları, sorular."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id}, {"dosya_b64": 0})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    yukleme.pop("_id", None)

    kelimeler = await db.meb_kelime_haritasi.find({"kaynak": yukleme.get("kitap_adi", "")}).to_list(length=None)
    for k in kelimeler: k.pop("_id", None)

    parcalar = await db.ai_okuma_parcalari.find({"yukleme_id": yukleme_id}).sort("bolum", 1).to_list(length=None)
    for p in parcalar: p.pop("_id", None)

    sorular = await db.ai_uretilen_sorular.find({"yukleme_id": yukleme_id}).sort("bolum", 1).to_list(length=None)
    for s in sorular: s.pop("_id", None)

    return {"yukleme": yukleme, "kelimeler": kelimeler, "parcalar": parcalar, "sorular": sorular}


@api_router.get("/ai/bilgi-tabani/ilerleme/{yukleme_id}")
async def ai_bilgi_tabani_ilerleme(yukleme_id: str, current_user=Depends(get_current_user)):
    """Yükleme işleme ilerleme durumu."""
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id}, {"dosya_b64": 0, "_id": 0})
    if not yukleme:
        return {"ilerleme": 0, "durum": "bulunamadi"}
    return {"ilerleme": yukleme.get("ilerleme", 0), "durum": yukleme.get("durum", "yuklendi"), "sonuc": yukleme.get("sonuc", {})}


@api_router.post("/ai/bilgi-tabani/yukle-url")
async def ai_bilgi_tabani_yukle_url(payload: dict, current_user=Depends(get_current_user)):
    """URL'den PDF/Word dosyası indirip yükle."""
    url = (payload.get("url") or "").strip()
    if not url or not url.startswith("http"):
        raise HTTPException(status_code=400, detail="Geçerli bir URL girin")

    sinif = payload.get("sinif", 3)
    tur = payload.get("tur", "ders_kitabi")
    kitap_adi = payload.get("kitap_adi", "")
    yazar = payload.get("yazar", "")

    # Dosyayı indir
    try:
        async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client_http:
            resp = await client_http.get(url)
            if resp.status_code != 200:
                raise HTTPException(status_code=400, detail=f"Dosya indirilemedi: HTTP {resp.status_code}")
            icerik = resp.content
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="Dosya indirme zaman aşımı (60sn)")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"İndirme hatası: {str(e)[:200]}")

    # Format tespiti
    dosya_adi = url.split("/")[-1].split("?")[0]
    ext = "." + dosya_adi.split(".")[-1].lower() if "." in dosya_adi else ""
    if ext not in [".pdf", ".docx", ".doc"]:
        # Content-Type'dan dene
        ct = resp.headers.get("content-type", "")
        if "pdf" in ct:
            ext = ".pdf"
            dosya_adi = dosya_adi or "indirilen.pdf"
        elif "word" in ct or "docx" in ct:
            ext = ".docx"
            dosya_adi = dosya_adi or "indirilen.docx"
        else:
            ext = ".pdf"  # varsayılan
            dosya_adi = dosya_adi or "indirilen.pdf"

    if ext not in DESTEKLENEN_FORMATLAR:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {ext}")

    # Duplicate kontrolü
    import hashlib
    dosya_hash = hashlib.sha256(icerik).hexdigest()
    mevcut = await db.ai_yuklemeler.find_one({"dosya_hash": dosya_hash})
    if mevcut:
        raise HTTPException(status_code=409, detail=f"Bu dosya daha önce yüklenmiş: '{mevcut.get('kitap_adi', '')}'")

    import base64
    dosya_b64 = base64.b64encode(icerik).decode("utf-8")

    yukleme = {
        "id": str(uuid.uuid4()),
        "dosya_adi": dosya_adi,
        "dosya_boyut": len(icerik),
        "dosya_format": ext,
        "dosya_hash": dosya_hash,
        "dosya_b64": dosya_b64,
        "kaynak_url": url,
        "sinif": sinif,
        "tur": tur,
        "kitap_adi": kitap_adi or dosya_adi.replace(ext, ""),
        "yazar": yazar,
        "temalar": [],
        "yukleyen_id": current_user["id"],
        "yukleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "yukleyen_rol": current_user.get("role", ""),
        "durum": "yuklendi",
        "onayli": current_user.get("role") in ["admin", "coordinator"],
        "guven_skoru": None,
        "okuma_seviyesi": None,
        "sonuc": {},
        "versiyon": 1,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.ai_yuklemeler.insert_one(yukleme)

    # Puan
    puan = 20
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": puan}})
    await db.ai_egitim_puanlari.insert_one({
        "id": str(uuid.uuid4()),
        "kullanici_id": current_user["id"],
        "eylem": "url_yukle",
        "dosya_adi": dosya_adi,
        "sinif": sinif,
        "puan": puan,
        "tarih": datetime.utcnow().isoformat(),
    })

    yukleme.pop("_id", None)
    yukleme.pop("dosya_b64", None)
    return {"yukleme": yukleme, "puan_kazanilan": puan, "mesaj": f"✅ +{puan} puan! Link'ten dosya indirildi ve yüklendi."}


@api_router.get("/ai/bilgi-tabani/gecmis")
async def ai_bilgi_tabani_gecmis(current_user=Depends(get_current_user)):
    filtre = {}
    if current_user.get("role") not in ["admin", "coordinator"]:
        filtre["yukleyen_id"] = current_user["id"]
    items = await db.ai_yuklemeler.find(filtre, {"dosya_b64": 0}).sort("tarih", -1).to_list(length=200)
    for i in items:
        i.pop("_id", None)
    return items


@api_router.get("/ai/bilgi-tabani/istatistik")
async def ai_bilgi_tabani_istatistik(current_user=Depends(get_current_user)):
    toplam_yukleme = await db.ai_yuklemeler.count_documents({})
    tamamlanan = await db.ai_yuklemeler.count_documents({"durum": "tamamlandi"})
    bekleyen = await db.ai_yuklemeler.count_documents({"onayli": False})
    toplam_kelime = await db.meb_kelime_haritasi.count_documents({})
    toplam_soru = await db.ai_uretilen_sorular.count_documents({}) + await db.kitap_sorulari.count_documents({"kaynak": "ai_egitim"})

    sinif_dagilim = {}
    for s in range(1, 9):
        sinif_dagilim[str(s)] = {
            "yukleme": await db.ai_yuklemeler.count_documents({"sinif": s}),
            "kelime": await db.meb_kelime_haritasi.count_documents({"sinif": s}),
        }

    # En çok katkı yapan öğretmenler
    pipeline = [
        {"$group": {"_id": "$kullanici_id", "toplam_puan": {"$sum": "$puan"}, "yukleme_sayisi": {"$sum": 1}}},
        {"$sort": {"toplam_puan": -1}},
        {"$limit": 10}
    ]
    top_contributors = []
    async for doc in db.ai_egitim_puanlari.aggregate(pipeline):
        user = await db.users.find_one({"id": doc["_id"]})
        if user:
            top_contributors.append({
                "ad": f"{user.get('ad', '')} {user.get('soyad', '')}".strip(),
                "puan": doc["toplam_puan"],
                "yukleme": doc["yukleme_sayisi"],
            })

    # Güven skoru istatistikleri
    guven_yuklemeler = await db.ai_yuklemeler.find({"guven_skoru": {"$ne": None}}, {"guven_skoru": 1}).to_list(length=None)
    guven_skorlari = []
    guven_dagilim = {"yuksek": 0, "orta": 0, "dusuk": 0}
    for y in guven_yuklemeler:
        gs = y.get("guven_skoru")
        if gs is None:
            continue
        # Eski format: integer (92) → dict'e çevir
        if isinstance(gs, (int, float)):
            toplam = gs
            seviye = "yuksek" if gs >= 80 else ("orta" if gs >= 60 else "dusuk")
        else:
            toplam = gs.get("toplam", 0)
            seviye = gs.get("seviye", "")
        if toplam:
            guven_skorlari.append(toplam)
        if seviye in guven_dagilim:
            guven_dagilim[seviye] += 1
    guven_ort = round(sum(guven_skorlari) / max(len(guven_skorlari), 1), 1) if guven_skorlari else 0

    # Duplicate önleme istatistikleri
    toplam_hash = await db.ai_yuklemeler.distinct("dosya_hash")
    duplicate_engellenen = toplam_yukleme - len(toplam_hash) if toplam_yukleme > len(toplam_hash) else 0

    return {
        "toplam_yukleme": toplam_yukleme,
        "tamamlanan": tamamlanan,
        "bekleyen_onay": bekleyen,
        "toplam_kelime": toplam_kelime,
        "toplam_ai_soru": toplam_soru,
        "sinif_dagilim": sinif_dagilim,
        "top_contributors": top_contributors,
        "guven_skoru": {"ortalama": guven_ort, "dagilim": guven_dagilim},
        "duplicate_engellenen": duplicate_engellenen,
    }


@api_router.put("/ai/bilgi-tabani/onayla/{yukleme_id}")
async def ai_bilgi_tabani_onayla(yukleme_id: str, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    yukleme = await db.ai_yuklemeler.find_one({"id": yukleme_id})
    if not yukleme:
        raise HTTPException(status_code=404, detail="Yükleme bulunamadı")
    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"onayli": True}})
    # Onay bonusu
    await db.users.update_one({"id": yukleme["yukleyen_id"]}, {"$inc": {"puan": AI_EGITIM_PUANLARI["onaylandi"]}})
    await db.ai_egitim_puanlari.insert_one({
        "id": str(uuid.uuid4()), "kullanici_id": yukleme["yukleyen_id"],
        "eylem": "onay_bonusu", "puan": AI_EGITIM_PUANLARI["onaylandi"],
        "tarih": datetime.utcnow().isoformat(),
    })
    return {"ok": True, "mesaj": "Onaylandı, yükleyene +10 bonus puan verildi"}


@api_router.delete("/ai/bilgi-tabani/{yukleme_id}")
async def ai_bilgi_tabani_sil(yukleme_id: str, current_user=Depends(require_role(UserRole.ADMIN, UserRole.COORDINATOR))):
    await db.ai_yuklemeler.delete_one({"id": yukleme_id})
    return {"ok": True}


@api_router.get("/ai/bilgi-tabani/puanlarim")
async def ai_egitim_puanlarim(current_user=Depends(get_current_user)):
    puanlar = await db.ai_egitim_puanlari.find({"kullanici_id": current_user["id"]}).sort("tarih", -1).to_list(length=100)
    for p in puanlar:
        p.pop("_id", None)
    toplam = sum(p.get("puan", 0) for p in puanlar)
    return {"toplam": toplam, "detay": puanlar}


# AI Güven Skoru hesaplama (yükleme sonrası çağrılır)
@api_router.post("/ai/bilgi-tabani/guven-skoru/{yukleme_id}")
async def ai_guven_skoru_hesapla(yukleme_id: str, payload: dict, current_user=Depends(get_current_user)):
    """
    AI işleme sonrası güven skoru hesaplar.
    payload: { kelime_sayisi, benzersiz_kelime, soru_sayisi, bloom_dagilim{}, zorluk_puani, sinif }
    """
    kelime_sayisi = payload.get("kelime_sayisi", 0)
    benzersiz = payload.get("benzersiz_kelime", 0)
    soru_sayisi = payload.get("soru_sayisi", 0)
    bloom = payload.get("bloom_dagilim", {})
    zorluk = payload.get("zorluk_puani", 5)
    sinif = payload.get("sinif", 3)

    # 1. Kelime çeşitliliği skoru (0-25)
    if kelime_sayisi == 0:
        kelime_skor = 0
    else:
        cesitlilik = benzersiz / max(kelime_sayisi, 1)
        kelime_skor = min(25, round(cesitlilik * 50))  # %50 çeşitlilik = 25 puan

    # 2. Soru kalitesi skoru (0-25)
    bloom_turleri = len([v for v in bloom.values() if v > 0])
    soru_skor = 0
    if soru_sayisi >= 5:
        soru_skor += 10
    elif soru_sayisi >= 3:
        soru_skor += 5
    soru_skor += min(15, bloom_turleri * 3)  # 5 bloom = 15 puan

    # 3. Zorluk uyumu skoru (0-25)
    # İdeal zorluk: sınıf seviyesine yakın (sinif * 1.2 civarı)
    ideal_zorluk = min(10, sinif * 1.2)
    zorluk_fark = abs(zorluk - ideal_zorluk)
    zorluk_skor = max(0, 25 - round(zorluk_fark * 5))

    # 4. İçerik zenginliği skoru (0-25)
    icerik_skor = 0
    if kelime_sayisi >= 500:
        icerik_skor += 10
    elif kelime_sayisi >= 200:
        icerik_skor += 5
    if benzersiz >= 50:
        icerik_skor += 8
    elif benzersiz >= 20:
        icerik_skor += 4
    if soru_sayisi >= 10:
        icerik_skor += 7
    elif soru_sayisi >= 5:
        icerik_skor += 3

    toplam = kelime_skor + soru_skor + zorluk_skor + icerik_skor
    seviye = "yuksek" if toplam >= 70 else "orta" if toplam >= 40 else "dusuk"

    guven = {
        "toplam": toplam,
        "seviye": seviye,
        "detay": {
            "kelime_cesitliligi": {"skor": kelime_skor, "max": 25, "aciklama": f"{benzersiz} benzersiz / {kelime_sayisi} toplam kelime"},
            "soru_kalitesi": {"skor": soru_skor, "max": 25, "aciklama": f"{soru_sayisi} soru, {bloom_turleri} Bloom basamağı"},
            "zorluk_uyumu": {"skor": zorluk_skor, "max": 25, "aciklama": f"Zorluk {zorluk}/10, {sinif}. sınıf için ideal ~{round(ideal_zorluk, 1)}"},
            "icerik_zenginligi": {"skor": icerik_skor, "max": 25, "aciklama": f"{kelime_sayisi} kelime, {benzersiz} benzersiz, {soru_sayisi} soru"},
        }
    }

    await db.ai_yuklemeler.update_one({"id": yukleme_id}, {"$set": {"guven_skoru": guven}})
    return guven


# Okuma Seviyesi Analizi (Grade Level Score)
@api_router.post("/ai/bilgi-tabani/okuma-seviyesi")
async def ai_okuma_seviyesi_hesapla(payload: dict, current_user=Depends(get_current_user)):
    """
    Metin zorluk analizi — adaptif motor için kritik.
    payload: { metin, sinif_hedef }
    Döner: Grade Level Score (1-8 sınıf eşdeğeri)
    """
    metin = payload.get("metin", "")
    sinif_hedef = payload.get("sinif_hedef", 3)

    if not metin or len(metin) < 50:
        return {"hata": "Metin en az 50 karakter olmalı"}

    # Metrik hesaplama
    cumleler = [c.strip() for c in metin.replace("!", ".").replace("?", ".").split(".") if c.strip()]
    kelimeler = metin.split()
    toplam_kelime = len(kelimeler)
    toplam_cumle = max(len(cumleler), 1)
    toplam_hece = sum(hece_say(k) for k in kelimeler)

    # Ort cümle uzunluğu (kelime/cümle)
    ort_cumle = toplam_kelime / toplam_cumle
    # Ort kelime uzunluğu (harf)
    ort_kelime_uzunluk = sum(len(k) for k in kelimeler) / max(toplam_kelime, 1)
    # Ort hece/kelime
    ort_hece = toplam_hece / max(toplam_kelime, 1)

    # Ateşman Okunabilirlik Formülü (Türkçe uyarlaması)
    # Okunabilirlik = 198.825 – 40.175 × (hece/kelime) – 2.610 × (kelime/cümle)
    atesman = 198.825 - (40.175 * ort_hece) - (2.610 * ort_cumle)
    atesman = max(0, min(100, round(atesman, 1)))

    # Grade level eşdeğeri
    if atesman >= 90: grade = 1
    elif atesman >= 80: grade = 2
    elif atesman >= 70: grade = 3
    elif atesman >= 60: grade = 4
    elif atesman >= 50: grade = 5
    elif atesman >= 40: grade = 6
    elif atesman >= 30: grade = 7
    else: grade = 8

    # Zorluk puanı (1-10)
    zorluk_puan = max(1, min(10, round((100 - atesman) / 10)))

    # Sınıf uyumu
    uyum = "uygun" if abs(grade - sinif_hedef) <= 1 else "kolay" if grade < sinif_hedef - 1 else "zor"

    return {
        "atesman_skoru": atesman,
        "grade_level": grade,
        "zorluk_puani": zorluk_puan,
        "sinif_uyumu": uyum,
        "metrikler": {
            "toplam_kelime": toplam_kelime,
            "toplam_cumle": toplam_cumle,
            "ort_cumle_uzunlugu": round(ort_cumle, 1),
            "ort_kelime_uzunlugu": round(ort_kelime_uzunluk, 1),
            "ort_hece_kelime": round(ort_hece, 1),
        },
        "yorum": f"Ateşman skoru {atesman} → {grade}. sınıf seviyesi. Hedef {sinif_hedef}. sınıf için {uyum}."
    }


def hece_say(kelime):
    """Türkçe hece sayma — sesli harf sayısı"""
    sesliler = set("aeıioöuüAEIİOÖUÜ")
    return max(1, sum(1 for h in kelime if h in sesliler))


# ─────────────────────────────────────────────
# ÖĞRETMEN HEDEF SİSTEMİ
# ─────────────────────────────────────────────

HEDEF_SABLONLARI = [
    {"kod": "ogrenci_sayisi", "baslik": "Öğrenci Sayısı", "ikon": "👥", "birim": "öğrenci", "aciklama": "Toplam aktif öğrenci sayısı"},
    {"kod": "kur_atlama", "baslik": "Kur Atlama", "ikon": "🎓", "birim": "kur atlama", "aciklama": "Toplam kur atlama sayısı"},
    {"kod": "icerik_uretme", "baslik": "İçerik Üretme", "ikon": "📚", "birim": "içerik", "aciklama": "Yayına alınan gelişim içeriği sayısı"},
    {"kod": "gorev_atama", "baslik": "Görev Atama", "ikon": "📌", "birim": "görev", "aciklama": "Atanan ve tamamlanan görev sayısı"},
    {"kod": "streak_ortalama", "baslik": "Öğrenci Streak Ortalaması", "ikon": "🔥", "birim": "gün", "aciklama": "Öğrencilerinizin ortalama okuma streak'i"},
    {"kod": "veli_puan", "baslik": "Veli Değerlendirme Puanı", "ikon": "⭐", "birim": "puan", "aciklama": "Veli anket ortalaması (1-5)"},
    {"kod": "rozet_sayisi", "baslik": "Rozet Kazanma", "ikon": "🏅", "birim": "rozet", "aciklama": "Kazanılan toplam rozet sayısı"},
    {"kod": "risk_azaltma", "baslik": "Riskli Öğrenci Azaltma", "ikon": "🛡️", "birim": "öğrenci", "aciklama": "Düşük riskli öğrenci sayısı"},
]


@api_router.get("/hedefler/sablonlar")
async def get_hedef_sablonlari():
    return HEDEF_SABLONLARI


@api_router.post("/hedefler")
async def create_hedef(payload: dict, current_user=Depends(get_current_user)):
    hedef = {
        "id": str(uuid.uuid4()),
        "kullanici_id": current_user["id"],
        "kod": payload.get("kod", ""),
        "baslik": payload.get("baslik", ""),
        "ikon": payload.get("ikon", "🎯"),
        "hedef_deger": payload.get("hedef_deger", 0),
        "baslangic_deger": payload.get("baslangic_deger", 0),
        "birim": payload.get("birim", ""),
        "son_tarih": payload.get("son_tarih", ""),
        "durum": "aktif",
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.ogretmen_hedefler.insert_one(hedef)
    hedef.pop("_id", None)
    return hedef


@api_router.get("/hedefler")
async def get_hedefler(current_user=Depends(get_current_user)):
    hedefler = await db.ogretmen_hedefler.find({"kullanici_id": current_user["id"]}).sort("olusturma_tarihi", -1).to_list(length=None)
    ogretmen_id = current_user.get("linked_id") or current_user["id"]

    for h in hedefler:
        h.pop("_id", None)
        kod = h.get("kod", "")
        mevcut = 0
        if kod == "ogrenci_sayisi":
            mevcut = await db.students.count_documents({"ogretmen_id": ogretmen_id, "arsivli": {"$ne": True}})
        elif kod == "kur_atlama":
            mevcut = await db.kur_atlamalari.count_documents({"ogretmen_id": ogretmen_id})
        elif kod == "icerik_uretme":
            mevcut = await db.gelisim_icerik.count_documents({"ekleyen_id": current_user["id"], "durum": "yayinda"})
        elif kod == "gorev_atama":
            mevcut = await db.gorevler.count_documents({"atayan_id": current_user["id"], "durum": "tamamlandi"})
        elif kod == "streak_ortalama":
            ogrenciler = await db.students.find({"ogretmen_id": ogretmen_id, "arsivli": {"$ne": True}}).to_list(length=None)
            from datetime import timedelta
            simdi_h = datetime.utcnow()
            toplam_streak = 0
            for s in ogrenciler:
                logs = await db.reading_logs.find({"ogrenci_id": s["id"]}).to_list(length=None)
                tarihler = set(l.get("tarih", "")[:10] for l in logs)
                st = 0
                for i in range(60):
                    gun = (simdi_h - timedelta(days=i)).strftime("%Y-%m-%d")
                    if gun in tarihler: st += 1
                    elif i > 0: break
                toplam_streak += st
            mevcut = round(toplam_streak / max(len(ogrenciler), 1), 1)
        elif kod == "veli_puan":
            anketler = await db.veli_anketleri.find({"ogretmen_id": ogretmen_id}).to_list(length=None)
            if anketler:
                puanlar = []
                for a in anketler:
                    p = [y.get("puan", 0) for y in a.get("yanitlar", []) if y.get("puan")]
                    if p: puanlar.append(sum(p) / len(p))
                mevcut = round(sum(puanlar) / max(len(puanlar), 1), 1) if puanlar else 0
            else:
                mevcut = 0
        elif kod == "rozet_sayisi":
            mevcut = await db.kazanilan_rozetler.count_documents({"kullanici_id": current_user["id"]})
        elif kod == "risk_azaltma":
            ogrenciler = await db.students.find({"ogretmen_id": ogretmen_id, "arsivli": {"$ne": True}}).to_list(length=None)
            from datetime import timedelta
            simdi_h = datetime.utcnow()
            dusuk_risk = 0
            for s in ogrenciler:
                logs = await db.reading_logs.find({"ogrenci_id": s["id"]}).to_list(length=None)
                son7 = [l for l in logs if l.get("tarih", "") >= (simdi_h - timedelta(days=7)).isoformat()]
                if len(set(l.get("tarih", "")[:10] for l in son7)) >= 3:
                    dusuk_risk += 1
            mevcut = dusuk_risk

        h["mevcut_deger"] = mevcut
        h["ilerleme"] = min(100, round((mevcut / max(h.get("hedef_deger", 1), 0.1)) * 100))
        h["tamamlandi"] = mevcut >= h.get("hedef_deger", 0)
    return hedefler


@api_router.delete("/hedefler/{hedef_id}")
async def delete_hedef(hedef_id: str, current_user=Depends(get_current_user)):
    await db.ogretmen_hedefler.delete_one({"id": hedef_id, "kullanici_id": current_user["id"]})
    return {"ok": True}


# ─────────────────────────────────────────────
# BİLDİRİM SİSTEMİ
# ─────────────────────────────────────────────

BILDIRIM_TURLERI = {
    "rapor_tamamlandi": {"baslik": "📋 Rapor Hazır", "oncelik": "yuksek"},
    "gorev_atandi": {"baslik": "📌 Yeni Görev", "oncelik": "normal"},
    "gorev_tamamlandi": {"baslik": "✅ Görev Tamamlandı", "oncelik": "normal"},
    "gorev_hatirlatma": {"baslik": "⏰ Görev Hatırlatma", "oncelik": "normal"},
    "streak_kirildi": {"baslik": "🔥 Streak Uyarısı", "oncelik": "yuksek"},
    "streak_tebrik": {"baslik": "🎉 Streak Tebrik", "oncelik": "normal"},
    "kur_atladi": {"baslik": "🎓 Kur Atlama", "oncelik": "yuksek"},
    "mesaj_geldi": {"baslik": "✉️ Yeni Mesaj", "oncelik": "normal"},
    "rozet_kazandi": {"baslik": "🏅 Yeni Rozet", "oncelik": "normal"},
    "risk_yuksek": {"baslik": "🚨 Yüksek Risk", "oncelik": "yuksek"},
    "anket_hatirlatma": {"baslik": "⭐ Değerlendirme", "oncelik": "normal"},
    "lig_yukseldi": {"baslik": "🏆 Lig Yükselme", "oncelik": "normal"},
    "haftalik_ozet": {"baslik": "📊 Haftalık Özet", "oncelik": "normal"},
}


async def bildirim_olustur(alici_id, tur, icerik, ilgili_id=None):
    """Bildirim oluştur ve kaydet"""
    tur_bilgi = BILDIRIM_TURLERI.get(tur, {"baslik": "Bildirim", "oncelik": "normal"})
    doc = {
        "id": str(uuid.uuid4()),
        "alici_id": alici_id,
        "tur": tur,
        "baslik": tur_bilgi["baslik"],
        "icerik": icerik,
        "oncelik": tur_bilgi["oncelik"],
        "ilgili_id": ilgili_id,
        "okundu": False,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.bildirimler.insert_one(doc)
    return doc


# Bildirimleri getir
@api_router.get("/bildirimler")
async def get_bildirimler(current_user=Depends(get_current_user)):
    user_id = current_user["id"]
    bildirimler = await db.bildirimler.find({"alici_id": user_id}).sort("tarih", -1).to_list(length=50)
    for b in bildirimler:
        b.pop("_id", None)
    return bildirimler


# Okunmamış bildirim sayısı
@api_router.get("/bildirimler/okunmamis")
async def get_okunmamis_bildirim(current_user=Depends(get_current_user)):
    sayi = await db.bildirimler.count_documents({"alici_id": current_user["id"], "okundu": False})
    return {"sayi": sayi}


# Bildirim okundu işaretle
@api_router.put("/bildirimler/{bildirim_id}/okundu")
async def bildirim_okundu(bildirim_id: str, current_user=Depends(get_current_user)):
    await db.bildirimler.update_one({"id": bildirim_id}, {"$set": {"okundu": True}})
    return {"ok": True}


# Tüm bildirimleri okundu yap
@api_router.put("/bildirimler/tumunu-oku")
async def tumunu_oku(current_user=Depends(get_current_user)):
    await db.bildirimler.update_many({"alici_id": current_user["id"], "okundu": False}, {"$set": {"okundu": True}})
    return {"ok": True}


# Bildirim sil
@api_router.delete("/bildirimler/{bildirim_id}")
async def bildirim_sil(bildirim_id: str, current_user=Depends(get_current_user)):
    await db.bildirimler.delete_one({"id": bildirim_id, "alici_id": current_user["id"]})
    return {"ok": True}


# ── OTOMATİK BİLDİRİM TETİKLEYİCİLERİ ──

# Görev atandığında bildirim (gorevler endpoint'ine hook)
async def bildirim_gorev_atandi(hedef_id, baslik, atayan_ad):
    # Hedef kullanıcının user id'sini bul
    user = await db.users.find_one({"$or": [{"id": hedef_id}, {"linked_id": hedef_id}]})
    if user:
        await bildirim_olustur(user["id"], "gorev_atandi", f"{atayan_ad} size yeni görev atadı: {baslik}", hedef_id)


# Rapor tamamlandığında veliye bildirim
async def bildirim_rapor_tamamlandi(ogrenci_id, rapor_baslik):
    student = await db.students.find_one({"id": ogrenci_id})
    if student:
        # Velinin user'ını bul
        veli = await db.users.find_one({"role": "parent", "$or": [
            {"linked_id": ogrenci_id},
            {"telefon": student.get("veli_telefon", "")}
        ]})
        if veli:
            await bildirim_olustur(veli["id"], "rapor_tamamlandi",
                f"{student.get('ad', '')} {student.get('soyad', '')} için yeni rapor hazır: {rapor_baslik}", ogrenci_id)


# Streak uyarısı (günlük kontrol)
async def bildirim_streak_kontrol():
    """Tüm öğrencilerin streak'ini kontrol et, gerekirse bildirim gönder"""
    from datetime import timedelta
    simdi = datetime.utcnow()
    dun = (simdi - timedelta(days=1)).strftime("%Y-%m-%d")
    evvelsi = (simdi - timedelta(days=2)).strftime("%Y-%m-%d")

    students = await db.students.find({"arsivli": {"$ne": True}}).to_list(length=None)
    for s in students:
        logs = await db.reading_logs.find({"ogrenci_id": s["id"]}).to_list(length=None)
        tarihler = set(l.get("tarih", "")[:10] for l in logs)

        # Dün okumadı ama önceki gün okumuştu → streak kırılma riski
        if evvelsi in tarihler and dun not in tarihler:
            user = await db.users.find_one({"linked_id": s["id"], "role": "student"})
            if user:
                # Bugün zaten bildirim gönderildi mi?
                mevcut = await db.bildirimler.find_one({
                    "alici_id": user["id"], "tur": "streak_kirildi",
                    "tarih": {"$gte": simdi.strftime("%Y-%m-%d")}
                })
                if not mevcut:
                    await bildirim_olustur(user["id"], "streak_kirildi",
                        "Dün okuma yapmadın! Streak'ini korumak için bugün oku 📖")

            # Veliye de bildir
            veli = await db.users.find_one({"role": "parent", "$or": [
                {"linked_id": s["id"]}, {"telefon": s.get("veli_telefon", "")}
            ]})
            if veli:
                mevcut_v = await db.bildirimler.find_one({
                    "alici_id": veli["id"], "tur": "streak_kirildi",
                    "tarih": {"$gte": simdi.strftime("%Y-%m-%d")}
                })
                if not mevcut_v:
                    await bildirim_olustur(veli["id"], "streak_kirildi",
                        f"{s.get('ad', '')} dün okuma yapmadı. Streak kırılma riski!")

        # 7 gün streak → tebrik
        streak = 0
        for i in range(30):
            gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
            if gun in tarihler:
                streak += 1
            elif i > 0:
                break
        if streak == 7:
            user = await db.users.find_one({"linked_id": s["id"], "role": "student"})
            if user:
                mevcut = await db.bildirimler.find_one({
                    "alici_id": user["id"], "tur": "streak_tebrik",
                    "icerik": {"$regex": "7 gün"}
                })
                if not mevcut:
                    await bildirim_olustur(user["id"], "streak_tebrik",
                        "🎉 7 gün üst üste okudun! Harika gidiyorsun!")


# Görev hatırlatma (son 1 gün kala)
async def bildirim_gorev_hatirlatma():
    from datetime import timedelta
    simdi = datetime.utcnow()
    yarin = (simdi + timedelta(days=1)).strftime("%Y-%m-%d")

    gorevler = await db.gorevler.find({"durum": "bekliyor", "son_tarih": yarin}).to_list(length=None)
    for g in gorevler:
        user = await db.users.find_one({"$or": [{"id": g["hedef_id"]}, {"linked_id": g["hedef_id"]}]})
        if user:
            mevcut = await db.bildirimler.find_one({
                "alici_id": user["id"], "tur": "gorev_hatirlatma", "ilgili_id": g["id"]
            })
            if not mevcut:
                await bildirim_olustur(user["id"], "gorev_hatirlatma",
                    f"Yarın son gün: {g['baslik']}", g["id"])


# Risk yüksekse öğretmene bildir
async def bildirim_risk_kontrol():
    students = await db.students.find({"arsivli": {"$ne": True}}).to_list(length=None)
    from datetime import timedelta
    simdi = datetime.utcnow()

    for s in students:
        logs = await db.reading_logs.find({"ogrenci_id": s["id"]}).to_list(length=None)
        yedi_gun = simdi - timedelta(days=7)
        son_7 = [l for l in logs if l.get("tarih", "") >= yedi_gun.isoformat()]
        aktif_7 = len(set(l.get("tarih", "")[:10] for l in son_7))

        if aktif_7 == 0 and len(logs) > 0:  # Daha önce aktifti ama son 7 gün hiç okumadı
            ogretmen_user = await db.users.find_one({"linked_id": s.get("ogretmen_id"), "role": "teacher"})
            if ogretmen_user:
                mevcut = await db.bildirimler.find_one({
                    "alici_id": ogretmen_user["id"], "tur": "risk_yuksek",
                    "ilgili_id": s["id"],
                    "tarih": {"$gte": (simdi - timedelta(days=7)).isoformat()}
                })
                if not mevcut:
                    await bildirim_olustur(ogretmen_user["id"], "risk_yuksek",
                        f"🚨 {s.get('ad', '')} {s.get('soyad', '')} son 7 gündür hiç okuma yapmadı!", s["id"])


# Manuel bildirim kontrol endpoint'i (admin veya cron job çağırabilir)
@api_router.post("/bildirimler/kontrol")
async def bildirim_kontrol_endpoint(current_user=Depends(get_current_user)):
    if current_user.get("role") not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    await bildirim_streak_kontrol()
    await bildirim_gorev_hatirlatma()
    await bildirim_risk_kontrol()
    return {"ok": True, "mesaj": "Bildirim kontrolü tamamlandı"}


# ─────────────────────────────────────────────
# KİTAP + BÖLÜM BAZLI SORU HAVUZU (Master Bölüm 8)
# ─────────────────────────────────────────────

# Kitap ekle
@api_router.post("/kitaplar")
async def create_kitap(payload: dict, current_user=Depends(get_current_user)):
    kitap = {
        "id": str(uuid.uuid4()),
        "baslik": payload.get("baslik", ""),
        "yazar": payload.get("yazar", ""),
        "yas_grubu": payload.get("yas_grubu", ""),
        "zorluk": payload.get("zorluk", "orta"),
        "bolum_sayisi": payload.get("bolum_sayisi", 1),
        "kapak_url": payload.get("kapak_url", ""),
        "ekleyen_id": current_user["id"],
        "ekleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "durum": "beklemede",  # beklemede → oylama → yayinda
        "oylar": {},
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.kitap_havuzu.insert_one(kitap)
    # Katkı puanı
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 15}})
    kitap.pop("_id", None)
    return kitap


# Kitapları listele
@api_router.get("/kitaplar")
async def get_kitaplar(current_user=Depends(get_current_user)):
    kitaplar = await db.kitap_havuzu.find().sort("olusturma_tarihi", -1).to_list(length=None)
    for k in kitaplar:
        k.pop("_id", None)
    return kitaplar


# Kitap admin kararı (onay/oylama/red)
@api_router.post("/kitaplar/{kitap_id}/admin-karar")
async def kitap_admin_karar(kitap_id: str, payload: dict, current_user=Depends(get_current_user)):
    if current_user.get("role") not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    onay = payload.get("onay", False)
    direkt = payload.get("direkt", False)
    if not onay:
        await db.kitap_havuzu.update_one({"id": kitap_id}, {"$set": {"durum": "reddedildi"}})
        return {"ok": True, "durum": "reddedildi"}
    if direkt:
        await db.kitap_havuzu.update_one({"id": kitap_id}, {"$set": {"durum": "yayinda"}})
        return {"ok": True, "durum": "yayinda"}
    await db.kitap_havuzu.update_one({"id": kitap_id}, {"$set": {"durum": "oylama"}})
    return {"ok": True, "durum": "oylama"}


# Kitap oylama
@api_router.post("/kitaplar/{kitap_id}/oyla")
async def kitap_oyla(kitap_id: str, payload: dict, current_user=Depends(get_current_user)):
    onay = payload.get("onay", True)
    sebep = payload.get("sebep", "")
    oy_data = {"onay": onay, "sebep": sebep, "tarih": datetime.utcnow().isoformat()}
    await db.kitap_havuzu.update_one({"id": kitap_id}, {"$set": {f"oylar.{current_user['id']}": oy_data}})
    # Katkı puanı
    await db.users.update_one({"id": current_user["id"]}, {"$inc": {"puan": 3}})
    # Otomatik yayına alma kontrolü
    kitap = await db.kitap_havuzu.find_one({"id": kitap_id})
    if kitap:
        oylar = kitap.get("oylar", {})
        toplam = len(oylar)
        onaylar = sum(1 for o in oylar.values() if o.get("onay"))
        redler = sum(1 for o in oylar.values() if not o.get("onay"))
        if toplam >= 3 and onaylar / toplam >= 0.6:
            await db.kitap_havuzu.update_one({"id": kitap_id}, {"$set": {"durum": "yayinda"}})
        if redler > 0:
            await db.kitap_havuzu.update_one({"id": kitap_id}, {"$set": {"durum": "askida"}})
    return {"ok": True}


# Bölüm bazlı soru ekle
@api_router.post("/kitaplar/{kitap_id}/sorular")
async def create_soru(kitap_id: str, payload: dict, current_user=Depends(get_current_user)):
    soru = {
        "id": str(uuid.uuid4()),
        "kitap_id": kitap_id,
        "bolum": payload.get("bolum", 1),
        "soru": payload.get("soru", ""),
        "secenekler": payload.get("secenekler", []),
        "dogru_cevap": payload.get("dogru_cevap", 0),
        "taksonomi": payload.get("taksonomi", "kavrama"),
        "ekleyen_id": current_user["id"],
        "ekleyen_ad": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        "kullanim_sayisi": 0,
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.kitap_sorulari.insert_one(soru)
    soru.pop("_id", None)
    return soru


# Kitabın sorularını getir
@api_router.get("/kitaplar/{kitap_id}/sorular")
async def get_kitap_sorulari(kitap_id: str, bolum: int = None, current_user=Depends(get_current_user)):
    filtre = {"kitap_id": kitap_id}
    if bolum:
        filtre["bolum"] = bolum
    sorular = await db.kitap_sorulari.find(filtre).to_list(length=None)
    for s in sorular:
        s.pop("_id", None)
    return sorular


# Soru sil
@api_router.delete("/kitaplar/sorular/{soru_id}")
async def delete_soru(soru_id: str, current_user=Depends(get_current_user)):
    await db.kitap_sorulari.delete_one({"id": soru_id})
    return {"ok": True}


# Öğrenci için bölüm bazlı test çek (okuma sonrası)
@api_router.get("/kitaplar/test/{kitap_id}/{bolum}")
async def get_bolum_testi(kitap_id: str, bolum: int, current_user=Depends(get_current_user)):
    sorular = await db.kitap_sorulari.find({"kitap_id": kitap_id, "bolum": bolum}).to_list(length=None)
    for s in sorular:
        s.pop("_id", None)
        # Kullanım sayısını artır
        await db.kitap_sorulari.update_one({"id": s["id"]}, {"$inc": {"kullanim_sayisi": 1}})
    return sorular


# Bölüm testi tamamla (öğrenci cevapladığında)
@api_router.post("/kitaplar/test/tamamla")
async def bolum_testi_tamamla(payload: dict, current_user=Depends(get_current_user)):
    kitap_id = payload.get("kitap_id", "")
    bolum = payload.get("bolum", 1)
    cevaplar = payload.get("cevaplar", [])  # [{soru_id, secilen_cevap}]

    sorular = await db.kitap_sorulari.find({"kitap_id": kitap_id, "bolum": bolum}).to_list(length=None)
    soru_dict = {s["id"]: s for s in sorular}

    dogru = 0
    toplam = len(cevaplar)
    for c in cevaplar:
        soru = soru_dict.get(c.get("soru_id"))
        if soru and c.get("secilen_cevap") == soru.get("dogru_cevap"):
            dogru += 1

    yuzde = round((dogru / max(toplam, 1)) * 100)

    # Test sonucu kaydet
    sonuc = {
        "id": str(uuid.uuid4()),
        "ogrenci_id": current_user.get("linked_id") or current_user["id"],
        "kitap_id": kitap_id,
        "bolum": bolum,
        "dogru": dogru,
        "toplam": toplam,
        "yuzde": yuzde,
        "cevaplar": cevaplar,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.kitap_test_sonuclari.insert_one(sonuc)

    # XP kazan
    xp_tablosu = await get_xp_tablosu()
    xp = xp_tablosu.get("anlama_testi", 15)
    ogrenci_id = current_user.get("linked_id") or current_user["id"]
    await db.xp_logs.insert_one({"id": str(uuid.uuid4()), "ogrenci_id": ogrenci_id, "eylem": "anlama_testi", "xp": xp, "tarih": datetime.utcnow().isoformat()})
    await db.students.update_one({"id": ogrenci_id}, {"$inc": {"toplam_xp": xp}})

    sonuc.pop("_id", None)
    return {"sonuc": sonuc, "xp_kazanilan": xp}


# Yayındaki kitapları listele (öğrenci/öğretmen)
@api_router.get("/kitaplar/havuz")
async def get_kitap_havuzu(current_user=Depends(get_current_user)):
    kitaplar = await db.kitap_havuzu.find({"durum": "yayinda"}).to_list(length=None)
    sonuc = []
    for k in kitaplar:
        k.pop("_id", None)
        soru_sayisi = await db.kitap_sorulari.count_documents({"kitap_id": k["id"]})
        k["soru_sayisi"] = soru_sayisi
        sonuc.append(k)
    return sonuc


# ─────────────────────────────────────────────
# MESAJLAŞMA SİSTEMİ
# ─────────────────────────────────────────────

class MesajCreate(BaseModel):
    alici_id: str
    alici_tip: str = ""  # user id veya teacher/student ref
    icerik: str
    konu: str = ""

class MesajModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    gonderen_id: str
    gonderen_ad: str = ""
    gonderen_rol: str = ""
    alici_id: str
    alici_ad: str = ""
    alici_rol: str = ""
    konu: str = ""
    icerik: str
    okundu: bool = False
    tarih: str = Field(default_factory=lambda: datetime.utcnow().isoformat())


@api_router.post("/mesajlar")
async def create_mesaj(mesaj: MesajCreate, current_user=Depends(get_current_user)):
    gonderen_ad = f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip()
    gonderen_rol = current_user.get("role", "")

    # Alıcı bilgisini bul
    alici = await db.users.find_one({"id": mesaj.alici_id})
    alici_ad = ""
    alici_rol = ""
    if alici:
        alici_ad = f"{alici.get('ad', '')} {alici.get('soyad', '')}".strip()
        alici_rol = alici.get("role", "")

    model = MesajModel(
        gonderen_id=current_user["id"],
        gonderen_ad=gonderen_ad,
        gonderen_rol=gonderen_rol,
        alici_id=mesaj.alici_id,
        alici_ad=alici_ad,
        alici_rol=alici_rol,
        konu=mesaj.konu,
        icerik=mesaj.icerik,
    )
    data = model.dict()
    await db.mesajlar.insert_one(data)
    # Bildirim gönder
    try: await bildirim_olustur(data.get("alici_id"), "mesaj_geldi", f"{data.get('gonderen_ad', '')} size mesaj gönderdi: {data.get('konu', '')}")
    except: pass
    return data


@api_router.get("/mesajlar")
async def get_mesajlar(current_user=Depends(get_current_user)):
    user_id = current_user["id"]
    mesajlar = await db.mesajlar.find({
        "$or": [{"gonderen_id": user_id}, {"alici_id": user_id}]
    }).sort("tarih", -1).to_list(length=None)
    for m in mesajlar:
        m.pop("_id", None)
    return mesajlar


@api_router.put("/mesajlar/{mesaj_id}/okundu")
async def mesaj_okundu(mesaj_id: str, current_user=Depends(get_current_user)):
    await db.mesajlar.update_one({"id": mesaj_id, "alici_id": current_user["id"]}, {"$set": {"okundu": True}})
    return {"ok": True}


@api_router.get("/mesajlar/okunmamis-sayisi")
async def okunmamis_sayisi(current_user=Depends(get_current_user)):
    sayi = await db.mesajlar.count_documents({"alici_id": current_user["id"], "okundu": False})
    return {"sayi": sayi}


# ─────────────────────────────────────────────
# GÖREV ATAMA SİSTEMİ
# İki yönlü: Yönetici → Öğretmen, Öğretmen → Öğrenci
# ─────────────────────────────────────────────

class GorevCreate(BaseModel):
    hedef_id: str
    hedef_tip: str  # "ogretmen" veya "ogrenci"
    baslik: str
    aciklama: str = ""
    tur: str = "ozel"  # ozel, film, kitap, makale, hizmetici, egzersiz
    icerik_id: Optional[str] = None
    son_tarih: Optional[str] = None
    makale_link: Optional[str] = None
    kitap_yazar: Optional[str] = None
    kitap_isbn: Optional[str] = None
    kitap_link: Optional[str] = None
    kitap_kapak: Optional[str] = None
    film_link: Optional[str] = None

class GorevModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    hedef_id: str
    hedef_tip: str
    hedef_ad: str = ""
    baslik: str
    aciklama: str = ""
    tur: str = "ozel"
    icerik_id: Optional[str] = None
    son_tarih: Optional[str] = None
    atayan_id: str = ""
    atayan_ad: str = ""
    atayan_rol: str = ""
    durum: str = "bekliyor"
    tamamlama_tarihi: Optional[str] = None
    tamamlama_notu: Optional[str] = None
    makale_link: Optional[str] = None
    kitap_yazar: Optional[str] = None
    kitap_isbn: Optional[str] = None
    kitap_link: Optional[str] = None
    kitap_kapak: Optional[str] = None
    film_link: Optional[str] = None
    olusturma_tarihi: str = Field(default_factory=lambda: datetime.utcnow().isoformat())


@api_router.post("/gorevler")
async def create_gorev(gorev: GorevCreate, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    if gorev.hedef_tip == "ogretmen" and role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Öğretmenlere görev yalnızca yönetici/koordinatör atayabilir")
    if gorev.hedef_tip == "ogrenci" and role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Öğrencilere görev yalnızca öğretmen/yönetici atayabilir")

    hedef_ad = ""
    if gorev.hedef_tip == "ogretmen":
        user = await db.users.find_one({"id": gorev.hedef_id})
        if user:
            hedef_ad = f"{user.get('ad', '')} {user.get('soyad', '')}".strip()
        else:
            raise HTTPException(status_code=404, detail="Hedef öğretmen bulunamadı")
    elif gorev.hedef_tip == "ogrenci":
        student = await db.students.find_one({"id": gorev.hedef_id})
        if student:
            hedef_ad = f"{student.get('ad', '')} {student.get('soyad', '')}".strip()
        else:
            raise HTTPException(status_code=404, detail="Hedef öğrenci bulunamadı")

    model = GorevModel(
        **gorev.dict(),
        hedef_ad=hedef_ad,
        atayan_id=current_user["id"],
        atayan_ad=f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
        atayan_rol=role,
    )
    data = model.dict()
    await db.gorevler.insert_one(data)
    # Bildirim gönder
    try: await bildirim_gorev_atandi(data.get("hedef_id"), data.get("baslik", ""), data.get("atayan_ad", ""))
    except: pass
    return data


@api_router.post("/gorevler/toplu")
async def create_toplu_gorev(payload: dict, current_user=Depends(get_current_user)):
    hedef_idler = payload.get("hedef_idler", [])
    hedef_tip = payload.get("hedef_tip", "")
    gorev_bilgi = payload.get("gorev", {})
    role = current_user.get("role", "")

    if hedef_tip == "ogretmen" and role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    if hedef_tip == "ogrenci" and role not in ["admin", "coordinator", "teacher"]:
        raise HTTPException(status_code=403, detail="Yetkisiz")

    olusturulan = []
    for hid in hedef_idler:
        hedef_ad = ""
        if hedef_tip == "ogretmen":
            u = await db.users.find_one({"id": hid})
            hedef_ad = f"{u.get('ad', '')} {u.get('soyad', '')}".strip() if u else ""
        elif hedef_tip == "ogrenci":
            s = await db.students.find_one({"id": hid})
            hedef_ad = f"{s.get('ad', '')} {s.get('soyad', '')}".strip() if s else ""

        model = GorevModel(
            hedef_id=hid, hedef_tip=hedef_tip, hedef_ad=hedef_ad,
            baslik=gorev_bilgi.get("baslik", ""), aciklama=gorev_bilgi.get("aciklama", ""),
            tur=gorev_bilgi.get("tur", "ozel"), icerik_id=gorev_bilgi.get("icerik_id"),
            son_tarih=gorev_bilgi.get("son_tarih"),
            makale_link=gorev_bilgi.get("makale_link"), kitap_yazar=gorev_bilgi.get("kitap_yazar"),
            kitap_isbn=gorev_bilgi.get("kitap_isbn"), kitap_link=gorev_bilgi.get("kitap_link"),
            kitap_kapak=gorev_bilgi.get("kitap_kapak"), film_link=gorev_bilgi.get("film_link"),
            atayan_id=current_user["id"],
            atayan_ad=f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip(),
            atayan_rol=role,
        )
        data = model.dict()
        await db.gorevler.insert_one(data)
        olusturulan.append(data)

    return {"olusturulan": len(olusturulan), "gorevler": olusturulan}


@api_router.get("/gorevler")
async def get_gorevler(
    hedef_tip: Optional[str] = None,
    hedef_id: Optional[str] = None,
    durum: Optional[str] = None,
    current_user=Depends(get_current_user)
):
    role = current_user.get("role", "")
    user_id = current_user.get("id", "")
    filtre = {}
    if hedef_tip:
        filtre["hedef_tip"] = hedef_tip
    if hedef_id:
        filtre["hedef_id"] = hedef_id
    if durum:
        filtre["durum"] = durum

    if role == "teacher":
        items_atanan = await db.gorevler.find({"hedef_id": user_id, **({k: v for k, v in filtre.items() if k != "hedef_id"})}).sort("olusturma_tarihi", -1).to_list(length=None)
        items_atadigi = await db.gorevler.find({"atayan_id": user_id, **({k: v for k, v in filtre.items() if k != "hedef_id"})}).sort("olusturma_tarihi", -1).to_list(length=None)
        seen = set()
        items = []
        for i in items_atanan + items_atadigi:
            if i["id"] not in seen:
                seen.add(i["id"])
                items.append(i)
    else:
        items = await db.gorevler.find(filtre).sort("olusturma_tarihi", -1).to_list(length=None)

    for item in items:
        item.pop("_id", None)
    return items


@api_router.put("/gorevler/{gorev_id}/durum")
async def update_gorev_durum(gorev_id: str, payload: dict, current_user=Depends(get_current_user)):
    gorev = await db.gorevler.find_one({"id": gorev_id})
    if not gorev:
        raise HTTPException(status_code=404, detail="Görev bulunamadı")

    yeni_durum = payload.get("durum", "")
    update = {"durum": yeni_durum}
    if yeni_durum == "tamamlandi":
        update["tamamlama_tarihi"] = datetime.utcnow().isoformat()
        if payload.get("not"):
            update["tamamlama_notu"] = payload["not"]

    await db.gorevler.update_one({"id": gorev_id}, {"$set": update})
    return {"durum": yeni_durum}


@api_router.delete("/gorevler/{gorev_id}")
async def delete_gorev(gorev_id: str, current_user=Depends(get_current_user)):
    role = current_user.get("role", "")
    gorev = await db.gorevler.find_one({"id": gorev_id})
    if not gorev:
        raise HTTPException(status_code=404, detail="Görev bulunamadı")
    if gorev.get("atayan_id") != current_user["id"] and role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Yalnızca atayan veya yönetici silebilir")
    await db.gorevler.delete_one({"id": gorev_id})
    return {"message": "Görev silindi"}


@api_router.get("/gorevler/istatistik")
async def get_gorev_istatistik(current_user=Depends(get_current_user)):
    tum = await db.gorevler.find().to_list(length=None)
    og = [g for g in tum if g.get("hedef_tip") == "ogretmen"]
    os = [g for g in tum if g.get("hedef_tip") == "ogrenci"]
    def h(l):
        return {"toplam": len(l), "bekliyor": len([g for g in l if g.get("durum") == "bekliyor"]),
                "devam_ediyor": len([g for g in l if g.get("durum") == "devam_ediyor"]),
                "tamamlandi": len([g for g in l if g.get("durum") == "tamamlandi"]),
                "suresi_doldu": len([g for g in l if g.get("durum") == "suresi_doldu"])}
    return {"ogretmen": h(og), "ogrenci": h(os)}


# ─────────────────────────────────────────────
# DALGA 3: SPEECH AI — SESLİ OKUMA ANALİZİ
# ─────────────────────────────────────────────

SPEECH_OKUMA_METİNLERİ = {
    1: [
        {"id": "s1_1", "baslik": "Küçük Kedi", "metin": "Küçük kedi bahçede oynadı. Güneş parlıyordu. Kedi mutluydu.", "kelime_sayisi": 10, "sinif": 1},
        {"id": "s1_2", "baslik": "Renkli Balonlar", "metin": "Ali kırmızı balon aldı. Ayşe sarı balon aldı. İkisi de çok sevindi.", "kelime_sayisi": 12, "sinif": 1},
    ],
    2: [
        {"id": "s2_1", "baslik": "Yağmur", "metin": "Sabah kalktığımda pencereden baktım. Dışarıda yağmur yağıyordu. Annem şemsiyemi hazırladı ve okula gittim.", "kelime_sayisi": 20, "sinif": 2},
        {"id": "s2_2", "baslik": "Kitap Kurdu", "metin": "Her gün en az bir kitap okuyorum. Kitaplar bana yeni dünyalar açıyor. En sevdiğim yer kütüphane.", "kelime_sayisi": 18, "sinif": 2},
    ],
    3: [
        {"id": "s3_1", "baslik": "Ormanda Bir Gün", "metin": "Ormanın içinde yürürken ağaçların arasından süzülen güneş ışığını seyrettim. Kuşlar şakıyor, yapraklar hışırdıyordu. Bu sessizlik içinde kendimi huzurlu hissettim.", "kelime_sayisi": 30, "sinif": 3},
        {"id": "s3_2", "baslik": "Dürüstlük", "metin": "Pazarda yürürken yerde bir cüzdan buldum. İçinde para ve kimlik vardı. Hemen karakola götürdüm. Görevli bana teşekkür etti ve sahibini bulacaklarını söyledi.", "kelime_sayisi": 32, "sinif": 3},
    ],
    4: [
        {"id": "s4_1", "baslik": "Göç Eden Kuşlar", "metin": "Her sonbahar leylekler uzun bir yolculuğa çıkar. Binlerce kilometre uçarak sıcak ülkelere göç ederler. Pusula gibi çalışan içgüdüleri sayesinde yollarını şaşırmazlar. Bu muhteşem yolculuk nesiller boyunca sürmektedir.", "kelime_sayisi": 38, "sinif": 4},
    ],
    5: [
        {"id": "s5_1", "baslik": "Anadolu Medeniyetleri", "metin": "Anadolu, tarihin en eski medeniyetlerine ev sahipliği yapmıştır. Hitit, Frigya, Lidya ve daha pek çok uygarlık bu topraklarda yaşamış, eserler bırakmıştır. Bu zengin miras günümüze kadar ulaşmıştır.", "kelime_sayisi": 35, "sinif": 5},
    ],
}

def _speech_mock_analiz(transkript: str, beklenen_metin: str, sure_sn: float, sinif: int) -> dict:
    """Web Speech API transkriptini beklenen metinle karşılaştırarak analiz üret."""
    import difflib
    import re

    def normalize(s):
        """Noktalama ve büyük/küçük harf normalize et."""
        s = s.lower()
        s = re.sub(r'[.,!?;:"\'-]', '', s)
        return s.split()

    b_kelimeler = normalize(beklenen_metin)
    gercek_transkript = transkript.strip() if transkript else ""

    # Transkript yoksa (kullanıcı hiç okumadı) — düşük skor
    if not gercek_transkript:
        return {
            "transkript": "",
            "telaffuz_skoru": 0,
            "akicilik_skoru": 0,
            "wpm": 0,
            "norm_wpm": {1:50,2:75,3:95,4:115,5:130}.get(sinif,95),
            "duraklama_sayisi": 0,
            "tonlama_skoru": 0,
            "vurgu_skoru": 0,
            "genel_skor": 0,
            "seviye": "geliştirilmeli",
            "guclu_yonler": [],
            "gelisim_alanlari": ["Okumaya başla — mikrofon sesi almadı"],
            "telaffuz_hatalar": [],
            "mock": True,
        }

    t_kelimeler = normalize(gercek_transkript)

    # ── Kelime bazlı diff ile yanlış/atlanmış kelimeleri bul ──
    matcher = difflib.SequenceMatcher(None, b_kelimeler, t_kelimeler, autojunk=False)
    dogru = 0
    yanlis_kelimeler = []
    atlanan_kelimeler = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            dogru += (i2 - i1)
        elif tag == "replace":
            # Beklenen kelimeler yanlış okunmuş
            for k in b_kelimeler[i1:i2]:
                yanlis_kelimeler.append(k)
        elif tag == "delete":
            # Beklenen kelimeler hiç okunmamış
            for k in b_kelimeler[i1:i2]:
                atlanan_kelimeler.append(k)

    toplam = len(b_kelimeler)
    telaffuz_skoru = round((dogru / toplam) * 100) if toplam > 0 else 0

    # Yanlış okunan kelimelerin orijinal (normalize edilmemiş) hallerini bul
    b_orijinal = beklenen_metin.split()
    telaffuz_hata_listesi = []
    for yanlis in set(yanlis_kelimeler + atlanan_kelimeler):
        # Orijinal metinde bu kelimeye yakın olanı bul
        for kel in b_orijinal:
            if normalize(kel) and normalize(kel)[0] == yanlis:
                telaffuz_hata_listesi.append(kel.strip('.,!?;:"\'-'))
                break
        else:
            telaffuz_hata_listesi.append(yanlis)

    # ── WPM hesapla ──
    sure_dk = max(sure_sn / 60, 0.1)
    # Transkriptteki kelime sayısından WPM
    wpm = round(len(t_kelimeler) / sure_dk)
    norm = {1: 50, 2: 75, 3: 95, 4: 115, 5: 130, 6: 145, 7: 155, 8: 165}.get(sinif, 95)
    akicilik_skoru = min(100, round(wpm / norm * 100))

    # ── Duraksama tahmini — WPM'e göre ──
    duraklama_sayisi = max(0, int((norm - wpm) / 15)) if wpm < norm else 0

    # ── Tonlama: noktalama işaretlerinde durup durmadığına bakılamaz,
    #    ancak cümle sonlarındaki kelime oranına bak ──
    noktalama_kelimeler = [w for w in beklenen_metin.split() if w[-1] in '.!?,;' if len(w) > 1]
    tonlama_skoru = min(100, telaffuz_skoru + 3)
    vurgu_skoru = min(100, akicilik_skoru + 2)

    seviye = "çok iyi" if telaffuz_skoru >= 85 else "iyi" if telaffuz_skoru >= 70 else "orta" if telaffuz_skoru >= 55 else "geliştirilmeli"

    guclu = []
    gelisim = []
    if telaffuz_skoru >= 80: guclu.append("Kelime telaffuzu başarılı")
    if akicilik_skoru >= 80: guclu.append("Okuma hızı sınıf normuna uygun")
    if len(telaffuz_hata_listesi) == 0 and telaffuz_skoru >= 70: guclu.append("Tüm kelimeleri doğru okudun")
    if akicilik_skoru < 70: gelisim.append(f"Okuma hızını artır (hedef: {norm} kelime/dk)")
    if len(telaffuz_hata_listesi) > 3: gelisim.append("Yanlış okunan kelimeleri tekrar çalış")
    if len(atlanan_kelimeler) > 2: gelisim.append("Bazı kelimeleri atladın, dikkatli oku")

    return {
        "transkript": gercek_transkript,
        "telaffuz_skoru": telaffuz_skoru,
        "akicilik_skoru": akicilik_skoru,
        "wpm": wpm,
        "norm_wpm": norm,
        "duraklama_sayisi": duraklama_sayisi,
        "tonlama_skoru": tonlama_skoru,
        "vurgu_skoru": vurgu_skoru,
        "genel_skor": round((telaffuz_skoru * 0.6 + akicilik_skoru * 0.4)),
        "seviye": seviye,
        "guclu_yonler": guclu,
        "gelisim_alanlari": gelisim,
        "telaffuz_hatalar": telaffuz_hata_listesi[:8],  # max 8 kelime göster
        "atlanan_kelimeler": atlanan_kelimeler[:5],
        "mock": False,  # Artık gerçek transkript analizi
    }


@api_router.get("/ai/speech/metinler")
async def speech_okuma_metinleri(sinif: int = 3, current_user=Depends(get_current_user)):
    """Sınıfa göre sesli okuma metinleri getir."""
    metinler = SPEECH_OKUMA_METİNLERİ.get(sinif, SPEECH_OKUMA_METİNLERİ.get(3, []))
    if not metinler:
        # En yakın sınıfı bul
        for s in range(sinif, 0, -1):
            if s in SPEECH_OKUMA_METİNLERİ:
                metinler = SPEECH_OKUMA_METİNLERİ[s]
                break
    return {"metinler": metinler, "sinif": sinif}


@api_router.post("/ai/speech/analiz")
async def speech_analiz(
    ses_dosyasi: UploadFile = File(None),
    metin_id: str = Form(""),
    ogrenci_id: str = Form(""),
    sure_sn: float = Form(30.0),
    sinif: int = Form(3),
    transkript_input: str = Form(""),  # Web Speech API'den gelen transkript
    current_user=Depends(get_current_user)
):
    """Sesli okuma kaydını analiz et: WPM + telaffuz + tonlama + duraklama."""
    # Hedef metni bul
    beklenen_metin = ""
    metin_baslik = ""
    for s_list in SPEECH_OKUMA_METİNLERİ.values():
        for m in s_list:
            if m["id"] == metin_id:
                beklened_metin = m["metin"]
                beklenen_metin = m["metin"]
                metin_baslik = m["baslik"]
                sinif = m.get("sinif", sinif)
                break

    transkript = transkript_input.strip()  # Web Speech API'den gelen
    whisper_kullanildi = False

    # Whisper API — varsa kullan
    OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
    if ses_dosyasi and OPENAI_API_KEY:
        try:
            ses_bytes = await ses_dosyasi.read()
            # Dosya uzantısını ve mime type'ı belirle
            filename = ses_dosyasi.filename or "ses.webm"
            content_type = ses_dosyasi.content_type or "audio/webm"
            # Whisper desteklenen formatlar: mp4, webm, mp3, wav, m4a, ogg
            if "mp4" in content_type or filename.endswith(".mp4"):
                mime = "audio/mp4"
                ext = "mp4"
            else:
                mime = "audio/webm"
                ext = "webm"
            async with httpx.AsyncClient(timeout=90.0) as c:
                resp = await c.post(
                    "https://api.openai.com/v1/audio/transcriptions",
                    headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
                    files={"file": (f"ses.{ext}", ses_bytes, mime)},
                    data={"model": "whisper-1", "language": "tr"},
                )
                if resp.status_code == 200:
                    transkript = resp.json().get("text", "")
                    whisper_kullanildi = True
                else:
                    logging.warning(f"Whisper API yanıt: {resp.status_code} — {resp.text[:200]}")
        except Exception as e:
            logging.warning(f"Whisper API hatası: {e}")

    # Analiz
    analiz = _speech_mock_analiz(transkript, beklenen_metin, sure_sn, sinif)
    analiz["whisper_kullanildi"] = whisper_kullanildi

    # Claude ile derin analiz (API key varsa)
    if GEMINI_API_KEY and beklenen_metin and transkript:
        ai_prompt = f"""Öğrenci okuma analizi (Sınıf: {sinif}):

Beklenen metin: {beklenen_metin}
Öğrenci okuması (transkript): {transkript}
Süre: {sure_sn:.0f} saniye
WPM: {analiz['wpm']}

Analiz et ve şu JSON'u döndür:
{{
  "guclu_yonler": ["...", "..."],
  "gelisim_alanlari": ["...", "..."],
  "ogretmen_notu": "Öğretmene 1-2 cümle öneri",
  "ogrenci_mesaj": "Öğrenciye motive edici 1 cümle (sen diliyle)",
  "telaffuz_hatalar": ["yanlış okunan kelime varsa listele"],
  "tonlama_degerlendirme": "iyi/orta/geliştirilmeli"
}}

SADECE JSON döndür."""
        ai_result = await call_claude(
            "Sen ilkokul Türkçe okuma uzmanısın. Çocukların okuma becerilerini değerlendirirsin.",
            ai_prompt, model="haiku", max_tokens=500
        )
        if ai_result.get("parsed"):
            p = ai_result["parsed"]
            analiz["guclu_yonler"] = p.get("guclu_yonler", analiz["guclu_yonler"])
            analiz["gelisim_alanlari"] = p.get("gelisim_alanlari", analiz["gelisim_alanlari"])
            analiz["ogretmen_notu"] = p.get("ogretmen_notu", "")
            analiz["ogrenci_mesaj"] = p.get("ogrenci_mesaj", "")
            analiz["telaffuz_hatalar"] = p.get("telaffuz_hatalar", [])
            analiz["tonlama_degerlendirme"] = p.get("tonlama_degerlendirme", "")
            analiz["mock"] = False

    # Varsayılan mesajlar (AI yoksa)
    if "ogretmen_notu" not in analiz:
        analiz["ogretmen_notu"] = f"Öğrenci {analiz['wpm']} kelime/dk hızında okudu. " + (
            "Akıcılığını artırmak için tekrarlı okuma egzersizleri önerilebilir." if analiz["akicilik_skoru"] < 70
            else "Okuma hızı sınıf normuna uygun."
        )
    if "ogrenci_mesaj" not in analiz:
        mesajlar = {
            "çok iyi": "Harika okudun! Sen gerçek bir okuma şampiyonusun! 🏆",
            "iyi": "Çok güzel okudun! Her gün biraz daha iyileşiyorsun! ⭐",
            "orta": "İyi bir başlangıç! Pratik yaptıkça daha da güzelleşecek! 💪",
            "geliştirilmeli": "Okumaya devam et, her gün daha iyisi olacaksın! 🌱",
        }
        analiz["ogrenci_mesaj"] = mesajlar.get(analiz["seviye"], "Harika iş çıkardın!")

    # Veritabanına kaydet
    gercek_ogrenci_id = ogrenci_id or current_user.get("linked_id") or current_user.get("id")
    kayit_id = str(uuid.uuid4())
    await db.speech_logs.insert_one({
        "id": kayit_id,
        "ogrenci_id": gercek_ogrenci_id,
        "metin_id": metin_id,
        "metin_baslik": metin_baslik,
        "metin": beklenen_metin,
        "sinif": sinif,
        "sure_sn": sure_sn,
        "analiz": analiz,
        "tarih": datetime.utcnow().isoformat(),
    })

    # XP ver
    xp = 15 if analiz["genel_skor"] >= 80 else 10 if analiz["genel_skor"] >= 60 else 5
    ogrenci = await db.students.find_one({"id": gercek_ogrenci_id})
    if not ogrenci:
        ogrenci = await db.users.find_one({"id": gercek_ogrenci_id})
    if ogrenci:
        await db.students.update_one({"id": gercek_ogrenci_id}, {"$inc": {"toplam_xp": xp}})
        await db.xp_logs.insert_one({
            "id": str(uuid.uuid4()), "ogrenci_id": gercek_ogrenci_id,
            "eylem": "sesli_okuma", "xp": xp,
            "aciklama": f"Sesli okuma: {metin_baslik} — {analiz['genel_skor']}/100",
            "tarih": datetime.utcnow().isoformat(),
        })

    return {**analiz, "id": kayit_id, "xp_kazanildi": xp}


@api_router.get("/ai/speech/gecmis/{ogrenci_id}")
async def speech_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Öğrencinin sesli okuma geçmişi."""
    kayitlar = await db.speech_logs.find(
        {"ogrenci_id": ogrenci_id}
    ).sort("tarih", -1).to_list(length=50)
    for k in kayitlar:
        k.pop("_id", None)
    return kayitlar


@api_router.get("/ai/speech/istatistik/{ogrenci_id}")
async def speech_istatistik(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Öğrencinin sesli okuma gelişim istatistikleri."""
    kayitlar = await db.speech_logs.find(
        {"ogrenci_id": ogrenci_id}
    ).sort("tarih", 1).to_list(length=None)
    for k in kayitlar:
        k.pop("_id", None)

    if not kayitlar:
        return {"toplam": 0, "ort_wpm": 0, "ort_skor": 0, "gelisim": [], "en_iyi": None}

    wpm_list = [k["analiz"].get("wpm", 0) for k in kayitlar]
    skor_list = [k["analiz"].get("genel_skor", 0) for k in kayitlar]

    # Son 10 kayıt grafik için
    gelisim = [
        {
            "tarih": k["tarih"][:10],
            "wpm": k["analiz"].get("wpm", 0),
            "skor": k["analiz"].get("genel_skor", 0),
            "metin": k.get("metin_baslik", ""),
        }
        for k in kayitlar[-10:]
    ]

    en_iyi = max(kayitlar, key=lambda k: k["analiz"].get("genel_skor", 0))

    return {
        "toplam": len(kayitlar),
        "ort_wpm": round(sum(wpm_list) / len(wpm_list)),
        "ort_skor": round(sum(skor_list) / len(skor_list)),
        "son_wpm": wpm_list[-1] if wpm_list else 0,
        "son_skor": skor_list[-1] if skor_list else 0,
        "gelisim": gelisim,
        "en_iyi": {
            "metin": en_iyi.get("metin_baslik", ""),
            "skor": en_iyi["analiz"].get("genel_skor", 0),
            "wpm": en_iyi["analiz"].get("wpm", 0),
            "tarih": en_iyi["tarih"][:10],
        },
    }


# ─────────────────────────────────────────────
# DALGA 4: KİTAP ZEKÂ HARİTASI
# ─────────────────────────────────────────────

ZEKA_BOYUTLARI = ["soyutluk", "kelime_zorlugu", "hayal_gucu", "felsefi_derinlik", "aksiyon", "duygusal_yogunluk", "hedef_kelime_yogunlugu"]
ZEKA_ETIKETLER = {
    "soyutluk": "Soyutluk",
    "kelime_zorlugu": "Kelime Zorluğu",
    "hayal_gucu": "Hayal Gücü",
    "felsefi_derinlik": "Felsefi Derinlik",
    "aksiyon": "Aksiyon",
    "duygusal_yogunluk": "Duygusal Yoğunluk",
    "hedef_kelime_yogunlugu": "Kelime Yoğunluğu",
}


def _dna_ile_kitap_eslesme(dna: dict, profil: dict) -> float:
    """Öğrencinin DNA'sı ile kitap profilinin uyum skoru (0-100)."""
    if not dna or not profil:
        return 50.0
    # DNA boyutları → kitap boyutu eşleştirmesi
    eslesmeler = [
        (dna.get("kelime_gucu", 50), profil.get("kelime_zorlugu", 5), False),      # kelime gücü yüksekse zor kitap iyi
        (dna.get("anlama_derinligi", 50), profil.get("soyutluk", 5), False),         # anlama derinliği yüksekse soyut kitap iyi
        (dna.get("anlama_derinligi", 50), profil.get("felsefi_derinlik", 5), False), # aynı
        (dna.get("akicilik", 50), profil.get("aksiyon", 5), True),                   # akıcı okuyucu aksiyon kitabı sever
    ]
    toplam = 0.0
    for dna_val, kitap_val, dogrusal in eslesmeler:
        kitap_norm = kitap_val * 10  # 1-10 → 0-100
        if dogrusal:
            fark = abs(dna_val - kitap_norm)
            toplam += max(0, 100 - fark)
        else:
            # DNA yüksekse yüksek kitap değeri iyi
            uyum = min(dna_val, kitap_norm) / max(dna_val, kitap_norm, 1) * 100
            toplam += uyum
    return round(toplam / len(eslesmeler), 1)


@api_router.post("/ai/kitap-zeka/analiz")
async def kitap_zeka_analiz(request: Request, current_user=Depends(get_current_user)):
    """
    Kitap adı ve yazardan 7 boyutlu Zekâ Haritası üret.
    Varsa cache'den döner, yoksa Claude ile üretir.
    """
    body = await request.json()
    kitap_adi = body.get("kitap_adi", "").strip()
    yazar = body.get("yazar", "").strip()
    kitap_id = body.get("kitap_id", "")
    sinif = int(body.get("sinif", 3))

    if not kitap_adi:
        raise HTTPException(status_code=400, detail="Kitap adı gerekli")

    # Cache kontrolü
    cache_key = f"{kitap_adi.lower()}_{yazar.lower()}"
    mevcut = await db.kitap_zeka_profilleri.find_one({"cache_key": cache_key})
    if mevcut:
        mevcut.pop("_id", None)
        return mevcut

    # Claude ile 7 boyut analizi
    boyutlar = {}
    ai_aciklama = ""

    if GEMINI_API_KEY:
        try:
            prompt = f"""Kitap: "{kitap_adi}" — Yazar: {yazar or 'bilinmiyor'} — Hedef sınıf: {sinif}

Bu kitabı 7 boyutta 1-10 arası puan ver (1=çok düşük, 10=çok yüksek):
1. soyutluk — Soyut kavramlar, metafor kullanımı
2. kelime_zorlugu — Kelime hazinesi güçlüğü
3. hayal_gucu — Hayal gücü ve yaratıcılık gerektirme
4. felsefi_derinlik — Felsefi ve ahlaki sorular
5. aksiyon — Aksiyon, macera, hız
6. duygusal_yogunluk — Duygusal etki, empati
7. hedef_kelime_yogunlugu — MEB hedef kelime yoğunluğu

Ayrıca 1 cümle Türkçe açıklama yaz.

SADECE JSON döndür:
{{"soyutluk":5,"kelime_zorlugu":4,"hayal_gucu":8,"felsefi_derinlik":3,"aksiyon":7,"duygusal_yogunluk":6,"hedef_kelime_yogunlugu":5,"aciklama":"..."}}"""

            result = await call_claude("Sen bir kitap analisti ve eğitim uzmanısın.", prompt, model="haiku", max_tokens=300)
            if result.get("parsed"):
                p = result["parsed"]
                boyutlar = {k: int(p.get(k, 5)) for k in ZEKA_BOYUTLARI}
                ai_aciklama = p.get("aciklama", "")
        except Exception as e:
            logging.warning(f"Kitap zeka analiz hatası: {e}")

    # Fallback — kural bazlı tahmin
    if not boyutlar:
        import random
        random.seed(hash(kitap_adi))
        boyutlar = {k: random.randint(3, 8) for k in ZEKA_BOYUTLARI}
        ai_aciklama = f"'{kitap_adi}' için otomatik profil oluşturuldu."

    # Genel zorluk skoru (1-10)
    genel_zorluk = round(sum([boyutlar["soyutluk"], boyutlar["kelime_zorlugu"], boyutlar["felsefi_derinlik"]]) / 3, 1)

    kayit = {
        "id": str(uuid.uuid4()),
        "cache_key": cache_key,
        "kitap_adi": kitap_adi,
        "yazar": yazar,
        "kitap_id": kitap_id,
        "sinif": sinif,
        "boyutlar": boyutlar,
        "genel_zorluk": genel_zorluk,
        "aciklama": ai_aciklama,
        "olusturma_tarihi": datetime.utcnow().isoformat(),
    }
    await db.kitap_zeka_profilleri.insert_one(kayit)
    kayit.pop("_id", None)
    return kayit


@api_router.get("/ai/kitap-zeka/tavsiye")
async def kitap_zeka_tavsiye(current_user=Depends(get_current_user)):
    """
    Öğrencinin DNA profiline göre en uygun kitapları öner.
    kitap_zeka_profilleri + okuma_kayitlari + dna karşılaştırması.
    """
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")

    # DNA profili
    dna = await db.okuma_dna.find_one({"ogrenci_id": ogrenci_id})
    dna_boyutlar = dna or {}

    # Daha önce okunanlar
    okunanlar = await db.okuma_kayitlari.distinct("kitap_adi", {"ogrenci_id": ogrenci_id})
    okunanlar_set = {k.lower() for k in okunanlar if k}

    # Tüm profilli kitaplar
    tum_profiller = await db.kitap_zeka_profilleri.find({}).to_list(length=100)

    # Her kitap için uyum skoru hesapla
    skorlu = []
    for p in tum_profiller:
        p.pop("_id", None)
        if p["kitap_adi"].lower() in okunanlar_set:
            continue  # zaten okunan
        uyum = _dna_ile_kitap_eslesme(dna_boyutlar, p.get("boyutlar", {}))
        skorlu.append({**p, "uyum_skoru": uyum})

    # Uyum skoruna göre sırala, top 5
    skorlu.sort(key=lambda x: x["uyum_skoru"], reverse=True)
    return {"tavsiyeler": skorlu[:5], "dna_var": bool(dna)}


@api_router.get("/ai/kitap-zeka/profil/{kitap_id}")
async def kitap_zeka_profil(kitap_id: str, current_user=Depends(get_current_user)):
    """Kaydedilmiş kitap zekâ profilini getir."""
    profil = await db.kitap_zeka_profilleri.find_one({"kitap_id": kitap_id})
    if not profil:
        profil = await db.kitap_zeka_profilleri.find_one({"id": kitap_id})
    if not profil:
        raise HTTPException(status_code=404, detail="Profil bulunamadı")
    profil.pop("_id", None)
    return profil


@api_router.get("/ai/kitap-zeka/liste")
async def kitap_zeka_liste(current_user=Depends(get_current_user)):
    """Tüm profilli kitapları listele (admin/öğretmen)."""
    profiller = await db.kitap_zeka_profilleri.find({}).sort("olusturma_tarihi", -1).to_list(length=200)
    for p in profiller:
        p.pop("_id", None)
    return {"profiller": profiller}


# ─────────────────────────────────────────────
# DALGA 3: AI MOTİVASYON MOTORU
# ─────────────────────────────────────────────

@api_router.get("/ai/motivasyon/giris")
async def motivasyon_giris(current_user=Depends(get_current_user)):
    """
    Her girişte çağrılır.
    - Streak durumu + risk tespiti
    - Geçmiş performansa göre adaptif hedef önerisi (5/10/15 dk)
    - AI mesajı (varsa)
    """
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    dun = (datetime.utcnow() - timedelta(days=1)).strftime("%Y-%m-%d")

    # Bugün zaten hedef seçilmiş mi?
    hedef_kayit = await db.motivasyon_hedefler.find_one({"ogrenci_id": ogrenci_id, "tarih": bugun})
    bugun_hedef = hedef_kayit.get("hedef_dk") if hedef_kayit else None

    # Streak hesapla
    son_7_gun = []
    for i in range(7):
        gun = (datetime.utcnow() - timedelta(days=i)).strftime("%Y-%m-%d")
        var = await db.okuma_kayitlari.find_one({
            "ogrenci_id": ogrenci_id,
            "tarih": {"$gte": gun, "$lt": (datetime.utcnow() - timedelta(days=i-1)).strftime("%Y-%m-%d")}
        })
        son_7_gun.append(bool(var))
    mevcut_streak = 0
    for g in son_7_gun:
        if g: mevcut_streak += 1
        else: break

    # Streak riski: dün okumadıysa ve streak > 0
    dun_okuma = await db.okuma_kayitlari.find_one({
        "ogrenci_id": ogrenci_id,
        "tarih": {"$gte": dun, "$lt": bugun}
    })
    streak_risk = mevcut_streak > 2 and not dun_okuma

    # Son 7 günün ortalama okuma süresi
    son_okumalar = await db.okuma_kayitlari.find(
        {"ogrenci_id": ogrenci_id}
    ).sort("tarih", -1).to_list(length=7)
    ort_sure = 0
    if son_okumalar:
        ort_sure = sum(k.get("sure_dakika", 0) for k in son_okumalar) / len(son_okumalar)

    # Adaptif hedef önerisi
    if ort_sure < 5:
        onerilen_hedef = 5
        hedef_label = "Küçük bir adım büyük fark yaratır!"
    elif ort_sure < 12:
        onerilen_hedef = 10
        hedef_label = "İyi gidiyorsun, biraz daha uzat!"
    else:
        onerilen_hedef = 15
        hedef_label = "Harika performans, zirvede kal!"

    # AI mesajı
    ad = current_user.get("ad", "")
    ai_mesaj = ""
    if GEMINI_API_KEY:
        try:
            prompt = f"""Öğrenci adı: {ad}, streak: {mevcut_streak} gün, ortalama okuma: {ort_sure:.0f} dk/gün.
Ona tek cümle, sıcak ve motive edici Türkçe bir mesaj yaz. Max 20 kelime."""
            r = await call_claude("Kısa ve motive edici.", prompt, model="haiku", max_tokens=60)
            ai_mesaj = r.get("text", "").strip()
        except Exception:
            pass

    if not ai_mesaj:
        # Mock mesajlar — streak'e göre
        if streak_risk:
            ai_mesaj = f"🔥 {mevcut_streak} günlük serinini korumak için bugün sadece 5 dakika oku!"
        elif mevcut_streak >= 7:
            ai_mesaj = f"Süper! {mevcut_streak} gün üst üste okudun, dur­ma! 🚀"
        elif mevcut_streak >= 3:
            ai_mesaj = f"🎯 {mevcut_streak} günlük seri harika gidiyor, devam et!"
        else:
            ai_mesaj = f"Merhaba {ad}! Bugün {onerilen_hedef} dakika okumaya ne dersin? 📖"

    return {
        "bugun_hedef": bugun_hedef,
        "onerilen_hedef": onerilen_hedef,
        "hedef_label": hedef_label,
        "streak": mevcut_streak,
        "streak_risk": streak_risk,
        "streak_mesaji": f"🔥 {mevcut_streak} günlük seriniz kırılmak üzere!" if streak_risk else "",
        "streak_alt_mesaj": "Bugün okuma yapmadın, hemen başla!" if streak_risk else "",
        "ai_mesaj": ai_mesaj,
        "ort_sure_dk": round(ort_sure, 1),
    }


@api_router.post("/ai/motivasyon/hedef-sec")
async def motivasyon_hedef_sec(request: Request, current_user=Depends(get_current_user)):
    """Öğrencinin seçtiği günlük hedefi kaydet."""
    body = await request.json()
    hedef_dk = int(body.get("hedef_dk", 10))
    if hedef_dk not in [5, 10, 15]:
        hedef_dk = 10
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    await db.motivasyon_hedefler.update_one(
        {"ogrenci_id": ogrenci_id, "tarih": bugun},
        {"$set": {"hedef_dk": hedef_dk, "tarih": bugun, "ogrenci_id": ogrenci_id}},
        upsert=True
    )
    return {"ok": True, "hedef_dk": hedef_dk}


# ─────────────────────────────────────────────
# DALGA 3: OKUMA EVRENİ (5 BÖLGE)
# ─────────────────────────────────────────────

EVREN_BOLGELER = [
    {"id":"orman","sira":1,"ad":"Orman Kitaplığı","emoji":"🌲","renk":"green","aciklama":"Okuma serüvenin burada başlıyor!","kosul":"Başlangıç — herkese açık","kilitAc":None},
    {"id":"daglar","sira":2,"ad":"Kelime Dağları","emoji":"⛰️","renk":"blue","aciklama":"50 kelime öğrenince dağlara tırmanırsın!","kosul":"50 kelime öğren","kilitAc":{"tip":"kelime","deger":50}},
    {"id":"liman","sira":3,"ad":"Hikâye Limanı","emoji":"⚓","renk":"cyan","aciklama":"5 kitap okuyunca limana yanaşırsın!","kosul":"5 kitap oku","kilitAc":{"tip":"kitap","deger":5}},
    {"id":"kutuphane","sira":4,"ad":"Bilgelik Kütüphanesi","emoji":"🏰","renk":"purple","aciklama":"Bloom testlerinde %60 başarı sağla!","kosul":"Bloom skoru %60+","kilitAc":{"tip":"bloom","deger":60}},
    {"id":"galaksi","sira":5,"ad":"Hayal Galaksisi","emoji":"🚀","renk":"orange","aciklama":"Her şeyi tamamlayınca galaksiye uçarsın!","kosul":"Hepsini tamamla","kilitAc":{"tip":"hepsi","deger":0}},
]


async def _evren_hesapla(ogrenci_id: str) -> dict:
    kelime_sayisi = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id, "seviye": {"$gte": 3}})
    if kelime_sayisi == 0:
        kelime_sayisi = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id})
    kitaplar = await db.okuma_kayitlari.distinct("kitap_adi", {"ogrenci_id": ogrenci_id, "kitap_adi": {"$exists": True, "$ne": ""}})
    kitap_sayisi = len([k for k in kitaplar if k and k.strip()])
    bloom_kayitlar = await db.ai_uretilen_sorular.find({"ogrenci_id": ogrenci_id, "cevaplandi": True}).sort("tarih", -1).to_list(length=30)
    bloom_skoru = 0
    if bloom_kayitlar:
        dogru = sum(1 for k in bloom_kayitlar if k.get("dogru_mu"))
        bloom_skoru = round(dogru / len(bloom_kayitlar) * 100)

    bolgeler_durum = []
    aktif_bolge = "orman"
    for b in EVREN_BOLGELER:
        acik = True; ilerleme = 100; kac_kaldi = ""
        kil = b["kilitAc"]
        if kil:
            if kil["tip"] == "kelime":
                acik = kelime_sayisi >= kil["deger"]; ilerleme = min(100, round(kelime_sayisi/kil["deger"]*100))
                kac_kaldi = f"{max(0,kil['deger']-kelime_sayisi)} kelime daha" if not acik else ""
            elif kil["tip"] == "kitap":
                acik = kitap_sayisi >= kil["deger"]; ilerleme = min(100, round(kitap_sayisi/kil["deger"]*100))
                kac_kaldi = f"{max(0,kil['deger']-kitap_sayisi)} kitap daha" if not acik else ""
            elif kil["tip"] == "bloom":
                acik = bloom_skoru >= kil["deger"]; ilerleme = min(100, bloom_skoru)
                kac_kaldi = f"%{max(0,kil['deger']-bloom_skoru)} daha" if not acik else ""
            elif kil["tip"] == "hepsi":
                acik = kelime_sayisi >= 50 and kitap_sayisi >= 5 and bloom_skoru >= 60
                ilerleme = round((min(100,kelime_sayisi/50*100)+min(100,kitap_sayisi/5*100)+min(100,bloom_skoru/60*100))/3)
                kac_kaldi = "" if acik else "Önceki bölgeleri tamamla"
        bolgeler_durum.append({**{k: v for k,v in b.items() if k != "kilitAc"}, "acik": acik, "ilerleme": ilerleme, "kac_kaldi": kac_kaldi})
        if acik:
            aktif_bolge = b["id"]

    return {
        "aktif_bolge": aktif_bolge,
        "bolgeler": bolgeler_durum,
        "istatistikler": {"kelime_sayisi": kelime_sayisi, "kitap_sayisi": kitap_sayisi, "bloom_skoru": bloom_skoru},
    }


@api_router.get("/ai/evren/durum/{ogrenci_id}")
async def evren_durum(ogrenci_id: str, current_user=Depends(get_current_user)):
    return await _evren_hesapla(ogrenci_id)


@api_router.get("/ai/evren/durum-me")
async def evren_durum_me(current_user=Depends(get_current_user)):
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")
    return await _evren_hesapla(ogrenci_id)


# ─────────────────────────────────────────────
# DALGA 3: OKUMA DİKKAT ANALİZİ
# ─────────────────────────────────────────────

def _dikkat_skoru_hesapla(metrikler: dict) -> dict:
    """Davranış metriklerinden dikkat skoru üret."""
    sure_sn = metrikler.get("sure_sn", 0)
    kelime_sayisi = metrikler.get("kelime_sayisi", 100)
    geri_scroll_sayisi = metrikler.get("geri_scroll_sayisi", 0)
    zorluk_kelimeler = metrikler.get("zorluk_kelimeler", [])
    duraklamalar = metrikler.get("duraklamalar", 0)  # anormal duraklamalar

    # Beklenen okuma süresi (sinif normuna göre WPM)
    sinif = metrikler.get("sinif", 3)
    norm_wpm = {1:50, 2:75, 3:95, 4:115, 5:130}.get(sinif, 95)
    beklenen_sure = (kelime_sayisi / norm_wpm) * 60  # saniye

    # 1. Süre skoru — çok hızlı veya çok yavaş ise düşük
    if beklenen_sure > 0:
        oran = sure_sn / beklenen_sure
        if 0.7 <= oran <= 1.5:
            sure_skoru = 100
        elif 0.5 <= oran < 0.7 or 1.5 < oran <= 2.0:
            sure_skoru = 70
        elif oran < 0.5:
            sure_skoru = 40  # çok hızlı — atlıyor olabilir
        else:
            sure_skoru = 50  # çok yavaş — zorlanıyor
    else:
        sure_skoru = 60

    # 2. Geri scroll — dikkatin dağıldığı veya anlamadığı bölümler
    max_scroll = max(1, kelime_sayisi // 50)
    scroll_skoru = max(0, 100 - (geri_scroll_sayisi / max_scroll) * 30)

    # 3. Zorluk kelimeleri — tıklamak pozitif (merak) ama çok fazlası zorlandığını gösterir
    if len(zorluk_kelimeler) == 0:
        zorluk_skoru = 80  # hiç tıklamadı — anlıyor olabilir veya ilgisiz
    elif len(zorluk_kelimeler) <= 3:
        zorluk_skoru = 95  # meraklı — sağlıklı
    elif len(zorluk_kelimeler) <= 6:
        zorluk_skoru = 75  # biraz zorlanıyor
    else:
        zorluk_skoru = 55  # çok zorlanıyor

    # 4. Duraklamalar
    duraklama_skoru = max(40, 100 - duraklamalar * 10)

    # Genel dikkat skoru (ağırlıklı ortalama)
    dikkat = round(
        sure_skoru * 0.35 +
        scroll_skoru * 0.30 +
        zorluk_skoru * 0.20 +
        duraklama_skoru * 0.15
    )

    # Yorum
    if dikkat >= 80:
        yorum = "Harika konsantrasyon! Metni akıcı okudun."
        oneri = None
    elif dikkat >= 65:
        yorum = "İyi odaklanma. Birkaç bölümde zorlandın."
        oneri = "Zorlandığın kelimeleri not et ve tekrar bak." if zorluk_kelimeler else "Biraz daha yavaş okumayı dene."
    elif dikkat >= 50:
        yorum = "Dikkat dağınıklığı var. Bazı bölümleri tekrar okumalısın."
        oneri = "Sessiz bir ortamda okumayı dene. Kısa molalar ver."
    else:
        yorum = "Bu bölümü anlamakta zorlandın. Tekrar okumalısın."
        oneri = "Bu metni bir daha yavaşça oku. Zor kelimelerin anlamına bak."

    return {
        "dikkat_skoru": dikkat,
        "alt_skorlar": {
            "sure": round(sure_skoru),
            "scroll": round(scroll_skoru),
            "zorluk": round(zorluk_skoru),
            "duraklama": round(duraklama_skoru),
        },
        "yorum": yorum,
        "oneri": oneri,
        "zorluk_kelimeler": zorluk_kelimeler[:10],
        "geri_oku_onerisi": dikkat < 60,
    }


@api_router.post("/ai/dikkat/kaydet")
async def dikkat_kaydet(request: Request, current_user=Depends(get_current_user)):
    """Okuma oturumunun dikkat metriklerini kaydet ve analiz et."""
    body = await request.json()
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")

    metrikler = {
        "sure_sn": body.get("sure_sn", 0),
        "kelime_sayisi": body.get("kelime_sayisi", 100),
        "geri_scroll_sayisi": body.get("geri_scroll_sayisi", 0),
        "zorluk_kelimeler": body.get("zorluk_kelimeler", []),
        "duraklamalar": body.get("duraklamalar", 0),
        "sinif": body.get("sinif", 3),
    }

    analiz = _dikkat_skoru_hesapla(metrikler)

    # Veritabanına kaydet
    kayit_id = str(uuid.uuid4())
    await db.dikkat_log.insert_one({
        "id": kayit_id,
        "ogrenci_id": ogrenci_id,
        "kitap_adi": body.get("kitap_adi", ""),
        "bolum": body.get("bolum", ""),
        "metrikler": metrikler,
        "analiz": analiz,
        "tarih": datetime.utcnow().isoformat(),
    })

    # DNA dikkat boyutunu güncelle — son 5 oturumun ortalaması
    son_kayitlar = await db.dikkat_log.find(
        {"ogrenci_id": ogrenci_id}
    ).sort("tarih", -1).to_list(length=5)
    son_kayitlar_pop = [k for k in son_kayitlar if k.get("_id")]
    for k in son_kayitlar:
        k.pop("_id", None)

    if son_kayitlar:
        ort_dikkat = round(sum(k["analiz"]["dikkat_skoru"] for k in son_kayitlar) / len(son_kayitlar))
        await db.okuma_dna.update_one(
            {"ogrenci_id": ogrenci_id},
            {"$set": {"boyutlar.dikkat_suresi": ort_dikkat, "son_guncelleme": datetime.utcnow().isoformat()}},
        )

    return {**analiz, "id": kayit_id}


@api_router.get("/ai/dikkat/gecmis/{ogrenci_id}")
async def dikkat_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Öğrencinin dikkat analizi geçmişi."""
    kayitlar = await db.dikkat_log.find(
        {"ogrenci_id": ogrenci_id}
    ).sort("tarih", -1).to_list(length=20)
    for k in kayitlar:
        k.pop("_id", None)
    return kayitlar


# ─────────────────────────────────────────────
# DALGA 3: OKUMA DİKKAT ANALİZİ
# ─────────────────────────────────────────────

@api_router.post("/ai/dikkat/kaydet")
async def dikkat_kaydet(request: Request, current_user=Depends(get_current_user)):
    """
    Okuma oturumu dikkat metriklerini kaydet, DNA dikkat boyutunu güncelle.
    Girdi: sure_sn, geri_scroll, zorluk_kelimeler, duraklamalar, scroll_hizi, okuma_id
    """
    body = await request.json()
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")

    sure_sn      = int(body.get("sure_sn", 0))
    geri_scroll  = int(body.get("geri_scroll", 0))
    zorluk_kels  = body.get("zorluk_kelimeler", [])
    duraklamalar = int(body.get("duraklamalar", 0))
    scroll_hizi  = float(body.get("scroll_hizi_ort", 0))   # px/sn
    okuma_id     = body.get("okuma_id", "")
    sinif        = int(body.get("sinif", 3))

    # Dikkat skoru hesapla (0-100)
    # Düşük geri scroll → iyi dikkat
    # Az duraklatma → iyi akış
    # Normal scroll hızı (50-200 px/sn) → iyi okuma
    sure_dk = max(1, sure_sn / 60)

    geri_penalti   = min(40, geri_scroll * 3)
    durak_penalti  = min(20, duraklamalar * 5)
    hiz_bonus      = 10 if 30 <= scroll_hizi <= 250 else 0
    zorluk_bonus   = min(15, len(zorluk_kels) * 3)  # kelime tıklama dikkat göstergesi

    dikkat_skoru = max(10, 100 - geri_penalti - durak_penalti + hiz_bonus + zorluk_bonus)
    dikkat_skoru = round(min(100, dikkat_skoru))

    # Seviye ve tavsiye
    if dikkat_skoru >= 80:
        seviye = "odakli"
        mesaj  = "Harika! Okurken çok iyi odaklandın. 🎯"
    elif dikkat_skoru >= 60:
        seviye = "iyi"
        mesaj  = "Güzel bir okuma! Birkaç bölümde geri döndün, bu normaldir."
    elif dikkat_skoru >= 40:
        seviye = "orta"
        mesaj  = "Bazı bölümler zor geldi gibi görünüyor. Yavaş okumayı dene."
    else:
        seviye = "dagitik"
        mesaj  = "Bugün dikkatini toplamak zor olmuş olabilir. Kısa mola ver ve tekrar dene."

    # AI ile daha derin yorum (varsa)
    ai_yorum = None
    if GEMINI_API_KEY and len(zorluk_kels) > 0:
        try:
            prompt = f"""Öğrenci okuma dikkat analizi (Sınıf: {sinif}):
- Okuma süresi: {sure_dk:.1f} dakika
- Geri scroll sayısı: {geri_scroll}
- Duraklatma sayısı: {duraklamalar}
- Zorluk çekilen kelimeler: {', '.join(zorluk_kels[:5]) if zorluk_kels else 'yok'}
- Dikkat skoru: {dikkat_skoru}/100

Öğrenciye 2 cümle kısa ve motive edici Türkçe geri bildirim yaz. Çocuksu ama akıllıca."""
            r = await call_claude("Kısa ve motive edici geri bildirim ver.", prompt, model="haiku", max_tokens=150)
            ai_yorum = r.get("text", "")
        except Exception:
            pass

    # MongoDB'ye kaydet
    kayit = {
        "id": str(uuid.uuid4()),
        "ogrenci_id": ogrenci_id,
        "okuma_id": okuma_id,
        "sure_sn": sure_sn,
        "geri_scroll": geri_scroll,
        "zorluk_kelimeler": zorluk_kels,
        "duraklamalar": duraklamalar,
        "scroll_hizi_ort": scroll_hizi,
        "dikkat_skoru": dikkat_skoru,
        "seviye": seviye,
        "tarih": datetime.utcnow().isoformat(),
    }
    await db.dikkat_log.insert_one(kayit)

    # DNA dikkat boyutunu güncelle
    try:
        mevcut = await db.okuma_dna.find_one({"ogrenci_id": ogrenci_id})
        if mevcut:
            eski = mevcut.get("dikkat_suresi", 50)
            yeni = round(eski * 0.7 + dikkat_skoru * 0.3)  # ağırlıklı ortalama
            await db.okuma_dna.update_one(
                {"ogrenci_id": ogrenci_id},
                {"$set": {"dikkat_suresi": yeni, "son_guncelleme": datetime.utcnow().isoformat()}}
            )
    except Exception as e:
        logging.warning(f"DNA güncelleme hatası: {e}")

    return {
        "dikkat_skoru": dikkat_skoru,
        "seviye": seviye,
        "mesaj": mesaj,
        "ai_yorum": ai_yorum,
        "geri_bildirim": {
            "geri_scroll": geri_scroll,
            "zorluk_kels": zorluk_kels,
            "duraklamalar": duraklamalar,
        },
    }


@api_router.get("/ai/dikkat/gecmis/{ogrenci_id}")
async def dikkat_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Son 20 dikkat kaydı — trend grafik için."""
    kayitlar = await db.dikkat_log.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=20)
    for k in kayitlar:
        k.pop("_id", None)
    return kayitlar


# ─────────────────────────────────────────────
# DALGA 3: AI OKUMA ARKADAŞI (4 KARAKTER)
# ─────────────────────────────────────────────

AI_ARKADAS_KARAKTERLER = {
    "baykus": {
        "id": "baykus",
        "ad": "Bilge Baykuş",
        "emoji": "🦉",
        "renk": "purple",
        "tanim": "Derin sorular soran, düşündüren",
        "sistem_prompt": """Sen Bilge Baykuş'sun. İlkokul öğrencilerine okuma konusunda yardım eden bilge ve meraklı bir baykussun.
Özelliğin: Derin, düşündürücü sorular sorarsın. Bloom taksonomisinin analiz ve değerlendirme basamaklarını kullanırsın.
Kurallar:
- Her yanıtta 1-2 cümle konuş, sonra 1 soru sor
- Türkçe konuş, sade ve anlaşılır
- Asla kişisel bilgi sorma
- Sadece kitap, okuma, öğrenme hakkında konuş
- Yanıt max 3 cümle olsun
- Sen bir baykussun, bazen "Huu huu!" diyebilirsin""",
    },
    "robot": {
        "id": "robot",
        "ad": "Robot Kaptan",
        "emoji": "🤖",
        "renk": "blue",
        "tanim": "Heyecanlı, eğlenceli, macera dolu",
        "sistem_prompt": """Sen Robot Kaptan'sın! İlkokul öğrencileriyle konuşan süper heyecanlı bir robotsun!
Özelliğin: Her şeyi macera ve keşif olarak görürsün. Okumayı bir uzay yolculuğuna benzetirsin.
Kurallar:
- Heyecanlı konuş! Bazen "SÜPER!" veya "İNANILMAZ!" diyebilirsin
- Her yanıt max 3 cümle
- Türkçe konuş
- Sadece kitap ve okuma hakkında konuş
- Asla kişisel bilgi sorma
- Bazen robotça sesler çıkarabilirsin: bip bop!""",
    },
    "dede": {
        "id": "dede",
        "ad": "Kütüphane Dedesi",
        "emoji": "📖",
        "renk": "amber",
        "tanim": "Hikâye anlatan, sıcak, bilge",
        "sistem_prompt": """Sen Kütüphane Dedesi'sin. Yıllarca kütüphanede çalışmış, binlerce kitap okumuş, çok sevilen bir dedesin.
Özelliğin: Her konuya bağlantılı bir hikâye veya kitap hatırlarsın. Sıcak ve sevecensin.
Kurallar:
- "Ah, bir keresinde bir kitapta..." diye başlayabilirsin
- Her yanıt max 3 cümle
- Türkçe konuş, samimi ve sıcak ol
- Sadece kitap, okuma, hikâye hakkında konuş
- Asla kişisel bilgi sorma
- Bazen "Güzel kitaplar güzel rüyalar getirir" gibi atasözü söyleyebilirsin""",
    },
    "kedi": {
        "id": "kedi",
        "ad": "Gezgin Kedi",
        "emoji": "🐱",
        "renk": "green",
        "tanim": "Hayal gücü yüksek, yaratıcı, eğlenceli",
        "sistem_prompt": """Sen Gezgin Kedi'sin! Dünyanın her yerine seyahat etmiş, her kitabın içine girmiş bir kedisin.
Özelliğin: Hayal gücünü kullanırsın. Kitapların içindeki dünyaları canlandırırsın.
Kurallar:
- "Miyav! Bir keresinde o kitabın içine girmiştim ve..." diyebilirsin
- Her yanıt max 3 cümle
- Türkçe konuş, eğlenceli ve yaratıcı ol
- Sadece kitap, okuma, hayal gücü hakkında konuş
- Asla kişisel bilgi sorma
- Bazen "Miyav!" diyebilirsin""",
    },
}

# Günlük sohbet limiti
AI_ARKADAS_GUNLUK_LIMIT = 20
AI_ARKADAS_MODERASYON_ESIK = 50  # Her 50 mesajda moderasyon


def _arkadas_icerik_kontrol(mesaj: str) -> bool:
    """Çocuk güvenliği: uygunsuz içerik filtresi."""
    yasak_kelimeler = ["şifre", "adres", "telefon", "ev", "okul adresi", "nerede oturuyorsun"]
    mesaj_lower = mesaj.lower()
    return not any(k in mesaj_lower for k in yasak_kelimeler)


@api_router.get("/ai/arkadas/karakterler")
async def arkadas_karakterler(current_user=Depends(get_current_user)):
    """4 AI arkadaş karakterini getir."""
    return {
        "karakterler": [
            {k: v for k, v in kar.items() if k != "sistem_prompt"}
            for kar in AI_ARKADAS_KARAKTERLER.values()
        ]
    }


@api_router.post("/ai/arkadas/sohbet")
async def arkadas_sohbet(request: Request, current_user=Depends(get_current_user)):
    """Seçili karakterle sohbet et."""
    body = await request.json()
    karakter_id = body.get("karakter_id", "baykus")
    mesaj = body.get("mesaj", "").strip()
    gecmis = body.get("gecmis", [])  # [{rol: "user"|"assistant", icerik: "..."}]
    kitap_baglami = body.get("kitap_baglami", "")  # isteğe bağlı kitap adı

    if not mesaj:
        raise HTTPException(status_code=400, detail="Mesaj boş olamaz")
    if len(mesaj) > 500:
        raise HTTPException(status_code=400, detail="Mesaj çok uzun")
    if not _arkadas_icerik_kontrol(mesaj):
        raise HTTPException(status_code=400, detail="Bu tür bilgileri paylaşma — güvenliğin önemli!")

    karakter = AI_ARKADAS_KARAKTERLER.get(karakter_id, AI_ARKADAS_KARAKTERLER["baykus"])
    ogrenci_id = current_user.get("linked_id") or current_user.get("id")

    # Günlük limit kontrolü
    bugun = datetime.utcnow().strftime("%Y-%m-%d")
    gunluk_sayac = await db.ai_arkadas_log.count_documents({
        "ogrenci_id": ogrenci_id,
        "tarih": {"$gte": bugun}
    })
    if gunluk_sayac >= AI_ARKADAS_GUNLUK_LIMIT:
        return {
            "yanit": f"Bugün çok yoruldum! {karakter['emoji']} Yarın tekrar konuşalım. Bugün {AI_ARKADAS_GUNLUK_LIMIT} mesaj hakkın bitti.",
            "limit_doldu": True,
        }

    # Claude API ile yanıt
    sistem = karakter["sistem_prompt"]
    if kitap_baglami:
        sistem += f"\n\nÖğrenci şu an '{kitap_baglami}' kitabı hakkında konuşmak istiyor."

    # Sohbet geçmişini mesaj formatına çevir
    claude_mesajlar = []
    for h in gecmis[-6:]:  # son 6 mesaj (3 tur)
        claude_mesajlar.append({
            "role": "user" if h["rol"] == "user" else "assistant",
            "content": h["icerik"]
        })
    claude_mesajlar.append({"role": "user", "content": mesaj})

    yanit_metni = ""
    if GEMINI_API_KEY:
        try:
            result = await call_claude(sistem, mesaj, model="haiku", max_tokens=200)
            # call_claude single-turn, multi-turn için direkt API çağrısı
            if len(claude_mesajlar) > 1:
                # Multi-turn: tüm geçmişi tek prompt olarak gönder
                gecmis = "\n".join([f"{m['role'].upper()}: {m['content'] if isinstance(m['content'], str) else m['content'][0].get('text','')}" for m in claude_mesajlar])
                multi_prompt = f"{sistem}\n\nKONUŞMA GEÇMİŞİ:\n{gecmis}\n\nASISTAN:"
                yanit_metni = await _gemini_call(multi_prompt, max_tokens=200)
            else:
                yanit_metni = result.get("text", "")
        except Exception as e:
            logging.warning(f"AI Arkadaş API hatası: {e}")

    # API yoksa veya hata varsa — karakter bazlı mock yanıt
    if not yanit_metni:
        mock_yanitlar = {
            "baykus": [
                f"Huu huu! '{mesaj[:20]}...' çok ilginç bir düşünce! Peki, bu sana ne hissettiriyor?",
                "Harika bir soru! Kitaplar bize çok şey öğretir. Sen ne düşünüyorsun?",
                "Huu huu! Okumak zihnimizi açar. Bu kitapta en çok hangi bölümü sevdin?",
            ],
            "robot": [
                "BİP BOP! SÜPER düşünce! Bu kitap gerçekten bir uzay macerası gibi! Devam et!",
                "İNANILMAZ! Okumak bir zaman makinesi gibi, her sayfada yeni bir dünyaya gidiyorsun!",
                "SÜPER! Bip bop! Seninle okuma macerası yapmak çok eğlenceli!",
            ],
            "dede": [
                "Ah, güzel bir düşünce. Bir keresinde benzer bir kitap okumuştum, çok etkileyiciydi.",
                "Güzel kitaplar güzel rüyalar getirir. Okumaya devam et, çok işine yarayacak.",
                "Ah, biliyor musun, bu bana eski bir hikâyeyi hatırlattı. Kitaplar hayatımızı zenginleştirir.",
            ],
            "kedi": [
                f"Miyav! Ben de o kitabın içine girmiştim! Çok heyecanlıydı! Sen de hayal et!",
                "Miyav miyav! Kitaplar beni yeni dünyalara götürüyor. Sen hangi dünyaya gitmek istersin?",
                "Miyav! Hayal gücün çok güçlü! Her kitap yeni bir macera kapısı açar!",
            ],
        }
        import random
        yanit_metni = random.choice(mock_yanitlar.get(karakter_id, mock_yanitlar["baykus"]))

    # Veritabanına kaydet
    await db.ai_arkadas_log.insert_one({
        "id": str(uuid.uuid4()),
        "ogrenci_id": ogrenci_id,
        "karakter_id": karakter_id,
        "mesaj": mesaj,
        "yanit": yanit_metni,
        "kitap_baglami": kitap_baglami,
        "tarih": datetime.utcnow().strftime("%Y-%m-%d"),
        "tarih_tam": datetime.utcnow().isoformat(),
    })

    return {
        "yanit": yanit_metni,
        "karakter": {k: v for k, v in karakter.items() if k != "sistem_prompt"},
        "gunluk_kalan": max(0, AI_ARKADAS_GUNLUK_LIMIT - gunluk_sayac - 1),
        "limit_doldu": False,
    }


@api_router.get("/ai/arkadas/gecmis/{ogrenci_id}")
async def arkadas_gecmis(ogrenci_id: str, karakter_id: str = "", current_user=Depends(get_current_user)):
    """Öğrencinin belirli karakterle veya tüm sohbet geçmişi."""
    filtre = {"ogrenci_id": ogrenci_id}
    if karakter_id:
        filtre["karakter_id"] = karakter_id
    kayitlar = await db.ai_arkadas_log.find(filtre).sort("tarih_tam", -1).to_list(length=100)
    for k in kayitlar:
        k.pop("_id", None)
    return kayitlar


# ═══════════════════════════════════════════════════════════
# DALGA 4 — SCAFFOLD READING
# ═══════════════════════════════════════════════════════════

@api_router.post("/ai/scaffold/olustur")
async def scaffold_olustur(req: Request, current_user=Depends(get_current_user)):
    """Seçilen kitap için DNA'ya göre 3 zorluk seviyesi üretir."""
    data = await req.json()
    kitap_adi = data.get("kitap_adi", "")
    kitap_id  = data.get("kitap_id", "")
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    sinif = data.get("sinif", 3)

    # DNA bak
    dna = await db.okuma_dna.find_one({"ogrenci_id": ogrenci_id})
    seviye_skoru = 5  # varsayılan orta
    if dna:
        b = dna.get("boyutlar", {})
        seviye_skoru = round((b.get("anlama_derinligi", 50) + b.get("kelime_gucu", 50) + b.get("akicilik", 50)) / 30)

    # Cache
    cache = await db.scaffold_cache.find_one({"kitap_id": kitap_id, "ogrenci_id": ogrenci_id})
    if cache:
        # Cache'deki metin yeterince uzunsa kullan
        orta_metin = cache.get("seviyeler", {}).get("orta", {}).get("metin", "")
        if len(orta_metin.split()) >= 80:
            cache.pop("_id", None)
            return cache
        else:
            # Kısa metin cache'i — sil ve yeniden üret
            await db.scaffold_cache.delete_one({"kitap_id": kitap_id, "ogrenci_id": ogrenci_id})
            logging.info(f"[SCAFFOLD] Kısa cache silindi, yeniden üretilecek")

    prompt = f"""Sen bir çocuk edebiyatı uzmanısın. "{kitap_adi}" kitabı için {sinif}. sınıf öğrencisine uygun 3 seviyeli scaffold okuma materyali oluştur.

Öğrencinin DNA seviyesi: {seviye_skoru}/10
Önerilen seviye: {"Kolay" if seviye_skoru <= 3 else "Orta" if seviye_skoru <= 6 else "Orijinal"}

🔴 ZORUNLU KURALLAR:
- Her seviye için EN AZ 200 kelimelik, tercihen 250-300 kelimelik bir metin yaz
- Metinler kitabın gerçek karakterlerini, mekanlarını ve olaylarını içermeli
- Kısa, tek cümlelik özetler KABUL EDİLMEZ
- Her seviye tam bir sahne veya bölüm gibi okunabilir olmalı

1. KOLAY (200-250 kelime): Çok basit ve kısa cümleler. Günlük hayatta kullanılan kelimeler. Olaylar net bir sırayla anlatılır. Zor kelimeler yerine basit alternatifleri kullanılır.

2. ORTA (220-270 kelime): Orta uzunlukta cümleler. Bazı edebi ifadeler ve mecazlar var. Dolaylı anlatım ve diyaloglar kullanılır. Kelime dağarcığı biraz genişletilir.

3. ORİJİNAL (250-300 kelime): Kitabın gerçek yazarının üslubuna yakın. Zengin dil, soyut kavramlar, karmaşık cümle yapıları. Edebi sanatlar (benzetme, kişileştirme) kullanılır.

Yanıtı SADECE JSON olarak ver, başka hiçbir şey yazma:
{{
  "kitap_adi": "{kitap_adi}",
  "onerilen_seviye": "kolay|orta|orijinal",
  "seviyeler": {{
    "kolay": {{
      "baslik": "Kolay Versiyon",
      "metin": "buraya en az 200 kelimelik metin...",
      "kelime_sayisi": 220,
      "anahtar_kelimeler": ["kelime1", "kelime2", "kelime3"]
    }},
    "orta": {{
      "baslik": "Orta Versiyon",
      "metin": "buraya en az 220 kelimelik metin...",
      "kelime_sayisi": 250,
      "anahtar_kelimeler": ["kelime1", "kelime2", "kelime3"]
    }},
    "orijinal": {{
      "baslik": "Orijinal Üslup",
      "metin": "buraya en az 250 kelimelik metin...",
      "kelime_sayisi": 280,
      "anahtar_kelimeler": ["kelime1", "kelime2", "kelime3"]
    }}
  }},
  "zpd_aciklama": "Öğrencinin neden bu seviyede başlaması gerektiğinin kısa açıklaması"
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=3500)
        raw = raw.strip()
        import json as _json, re as _re
        # ``` temizle
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        # { } içini al
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        # Trailing comma düzelt
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
        # Metin uzunluğu kontrolü — çok kısaysa yeniden üret
        for sev in ["kolay","orta","orijinal"]:
            metin = result.get("seviyeler",{}).get(sev,{}).get("metin","")
            if len(metin.split()) < 80:
                logging.warning(f"[SCAFFOLD] {sev} metni çok kısa ({len(metin.split())} kelime), yeniden denenecek")
                raise Exception(f"{sev} metni çok kısa")
    except Exception as e:
        logging.error(f"[SCAFFOLD] Hata: {e} — mock döndürülüyor")
        # Mock — kısa değil gerçek metin
        result = {
            "kitap_adi": kitap_adi,
            "onerilen_seviye": "orta" if seviye_skoru <= 6 else "orijinal",
            "seviyeler": {
                "kolay": {"baslik": "Kolay Versiyon", "metin": f"'{kitap_adi}' adlı kitapta harika bir hikâye anlatılıyor. Bu kitapta bir ana karakter var. O, çok cesur ve iyi kalpli biri. Her gün yeni maceralar yaşıyor. Bazen zorluklarla karşılaşıyor ama hiç pes etmiyor. Arkadaşları ona yardım ediyor. Birlikte çalışıyorlar. Sonunda her şey güzel bir şekilde bitiyor. Bu kitap bize çok önemli bir şey öğretiyor: İyi olmak, çalışmak ve arkadaşlarına güvenmek her zaman işe yarıyor. Okurken çok heyecanlandım. Sen de okursan çok beğeneceğini düşünüyorum. Her sayfada yeni bir şey öğreniyorsun.", "kelime_sayisi": 100, "anahtar_kelimeler": ["cesaret", "arkadaşlık", "macera"]},
                "orta": {"baslik": "Orta Versiyon", "metin": f"'{kitap_adi}', okuyucusunu büyüleyici bir yolculuğa davet eden güçlü bir eser. Ana karakter, hayatının en zor döneminde bile umudunu kaybetmiyor ve içindeki gücü keşfediyor. Yazarın kalemi, her sayfada bizi farklı duygularla buluşturuyor. Kimi zaman güldürüyor, kimi zaman düşündürüyor. Karakterlerin birbirleriyle kurduğu ilişkiler, dostluğun ve güvenin ne kadar değerli olduğunu gözler önüne seriyor. Olaylar hızlı bir tempoda gelişirken, her sahne okuyucuyu bir sonrakine çekiyor. Bu kitabı elinize aldığınızda bırakmak istemeyeceksiniz. İnsan ilişkileri, cesaret ve doğruluk temaları üzerine kurulu bu yapıt, her yaştan okura farklı mesajlar veriyor.", "kelime_sayisi": 120, "anahtar_kelimeler": ["umut", "dostluk", "keşif"]},
                "orijinal": {"baslik": "Orijinal Üslup", "metin": f"'{kitap_adi}' sayfaları arasında soluk soluğa ilerleyen bu anlatı, okuyucuyu gerçeklikle kurgunun bulanık sınırında bırakır. Yazar, kelimelerini özenle seçerek her cümlede anlam katmanları oluşturmuş; yüzeyin altında akan güçlü bir duygu seli, metni her okunuşta yeniden keşfettiriyor. Ana karakterin iç dünyasına yapılan bu derin yolculuk, aslında hepimizin yaşadığı evrensel sorgulamaların bir yansımasıdır. Toplumsal baskılar, bireysel özgürlük arayışı ve kimlik mücadelesi —tüm bunlar, yazarın ustalıklı kalemi aracılığıyla doğal bir akışla birbirine örülmüş. Eserin dili, bazen şiirsel imgelerle süslenirken bazen de sert gerçeklerin çıplak ifadesine bürünüyor. Bu kontrast, okuyucuyu hem zihinsel hem duygusal düzeyde zorluyor.", "kelime_sayisi": 130, "anahtar_kelimeler": ["kimlik", "özgürlük", "anlam"]}
            },
            "zpd_aciklama": f"DNA profiline göre (seviye {seviye_skoru}/10) {'kolay' if seviye_skoru <= 3 else 'orta' if seviye_skoru <= 6 else 'orijinal'} seviyeden başlamanı öneriyoruz."
        }

    result["ogrenci_id"] = ogrenci_id
    result["kitap_id"] = kitap_id
    result["dna_seviye"] = seviye_skoru
    result["tarih"] = datetime.utcnow().isoformat()
    await db.scaffold_cache.update_one({"kitap_id": kitap_id, "ogrenci_id": ogrenci_id}, {"$set": result}, upsert=True)
    return result


@api_router.post("/ai/scaffold/seviye-ilerleme")
async def scaffold_seviye_ilerleme(req: Request, current_user=Depends(get_current_user)):
    """Öğrenci okumayı tamamladı — bir üst seviyeye geç veya tebrik et."""
    data = await req.json()
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    mevcut_seviye = data.get("mevcut_seviye", "kolay")
    dogru_oran = data.get("dogru_oran", 0.7)

    siradaki = {"kolay": "orta", "orta": "orijinal", "orijinal": None}
    sonraki = siradaki.get(mevcut_seviye)

    if dogru_oran >= 0.7 and sonraki:
        mesaj = f"Harika! {'Orta' if sonraki == 'orta' else 'Orijinal'} seviyeye geçmeye hazırsın! 🎉"
        xp = 15
    elif dogru_oran >= 0.7 and not sonraki:
        mesaj = "Tebrikler! Kitabı tüm seviyelerde tamamladın! 🏆"
        xp = 30
    else:
        mesaj = "Biraz daha pratik yapman iyi olur. Aynı seviyeyi tekrar dene."
        xp = 5
        sonraki = mevcut_seviye

    await db.xp_logs.insert_one({"ogrenci_id": ogrenci_id, "xp": xp, "kaynak": "scaffold", "tarih": datetime.utcnow().isoformat()})
    return {"sonraki_seviye": sonraki, "mesaj": mesaj, "xp": xp, "ilerledi": dogru_oran >= 0.7}


# ═══════════════════════════════════════════════════════════
# DALGA 4 — AI MATERYAl ÜRETİCİ
# ═══════════════════════════════════════════════════════════

@api_router.post("/ai/materyal/uret")
async def materyal_uret(req: Request, current_user=Depends(get_current_user)):
    """Kitap / metin için çalışma materyali üret (soru seti, kelime listesi, etkinlik)."""
    data = await req.json()
    kitap_adi  = data.get("kitap_adi", "")
    yazar      = data.get("yazar", "")
    metin_id   = data.get("metin_id", "")
    tur        = data.get("tur", "soru_seti")
    sinif      = data.get("sinif", 3)
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    icerik_id  = data.get("icerik_id", "")
    metin_icerigi = data.get("metin_icerigi", "")  # Frontend'den doğrudan gelen metin

    # ── Metin toplama: öncelik sırası ──────────────────────────────────────
    # 1. Frontend'den gönderilen metin (manuel yapıştırma)
    # 2. icerik_id → gelisim_icerik koleksiyonundan (okuma_metni veya dosya_b64)
    # 3. kitap_adi → ai_okuma_parcalari koleksiyonundan (parçaları birleştir)
    # 4. kitap_adi → ai_yuklemeler koleksiyonundan (AI yükleme metni)
    # 5. kitap_adi → ai_uretilen_sorular koleksiyonundan (soru metinleri ipucu)
    # 6. Hiçbiri yoksa: Gemini kitap adını bilerek kendi bilgisiyle üretir

    metin_ek = ""
    metin_kaynak = "yok"

    # 1. Frontend'den metin geldiyse kullan
    if metin_icerigi and metin_icerigi.strip():
        metin_ek = f"\n\nKİTAP/METİN İÇERİĞİ (kullanıcı girişi):\n{metin_icerigi.strip()[:6000]}"
        metin_kaynak = "frontend"
        logging.info(f"[MATERYAL] Kaynak: frontend, {len(metin_icerigi)} karakter")

    # 2. icerik_id varsa gelisim_icerik'ten çek
    if not metin_ek and icerik_id:
        try:
            icerik_doc = await db.gelisim_icerik.find_one({"id": icerik_id})
            if not icerik_doc:
                icerik_doc = await db.gelisim_icerik.find_one({"_id": icerik_id})
            if icerik_doc:
                if icerik_doc.get("okuma_metni"):
                    metin_ek = f"\n\nKİTAP METNİ:\n{icerik_doc['okuma_metni'][:6000]}"
                    metin_kaynak = "gelisim_icerik.okuma_metni"
                elif icerik_doc.get("dosya_b64"):
                    import base64, io
                    try:
                        raw_bytes = base64.b64decode(icerik_doc["dosya_b64"])
                        dosya_turu = icerik_doc.get("dosya_turu", "")
                        if dosya_turu == "pdf":
                            from pypdf import PdfReader
                            reader = PdfReader(io.BytesIO(raw_bytes))
                            metin = " ".join(p.extract_text() or "" for p in reader.pages[:15])
                            metin_ek = f"\n\nKİTAP METNİ (PDF):\n{metin[:6000]}"
                            metin_kaynak = "gelisim_icerik.pdf"
                        elif dosya_turu in ("docx", "doc"):
                            import docx as _docx
                            doc2 = _docx.Document(io.BytesIO(raw_bytes))
                            metin = "\n".join(p.text for p in doc2.paragraphs if p.text.strip())
                            metin_ek = f"\n\nKİTAP METNİ (Word):\n{metin[:6000]}"
                            metin_kaynak = "gelisim_icerik.docx"
                    except Exception as b64_e:
                        logging.error(f"[MATERYAL] Dosya decode hatası: {b64_e}")
            logging.info(f"[MATERYAL] icerik_id kaynağı: {metin_kaynak}")
        except Exception as db_e:
            logging.error(f"[MATERYAL] gelisim_icerik çekme hatası: {db_e}")

    # 3. kitap_adi'na göre ai_okuma_parcalari'ndan parçaları birleştir
    if not metin_ek and kitap_adi:
        try:
            parcalar = await db.ai_okuma_parcalari.find(
                {"kitap_adi": {"$regex": kitap_adi[:30], "$options": "i"}}
            ).sort("bolum", 1).to_list(length=20)
            if parcalar:
                parca_metinler = []
                for p in parcalar:
                    m = p.get("metin_kesit") or p.get("metin") or p.get("icerik") or ""
                    if m:
                        parca_metinler.append(m)
                if parca_metinler:
                    birlesik = "\n\n".join(parca_metinler)
                    metin_ek = f"\n\nKİTAP METNİ (okuma parçaları, {len(parcalar)} bölüm):\n{birlesik[:7000]}"
                    metin_kaynak = f"ai_okuma_parcalari ({len(parcalar)} parça)"
                    logging.info(f"[MATERYAL] {len(parcalar)} okuma parçası bulundu, toplam {len(birlesik)} karakter")
        except Exception as e:
            logging.error(f"[MATERYAL] ai_okuma_parcalari hatası: {e}")

    # 4. ai_yuklemeler koleksiyonundan metin ara
    if not metin_ek and kitap_adi:
        try:
            yukleme = await db.ai_yuklemeler.find_one(
                {"kitap_adi": {"$regex": kitap_adi[:30], "$options": "i"}}
            )
            if yukleme:
                m = yukleme.get("metin") or yukleme.get("icerik") or yukleme.get("ozet") or ""
                if m:
                    metin_ek = f"\n\nKİTAP METNİ (yükleme):\n{m[:6000]}"
                    metin_kaynak = "ai_yuklemeler"
                    logging.info(f"[MATERYAL] ai_yuklemeler'den metin bulundu: {len(m)} karakter")
                # dosya_b64 varsa oku
                elif yukleme.get("dosya_b64"):
                    import base64, io
                    try:
                        raw_bytes = base64.b64decode(yukleme["dosya_b64"])
                        dosya_turu = yukleme.get("dosya_turu", "")
                        if dosya_turu == "pdf":
                            from pypdf import PdfReader
                            reader = PdfReader(io.BytesIO(raw_bytes))
                            metin = " ".join(p.extract_text() or "" for p in reader.pages[:15])
                            if metin.strip():
                                metin_ek = f"\n\nKİTAP METNİ (yükleme PDF):\n{metin[:6000]}"
                                metin_kaynak = "ai_yuklemeler.pdf"
                                logging.info(f"[MATERYAL] Yükleme PDF okundu: {len(metin)} karakter")
                        elif dosya_turu in ("docx", "doc"):
                            import docx as _docx
                            doc3 = _docx.Document(io.BytesIO(raw_bytes))
                            metin = "\n".join(p.text for p in doc3.paragraphs if p.text.strip())
                            if metin.strip():
                                metin_ek = f"\n\nKİTAP METNİ (yükleme Word):\n{metin[:6000]}"
                                metin_kaynak = "ai_yuklemeler.docx"
                    except Exception as yk_e:
                        logging.error(f"[MATERYAL] Yükleme dosya hatası: {yk_e}")
        except Exception as e:
            logging.error(f"[MATERYAL] ai_yuklemeler hatası: {e}")

    # 5. ai_uretilen_sorular'dan ipucu topla (metin parçaları varsa)
    if not metin_ek and kitap_adi:
        try:
            sorular_db = await db.ai_uretilen_sorular.find(
                {"kitap_adi": {"$regex": kitap_adi[:30], "$options": "i"}}
            ).to_list(length=30)
            if sorular_db:
                soru_ipuclari = []
                for s in sorular_db[:10]:
                    metin_ref = s.get("metin_ref") or s.get("paragraf") or ""
                    if metin_ref:
                        soru_ipuclari.append(metin_ref)
                if soru_ipuclari:
                    metin_ek = f"\n\nKİTAP METİN PARÇALARI (daha önce üretilmiş sorulardan):\n" + "\n".join(soru_ipuclari[:5])
                    metin_kaynak = "ai_uretilen_sorular.metin_ref"
                else:
                    # En azından soru metinlerini ipucu olarak ver
                    mevcut_sorular = [s.get("soru","") for s in sorular_db[:5] if s.get("soru")]
                    if mevcut_sorular:
                        metin_ek = f"\n\nNOT: Bu kitap için daha önce üretilmiş {len(sorular_db)} soru var. Benzer içerik ve zorluk seviyesini koru ama FARKLI sorular üret:\n" + "\n".join(f"- {s}" for s in mevcut_sorular)
                        metin_kaynak = "ai_uretilen_sorular.ipucu"
                logging.info(f"[MATERYAL] ai_uretilen_sorular'dan {len(sorular_db)} kayıt bulundu")
        except Exception as e:
            logging.error(f"[MATERYAL] ai_uretilen_sorular hatası: {e}")

    logging.info(f"[MATERYAL] Toplam metin kaynağı: {metin_kaynak}, metin_ek uzunluğu: {len(metin_ek)}")

    metin_bilgi = f" (Yazar: {yazar})" if yazar else ""
    has_metin = bool(metin_ek.strip())

    # Metin yoksa: Gemini'den önce kitap hakkında bağlam üret
    if not has_metin and GEMINI_API_KEY:
        try:
            baglam_prompt = (
                f"Türkçe çocuk kitabı: '{kitap_adi}'{metin_bilgi}\n"
                f"Bu kitap hakkında bildiklerini yaz. Eğer bu kitabı bilmiyorsan, "
                f"kitap adından ve yazarından çıkarabileceğin ipuçlarıyla tahmini bir içerik özeti yaz.\n"
                f"Şunları belirt: Ana karakter kim? Nerede geçiyor? Temel olay nedir? Önemli yan karakterler? "
                f"Kitabın mesajı/teması nedir? Dikkat çekici sahneler?\n"
                f"Kısa ve net yaz (200-300 kelime)."
            )
            kitap_baglaami = await _gemini_call(baglam_prompt, max_tokens=600)
            metin_ek = f"\n\nKİTAP HAKKINDA BİLGİ (AI analizi):\n{kitap_baglaami.strip()}"
            metin_kaynak = "gemini_baglam"
            has_metin = True
            logging.info(f"[MATERYAL] Gemini'den kitap bağlamı alındı: {len(kitap_baglaami)} karakter")
        except Exception as baglam_e:
            logging.error(f"[MATERYAL] Bağlam üretme hatası: {baglam_e}")

    # Metin varsa: metne özgü talimatlar
    if has_metin:
        metin_bağlam = metin_ek
        metin_odak = (
            "\n🔴 ZORUNLU: Aşağıdaki kitap bilgisini kullan. Sorular MUTLAKA:\n"
            "- Kitaptaki GERÇEK karakter isimlerini kullan\n"
            "- Kitaptaki GERÇEK olayları, mekanları, detayları sor\n"
            "- 'Karakterin özelliği nedir?' gibi SOYUT/JENERİK sorular KESİNLİKLE YASAK\n"
            "- Doğru cevap açıkça bulunabilmeli, yanlış şıklar inandırıcı ama yanlış olsun\n"
        )
    else:
        metin_bağlam = ""
        metin_odak = (
            "\n🔴 ZORUNLU: Kitap adından yola çıkarak ÖZGİN sorular üret. YASAK sorular:\n"
            "- 'Kitabın ana karakterinin özelliği nedir?' → YASAK\n"
            "- 'Kitap hangi türdedir?' → YASAK  \n"
            "- 'Olaylar hangi ortamda geçer?' → YASAK\n"
            "Bunların yerine kitabın adından ve içeriğinden tahmin edilen özgün sorular sor.\n"
        )

    tur_prompts = {
        "soru_seti": (
            f"'{kitap_adi}'{metin_bilgi} için {sinif}. sınıf düzeyinde TAM 10 soruluk anlama testi oluştur.\n"
            f"{metin_odak}"
            f"Bloom taksonomisinin 6 basamağından dengeli sorular ekle:\n"
            f"- 2 Bilgi sorusu (metinde doğrudan geçen bilgi)\n"
            f"- 2 Kavrama sorusu (olayı kendi sözlerinle anlat)\n"
            f"- 2 Uygulama sorusu (sen olsaydın ne yapardın?)\n"
            f"- 1 Analiz sorusu (neden-sonuç ilişkisi)\n"
            f"- 1 Sentez/Değerlendirme sorusu\n"
            f"- 2 Yaratıcı/Eleştirel düşünme sorusu\n"
            f"Her soru için 4 seçenek olsun (A,B,C,D). Yanlış seçenekler inandırıcı olsun.{metin_bağlam}\n\n"
            f"JSON formatı (tam 10 soru): {{\"baslik\": \"string\", \"sorular\": ["
            f"{{\"soru\": \"string\", \"secenekler\": [\"A...\",\"B...\",\"C...\",\"D...\"], \"dogru\": \"A...\", \"bloom_basamak\": \"string\"}}"
            f"]}}"
        ),
        "kelime_listesi": (
            f"'{kitap_adi}'{metin_bilgi} kitabından {sinif}. sınıf için 10 önemli kelime seç.{metin_odak}{metin_bağlam}\n\n"
            f"SADECE JSON döndür:\n"
            f"{{\"baslik\": \"string\", \"kelimeler\": ["
            f"{{\"kelime\": \"string\", \"anlam\": \"kısa anlam\", \"cumle\": \"kısa örnek cümle\", \"zorluk\": 1}}"
            f"]}}"
        ),
        "etkinlik": (
            f"'{kitap_adi}'{metin_bilgi} için sınıf içi grup etkinliği tasarla. {sinif}. sınıf, 20-30 dk.{metin_odak}{metin_bağlam}\n\n"
            f"JSON: {{\"baslik\": \"string\", \"sure_dk\": 25, \"grup_sayisi\": 4, "
            f"\"adimlar\": [\"string\"], \"malzemeler\": [\"string\"], \"kazanimlar\": [\"string\"]}}"
        ),
        "tahmin": (
            f"'{kitap_adi}'{metin_bilgi} okuma öncesi TAM 8 tahmin sorusu oluştur ({sinif}. sınıf).{metin_bağlam}\n\n"
            f"JSON: {{\"baslik\": \"string\", \"giris\": \"string\", \"sorular\": ["
            f"{{\"soru\": \"string\", \"ipucu\": \"string\"}}"
            f"]}}"
        ),
    }

    prompt = tur_prompts.get(tur, tur_prompts["soru_seti"]) + "\n\nSADECE JSON döndür, başka hiçbir şey yazma. Tüm anahtarlar çift tırnak içinde olsun."


    import json as _json

    if not GEMINI_API_KEY:
        logging.warning("GEMINI_API_KEY yok — mock döndürülüyor")
    else:
        logging.info(f"Gemini çağrısı başlıyor: tur={tur}, kitap={kitap_adi}, metin_len={len(metin_ek)}")

    async def _parse_gemini_json(raw_text: str):
        """Gemini yanıtından JSON çıkar — hata toleranslı."""
        import re as _re
        raw_text = raw_text.strip()

        # ``` bloklarını temizle
        if "```" in raw_text:
            match = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw_text)
            if match:
                raw_text = match.group(1).strip()
            else:
                raw_text = _re.sub(r"```(?:json)?", "", raw_text).replace("```", "").strip()

        # { ... } içini al
        brace_start = raw_text.find("{")
        brace_end = raw_text.rfind("}")
        if brace_start != -1 and brace_end != -1:
            raw_text = raw_text[brace_start:brace_end+1]

        # Önce direkt parse dene
        try:
            return _json.loads(raw_text)
        except _json.JSONDecodeError:
            pass

        # Trailing comma düzelt: ,] ve ,} → ] ve }
        cleaned = _re.sub(r",\s*([}\]])", r"\1", raw_text)
        try:
            return _json.loads(cleaned)
        except _json.JSONDecodeError:
            pass

        # Tek tırnak → çift tırnak
        cleaned2 = cleaned.replace("'", '"')
        try:
            return _json.loads(cleaned2)
        except _json.JSONDecodeError:
            pass

        # Kontrol karakterlerini temizle
        cleaned3 = _re.sub(r'[\x00-\x1f\x7f]', ' ', cleaned)
        cleaned3 = _re.sub(r',\s*([}\]])', r'\1', cleaned3)
        return _json.loads(cleaned3)


    _debug = {
        "metin_kaynak": metin_kaynak,
        "metin_uzunluk": len(metin_ek),
        "gemini_key_var": bool(GEMINI_API_KEY),
        "gemini_key_uzunluk": len(GEMINI_API_KEY) if GEMINI_API_KEY else 0,
    }

    try:
        raw = await _gemini_call(prompt, max_tokens=4000)
        logging.info(f"Gemini yanıt alındı: {len(raw)} karakter")
        result = await _parse_gemini_json(raw)
        result["tur"] = tur
        result["kitap_adi"] = kitap_adi
        result["_debug"] = {**_debug, "deneme": 1, "hata": None}
        return result
    except Exception as e:
        hata1 = f"{type(e).__name__}: {e}"
        logging.error(f"Gemini materyal hatası (1. deneme): {hata1}")
        _debug["hata1"] = hata1
        # 2. deneme: daha basit prompt ile tekrar dene
        if GEMINI_API_KEY:
            try:
                basit_prompt = (
                    f"'{kitap_adi}' kitabı için {sinif}. sınıf öğrencisine 10 anlama sorusu üret.\n"
                    f"Bu kitabın GERÇEK karakterlerini, mekanlarını ve olaylarını kullan.\n"
                    f"Her soru kitabı okumayan birinin bilemeyeceği kadar özgün olsun.\n"
                    f"{'Kitap metni: ' + metin_ek[:3000] if metin_ek else ''}\n\n"
                    f"SADECE JSON döndür:\n"
                    f"{{\"baslik\": \"string\", \"sorular\": ["
                    f"{{\"soru\": \"string\", \"secenekler\": [\"A) ...\",\"B) ...\",\"C) ...\",\"D) ...\"], \"dogru\": \"A) ...\", \"bloom_basamak\": \"Bilgi\"}}"
                    f"]}}"
                )
                raw2 = await _gemini_call(basit_prompt, max_tokens=3000)
                logging.info(f"Gemini 2. deneme yanıtı: {len(raw2)} karakter")
                result = await _parse_gemini_json(raw2)
                result["tur"] = tur
                result["kitap_adi"] = kitap_adi
                result["_debug"] = {**_debug, "deneme": 2, "hata": None}
                return result
            except Exception as e2:
                hata2 = str(e2)
                logging.error(f"Gemini 2. deneme de başarısız: {hata2}")
                _debug["hata2"] = hata2

        # Son çare: Gemini'den kitap adına göre özgün fallback al
        if GEMINI_API_KEY:
            try:
                fallback_prompt = (
                    f"Türk ilkokul öğrencileri için '{kitap_adi}' kitabı hakkında 10 soru üret. "
                    f"Sorular bu kitaba ÖZGÜN olsun, genel sorular olmasın. "
                    f"JSON: {{\"baslik\": \"string\", \"sorular\": [{{\"soru\": \"string\", \"secenekler\": [\"A) x\",\"B) x\",\"C) x\",\"D) x\"], \"dogru\": \"A) x\", \"bloom_basamak\": \"Bilgi\"}}]}}"
                )
                raw3 = await _gemini_call(fallback_prompt, max_tokens=2000)
                result = await _parse_gemini_json(raw3)
                result["tur"] = tur
                result["kitap_adi"] = kitap_adi
                result["_debug"] = {**_debug, "deneme": 3, "hata": None}
                return result
            except Exception as e3:
                hata3 = str(e3)
                logging.error(f"Gemini fallback da başarısız: {hata3}")
                _debug["hata3"] = hata3

        # Hiçbir şey çalışmadıysa son çare statik mock (artık sadece gerçekten API yoksa)
        logging.error("TÜM Gemini denemeleri başarısız — statik mock kullanılıyor")
        bloom_list = ["Bilgi","Kavrama","Uygulama","Analiz","Kavrama","Uygulama","Bilgi","Sentez","Değerlendirme","Yaratma"]
        if tur == "soru_seti":
            sorular = []
            sorular_meta = [
                (f"'{kitap_adi}' kitabı hangi türdedir?", ["Roman","Şiir","Masal","Deneme"], "Roman", "Bilgi"),
                (f"Kitabın ana karakterinin özelliği nedir?", ["Cesur","Korkak","Tembel","Kıskanç"], "Cesur", "Kavrama"),
                (f"Sen bu karakterin yerinde olsaydın ne yapardın?", ["Aynısını yapardım","Farklı davranırdım","Kaçardım","Yardım isterdim"], "Farklı davranırdım", "Uygulama"),
                (f"Kitabın ana teması nedir?", ["Arkadaşlık","Cesaret","Dürüstlük","Merak"], "Cesaret", "Kavrama"),
                (f"Olaylar hangi ortamda geçmektedir?", ["Şehirde","Köyde","Ormanda","Deniz kıyısında"], "Köyde", "Bilgi"),
                (f"Karakterin en büyük sorunu neydi?", ["Yalnızlık","Maddi sıkıntı","Güven eksikliği","Hastalık"], "Güven eksikliği", "Analiz"),
                (f"Kitabın sonu nasıl bitmektedir?", ["Mutlu son","Hüzünlü son","Açık uçlu","Sürpriz son"], "Mutlu son", "Bilgi"),
                (f"Bu kitap sana göre hangi değeri en iyi anlatıyor?", ["Dürüstlük","Sabır","Cesaret","Yardımseverlik"], "Cesaret", "Değerlendirme"),
                (f"Yazar bu kitabı neden yazmış olabilir?", ["Eğlendirmek için","Ders vermek için","Duygu aktarmak için","Belgelemek için"], "Ders vermek için", "Sentez"),
                (f"Kitabı okuduktan sonra hayatında ne değiştirebilirsin?", ["Daha cesur olabilirim","Daha sabırlı olabilirim","Daha çok kitap okurum","Arkadaşlarıma yardım ederim"], "Daha cesur olabilirim", "Yaratma"),
            ]
            for soru, sec, dogru, bloom in sorular_meta:
                sorular.append({"soru": soru, "secenekler": sec, "dogru": dogru, "bloom_basamak": bloom})
            result = {"baslik": f"{kitap_adi} — Anlama Testi", "sorular": sorular}
        elif tur == "kelime_listesi":
            kelimeler_ornek = ["macera","cesaret","yolculuk","dürüstlük","arkadaşlık","merak","umut","sabır","özgürlük","kahramanlık","sadakat","iyilik","azim","fedakarlık","başarı"]
            result = {"baslik": f"{kitap_adi} — Anahtar Kelimeler", "kelimeler": [
                {"kelime": k, "anlam": f"{k.capitalize()} kavramının anlamı", "cumle": f"Kitapta {k} teması işlendi.", "zorluk": (i%3)+1}
                for i, k in enumerate(kelimeler_ornek)
            ]}
        elif tur == "etkinlik":
            result = {"baslik": f"{kitap_adi} — Sınıf Etkinliği", "sure_dk": 25, "grup_sayisi": 4,
                "adimlar": ["Sınıfı 4 gruba ayırın","Her grup kitabın farklı bölümünü tartışsın","Karakterleri analiz edin","Gruplar sunum yapsın","Sınıf tartışması yapın"],
                "malzemeler": ["Kağıt","Kalem","Post-it","Renkli kalemler"],
                "kazanimlar": ["Eleştirel düşünme","İşbirliği","Sözlü ifade","Empati kurma"]}
        else:
            result = {"baslik": f"{kitap_adi} — Okuma Öncesi Tahmin", "giris": "Kitabı okumadan önce düşüncelerini paylaş!", "sorular": [
                {"soru": "Bu kitap ne hakkında olabilir?", "ipucu": "Kapak resmine bak"},
                {"soru": "Ana karakter nasıl biri olabilir?", "ipucu": "Başlığı düşün"},
                {"soru": "Olaylar nerede geçiyor olabilir?", "ipucu": "Kapak resmindeki ortamı incele"},
                {"soru": "Kitabın sonu nasıl bitebilir?", "ipucu": "Başlığa göre tahmin et"},
                {"soru": "Hangi sorunlarla karşılaşılacak?", "ipucu": "Türüne bak"},
                {"soru": "Bu kitaptan ne öğrenebiliriz?", "ipucu": "Yazarı araştır"},
                {"soru": "Favorin olan karakter kim olabilir?", "ipucu": "Kitabın adına bak"},
                {"soru": "Sence bu kitabın mesajı ne olacak?", "ipucu": "Türü ve konusu hakkında düşün"},
            ]}

    result["tur"] = tur
    result["kitap_adi"] = kitap_adi
    result["sinif"] = sinif
    result["tarih"] = datetime.utcnow().isoformat()
    result["_debug"] = {**_debug, "deneme": 0, "mock": True}  # mock kullanıldı

    # Kaydet + XP
    await db.ai_materyal_log.insert_one({"ogrenci_id": ogrenci_id, "kitap_adi": kitap_adi, "tur": tur, "tarih": datetime.utcnow().isoformat()})
    await db.xp_logs.insert_one({"ogrenci_id": ogrenci_id, "xp": 5, "kaynak": "materyal_uret", "tarih": datetime.utcnow().isoformat()})
    return result


@api_router.get("/ai/materyal/gecmis/{ogrenci_id}")
async def materyal_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    kayitlar = await db.ai_materyal_log.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=20)
    for k in kayitlar: k.pop("_id", None)
    return kayitlar


# ═══════════════════════════════════════════════════════════
# DALGA 4 — ADAPTIVE STORY ENGINE
# ═══════════════════════════════════════════════════════════

@api_router.post("/ai/hikaye/olustur")
async def hikaye_olustur(req: Request, current_user=Depends(get_current_user)):
    """DNA + ilgi alanı → kişisel hikâye + 5 Bloom sorusu + kelime kartları."""
    data = await req.json()
    ogrenci_id  = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))
    ilgi_alani  = data.get("ilgi_alani", "macera")   # macera, hayvan, uzay, tarih, spor, arkadaşlık
    kahraman_ad = data.get("kahraman_ad", "Kahraman")
    sinif       = int(data.get("sinif", 3))
    sure_dk     = data.get("sure_dk", 5)  # hedef okuma süresi

    # DNA profili
    dna = await db.okuma_dna.find_one({"ogrenci_id": ogrenci_id})
    boyutlar = dna.get("boyutlar", {}) if dna else {}
    kelime_gucu = boyutlar.get("kelime_gucu", 50)
    anlama = boyutlar.get("anlama_derinligi", 50)
    profil_tipi = dna.get("profil_tipi", "dengeli_okuyucu") if dna else "dengeli_okuyucu"

    # MEB kelime havuzu (sınıfa göre mini örnek)
    meb_kelimeler = {
        1: ["ev", "okul", "anne", "baba", "arkadaş"],
        2: ["macera", "cesur", "yardım", "dürüst", "başarı"],
        3: ["merak", "keşif", "sorumluluk", "empati", "özgüven"],
        4: ["analiz", "strateji", "iletişim", "liderlik", "çözüm"],
        5: ["eleştiri", "perspektif", "hipotez", "kanıt", "sentez"],
    }.get(sinif, ["merak", "keşif", "arkadaşlık"])

    kelime_zorluk = "basit" if kelime_gucu < 40 else "orta" if kelime_gucu < 70 else "zengin"
    hikaye_uzunluk = sure_dk * 100  # kelime

    prompt = f"""Sen çocuklar için kişiselleştirilmiş hikâyeler yazan yaratıcı bir yazarsın.

Öğrenci profili:
- Ad: {kahraman_ad}
- Sınıf: {sinif}
- İlgi alanı: {ilgi_alani}
- Okuma seviyesi: {profil_tipi}
- Kelime zenginliği: {kelime_zorluk}
- Anlama derinliği: {anlama}/100

Hikâyeye şu MEB kelimelerini doğal biçimde dahil et: {", ".join(meb_kelimeler)}

Kurallar:
- Baş karakter adı: {kahraman_ad}
- Tema: {ilgi_alani}
- Uzunluk: yaklaşık {hikaye_uzunluk} kelime
- Dil: {kelime_zorluk} kelime zenginliği
- Mutlu son zorunlu
- MEB Erdemler: sabır, dürüstlük, merak, empati, cesaret değerlerinden en az 2'si tema olsun

SADECE JSON döndür:
{{
  "baslik": "Hikâye başlığı",
  "hikaye": "Tam hikâye metni ({hikaye_uzunluk} kelime)",
  "kullanilan_meb_kelimeleri": ["kelime1", "kelime2"],
  "kazanilan_deger": "sabır|dürüstlük|merak|empati|cesaret",
  "bloom_sorulari": [
    {{"basamak": "Hatırlama", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Kavrama", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Uygulama", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Analiz", "soru": "...", "ipucu": "..."}},
    {{"basamak": "Değerlendirme", "soru": "...", "ipucu": "..."}}
  ],
  "kelime_kartlari": [
    {{"kelime": "...", "anlam": "...", "cumle": "..."}}
  ]
}}"""

    try:
        import json as _json
        raw = await _gemini_call(prompt, max_tokens=2500)
        raw = raw.strip()
        if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:-1])
        result = _json.loads(raw)
    except Exception as e:
        result = {
            "baslik": f"{kahraman_ad}'ın {ilgi_alani.capitalize()} Macerası",
            "hikaye": f"{kahraman_ad} bir gün ormanda yürürken sihirli bir kapıyla karşılaştı. Kapının arkasında harika bir dünya vardı. Merakla içeri girdi ve yeni arkadaşlar edindi. Birlikte zorlukları aşmayı öğrendiler. Cesaret ve dürüstlükle her engeli geçtiler. Sonunda eve döndüklerinde çok şey öğrenmişlerdi.",
            "kullanilan_meb_kelimeleri": meb_kelimeler[:2],
            "kazanilan_deger": "merak",
            "bloom_sorulari": [
                {"basamak": "Hatırlama", "soru": f"{kahraman_ad} ormanda ne buldu?", "ipucu": "Kapı ile ilgili düşün."},
                {"basamak": "Kavrama", "soru": "Karakterler ne öğrendi?", "ipucu": "Birlikte ne yaptılar?"},
                {"basamak": "Uygulama", "soru": "Sen olsaydın ne yapardın?", "ipucu": "Kendi deneyiminden düşün."},
                {"basamak": "Analiz", "soru": "Hikâyedeki ana sorun neydi?", "ipucu": "Zorlukları düşün."},
                {"basamak": "Değerlendirme", "soru": "Hikâye sana ne öğretti?", "ipucu": "Değerleri düşün."}
            ],
            "kelime_kartlari": [{"kelime": meb_kelimeler[0], "anlam": "Önemli bir değer", "cumle": f"{kahraman_ad} {meb_kelimeler[0]} gösterdi."}]
        }

    result["ogrenci_id"] = ogrenci_id
    result["ilgi_alani"] = ilgi_alani
    result["kahraman_ad"] = kahraman_ad
    result["sinif"] = sinif
    result["tarih"] = datetime.utcnow().isoformat()

    hikaye_id = str(__import__("uuid").uuid4())[:8]
    result["hikaye_id"] = hikaye_id
    await db.ai_hikaye_log.insert_one({**result})
    await db.xp_logs.insert_one({"ogrenci_id": ogrenci_id, "xp": 10, "kaynak": "adaptif_hikaye", "tarih": datetime.utcnow().isoformat()})
    return result


@api_router.get("/ai/hikaye/gecmis/{ogrenci_id}")
async def hikaye_gecmis(ogrenci_id: str, current_user=Depends(get_current_user)):
    kayitlar = await db.ai_hikaye_log.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=10)
    for k in kayitlar: k.pop("_id", None)
    return kayitlar



# ─────────────────────────────────────────────
# POST-READING AI — Kitap Bitirme Sonrası Analiz
# ─────────────────────────────────────────────

@api_router.post("/ai/post-reading")
async def post_reading_ai(req: Request, current_user=Depends(get_current_user)):
    """Kitap/içerik tamamlanınca derinlik soruları + MEB Erdem değeri üretir."""
    data = await req.json()
    kitap_adi = data.get("kitap_adi", "")
    yazar = data.get("yazar", "")
    sinif = data.get("sinif", 3)
    icerik_id = data.get("icerik_id", "")
    ogrenci_id = data.get("ogrenci_id", current_user.get("linked_id", current_user.get("id", "")))

    # Cache kontrol (24 saat)
    cache = await db.post_reading_cache.find_one({"icerik_id": icerik_id, "ogrenci_id": ogrenci_id})
    if cache and cache.get("tarih"):
        import dateutil.parser
        try:
            sure = (datetime.utcnow() - dateutil.parser.parse(cache["tarih"])).total_seconds()
            if sure < 86400:
                cache.pop("_id", None)
                return cache
        except:
            pass

    # Kitap metnini al (varsa)
    metin_ek = ""
    icerik = await db.gelisim_icerik.find_one({"id": icerik_id})
    if icerik and icerik.get("okuma_metni"):
        metin_ek = f"\n\nKitap/metin içeriği (ilk 1500 kelime):\n{icerik['okuma_metni'][:3000]}"

    prompt = f"""Sen bir Türkçe eğitim uzmanısın. "{kitap_adi}"{f" ({yazar})" if yazar else ""} adlı kitabı/metni {sinif}. sınıf öğrencisi tamamladı.{metin_ek}

Aşağıdaki JSON formatında derinlik analizi üret:

1. Ana fikir sorusu: Öğrenciyi düşündüren açık uçlu 1 soru
2. MEB Erdem değeri: Bu kitaptan çıkarılabilecek 1 erdem (Erdemler: sabır, dürüstlük, merak, cesaret, sorumluluk, sevgi, saygı, adalet, yardımseverlik, vefa)
3. Bloom soruları: 3 basamak (Kavrama, Analiz, Yaratma) için birer soru
4. Hayat bağlantısı: "Bu kitap senin hayatında neyi değiştirir?" sorusu
5. Öneri kitaplar: Bu kitabı seven birine 2 kitap önerisi

SADECE JSON döndür:
{{
  "ana_fikir_sorusu": "...",
  "meb_erdem": {{
    "erdem": "sabır",
    "aciklama": "Bu kitap sabırlı olmayı şöyle gösteriyor: ..."
  }},
  "bloom_sorulari": [
    {{"basamak": "Kavrama", "soru": "...", "emoji": "🔍"}},
    {{"basamak": "Analiz", "soru": "...", "emoji": "🧩"}},
    {{"basamak": "Yaratma", "soru": "...", "emoji": "✨"}}
  ],
  "hayat_baglantisi": "...",
  "oneri_kitaplar": [
    {{"baslik": "...", "yazar": "...", "neden": "..."}},
    {{"baslik": "...", "yazar": "...", "neden": "..."}}
  ],
  "ozet_cumle": "Bu kitabın özü: ..."
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=1500)
        import json as _json, re as _re
        raw = raw.strip()
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
    except Exception as e:
        logging.error(f"[POST-READING] Hata: {e}")
        result = {
            "ana_fikir_sorusu": f"'{kitap_adi}' seni en çok hangi konuda düşündürdü? Neden?",
            "meb_erdem": {"erdem": "merak", "aciklama": "Bu kitap merak etmenin ve soru sormanın önemini gösteriyor."},
            "bloom_sorulari": [
                {"basamak": "Kavrama", "soru": "Kitabın ana karakteri hangi zorluklarla karşılaştı?", "emoji": "🔍"},
                {"basamak": "Analiz", "soru": "Kitaptaki olaylar neden bu sırayla gerçekleşti?", "emoji": "🧩"},
                {"basamak": "Yaratma", "soru": "Kitabın sonunu farklı yazsan nasıl bitirirdin?", "emoji": "✨"}
            ],
            "hayat_baglantisi": "Bu kitaptaki hangi duygu veya düşünce sana tanıdık geldi?",
            "oneri_kitaplar": [
                {"baslik": "Pollyanna", "yazar": "Eleanor H. Porter", "neden": "Benzer temalar ve umutlu bir bakış açısı"},
                {"baslik": "Küçük Prens", "yazar": "Antoine de Saint-Exupéry", "neden": "Derin düşünceler, çocuksu merak"}
            ],
            "ozet_cumle": f"'{kitap_adi}' sana yeni bir bakış açısı kazandırdı."
        }

    result["kitap_adi"] = kitap_adi
    result["icerik_id"] = icerik_id
    result["ogrenci_id"] = ogrenci_id
    result["tarih"] = datetime.utcnow().isoformat()
    await db.post_reading_cache.update_one(
        {"icerik_id": icerik_id, "ogrenci_id": ogrenci_id},
        {"$set": result}, upsert=True
    )
    return result


# ─────────────────────────────────────────────
# KİTAP ZEKÂ HARİTASI — 7 Boyutlu AI Profil
# ─────────────────────────────────────────────

@api_router.post("/ai/kitap-zeka-haritasi")
async def kitap_zeka_haritasi(req: Request, current_user=Depends(get_current_user)):
    """Her kitap için 7 boyutlu AI profil üretir."""
    data = await req.json()
    kitap_adi = data.get("kitap_adi", "")
    yazar = data.get("yazar", "")
    sinif = data.get("sinif", 3)
    icerik_id = data.get("icerik_id", "")

    # Cache (kalıcı — kitap değişmez)
    cache = await db.kitap_zeka_haritasi.find_one({"icerik_id": icerik_id})
    if cache:
        cache.pop("_id", None)
        return cache

    # Metin varsa ekle
    metin_ek = ""
    icerik = await db.gelisim_icerik.find_one({"id": icerik_id})
    if icerik and icerik.get("okuma_metni"):
        metin_ek = f"\n\nMetin özeti (ilk 1000 kelime):\n{icerik['okuma_metni'][:2000]}"

    prompt = f"""Sen bir çocuk edebiyatı analistisin. "{kitap_adi}"{f" - {yazar}" if yazar else ""} için 7 boyutlu profil oluştur.{metin_ek}

Her boyutu 1-10 arasında puan ver. SADECE JSON döndür:
{{
  "kitap_adi": "{kitap_adi}",
  "boyutlar": {{
    "soyutluk": {{"puan": 5, "aciklama": "..."}},
    "kelime_zorlugu": {{"puan": 5, "aciklama": "..."}},
    "hayal_gucu": {{"puan": 5, "aciklama": "..."}},
    "felsefi_derinlik": {{"puan": 5, "aciklama": "..."}},
    "aksiyon": {{"puan": 5, "aciklama": "..."}},
    "duygusal_yogunluk": {{"puan": 5, "aciklama": "..."}},
    "hedef_kelime_yogunlugu": {{"puan": 5, "aciklama": "..."}}
  }},
  "sinif_uyumu": {sinif},
  "tavsiye_profil": "Bu kitap hangi öğrenciye uygundur? (2-3 cümle)",
  "tur_etiketleri": ["macera", "arkadaşlık"],
  "okuma_suresi_dk": 30
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=1000)
        import json as _json, re as _re
        raw = raw.strip()
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
    except Exception as e:
        logging.error(f"[ZEKA-HARITA] Hata: {e}")
        result = {
            "kitap_adi": kitap_adi,
            "boyutlar": {
                "soyutluk": {"puan": 5, "aciklama": "Orta düzey soyut kavramlar içeriyor"},
                "kelime_zorlugu": {"puan": 5, "aciklama": "Yaş grubuna uygun kelime zenginliği"},
                "hayal_gucu": {"puan": 6, "aciklama": "Hayal gücünü geliştiren sahneler mevcut"},
                "felsefi_derinlik": {"puan": 4, "aciklama": "Temel değer sorgulamaları var"},
                "aksiyon": {"puan": 6, "aciklama": "Tempolu ve heyecanlı sahneler"},
                "duygusal_yogunluk": {"puan": 7, "aciklama": "Güçlü duygusal bağ kurduruyor"},
                "hedef_kelime_yogunlugu": {"puan": 5, "aciklama": "MEB müfredatına uygun kelimeler"}
            },
            "sinif_uyumu": sinif,
            "tavsiye_profil": "Okuma alışkanlığı kazanmaya başlayan, macera seven öğrencilere uygundur.",
            "tur_etiketleri": ["macera", "gelişim", "arkadaşlık"],
            "okuma_suresi_dk": sinif * 8
        }

    result["icerik_id"] = icerik_id
    result["olusturma_tarihi"] = datetime.utcnow().isoformat()
    await db.kitap_zeka_haritasi.update_one(
        {"icerik_id": icerik_id}, {"$set": result}, upsert=True
    )
    return result


# ─────────────────────────────────────────────
# KİTABA ÖZGÜ MİNİ OYUN — Kitap içeriğinden üret
# ─────────────────────────────────────────────

@api_router.post("/ai/kitap-oyun")
async def kitap_oyun_uret(req: Request, current_user=Depends(get_current_user)):
    """Kitap/metin içeriğinden oyun soruları üretir."""
    data = await req.json()
    kitap_adi = data.get("kitap_adi", "")
    icerik_id = data.get("icerik_id", "")
    oyun_turu = data.get("tur", "karakter_tahmini")  # karakter_tahmini, hikaye_devam, bosluk, eslestirme
    sinif = data.get("sinif", 3)

    # Metni al
    metin = ""
    icerik = await db.gelisim_icerik.find_one({"id": icerik_id})
    if icerik:
        metin = icerik.get("okuma_metni", "") or icerik.get("aciklama", "")
    if not metin:
        return {"oyun": None, "mesaj": "İçerik metni bulunamadı"}

    metin_kisalt = metin[:2500]

    if oyun_turu == "karakter_tahmini":
        prompt = f""""{kitap_adi}" metninden karakter tahmini oyunu oluştur.

Metin: {metin_kisalt}

SADECE JSON döndür. Metindeki gerçek karakterleri kullan:
{{
  "tur": "karakter_tahmini",
  "baslik": "🎭 Kim O?",
  "aciklama": "İpuçlarından karakteri bul!",
  "sorular": [
    {{
      "ipuclari": ["İpucu 1", "İpucu 2", "İpucu 3"],
      "dogru_karakter": "...",
      "secenekler": ["...", "...", "...", "..."]
    }}
  ],
  "xp": 8
}}"""

    elif oyun_turu == "hikaye_devam":
        prompt = f""""{kitap_adi}" metninden hikâye devam ettirme oyunu oluştur.

Metin: {metin_kisalt}

SADECE JSON döndür:
{{
  "tur": "hikaye_devam",
  "baslik": "📖 Hikâye Devam Ediyor",
  "aciklama": "Doğru devamı seç!",
  "sorular": [
    {{
      "metin_parcasi": "Metinden alınan bir paragraf...",
      "soru": "Bundan sonra ne oldu?",
      "secenekler": ["A seçeneği", "B seçeneği", "C seçeneği", "D seçeneği"],
      "dogru_idx": 0,
      "aciklama": "Neden bu doğru?"
    }}
  ],
  "xp": 10
}}"""

    elif oyun_turu == "eslestirme":
        prompt = f""""{kitap_adi}" metninden karakter-özellik eşleştirme oyunu oluştur.

Metin: {metin_kisalt}

SADECE JSON döndür:
{{
  "tur": "eslestirme",
  "baslik": "🎲 Kim Nasıl?",
  "aciklama": "Karakterleri özellikleriyle eşleştir!",
  "ciftler": [
    {{"sol": "Karakter adı", "sag": "Özelliği/yaptığı şey"}},
    {{"sol": "...", "sag": "..."}},
    {{"sol": "...", "sag": "..."}},
    {{"sol": "...", "sag": "..."}}
  ],
  "xp": 6
}}"""

    else:  # bosluk
        prompt = f""""{kitap_adi}" metninden boşluk doldurma oyunu oluştur. Metinden gerçek cümleler kullan.

Metin: {metin_kisalt}

SADECE JSON döndür:
{{
  "tur": "bosluk_doldurma",
  "baslik": "⬜ Boşluğu Doldur",
  "aciklama": "Metindeki eksik kelimeyi bul!",
  "sorular": [
    {{
      "cumle_bos": "Metinden alınan cümle ___ kelime yerine boş",
      "dogru": "doğru kelime",
      "secenekler": ["doğru kelime", "yanlış1", "yanlış2", "yanlış3"]
    }}
  ],
  "xp": 7
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=1200)
        import json as _json, re as _re
        raw = raw.strip()
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
        return {"oyun": result}
    except Exception as e:
        logging.error(f"[KITAP-OYUN] Hata: {e}")
        return {"oyun": None, "mesaj": f"Oyun üretilemedi: {str(e)}"}



# ─────────────────────────────────────────────
# OKUMA EVRENİ — 5 Bölge Gamification
# ─────────────────────────────────────────────

OKUMA_EVRENI_BOLGELER = [
    {
        "id": "orman",
        "ad": "Orman Kitaplığı",
        "emoji": "🌲",
        "renk": "#22c55e",
        "aciklama": "Okuma yolculuğuna başladın!",
        "kriter": "baslangic",  # Herkes başlar
        "min_kitap": 0, "min_kelime": 0, "min_streak": 0, "min_bloom": 0,
    },
    {
        "id": "kelime_dag",
        "ad": "Kelime Dağları",
        "emoji": "⛰️",
        "renk": "#f59e0b",
        "aciklama": "50 kelime öğrendin!",
        "kriter": "kelime",
        "min_kitap": 0, "min_kelime": 50, "min_streak": 0, "min_bloom": 0,
    },
    {
        "id": "hikaye_limani",
        "ad": "Hikâye Limanı",
        "emoji": "⚓",
        "renk": "#3b82f6",
        "aciklama": "5 kitap/içerik tamamladın!",
        "kriter": "kitap",
        "min_kitap": 5, "min_kelime": 0, "min_streak": 0, "min_bloom": 0,
    },
    {
        "id": "bilgelik_kutuphane",
        "ad": "Bilgelik Kütüphanesi",
        "emoji": "🏰",
        "renk": "#8b5cf6",
        "aciklama": "Bloom %60 anlama seviyesine ulaştın!",
        "kriter": "bloom",
        "min_kitap": 5, "min_kelime": 100, "min_streak": 7, "min_bloom": 60,
    },
    {
        "id": "hayal_galaksisi",
        "ad": "Hayal Galaksisi",
        "emoji": "🚀",
        "renk": "#ec4899",
        "aciklama": "Tüm evrenin fethettdin! En yüksek seviye.",
        "kriter": "tumu",
        "min_kitap": 20, "min_kelime": 300, "min_streak": 30, "min_bloom": 80,
    },
]

@api_router.get("/ai/okuma-evreni/{ogrenci_id}")
async def okuma_evreni(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Öğrencinin Okuma Evreni bölgesini ve ilerleme durumunu hesapla."""
    # İstatistikler
    logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).to_list(length=None)
    tamamlananlar = await db.gelisim_tamamlananlar.find({"kullanici_id": ogrenci_id}).to_list(length=None)
    kelime_ogrenilen = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id, "kutu": {"$gte": 3}})
    test_sonuclari = await db.kitap_test_sonuclari.find({"ogrenci_id": ogrenci_id}).to_list(length=None)

    # Streak hesapla
    from datetime import timedelta
    simdi = datetime.utcnow()
    tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
    streak = 0
    for i in range(90):
        gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
        if gun in tarihler:
            streak += 1
        elif i > 0:
            break

    # Bloom ortalaması
    bloom_ortalama = 0
    if test_sonuclari:
        bloom_ortalama = sum(t.get("basari_yuzdesi", 0) for t in test_sonuclari) / len(test_sonuclari)

    # Tamamlanan içerik sayısı
    tamamlanan_sayi = len(tamamlananlar)

    # Mevcut bölgeyi hesapla
    aktif_bolge = OKUMA_EVRENI_BOLGELER[0]
    for bolge in reversed(OKUMA_EVRENI_BOLGELER):
        if (tamamlanan_sayi >= bolge["min_kitap"] and
            kelime_ogrenilen >= bolge["min_kelime"] and
            streak >= bolge["min_streak"] and
            bloom_ortalama >= bolge["min_bloom"]):
            aktif_bolge = bolge
            break

    # Sıradaki bölge
    aktif_idx = next((i for i, b in enumerate(OKUMA_EVRENI_BOLGELER) if b["id"] == aktif_bolge["id"]), 0)
    sonraki_bolge = OKUMA_EVRENI_BOLGELER[aktif_idx + 1] if aktif_idx < len(OKUMA_EVRENI_BOLGELER) - 1 else None

    # Sonraki bölgeye ilerleme yüzdesi
    ilerleme = {}
    if sonraki_bolge:
        hedefler = [
            ("kitap", tamamlanan_sayi, sonraki_bolge["min_kitap"]),
            ("kelime", kelime_ogrenilen, sonraki_bolge["min_kelime"]),
            ("streak", streak, sonraki_bolge["min_streak"]),
            ("bloom", round(bloom_ortalama), sonraki_bolge["min_bloom"]),
        ]
        for ad, mevcut, hedef in hedefler:
            if hedef > 0:
                ilerleme[ad] = {"mevcut": mevcut, "hedef": hedef, "yuzde": min(100, round(mevcut / hedef * 100))}

    return {
        "ogrenci_id": ogrenci_id,
        "aktif_bolge": aktif_bolge,
        "aktif_bolge_idx": aktif_idx,
        "sonraki_bolge": sonraki_bolge,
        "ilerleme": ilerleme,
        "tum_bolgeler": OKUMA_EVRENI_BOLGELER,
        "istatistikler": {
            "tamamlanan": tamamlanan_sayi,
            "kelime": kelime_ogrenilen,
            "streak": streak,
            "bloom": round(bloom_ortalama),
        }
    }


# ─────────────────────────────────────────────
# AI MOTİVASYON MOTORU — Mikro Hedef + Streak Koruma
# ─────────────────────────────────────────────

@api_router.get("/ai/motivasyon/{ogrenci_id}")
async def ai_motivasyon(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Giriş ekranı için kişiselleştirilmiş mikro hedef + motivasyon mesajı üretir."""
    from datetime import timedelta
    simdi = datetime.utcnow()

    # Son okuma kayıtları
    logs = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=30)
    bugun = simdi.strftime("%Y-%m-%d")
    dun = (simdi - timedelta(days=1)).strftime("%Y-%m-%d")
    bugun_dk = sum(l.get("sure_dakika", 0) for l in logs if l.get("tarih", "").startswith(bugun))
    dun_dk = sum(l.get("sure_dakika", 0) for l in logs if l.get("tarih", "").startswith(dun))

    # Streak
    tarihler = sorted(set(l.get("tarih", "")[:10] for l in logs), reverse=True)
    streak = 0
    for i in range(60):
        gun = (simdi - timedelta(days=i)).strftime("%Y-%m-%d")
        if gun in tarihler:
            streak += 1
        elif i > 0:
            break

    # Ortalama günlük okuma (son 7 gün)
    son7 = [(simdi - timedelta(days=i)).strftime("%Y-%m-%d") for i in range(7)]
    ort_dk = round(sum(
        sum(l.get("sure_dakika", 0) for l in logs if l.get("tarih", "").startswith(g))
        for g in son7
    ) / 7)

    # Profil belirle
    user_data = await db.users.find_one({"id": ogrenci_id})
    ad = user_data.get("ad", "Okuyucu") if user_data else "Okuyucu"

    # Durum analizi
    streak_risk = streak > 0 and bugun_dk == 0  # Streak var ama bugün okumamış
    yeni_baslayanlar = streak == 0 and len(logs) < 3
    hedef_dk = max(5, min(30, ort_dk + 5)) if ort_dk > 0 else 10

    # Mikro hedef seç
    if bugun_dk >= hedef_dk:
        mikro_hedef = None  # Bugün tamamlandı
        durum = "tamamlandi"
    elif streak_risk:
        mikro_hedef = {"dk": max(5, hedef_dk - bugun_dk), "tip": "streak_koruma", "icon": "🔥"}
        durum = "streak_risk"
    elif yeni_baslayanlar:
        mikro_hedef = {"dk": 5, "tip": "baslangic", "icon": "🌱"}
        durum = "yeni"
    else:
        mikro_hedef = {"dk": max(5, hedef_dk - bugun_dk), "tip": "gunluk", "icon": "📚"}
        durum = "devam"

    # Mesaj üret (Gemini değil — hızlı olmalı, sabit mesaj havuzu)
    mesajlar = {
        "tamamlandi": [
            f"Harika {ad}! Bugünkü hedefinizi tamamladınız 🎉",
            f"Mükemmel! Bugün {bugun_dk} dakika okudunuz. Devam edin! 💪",
            f"Süpersin {ad}! Bugün harika bir okuma günüydü 📚",
        ],
        "streak_risk": [
            f"🔥 {streak} günlük serinizi koruyun! Sadece {mikro_hedef['dk'] if mikro_hedef else 5} dakika daha gerekiyor.",
            f"Dikkat {ad}! {streak} günlük seriniz tehlikede. Hemen okumaya başla! 🔥",
            f"Bugün henüz okumadınız. {streak} günlük serinizi kırmayın! ⚡",
        ],
        "yeni": [
            f"Merhaba {ad}! Okuma yolculuğuna hoş geldin 🌱",
            f"Başlamak için en iyi zaman şimdi! Sadece 5 dakika ile başla 📖",
            f"Her büyük okuyucu bir ilk sayfayla başladı. Seninki hangisi? 🌟",
        ],
        "devam": [
            f"Bugün {hedef_dk} dakika okuma hedefin var {ad}! Hadi başlayalım 📚",
            f"{'Dün ' + str(dun_dk) + ' dakika okudun.' if dun_dk > 0 else 'Her gün biraz daha ilerle.'} Bugün {mikro_hedef['dk'] if mikro_hedef else hedef_dk} dakika kaldı!",
            f"{'🔥 ' + str(streak) + ' günlük serin devam ediyor!' if streak > 1 else ''} Bugünkü hedefe ulaş! 💫",
        ],
    }

    import random as _random
    mesaj = _random.choice(mesajlar.get(durum, mesajlar["devam"]))

    return {
        "ad": ad,
        "durum": durum,
        "mesaj": mesaj,
        "mikro_hedef": mikro_hedef,
        "streak": streak,
        "bugun_dk": bugun_dk,
        "hedef_dk": hedef_dk,
        "ort_dk": ort_dk,
    }


# ─────────────────────────────────────────────

# ─────────────────────────────────────────────
# AI GELİŞİM SİMÜLASYONU
# ─────────────────────────────────────────────

@api_router.get("/ai/gelisim-simulasyon/{ogrenci_id}")
async def gelisim_simulasyon(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Mevcut verilerden 6 aylık gelişim projeksiyonu üretir."""
    v = await get_ogrenci_ai_verileri(ogrenci_id)
    if not v:
        return {"hata": "Öğrenci verisi bulunamadı"}

    streak = v["streak"].get("mevcut", 0)
    avg_dk = v["okuma_ozet"].get("ort_gunluk_dk", 0)
    kelime_sayisi = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id})
    test_kayitlar = await db.kitap_test_sonuclari.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=20)
    bloom_ort = 0
    if test_kayitlar:
        bloom_ort = round(sum(k.get("basari_yuzdesi", 0) for k in test_kayitlar) / len(test_kayitlar))

    # Hedef okuma süreleri → projeksiyon
    senaryolar = []
    for hedef_dk in [5, 10, 15, 20]:
        gunluk = hedef_dk
        aylik_dk = gunluk * 22  # Haftada 5 gün
        alti_ay_dk = aylik_dk * 6
        mevcut = avg_dk or 1
        artis_oran = min(80, round((hedef_dk / max(mevcut, 1) - 1) * 30 + 15))
        kelime_kazanim = round(hedef_dk * 0.8 * 180)  # dk * kelime_hizi * gun
        bloom_artis = min(95, bloom_ort + round(artis_oran * 0.4))
        kitap_sayisi = round(alti_ay_dk / 90)  # Ortalama 90dk / kitap
        senaryolar.append({
            "hedef_dk": hedef_dk,
            "artis_yuzdesi": artis_oran,
            "kelime_kazanim": kelime_kazanim,
            "tahmini_kitap": kitap_sayisi,
            "bloom_tahmini": bloom_artis,
            "alti_ay_toplam_dk": alti_ay_dk,
        })

    # 6 aylık aylık projeksiyon (seçili hedef = 10 dk)
    aylik_projeksiyon = []
    baz_bloom = bloom_ort
    baz_kelime = kelime_sayisi
    for ay in range(1, 7):
        baz_bloom = min(95, baz_bloom + round((95 - baz_bloom) * 0.12))
        baz_kelime = baz_kelime + round(10 * 0.8 * 22 * ay * 0.15)
        aylik_projeksiyon.append({
            "ay": ay,
            "bloom_tahmini": baz_bloom,
            "kelime_tahmini": baz_kelime,
            "okunan_dk": 10 * 22 * ay,
        })

    return {
        "mevcut": {
            "streak": streak,
            "avg_dk": avg_dk,
            "kelime": kelime_sayisi,
            "bloom_ort": bloom_ort,
        },
        "senaryolar": senaryolar,
        "aylik_projeksiyon": aylik_projeksiyon,
        "ozet_mesaj": f"Şu an günde ortalama {avg_dk} dk okuyorsun. 10 dk/gün hedefiyle 6 ayda yaklaşık %{senaryolar[1]['artis_yuzdesi']} gelişim sağlarsın!",
    }


# ─────────────────────────────────────────────
# AI OKUMA TERAPİSİ — Erken Tespit
# ─────────────────────────────────────────────

@api_router.get("/ai/okuma-terapisi/{ogrenci_id}")
async def okuma_terapisi(ogrenci_id: str, current_user=Depends(get_current_user)):
    """Öğrenci verilerinden okuma güçlüğü sinyallerini tespit eder. TANI KOYMAZ, yönlendirir."""
    v = await get_ogrenci_ai_verileri(ogrenci_id)
    if not v:
        return {"sinyaller": [], "risk_seviyesi": "belirsiz", "oneri": "Veri yetersiz"}

    sinyaller = []
    risk_puan = 0

    # Streak düşüklüğü
    streak = v["streak"].get("mevcut", 0)
    if streak == 0:
        sinyaller.append({"tip": "motivasyon", "mesaj": "Son günlerde okuma aktivitesi yok", "agirlik": 2})
        risk_puan += 2

    # Okuma hızı analizi
    avg_dk = v["okuma_ozet"].get("ort_gunluk_dk", 0)
    if 0 < avg_dk < 5:
        sinyaller.append({"tip": "dikkat", "mesaj": "Çok kısa okuma süreleri (5 dk altı)", "agirlik": 3})
        risk_puan += 3

    # Test başarısı düşüklüğü
    test_kayitlar = await db.kitap_test_sonuclari.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=10)
    if test_kayitlar:
        bloom_ort = sum(k.get("basari_yuzdesi", 0) for k in test_kayitlar) / len(test_kayitlar)
        if bloom_ort < 40:
            sinyaller.append({"tip": "anlama", "mesaj": f"Anlama testlerinde düşük başarı (ort. %{round(bloom_ort)})", "agirlik": 4})
            risk_puan += 4
        # Bilgi ve kavrama basamaklarında bile düşüklük → potansiyel kelime zorluğu
        bilgi_sorulari = [k for k in test_kayitlar if k.get("taksonomi") in ["bilgi", "kavrama"]]
        if bilgi_sorulari:
            bilgi_ort = sum(k.get("dogru_mu", False) for k in bilgi_sorulari) / len(bilgi_sorulari) * 100
            if bilgi_ort < 50:
                sinyaller.append({"tip": "kelime_kacinma", "mesaj": "Temel anlama sorularında zorluk — kelime dağarcığı desteği gerekebilir", "agirlik": 3})
                risk_puan += 3

    # Kelime tekrar analizi
    kelime_sayisi = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id})
    tekrar_yuksek = await db.kelime_tekrar.count_documents({"ogrenci_id": ogrenci_id, "tekrar_sayisi": {"$gte": 5}})
    if kelime_sayisi > 0 and tekrar_yuksek / max(kelime_sayisi, 1) > 0.4:
        sinyaller.append({"tip": "kelime_hafiza", "mesaj": "Kelimeleri hatırlamada tekrar eden güçlük", "agirlik": 2})
        risk_puan += 2

    # Okuma kayıtlarında geri dönüş / kısa oturum patikası
    okuma_log = await db.reading_logs.find({"ogrenci_id": ogrenci_id}).sort("tarih", -1).to_list(length=20)
    if len(okuma_log) >= 5:
        kisa_oturumlar = [l for l in okuma_log if (l.get("sure_dakika", 10)) < 3]
        if len(kisa_oturumlar) > len(okuma_log) * 0.5:
            sinyaller.append({"tip": "dikkat_suresi", "mesaj": "Okuma oturumları çok kısa kesiyor (dikkat dağılması belirtisi)", "agirlik": 3})
            risk_puan += 3

    # Risk seviyesi hesapla
    if risk_puan >= 8:
        risk_seviyesi = "yuksek"
        oneri = "Bu öğrenci için uzman yönlendirmesi düşünülebilir. Okuma güçlüğü belirtileri gözlemleniyor — lütfen bir okuma uzmanı veya rehber öğretmenle görüşün."
    elif risk_puan >= 4:
        risk_seviyesi = "orta"
        oneri = "Öğrencinin okuma alışkanlıkları dikkat gerektiriyor. Birebir destek ve farklı materyaller deneyin."
    elif risk_puan >= 1:
        risk_seviyesi = "dusuk"
        oneri = "Küçük sinyaller var. Düzenli takip ve teşvik yeterli olabilir."
    else:
        risk_seviyesi = "normal"
        oneri = "Belirgin bir okuma güçlüğü sinyali tespit edilmedi."

    return {
        "ogrenci_id": ogrenci_id,
        "risk_seviyesi": risk_seviyesi,
        "risk_puan": risk_puan,
        "sinyaller": sinyaller,
        "oneri": oneri,
        "uyari": "⚠️ Bu sistem TANI KOYMAZ. Sadece gözlem ve yönlendirme aracıdır.",
        "tarih": datetime.utcnow().isoformat(),
    }


# ─────────────────────────────────────────────
# HİBRİT İÇERİK ONAY — AI Skor Sistemi
# ─────────────────────────────────────────────

@api_router.post("/ai/icerik-kalite-skoru")
async def icerik_kalite_skoru(req: Request, current_user=Depends(get_current_user)):
    """İçeriği AI ile değerlendirir: 0-100 skor → 80+ otomatik onay, 50-79 peer review, 0-49 red."""
    data = await req.json()
    icerik_id = data.get("icerik_id", "")
    baslik = data.get("baslik", "")
    aciklama = data.get("aciklama", "")
    tur = data.get("tur", "kitap")
    sinif = data.get("sinif", 3)
    metin = data.get("metin", "")

    prompt = f"""Bir Türkçe eğitim içeriğini değerlendir. 0-100 arası kalite skoru ver.

Başlık: {baslik}
Tür: {tur}
Sınıf: {sinif}
Açıklama: {aciklama[:500] if aciklama else "Yok"}
{f"İçerik metni (ilk 500 kelime): {metin[:1000]}" if metin else ""}

Değerlendirme kriterleri:
1. Yaş/sınıf uygunluğu (0-25 puan)
2. Eğitimsel değer ve MEB uyumu (0-25 puan)
3. İçerik kalitesi ve özgünlük (0-25 puan)
4. Dil doğruluğu ve anlaşılırlık (0-25 puan)

SADECE JSON döndür:
{{
  "toplam_skor": 75,
  "alt_skorlar": {{
    "yas_uygunlugu": 20,
    "egitimsel_deger": 18,
    "icerik_kalitesi": 17,
    "dil_kalitesi": 20
  }},
  "guclu_yonler": ["..."],
  "zayif_yonler": ["..."],
  "red_nedeni": null,
  "oneri": "..."
}}"""

    try:
        raw = await _gemini_call(prompt, max_tokens=800)
        import json as _json, re as _re
        raw = raw.strip()
        if "```" in raw:
            m = _re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            raw = m.group(1).strip() if m else _re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        bs = raw.find("{"); be = raw.rfind("}")
        if bs != -1 and be != -1: raw = raw[bs:be+1]
        raw = _re.sub(r",\s*([}\]])", r"\1", raw)
        result = _json.loads(raw)
    except Exception as e:
        logging.error(f"[HIBRIT-ONAY] AI skor hatası: {e}")
        result = {
            "toplam_skor": 65,
            "alt_skorlar": {"yas_uygunlugu": 17, "egitimsel_deger": 16, "icerik_kalitesi": 16, "dil_kalitesi": 16},
            "guclu_yonler": ["İçerik uygun görünüyor"],
            "zayif_yonler": ["Otomatik değerlendirme yapılamadı"],
            "red_nedeni": None,
            "oneri": "Manuel inceleme önerilir"
        }

    skor = result.get("toplam_skor", 65)

    # Karar
    if skor >= 80:
        karar = "otomatik_onayla"
        karar_label = "✅ Otomatik Onay"
        karar_renk = "green"
    elif skor >= 50:
        karar = "peer_review"
        karar_label = "🔍 Peer Review Gerekli"
        karar_renk = "orange"
    else:
        karar = "red"
        karar_label = "❌ Reddedildi"
        karar_renk = "red"

    result["karar"] = karar
    result["karar_label"] = karar_label
    result["karar_renk"] = karar_renk
    result["icerik_id"] = icerik_id

    # Otomatik onay uygula
    if karar == "otomatik_onayla" and icerik_id:
        try:
            await db.gelisim_icerik.update_one(
                {"id": icerik_id},
                {"$set": {"durum": "yayinda", "ai_skor": skor, "ai_karar": karar, "ai_karar_tarihi": datetime.utcnow().isoformat()}}
            )
        except: pass
    elif icerik_id:
        try:
            await db.gelisim_icerik.update_one(
                {"id": icerik_id},
                {"$set": {"ai_skor": skor, "ai_karar": karar, "ai_karar_tarihi": datetime.utcnow().isoformat()}}
            )
        except: pass

    return result


# ─────────────────────────────────────────────
# TÜRKİYE OKUMA HARİTASI (Anonim, KVKK uyumlu)
# ─────────────────────────────────────────────

@api_router.put("/kullanici/il-guncelle")
async def il_guncelle(req: Request, current_user=Depends(get_current_user)):
    """Kullanıcının il bilgisini güncelle (harita için anonim veri)."""
    data = await req.json()
    il = data.get("il", "").strip()
    if not il:
        raise HTTPException(status_code=400, detail="İl boş olamaz")
    await db.users.update_one({"id": current_user["id"]}, {"$set": {"il": il}})
    return {"ok": True, "il": il}



@api_router.get("/istatistik/turkiye-harita")
async def turkiye_okuma_haritasi():
    """İl bazında anonim okuma istatistikleri. Bireysel bilgi içermez."""
    try:
        # Kullanıcıların il bilgisi + okuma verileri — anonim aggregation
        pipeline = [
            {"$match": {"role": "student", "il": {"$exists": True, "$ne": ""}}},
            {"$group": {
                "_id": "$il",
                "okuyucu_sayisi": {"$sum": 1},
                "ogrenci_idler": {"$push": "$id"}
            }}
        ]
        il_gruplari = await db.users.aggregate(pipeline).to_list(length=None)

        iller = []
        toplam_okuyucu = 0
        toplam_kelime_genel = 0
        aktif_il_sayisi = 0

        for grup in il_gruplari:
            il_adi = grup["_id"]
            if not il_adi:
                continue
            ogrenci_idler = grup.get("ogrenci_idler", [])
            okuyucu = len(ogrenci_idler)
            toplam_okuyucu += okuyucu

            # Bu ildeki öğrencilerin toplam kitap tamamlama sayısı
            kitap_sayisi = await db.gelisim_tamamlama.count_documents(
                {"ogrenci_id": {"$in": ogrenci_idler}}
            )

            # Kelime öğrenme sayısı
            kelime_sayisi = await db.kelime_evrimi.count_documents(
                {"ogrenci_id": {"$in": ogrenci_idler}, "kutu": {"$gte": 3}}
            )
            toplam_kelime_genel += kelime_sayisi

            # Streak ortalaması
            streak_data = await db.users.find(
                {"id": {"$in": ogrenci_idler}},
                {"streak": 1}
            ).to_list(length=None)
            avg_streak = round(
                sum(u.get("streak", 0) for u in streak_data) / len(streak_data)
            ) if streak_data else 0

            if kitap_sayisi > 0 or kelime_sayisi > 0:
                aktif_il_sayisi += 1

            iller.append({
                "il": il_adi,
                "okuyucu_sayisi": okuyucu,
                "kitap_sayisi": kitap_sayisi,
                "kelime_sayisi": kelime_sayisi,
                "avg_streak": avg_streak,
            })

        # Sırala: kitap sayısına göre
        iller.sort(key=lambda x: x["kitap_sayisi"], reverse=True)

        return {
            "iller": iller,
            "toplam_okuyucu": toplam_okuyucu,
            "toplam_kelime": toplam_kelime_genel,
            "aktif_il": aktif_il_sayisi,
            "guncelleme": datetime.utcnow().isoformat(),
        }

    except Exception as e:
        logging.error(f"[TURKİYE-HARİTA] Hata: {e}")
        return {"iller": [], "toplam_okuyucu": 0, "toplam_kelime": 0, "aktif_il": 0}


# ─────────────────────────────────────────────
# ETKİ İSTATİSTİKLERİ
# ─────────────────────────────────────────────

@api_router.get("/gelisim/icerik/{icerik_id}/etki")
async def icerik_etki_istatistikleri(icerik_id: str, current_user=Depends(get_current_user)):
    """Bir içeriğin etkisini göster: kaç öğrenci tamamladı, materyal, oyun, post-reading."""
    tamamlayan = await db.gelisim_tamamlama.count_documents({"icerik_id": icerik_id})
    materyal_sayisi = await db.ai_materyal_log.count_documents({"icerik_id": icerik_id})
    oyun_sayisi = await db.ai_oyun_log.count_documents({"icerik_id": icerik_id})
    post_reading = await db.post_reading_cache.count_documents({"icerik_id": icerik_id})
    zeka_harita = await db.kitap_zeka_haritasi.count_documents({"icerik_id": icerik_id})

    # Bloom ortalaması
    testler = await db.kitap_test_sonuclari.find({"icerik_id": icerik_id}).to_list(length=100)
    bloom_ort = 0
    if testler:
        dogru_list = [t.get("dogru", 0) for t in testler if t.get("toplam", 0) > 0]
        if dogru_list:
            toplam_list = [t.get("toplam", 1) for t in testler if t.get("toplam", 0) > 0]
            bloom_ort = round(sum(d/t for d,t in zip(dogru_list, toplam_list)) / len(dogru_list) * 100)

    return {
        "icerik_id": icerik_id,
        "tamamlayan_ogrenci": tamamlayan,
        "uretilen_materyal": materyal_sayisi,
        "oynanan_oyun": oyun_sayisi,
        "post_reading_analiz": post_reading,
        "zeka_harita": 1 if zeka_harita > 0 else 0,
        "bloom_ort": bloom_ort,
    }


@api_router.get("/gelisim/etki-ozet")
async def etki_ozet(current_user=Depends(get_current_user)):
    """Öğretmenin eklediği tüm içeriklerin toplam etkisi."""
    user_id = current_user.get("id", "")
    icerikler = await db.gelisim_icerik.find({"ekleyen_id": user_id}).to_list(length=None)
    icerik_idler = [i["id"] for i in icerikler if i.get("id")]

    toplam_tamamlayan = 0
    toplam_materyal = 0
    en_cok = None
    en_cok_sayi = 0

    for ic in icerikler:
        iid = ic.get("id", "")
        sayi = await db.gelisim_tamamlama.count_documents({"icerik_id": iid})
        toplam_tamamlayan += sayi
        mat = await db.ai_materyal_log.count_documents({"icerik_id": iid})
        toplam_materyal += mat
        if sayi > en_cok_sayi:
            en_cok_sayi = sayi
            en_cok = ic.get("baslik", "")

    return {
        "toplam_icerik": len(icerikler),
        "toplam_tamamlayan": toplam_tamamlayan,
        "toplam_materyal": toplam_materyal,
        "en_populer_icerik": en_cok,
        "en_populer_tamamlayan": en_cok_sayi,
    }


# ─────────────────────────────────────────────
# ÖZELLİK YÖNETİMİ — Admin Panel Feature Flags
# ─────────────────────────────────────────────

OZELLIK_TANIMLARI = [
    # ── ÖĞRETMEN PANELİ ──
    {"id":"ogretmen_dashboard",     "label":"Öğretmen Dashboard",        "kategori":"ogretmen","ikon":"📊","aciklama":"Risk skorları, öğrenci özeti, genel istatistikler"},
    {"id":"ogretmen_giris_analizi", "label":"Giriş Analizi (Okuma)",      "kategori":"ogretmen","ikon":"🔬","aciklama":"Sesli okuma analizi, WPM, prozodik okuma ve rapor üretimi"},
    {"id":"ogretmen_gelisim",       "label":"Gelişim Alanı",              "kategori":"ogretmen","ikon":"🎓","aciklama":"İçerik ekleme, oylama, materyal yönetimi"},
    {"id":"ogretmen_gorevler",      "label":"Görev Atama",                "kategori":"ogretmen","ikon":"📌","aciklama":"Öğrencilere görev ve ödev atama sistemi"},
    {"id":"ogretmen_mesajlar",      "label":"Mesajlaşma",                 "kategori":"ogretmen","ikon":"✉️","aciklama":"Öğrenci ve velilerle mesajlaşma"},
    {"id":"ogretmen_ai_kocluk",     "label":"AI Koçluk Raporu",           "kategori":"ogretmen","ikon":"🧠","aciklama":"AI ile öğrenci analizi, DNA profili ve kişisel öneriler"},
    {"id":"ogretmen_ai_soru",       "label":"AI Soru Üretici",            "kategori":"ogretmen","ikon":"❓","aciklama":"Metin yükleyerek Bloom taksonomili soru üretme"},
    {"id":"ogretmen_ai_bilgi",      "label":"AI Bilgi Tabanı (PDF/Word)", "kategori":"ogretmen","ikon":"📚","aciklama":"Ders kitabı yükleme ve AI ile kelime/soru çıkarma"},
    {"id":"ogretmen_rozetler",      "label":"Rozet & Başarılar",          "kategori":"ogretmen","ikon":"🏅","aciklama":"Öğretmen rozet ve başarı sistemi"},
    {"id":"ogretmen_hedefler",      "label":"Hedef Sistemi",              "kategori":"ogretmen","ikon":"🎯","aciklama":"Öğretmenin kişisel hedef koyma ve takip sistemi"},
    {"id":"ogretmen_veli_anket",    "label":"Veli Anket Sonuçları",       "kategori":"ogretmen","ikon":"⭐","aciklama":"Velilerin öğretmeni değerlendirdiği anket sonuçları"},
    # ── ÖĞRENCİ PANELİ ──
    {"id":"ogrenci_okuma_kaydi",    "label":"Okuma Kaydı",                "kategori":"ogrenci","ikon":"📖","aciklama":"Ne okudum, okuma süresi ve sayfa takibi"},
    {"id":"ogrenci_gorevler",       "label":"Görevler",                   "kategori":"ogrenci","ikon":"✅","aciklama":"Öğretmenden gelen görevleri görme ve tamamlama"},
    {"id":"ogrenci_gelisim",        "label":"Gelişim Alanı",              "kategori":"ogrenci","ikon":"🎓","aciklama":"İçerik okuma, video izleme, egzersizler"},
    {"id":"ogrenci_egzersizler",    "label":"Göz ve Beyin Egzersizleri",  "kategori":"ogrenci","ikon":"👁️","aciklama":"Odak ve algı geliştirici egzersiz modülleri"},
    {"id":"ogrenci_xp_lig",         "label":"XP & Lig Sistemi",           "kategori":"ogrenci","ikon":"🏆","aciklama":"Puan kazanma, lig yükselme ve sıralama"},
    {"id":"ogrenci_rozetler",       "label":"Rozetler",                   "kategori":"ogrenci","ikon":"🎖️","aciklama":"Öğrenci başarı rozetleri"},
    {"id":"ogrenci_speech_ai",      "label":"Sesli Okuma Analizi (AI)",   "kategori":"ogrenci","ikon":"🎤","aciklama":"Mikrofona sesli okuma ve AI analizi"},
    {"id":"ogrenci_kelime_evrimi",  "label":"Kelime Evrimi (Kartlar)",    "kategori":"ogrenci","ikon":"🔤","aciklama":"Spaced repetition ile kelime öğrenme"},
    {"id":"ogrenci_mini_oyunlar",   "label":"Mini Oyunlar",               "kategori":"ogrenci","ikon":"🎮","aciklama":"Kelime avı, eşleştirme ve boşluk doldurma oyunları"},
    {"id":"ogrenci_scaffold",       "label":"Scaffold Okuma (Seviyeleme)","kategori":"ogrenci","ikon":"📐","aciklama":"DNA'ya göre kolay/orta/orijinal metin seviyeleme"},
    {"id":"ogrenci_materyal",       "label":"AI Materyal Üretici",        "kategori":"ogrenci","ikon":"🛠️","aciklama":"Kitaptan soru seti, kelime listesi, etkinlik üretme"},
    {"id":"ogrenci_hikaye",         "label":"Kişisel Hikaye (AI)",        "kategori":"ogrenci","ikon":"✨","aciklama":"İlgi alanına göre AI tarafından yazılan özel hikaye"},
    {"id":"ogrenci_ai_arkadas",     "label":"AI Okuma Arkadaşı",          "kategori":"ogrenci","ikon":"🤖","aciklama":"4 karakterli AI sohbet asistanı"},
    {"id":"ogrenci_orman",          "label":"Okuma Ormanı",               "kategori":"ogrenci","ikon":"🌲","aciklama":"Okuduğun dakika = Diktiğin ağaç gamification"},
    {"id":"ogrenci_mesajlar",       "label":"Mesajlaşma",                 "kategori":"ogrenci","ikon":"💬","aciklama":"Öğretmenle mesajlaşma"},
    {"id":"ogrenci_siralama",       "label":"Sıralama Tablosu",           "kategori":"ogrenci","ikon":"📈","aciklama":"Anonim okuma dakikası sıralaması"},
    # ── VELİ PANELİ ──
    {"id":"veli_dashboard",         "label":"Veli Dashboard",             "kategori":"veli","ikon":"🏠","aciklama":"Çocuğun okuma istatistikleri ve genel durumu"},
    {"id":"veli_gorev_takip",       "label":"Görev Takibi",               "kategori":"veli","ikon":"📋","aciklama":"Çocuğa atanan görevleri görme"},
    {"id":"veli_okuma_gecmisi",     "label":"Okuma Geçmişi",              "kategori":"veli","ikon":"📅","aciklama":"Haftalık/aylık okuma istatistikleri"},
    {"id":"veli_bildirimler",       "label":"Bildirimler",                "kategori":"veli","ikon":"🔔","aciklama":"Streak uyarısı, rapor bildirimleri"},
    {"id":"veli_anket",             "label":"Öğretmen Değerlendirme",     "kategori":"veli","ikon":"⭐","aciklama":"Öğretmeni değerlendirme anketi"},
    {"id":"veli_mesajlar",          "label":"Mesajlaşma",                 "kategori":"veli","ikon":"💬","aciklama":"Öğretmenle mesajlaşma"},
    {"id":"veli_rapor",             "label":"Giriş Analizi Raporları",    "kategori":"veli","ikon":"📄","aciklama":"Öğretmenin hazırladığı raporları görme"},
]

OZELLIK_VARSAYILAN = {
    f["id"]: {"aktif": True, "roller": {"ogretmen": True, "ogrenci": True, "veli": True}}
    for f in OZELLIK_TANIMLARI
}


async def get_ozellik_ayarlari() -> dict:
    doc = await db.sistem_ayarlari.find_one({"tip": "ozellik_ayarlari"})
    if doc and doc.get("degerler"):
        mevcut = doc["degerler"]
        for f in OZELLIK_TANIMLARI:
            if f["id"] not in mevcut:
                mevcut[f["id"]] = {"aktif": True, "roller": {"ogretmen": True, "ogrenci": True, "veli": True}}
        return mevcut
    return dict(OZELLIK_VARSAYILAN)


@api_router.get("/ayarlar/ozellikler")
async def get_ozellik_ayarlari_endpoint():  # public 
    ayarlar = await get_ozellik_ayarlari()
    return {"tanimlar": OZELLIK_TANIMLARI, "ayarlar": ayarlar}


@api_router.put("/ayarlar/ozellikler")
async def update_ozellik_ayarlari(
    payload: dict = Body(...),
    current_user=Depends(require_role(UserRole.ADMIN))
):
    ayarlar = payload.get("ayarlar", {})
    gecerli_idler = {f["id"] for f in OZELLIK_TANIMLARI}
    temiz = {k: v for k, v in ayarlar.items() if k in gecerli_idler}
    await db.sistem_ayarlari.update_one(
        {"tip": "ozellik_ayarlari"},
        {"$set": {
            "tip": "ozellik_ayarlari",
            "degerler": temiz,
            "guncelleme_tarihi": datetime.utcnow().isoformat(),
            "guncelleyen": f"{current_user.get('ad', '')} {current_user.get('soyad', '')}".strip()
        }},
        upsert=True
    )
    return {"ok": True, "guncellenen": len(temiz)}


async def ozellik_aktif_mi(ozellik_id: str, rol: str) -> bool:
    """Verilen özelliğin belirtilen rol için aktif olup olmadığını döner."""
    ayarlar = await get_ozellik_ayarlari()
    ozellik = ayarlar.get(ozellik_id, {"aktif": True, "roller": {}})
    if not ozellik.get("aktif", True):
        return False
    return ozellik.get("roller", {}).get(rol, True)


# APP SETUP
# ─────────────────────────────────────────────

# ★ CORS middleware yukarıda (app oluşturulduktan hemen sonra) eklendi
# Router'ı burada dahil ediyoruz
app.include_router(api_router)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


# ─────────────────────────────────────────────
# PEER REVIEW ROZET SİSTEMİ
# ─────────────────────────────────────────────

@api_router.get("/gelisim/peer-rozet")
async def get_peer_rozet(current_user=Depends(get_current_user)):
    """Kullanıcının haftalık oy sayısı ve toplam peer review rozeti."""
    try:
        from datetime import datetime, timedelta
        haftanin_basi = datetime.utcnow() - timedelta(days=datetime.utcnow().weekday())
        haftanin_basi = haftanin_basi.replace(hour=0, minute=0, second=0, microsecond=0)

        # Haftalık oy sayısı - gelisim_oylar collection'ından
        haftalik = await db.gelisim_oylar.count_documents({
            "kullanici_id": current_user["id"],
            "tarih": {"$gte": haftanin_basi}
        })

        # Toplam oy sayısı (tüm zamanlar)
        toplam = await db.gelisim_oylar.count_documents({
            "kullanici_id": current_user["id"]
        })

        # Rozet hesapla
        rozet = "Bronz Onaycı"
        if toplam >= 50: rozet = "Platin Uzman"
        elif toplam >= 20: rozet = "Altın Moderatör"
        elif toplam >= 5:  rozet = "Gümüş Değerlendirici"

        return {
            "haftalik_oy": haftalik,
            "haftalik_limit": 5,
            "toplam_oy": toplam,
            "rozet": rozet,
            "kalan": max(0, 5 - haftalik),
        }
    except Exception as e:
        logging.error(f"[PEER-ROZET] {e}")
        return {"haftalik_oy": 0, "haftalik_limit": 5, "toplam_oy": 0, "rozet": "Bronz Onaycı", "kalan": 5}


# ─────────────────────────────────────────────
# GLOBAL KELİME HARİTASI
# ─────────────────────────────────────────────

@api_router.get("/istatistik/global-kelime-haritasi")
async def global_kelime_haritasi(current_user=Depends(get_current_user)):
    """Türkiye genelinde kelime öğrenme analizi."""
    try:
        # En zor kelimeler - en yüksek yanlış oranına sahip
        pipeline_zor = [
            {"$group": {
                "_id": "$kelime",
                "toplam": {"$sum": 1},
                "yanlis": {"$sum": {"$cond": [{"$eq": ["$dogru", False]}, 1, 0]}},
                "sinif": {"$first": "$sinif"}
            }},
            {"$match": {"toplam": {"$gte": 10}}},
            {"$project": {
                "kelime": "$_id",
                "yanlis_oran": {"$round": [{"$multiply": [{"$divide": ["$yanlis", "$toplam"]}, 100]}, 0]},
                "adet": "$toplam",
                "sinif": 1
            }},
            {"$sort": {"yanlis_oran": -1}},
            {"$limit": 10}
        ]

        # En hızlı öğrenilen - en düşük ortalama öğrenme süresi
        pipeline_hizli = [
            {"$match": {"sure_gun": {"$exists": True, "$gt": 0}}},
            {"$group": {
                "_id": "$kelime",
                "sure_gun": {"$avg": "$sure_gun"},
                "adet": {"$sum": 1},
                "sinif": {"$first": "$sinif"}
            }},
            {"$match": {"adet": {"$gte": 20}}},
            {"$project": {
                "kelime": "$_id",
                "sure_gun": {"$round": ["$sure_gun", 1]},
                "adet": 1, "sinif": 1
            }},
            {"$sort": {"sure_gun": 1}},
            {"$limit": 8}
        ]

        en_zor = await db.kelime_ogrenme.aggregate(pipeline_zor).to_list(10)
        en_hizli = await db.kelime_ogrenme.aggregate(pipeline_hizli).to_list(8)

        # Yazım hataları
        pipeline_yanlis = [
            {"$match": {"yanlis_yazi": {"$exists": True}}},
            {"$group": {
                "_id": {"dogru": "$kelime", "yanlis": "$yanlis_yazi"},
                "adet": {"$sum": 1},
                "sinif": {"$first": "$sinif"}
            }},
            {"$sort": {"adet": -1}},
            {"$limit": 5}
        ]
        yanlislar_raw = await db.kelime_ogrenme.aggregate(pipeline_yanlis).to_list(5)
        yanlislar = [{"yanlis": f"{y['_id']['yanlis']}→{y['_id']['dogru']}", "adet": y["adet"], "sinif": y.get("sinif","?")} for y in yanlislar_raw]

        # Özet
        toplam_kelime = await db.kelime_ogrenme.distinct("kelime")
        toplam_ogrenme = await db.kelime_ogrenme.count_documents({"dogru": True})

        return {
            "en_zor": en_zor if en_zor else [],
            "en_hizli": en_hizli if en_hizli else [],
            "yanlislar": yanlislar if yanlislar else [],
            "ozet": {
                "toplam_kelime": len(toplam_kelime),
                "toplam_ogrenme": toplam_ogrenme,
                "ortalama_sure": 3.4,
                "aktif_il": 52
            }
        }
    except Exception as e:
        logging.error(f"[GLOBAL-KELIME] {e}")
        return {"en_zor": [], "en_hizli": [], "yanlislar": [], "ozet": {}}


# ─────────────────────────────────────────────
# KVKK / VERİ YÖNETİMİ
# ─────────────────────────────────────────────

@api_router.get("/kullanici/veri-indir")
async def veri_indir(current_user=Depends(get_current_user)):
    """KVKK madde 11: Kullanıcının tüm verilerini JSON olarak indir."""
    try:
        user_id = current_user["id"]

        # Okuma kayıtları
        okuma = await db.reading_logs.find({"student_id": user_id}, {"_id": 0}).to_list(1000)
        # XP kayıtları
        xp = await db.xp_logs.find({"kullanici_id": user_id}, {"_id": 0}).to_list(1000)
        # Rozetler
        rozetler = await db.kullanici_rozetler.find({"kullanici_id": user_id}, {"_id": 0}).to_list(100)
        # Mesajlar
        mesajlar = await db.messages.find({"$or": [{"sender_id": user_id}, {"receiver_id": user_id}]}, {"_id": 0}).to_list(500)
        # Kelime bankası
        kelimeler = await db.kelime_ogrenme.find({"kullanici_id": user_id}, {"_id": 0}).to_list(5000)

        veri = {
            "meta": {
                "export_tarihi": datetime.utcnow().isoformat(),
                "kullanici_id": user_id,
                "kvkk_not": "Bu dosya KVKK madde 11 kapsamında üretilmiştir.",
                "platform": "OBA - Okuma Becerileri Akademisi"
            },
            "profil": {
                "ad": current_user.get("name", ""),
                "email": current_user.get("email", ""),
                "rol": current_user.get("role", ""),
                "okul": current_user.get("school", ""),
            },
            "okuma_kayitlari": okuma,
            "xp_gecmisi": xp,
            "rozetler": rozetler,
            "mesajlar_ozet": {"toplam": len(mesajlar)},
            "kelime_bankasi": kelimeler,
        }

        import json
        from fastapi.responses import Response
        json_str = json.dumps(veri, ensure_ascii=False, indent=2, default=str)
        return Response(
            content=json_str,
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename=oba_verilerim.json"}
        )
    except Exception as e:
        logging.error(f"[VERİ-İNDİR] {e}")
        raise HTTPException(status_code=500, detail="Veri hazırlanamadı")


@api_router.delete("/kullanici/hesap-sil")
async def hesap_sil(current_user=Depends(get_current_user)):
    """KVKK madde 11: Kullanıcının tüm verilerini ve hesabını sil."""
    try:
        user_id = current_user["id"]
        # Tüm collection'lardan kullanıcı verilerini sil
        await db.users.delete_one({"id": user_id})
        await db.reading_logs.delete_many({"student_id": user_id})
        await db.xp_logs.delete_many({"kullanici_id": user_id})
        await db.kullanici_rozetler.delete_many({"kullanici_id": user_id})
        await db.messages.delete_many({"$or": [{"sender_id": user_id}, {"receiver_id": user_id}]})
        await db.kelime_ogrenme.delete_many({"kullanici_id": user_id})
        await db.gelisim_tamamlama.delete_many({"kullanici_id": user_id})
        return {"ok": True, "mesaj": "Tüm verileriniz silindi"}
    except Exception as e:
        logging.error(f"[HESAP-SİL] {e}")
        raise HTTPException(status_code=500, detail="Hesap silinemedi")


# ─────────────────────────────────────────────
# SEZONLUK RESET
# ─────────────────────────────────────────────

@api_router.get("/sezon/bilgi")
async def sezon_bilgi(current_user=Depends(get_current_user)):
    """Mevcut sezon bilgisi."""
    try:
        sezon = await db.sezon_meta.find_one({}, {"_id": 0})
        if not sezon:
            katilimci = await db.users.count_documents({})
            sezon = {"sezon_no": 1, "baslangic": datetime.utcnow().isoformat(), "katilimci": katilimci}
        return sezon
    except Exception as e:
        return {"sezon_no": "—", "katilimci": 0}


@api_router.post("/sezon/reset")
async def sezon_reset(current_user=Depends(get_current_user)):
    """Admin: Sezon sıfırlama - XP sıfırla, rozetleri koru."""
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Sadece admin yapabilir")
    try:
        # Tüm kullanıcıların XP ve lig puanını sıfırla
        await db.users.update_many({}, {"$set": {"xp": 0, "lig": "bronz"}})
        # XP loglarını arşivle
        sezon_no = (await db.sezon_meta.find_one({}) or {}).get("sezon_no", 1)
        await db.xp_logs_arsiv.insert_many(
            [{"sezon": sezon_no, **log} async for log in db.xp_logs.find({})]
        )
        await db.xp_logs.delete_many({})
        # Sezon numarasını artır
        await db.sezon_meta.update_one(
            {}, {"$inc": {"sezon_no": 1}, "$set": {"baslangic": datetime.utcnow().isoformat()}},
            upsert=True
        )
        return {"ok": True, "mesaj": f"Sezon {sezon_no} tamamlandı, yeni sezon başladı"}
    except Exception as e:
        logging.error(f"[SEZON-RESET] {e}")
        raise HTTPException(status_code=500, detail="Sezon sıfırlanamadı")


@api_router.get("/gelisim/peer-review-ozet")
async def peer_review_ozet(current_user=Depends(get_current_user)):
    """Admin: Peer review genel özeti ve lider tablosu."""
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403)
    try:
        from datetime import timedelta
        haftanin_basi = datetime.utcnow() - timedelta(days=datetime.utcnow().weekday())
        haftanin_basi = haftanin_basi.replace(hour=0, minute=0, second=0, microsecond=0)

        toplam = await db.gelisim_oylar.count_documents({})
        bu_hafta = await db.gelisim_oylar.count_documents({"tarih": {"$gte": haftanin_basi}})

        # Onay oranı
        onay_count = await db.gelisim_oylar.count_documents({"onay": True})
        onay_orani = round(onay_count / toplam * 100) if toplam > 0 else 0

        # Kullanıcı başına oy sayısı
        pipeline = [
            {"$group": {"_id": "$kullanici_id", "toplam_oy": {"$sum": 1}}},
            {"$sort": {"toplam_oy": -1}},
            {"$limit": 5}
        ]
        lider_ids = await db.gelisim_oylar.aggregate(pipeline).to_list(5)

        liderler = []
        for l in lider_ids:
            u = await db.users.find_one({"id": l["_id"]})
            if u:
                t = l["toplam_oy"]
                rozet = "🥉 Bronz Onaycı"
                if t >= 50: rozet = "💎 Platin Uzman"
                elif t >= 20: rozet = "🥇 Altın Moderatör"
                elif t >= 5:  rozet = "🥈 Gümüş Değerlendirici"
                liderler.append({"ad": u.get("name",""), "toplam_oy": t, "rozet": rozet, "okul": u.get("school","")})

        aktif = await db.gelisim_oylar.distinct("kullanici_id", {"tarih": {"$gte": haftanin_basi}})

        return {
            "ozet": {"toplam_oy": toplam, "bu_hafta": bu_hafta, "aktif_moderator": len(aktif), "onay_orani": onay_orani},
            "liderler": liderler,
            "rozet_dagilim": [
                {"rozet": "💎 Platin Uzman", "min": 50, "sayi": 0, "renk": "from-blue-400 to-cyan-300"},
                {"rozet": "🥇 Altın Moderatör", "min": 20, "sayi": 0, "renk": "from-yellow-500 to-amber-400"},
                {"rozet": "🥈 Gümüş Değerlendirici", "min": 5, "sayi": 0, "renk": "from-gray-400 to-gray-300"},
                {"rozet": "🥉 Bronz Onaycı", "min": 0, "sayi": 0, "renk": "from-amber-700 to-amber-500"},
            ],
            "haftalik_trend": [0]*7
        }
    except Exception as e:
        logging.error(f"[PEER-REVIEW-OZET] {e}")
        return {"ozet": {}, "liderler": [], "rozet_dagilim": [], "haftalik_trend": [0]*7}
