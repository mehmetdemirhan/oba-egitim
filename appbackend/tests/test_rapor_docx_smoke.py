"""Giriş Analizi Raporu — DOCX (Word) export + 4-bant prozodik + doğruluk düzey smoke.

Kapsam:
- GET /diagnostic/rapor/{id}/docx → geçerli .docx (grafik görselleri gömülü)
- Prozodik 4 bant: 14/20 → "Çok İyi" (eski 3-bant "İyi" DEĞİL)
- Doğru okuma oranı düzeyi: %95 → "Geliştirilmeli" + açıklama
- prozodik_esikleri rapor-ayarlari'nda düzenlenebilir

İzole test DB. cd appbackend && .venv/Scripts/python.exe tests/test_rapor_docx_smoke.py
"""
import asyncio
import io
import os
import sys
import uuid

TEST_DB = "oba_test_rapor_docx"
os.environ["MONGO_URL"] = os.environ.get("MONGO_URL_TEST", "mongodb://localhost:27017")
os.environ["DB_NAME"] = TEST_DB
os.environ.setdefault("SECRET_KEY", "test-secret-key")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_gecen = 0
_kalan = 0


def check(kosul, mesaj):
    global _gecen, _kalan
    if kosul:
        _gecen += 1; print(f"  [GECTI] {mesaj}")
    else:
        _kalan += 1; print(f"  [KALDI] {mesaj}")


async def run():
    import docx
    import server
    from core.auth import create_access_token
    from httpx import AsyncClient, ASGITransport

    await server.client.drop_database(TEST_DB)
    tid = str(uuid.uuid4()); sid = str(uuid.uuid4()); rid = str(uuid.uuid4())
    await server.db.users.insert_one({"id": tid, "role": "admin", "ad": "Yön", "soyad": "E"})
    await server.db.students.insert_one({"id": sid, "ad": "Ayşe", "soyad": "Yıl", "sinif": "3"})
    await server.db.diagnostic_raporlar.insert_one({
        "id": rid, "ogrenci_id": sid, "ogrenci_ad": "Ayşe Yıl", "ogretmen_id": tid, "ogretmen_ad": "Öğ Rt",
        "ogrenci_sinif": "3", "metin_turu": "olcum", "metin_baslik": "Kırmızı Balık", "kelime_sayisi": 168,
        "dogru_kelime": 160, "yanlis_kelime": 8, "sure_saniye": 100, "wpm": 96, "hiz_deger": "yeterli",
        "dogruluk_yuzde": 95, "anlama": {"ana_fikir": "iyi", "konu": "orta", "cumle_anlama": "iyi", "bilgi": "orta"},
        "anlama_yuzde": 78, "prozodik": {"noktalama": 3, "vurgu": 3, "tonlama": 3, "akicilik": 3, "anlamli_gruplama": 2},
        "prozodik_toplam": 14, "hata_sayilari": [{"tur": "harf_atlama", "sayi": 3}, {"tur": "tekrar_etme", "sayi": 2}],
        "ogretmen_notu": "Okumaya istekli.", "aktif_anlama_gruplari": ["4.1", "4.2", "4.3", "4.4", "4.5"],
        "rapor_tipi": "olcum", "olusturma_tarihi": "2026-07-21T00:00:00"})
    H = {"Authorization": f"Bearer {create_access_token({'sub': tid})}"}

    async with AsyncClient(transport=ASGITransport(app=server.app), base_url="http://test") as ac:
        # DOCX
        r = await ac.get(f"/api/diagnostic/rapor/{rid}/docx", headers=H)
        check(r.status_code == 200, f"DOCX üretildi (status={r.status_code})")
        check(r.content[:2] == b"PK", "DOCX ZIP imzası (PK) doğru")
        d = docx.Document(io.BytesIO(r.content))
        txt = "\n".join(p.text for p in d.paragraphs) + "\n" + "\n".join(
            c.text for tb in d.tables for row in tb.rows for c in row.cells)
        check("Çok İyi" in txt, "prozodik 14/20 → 'Çok İyi' (4-bant)")
        check("Geliştirilmeli" in txt, "doğruluk %95 → 'Geliştirilmeli' düzey etiketi")
        check("anlam kopukluk" in txt.lower(), "doğruluk düzey açıklaması var")
        check("Eğitimci Notu" in txt, "Eğitimci Notu bölümü var")
        check(len(d.inline_shapes) >= 5, f"grafik görselleri gömülü ({len(d.inline_shapes)} adet)")

        # PDF regresyon: aynı raporda prozodik 'Çok İyi'
        import fitz
        rp = await ac.get(f"/api/diagnostic/rapor/{rid}/pdf", headers=H)
        pm = "\n".join(pg.get_text() for pg in fitz.open(stream=rp.content, filetype="pdf"))
        check(rp.status_code == 200 and "Çok İyi" in pm, "PDF'te de prozodik 'Çok İyi'")

        # prozodik_esikleri ayarı okunur + güncellenir
        r = await ac.get("/api/admin/rapor-ayarlari/prozodik_esikleri", headers=H)
        check(r.status_code == 200 and r.json()["degerler"].get("cokiyi") == 13,
              f"prozodik_esikleri varsayılan cokiyi=13 ({r.json().get('degerler')})")
        r = await ac.put("/api/admin/rapor-ayarlari/prozodik_esikleri", headers=H,
                         json={"degerler": {"cokiyi": 15, "iyi": 10, "orta": 7}})
        check(r.status_code == 200, "prozodik_esikleri güncellendi")
        # eşik 15 olunca 14 → artık 'İyi'
        rp2 = await ac.get(f"/api/diagnostic/rapor/{rid}/pdf", headers=H)
        pm2 = "\n".join(pg.get_text() for pg in fitz.open(stream=rp2.content, filetype="pdf"))
        check("Prozodik okuma performansı: İyi" in pm2 or "performansı: İyi" in pm2,
              "eşik 15 → prozodik 14 artık 'İyi' (ayar etkili)")

    await server.client.drop_database(TEST_DB)


def main():
    asyncio.run(run())
    print(f"\nSONUC: {_gecen}/{_gecen + _kalan} kontrol gecti")
    sys.exit(0 if _kalan == 0 else 1)


if __name__ == "__main__":
    main()
