"""Giriş Analizi Raporu — .docx (Word) render'ı.

PDF ile AYNI veri modelinden (ctx) üretilir; grafikler PDF'te kullanılan aynı
reportlab.graphics çizimlerinin PNG'ye dönüştürülmüş HÂLİ olarak gömülür — böylece
iki format arasında içerik/görsel tutarsızlığı olmaz (kullanıcı isteği #4).

Grafik PNG dönüşümü renderPM ile yapılır (matplotlib gerekmez). python-docx tablo/
paragraf yapısı PDF bölümleriyle birebir eşlenir.
"""
from __future__ import annotations
import io

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import core.rapor_grafik as grafik
from core.giris_rapor import (
    metin_turu_ad, hata_turu_ad, hata_turu_aciklama,
    dogruluk_duzey_ad, dogruluk_duzey_aciklama, prozodik_duzey_ad,
    sonuc_paragrafi_uret,
)

LACIVERT = RGBColor(0x1F, 0x4E, 0x79)
GRI = RGBColor(0x66, 0x66, 0x66)
YESIL = RGBColor(0x2E, 0x9E, 0x5B)
KOYU_YESIL = RGBColor(0x1E, 0x7A, 0x44)
SARI = RGBColor(0xB8, 0x8A, 0x00)
KIRMIZI = RGBColor(0xD9, 0x53, 0x4F)

DUZEY_RENK = {
    "Bağımsız Düzey": YESIL, "Geliştirilmeli": SARI, "Yetersiz": KIRMIZI,
    "Çok İyi": KOYU_YESIL, "İyi": YESIL, "Orta": SARI, "Zayıf": KIRMIZI,
}


def _png(drawing, olcek=3) -> bytes | None:
    """reportlab Drawing → PNG bytes. renderPM'in C/cairo backend'i ortamda yok;
    bu yüzden çizim önce tek sayfalık PDF'e (renderPDF), ardından PyMuPDF (fitz) ile
    yüksek çözünürlüklü PNG'ye rasterize edilir (her ikisi de kurulu)."""
    if drawing is None:
        return None
    try:
        from reportlab.graphics import renderPDF
        import fitz
        pdfb = renderPDF.drawToString(drawing)
        d = fitz.open(stream=pdfb, filetype="pdf")
        pix = d[0].get_pixmap(matrix=fitz.Matrix(olcek, olcek), alpha=False)
        return pix.tobytes("png")
    except Exception:
        return None


def _resim(hucre_veya_doc, png: bytes, genislik_cm: float):
    if not png:
        return
    p = hucre_veya_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(png), width=Cm(genislik_cm))


def _baslik(doc, metin, boyut=13, renk=LACIVERT, before=10, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(metin)
    r.bold = True
    r.font.size = Pt(boyut)
    r.font.color.rgb = renk
    return p


def _para(doc, parcalar, boyut=10, after=4, italik=False):
    """parcalar: str veya [(metin, {bold, color, italic}), ...]."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if isinstance(parcalar, str):
        parcalar = [(parcalar, {})]
    for metin, ops in parcalar:
        r = p.add_run(metin)
        r.font.size = Pt(ops.get("size", boyut))
        r.bold = ops.get("bold", False)
        r.italic = ops.get("italic", italik)
        if ops.get("color"):
            r.font.color.rgb = ops["color"]
    return p


def _tablo_stil(t):
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)


def _baslik_satiri(t, hucreler):
    hdr = t.rows[0].cells
    for i, h in enumerate(hucreler):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = LACIVERT


def rapor_docx_uret(rapor: dict, ctx: dict) -> bytes:
    """Tam Giriş Analizi Raporu'nu .docx olarak üretir; bytes döner."""
    doc = Document()
    doc.core_properties.title = "Okuma Becerileri Ölçüm Raporu"

    # ── BAŞLIK ──
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Okuma Becerileri Akademisi"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = LACIVERT
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Giriş Analizi Raporu"); r.font.size = Pt(11); r.font.color.rgb = GRI

    # ── 1. ÖĞRENCİ BİLGİLERİ ──
    _baslik(doc, "1. Öğrenci Bilgileri")
    tb = doc.add_table(rows=2, cols=4)
    tb.style = "Table Grid"
    veriler = [
        ("Adı Soyadı:", rapor.get("ogrenci_ad", "-"), "Sınıfı:", str(rapor.get("ogrenci_sinif", "-"))),
        ("Eğitimci:", rapor.get("ogretmen_ad", "-"), "Tarih:", str(rapor.get("olusturma_tarihi", ""))[:10]),
    ]
    for ri, (a, b, c, d) in enumerate(veriler):
        cells = tb.rows[ri].cells
        for ci, (txt, bold) in enumerate([(a, True), (b, False), (c, True), (d, False)]):
            cells[ci].text = ""
            rr = cells[ci].paragraphs[0].add_run(str(txt))
            rr.bold = bold; rr.font.size = Pt(9)
            if bold:
                rr.font.color.rgb = LACIVERT

    # ── 2. METİN BİLGİLERİ ── (PDF ile birebir aynı 6 alan — Word Tablo 1 paritesi)
    _baslik(doc, "2. Metin Bilgileri")
    kelime_s = ctx["kelime_s"]; dogru_k = ctx["dogru_k"]; yanlis_k = ctx["yanlis_k"]
    _sst = int(rapor.get("sure_saniye") or 0)
    metin_satir = [
        ("Metnin Adı", (rapor.get("metin_adi") or rapor.get("metin_baslik") or "-")),
        ("Metnin Türü", metin_turu_ad(rapor.get("metin_turu"))),
        ("Toplam Kelime Sayısı", str(kelime_s)),
        ("Doğru Okunan Kelime", str(dogru_k)),
        ("Yanlış Okunan Kelime", str(yanlis_k)),
        ("Tamamlama Süresi", f"{_sst // 60}:{str(_sst % 60).zfill(2)} ({_sst} sn)"),
    ]
    tb2 = doc.add_table(rows=len(metin_satir), cols=2); tb2.style = "Table Grid"
    for ri, (a, b) in enumerate(metin_satir):
        cells = tb2.rows[ri].cells
        cells[0].text = ""; rr = cells[0].paragraphs[0].add_run(a); rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = LACIVERT
        cells[1].text = ""; cells[1].paragraphs[0].add_run(str(b)).font.size = Pt(9)

    # ── 3. OKUMA HIZI ──
    _baslik(doc, "3. Okuma Hızı")
    _para(doc, [("Okuma Hızı = Doğru okunan kelime sayısı / Metnin tamamının okunduğu süre (sn) × 60", {"italic": True, "size": 8, "color": GRI})])
    sinif_norm = ctx.get("sinif_norm") or {}
    if sinif_norm:
        _dd, _oo, _yy = int(sinif_norm.get("dusuk", 0)), int(sinif_norm.get("orta", 0)), int(sinif_norm.get("yeterli", 0))
        _para(doc, [(f"Referans (Sınıf {rapor.get('ogrenci_sinif','')}): ", {"bold": True, "size": 8}),
                    (f"Düşük 0–{_dd} · Orta {_dd+1}–{_oo} · Yeterli {_oo+1}–{_yy} · İleri {_yy+1}+ kelime/dk", {"size": 8, "color": GRI})])
    wpm = ctx["wpm"]
    if rapor.get("veri_anomali"):
        _para(doc, [("⚠ Veri Anomalisi: ", {"bold": True, "color": KIRMIZI}),
                    ((rapor.get("anomali_notu") or "Okuma hızı olağandışı, kontrol edilmeli.") +
                     " Bu ölçüm güvenilir kabul edilmemeli; okuma yeniden ölçülmelidir.", {"color": KIRMIZI})])
        _para(doc, [(f"Kaydedilen okuma hızı: {wpm} kelime/dakika (kontrol edilmeli — düzey belirlenmedi).", {})])
    else:
        _para(doc, [(f"Öğrencinin okuma hızı dakikada ", {}), (f"{wpm} kelime", {"bold": True}),
                    ("dir. Bu okuma hızı, öğrencinin bulunduğu sınıf düzeyi normlarına göre ", {}),
                    (f"{ctx['hiz_label'].lower()} düzeydedir", {"bold": True}), (".", {})])
        # gauge + norm karşılaştırma (yan yana)
        gt = doc.add_table(rows=1, cols=2); gt.alignment = WD_TABLE_ALIGNMENT.CENTER
        _resim(gt.rows[0].cells[0], _png(grafik.hiz_gauge(wpm, ctx.get("norm_wpm", 0))), 6.5)
        _resim(gt.rows[0].cells[1], _png(grafik.norm_karsilastirma(wpm, ctx.get("norm_wpm", 0))), 7.5)
    if rapor.get("ogretmen_onayli"):
        _para(doc, [("✓ Öğretmen onaylı ölçüm.", {"bold": True, "size": 8})])
        if rapor.get("olcum_farki_notu"):
            _para(doc, [(rapor.get("olcum_farki_notu"), {"italic": True, "size": 8, "color": GRI})])
        if rapor.get("tutarlilik_notu"):
            _para(doc, [("⚠ " + rapor.get("tutarlilik_notu"), {"size": 8, "color": KIRMIZI})])

    # ── 3.1 DOĞRU OKUMA ORANI ──
    _baslik(doc, "3.1. Doğru Okuma Oranı", boyut=10, renk=RGBColor(0x33, 0x33, 0x33))
    _para(doc, [("Doğru Okuma Oranı = Doğru okunan kelime sayısı / Metnin tamamındaki kelime sayısı × 100", {"italic": True, "size": 8, "color": GRI})])
    dogruluk = ctx["dogruluk"]; dog_duzey = ctx["dog_duzey"]; dog_acikla = ctx["dog_acikla"]
    dt = doc.add_table(rows=1, cols=2); dt.alignment = WD_TABLE_ALIGNMENT.CENTER
    _resim(dt.rows[0].cells[0], _png(grafik.dogruluk_donut(dogruluk)), 3.6)
    hc = dt.rows[0].cells[1]
    p = hc.paragraphs[0]
    p.add_run(f"Doğru okuma oranı %{round(dogruluk)} — ").font.size = Pt(9)
    rr = p.add_run(dog_duzey); rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = DUZEY_RENK.get(dog_duzey, LACIVERT)
    p.add_run(f". Toplam {kelime_s} kelimeden {dogru_k} doğru, {yanlis_k} hatalı okunmuştur.").font.size = Pt(9)
    p2 = hc.add_paragraph(); rr = p2.add_run(dog_acikla); rr.font.size = Pt(8.5); rr.font.color.rgb = GRI

    # Hata türleri tablosu
    hatalar = ctx.get("hatalar") or []
    if hatalar:
        th = doc.add_table(rows=len(hatalar) + 2, cols=3)
        _baslik_satiri(th, ["Hata Türü", "Açıklama", "Sayı"])
        toplam = 0
        for i, h in enumerate(hatalar):
            tur = h.get("tur") if isinstance(h, dict) else h
            sayi = int(h.get("sayi", 0)) if isinstance(h, dict) else 0
            toplam += sayi
            cells = th.rows[i + 1].cells
            cells[0].text = hata_turu_ad(tur); cells[1].text = hata_turu_aciklama(tur); cells[2].text = str(sayi)
        son = th.rows[-1].cells
        son[0].text = "Toplam"; son[2].text = str(toplam)
        son[0].paragraphs[0].runs[0].bold = True
        _tablo_stil(th)

    # ── 4. OKUDUĞUNU ANLAMA ──
    anlama = rapor.get("anlama") or {}
    anlama_pct = rapor.get("anlama_yuzde") or 0
    aktif = ctx["aktif_gruplar"]
    _baslik(doc, f"4. Okuduğunu Anlama Becerileri — %{anlama_pct}")
    _hepsi = set(aktif) >= {"4.1", "4.2", "4.3", "4.4", "4.5"}
    for gk, ad, olcutler in ctx["anlama_ham"]:
        if gk not in aktif:
            continue
        _baslik(doc, (f"{gk} {ad}" if _hepsi else ad), boyut=10, renk=RGBColor(0x33, 0x33, 0x33), before=6, after=2)
        ta = doc.add_table(rows=len(olcutler) + 1, cols=4)
        _baslik_satiri(ta, ["Ölçüt", "Zayıf", "Orta", "İyi"])
        for i, (label, key) in enumerate(olcutler):
            val = anlama.get(key, "orta")
            cells = ta.rows[i + 1].cells
            cells[0].text = label
            for ci, sv in enumerate(["zayif", "orta", "iyi"], start=1):
                cells[ci].text = ""
                if val == sv:
                    rr = cells[ci].paragraphs[0].add_run("✔")
                    rr.bold = True; rr.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tablo_stil(ta)
    # radar görselleri
    rt = doc.add_table(rows=1, cols=2); rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    _resim(rt.rows[0].cells[0], ctx.get("radar1_png"), 7.5)
    _resim(rt.rows[0].cells[1], ctx.get("radar2_png"), 7.5)

    # ── 5. PROZODİK OKUMA ──
    proz = rapor.get("prozodik") or {}
    proz_toplam = rapor.get("prozodik_toplam") or 0
    proz_sev = ctx["proz_sev"]
    _baslik(doc, "5. Prozodik Okuma Ölçeği")
    proz_desc = ctx["proz_desc"]; proz_labels = ctx["proz_labels"]
    tp = doc.add_table(rows=len(proz_labels) + 2, cols=6)
    _baslik_satiri(tp, ["Ölçüt", "1 puan", "2 puan", "3 puan", "4 puan", "Puan"])
    for i, key in enumerate(["noktalama", "vurgu", "tonlama", "akicilik", "anlamli_gruplama"]):
        puan = proz.get(key, 0)
        cells = tp.rows[i + 1].cells
        cells[0].text = proz_labels.get(key, key)
        for pi in range(4):
            d = (proz_desc.get(key, ["", "", "", ""])[pi]) or ""
            cells[pi + 1].text = ""
            rr = cells[pi + 1].paragraphs[0].add_run(d)
            if pi + 1 == puan:
                rr.bold = True; rr.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
            cells[pi + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[5].text = str(puan); cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    son = tp.rows[-1].cells
    son[4].text = "Toplam"; son[4].paragraphs[0].runs[0].bold = True
    son[5].text = str(proz_toplam); son[5].paragraphs[0].runs[0].bold = True
    _tablo_stil(tp)
    _para(doc, [("Prozodik okuma performansı: ", {}), (proz_sev, {"bold": True, "color": DUZEY_RENK.get(proz_sev, LACIVERT)}),
                (f" (Toplam {proz_toplam}/20)", {})])
    _resim(doc, _png(grafik.prozodik_bar(proz_toplam)), 11.0)

    # ── GELİŞİM (varsa) ──
    for blok in ctx.get("gelisim_satirlari") or []:
        pass  # gelişim ctx'te önceden hazırlanır (aşağıda)
    if ctx.get("gelisim"):
        g = ctx["gelisim"]
        _baslik(doc, f"{g['no']}. Gelişim (Önceki Rapora Göre)")
        _para(doc, [(f"Önceki ölçüm: {g['onceki_tarih']} tarihli rapor ile karşılaştırma:", {})])
        tg = doc.add_table(rows=len(g["satirlar"]) + 1, cols=4)
        _baslik_satiri(tg, ["Ölçüt", "Önceki", "Şimdi", "Değişim"])
        for i, sat in enumerate(g["satirlar"]):
            cells = tg.rows[i + 1].cells
            for ci, v in enumerate(sat):
                cells[ci].text = str(v)
        _tablo_stil(tg)

    # ── SONUÇ VE GENEL YORUM ──
    _baslik(doc, f"{ctx['sonuc_no']}. Sonuç ve Genel Yorum")
    for cumle in ctx["sonuc_cumleler"]:
        _para(doc, [(cumle, {})])
    m = ctx["metinler"]
    _baslik(doc, m.get("oneriler_baslik", "Eğitsel ve Ev Temelli Gelişim Önerileri"), boyut=11, before=8)
    _para(doc, [(m.get("oneriler_okul_baslik", "Eğitsel Ortamda Yapılacak Çalışmalar"), {"bold": True, "size": 9})], after=2)
    for o in m.get("oneriler_okul", []):
        doc.add_paragraph(o, style="List Bullet")
    _para(doc, [(m.get("oneriler_ev_baslik", "Evde Yapılması Önerilen Çalışmalar"), {"bold": True, "size": 9})], after=2)
    for o in m.get("oneriler_ev", []):
        doc.add_paragraph(o, style="List Bullet")

    # Eğitimci notu
    if rapor.get("ogretmen_notu"):
        _baslik(doc, "Eğitimci Notu", boyut=10, renk=RGBColor(0x33, 0x33, 0x33))
        _para(doc, [(rapor.get("ogretmen_notu"), {})])

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
